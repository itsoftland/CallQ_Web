import logging
import traceback
from rest_framework import viewsets, status

logger = logging.getLogger('actions')
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db import models, transaction
from django.conf import settings
from django.http import JsonResponse
from django.contrib import messages
from django.contrib.auth import logout
from datetime import date, timedelta
from .models import Company, AuthenticationLog, Branch, DealerCustomer, ActivityLog, Country, State, District
from callq_core.services import LicenseManagementService, log_activity
from callq_core.permissions import superadmin_required, company_required, dealer_required, branch_required
from django.utils.decorators import method_decorator
from django.contrib.auth.views import LoginView, PasswordResetView
from .forms import CustomPasswordResetForm
from django.db.models import Q, Sum
from django.utils.timezone import localtime
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

class CustomPasswordResetView(PasswordResetView):
    form_class = CustomPasswordResetForm
    # We don't set template_name here because it's set in urls.py or defaults to registration/password_reset_form.html
    # But we can if we want to customize the form page itself.
    
    def form_valid(self, form):
        email = form.cleaned_data.get('email')
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = list(User.objects.filter(email__iexact=email, is_active=True))
        
        logger.info(f"Password reset requested for email: {email}")
        logger.info(f"Custom View: Found {len(users)} active users with this email.")
        for u in users:
            logger.info(f"User ID: {u.id}, Username: {u.username}, Has Usable Password: {u.has_usable_password()}")

        try:
            response = super().form_valid(form)
            logger.info(f"Password reset email sent (or skipped by Django) for: {email}")
            return response
        except Exception as e:
            logger.error(f"Failed to send password reset email to {email}: {str(e)}")
            logger.error(traceback.format_exc())
            messages.error(self.request, "There was an error sending the password reset email. Please contact the administrator.")
            return self.form_invalid(form)

    def form_invalid(self, form):
        """
        Override to check if the error is due to email not found.
        If so, render the specific 'email not found' template/page.
        """
        email_errors = form.errors.get('email')
        if email_errors and "This email is not linked with any account." in email_errors:
            return render(self.request, 'registration/password_reset_email_not_found.html', {'hide_sidebar': True})
            
        return super().form_invalid(form)

import csv
from django.http import HttpResponse
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
# --- SERIALIZERS (Defining here for brevity, ideally in serializers.py) ---
from rest_framework import serializers

class CompanySerializer(serializers.ModelSerializer):
    class Meta:
        model = Company
        fields = '__all__'
        read_only_fields = ['is_dealer_created', 'created_at', 'updated_at']

# --- API VIEWS ---

class CompanyViewSet(viewsets.ModelViewSet):
    queryset = Company.objects.all()
    serializer_class = CompanySerializer
    # permission_classes = [IsAuthenticated, IsAdminUser] # Ensure strict permissions

    @action(detail=False, methods=['post'], url_path='customer_register')
    def customer_register_external(self, request):
        """Step 1: Register customer with external Licence Portal"""
        # If no PK, we expect data in payload
        data = request.data
        details = {
            "company_name": data.get('company_name'),
            "company_email": data.get('company_email'),
            "contact_person": data.get('contact_person'),
            "contact_number": data.get('contact_number'),
            "address": data.get('address'),
            "city": data.get('city'),
            "state": data.get('state'),
            "zip_code": data.get('zip_code'),
            "gst_number": data.get('gst_number', ''),
            "number_of_licence": data.get('number_of_licence', 1)
        }
        response_data = LicenseManagementService.register_product(details)
        return Response(response_data)

    @action(detail=True, methods=['patch'], url_path='customer_registration/save')
    def save_registration(self, request, pk=None):
        """Step 2: Save CustomerId returned from external API"""
        company = self.get_object()
        company_id = request.data.get('company_id')
        if company_id:
            company.company_id = str(company_id)
            company.save()
        
        serializer = self.get_serializer(company)
        return Response(serializer.data)

    @action(detail=False, methods=['post'], url_path='customer_authentication')
    def customer_authentication_external_standalone(self, request):
        """Step 3: Validate license status and retrieve license details (Standalone)"""
        customer_id = request.data.get('CustomerId')
        if not customer_id:
            return Response({"error": "CustomerId is required"}, status=status.HTTP_400_BAD_REQUEST)
        
        response_data = LicenseManagementService.authenticate_product(customer_id)
        return Response(response_data)

    @action(detail=True, methods=['patch'], url_path='save_product_authentication')
    def save_authentication(self, request, pk=None):
        """Step 4: Persist license authentication data in local database"""
        company = self.get_object()
        data = request.data
        
        # Helper to parse date or return None
        def parse_date(date_str):
            if not date_str: return None
            try:
                # API returns "2025-12-18 00:00:00"
                from datetime import datetime
                return datetime.strptime(date_str.split(' ')[0], '%Y-%m-%d').date()
            except (ValueError, IndexError):
                return None

        # Helper to parse int or return default
        def parse_int(val, default=0):
            try:
                return int(val)
            except (ValueError, TypeError):
                return default

        company.authentication_status = data.get('Authenticationstatus')
        company.product_registration_id = data.get('ProductRegistrationId')
        company.unique_identifier = data.get('UniqueIDentifier')
        company.product_from_date = parse_date(data.get('ProductFromDate'))
        company.product_to_date = parse_date(data.get('ProductToDate'))
        
        # Update device counts
        company.noof_broker_devices = parse_int(data.get('NoofBrokerdevices'), company.noof_broker_devices)
        company.noof_token_dispensors = parse_int(data.get('NoofTokenDispensors'), company.noof_token_dispensors)
        company.noof_keypad_devices = parse_int(data.get('NoofKeypaddevices'), company.noof_keypad_devices)
        company.noof_television_devices = parse_int(data.get('NoofTelevisiondevices'), company.noof_television_devices)
        company.noof_led_devices = parse_int(data.get('NoofLeddevices'), company.noof_led_devices)
        
        # Safely handle license count
        company.number_of_licence = parse_int(data.get('NumberOfLicence'), company.number_of_licence)
        company.save()
        
        try:
            from licenses.models import Batch
            from django.db.models import Sum
            
            # API totals from current payload
            api_tvs = parse_int(data.get('NoofTelevisiondevices'))
            api_dispensers = parse_int(data.get('NoofTokenDispensors'))
            api_keypads = parse_int(data.get('NoofKeypaddevices'))
            api_brokers = parse_int(data.get('NoofBrokerdevices'))
            api_leds = parse_int(data.get('NoofLeddevices'))

            # Sum of all existing batches
            existing_sums = Batch.objects.filter(customer=company).aggregate(
                t=Sum('max_tvs'), d=Sum('max_dispensers'), k=Sum('max_keypads'),
                b=Sum('max_brokers'), l=Sum('max_leds')
            )
            
            # Difference is what goes into the new batch
            diff_t = max(0, api_tvs - (existing_sums['t'] or 0))
            diff_d = max(0, api_dispensers - (existing_sums['d'] or 0))
            diff_k = max(0, api_keypads - (existing_sums['k'] or 0))
            diff_b = max(0, api_brokers - (existing_sums['b'] or 0))
            diff_l = max(0, api_leds - (existing_sums['l'] or 0))

            if any([diff_t, diff_d, diff_k, diff_b, diff_l]):
                 existing_batches_count = Batch.objects.filter(customer=company).count()
                 new_batch_name = f"B{existing_batches_count + 1}"
                 
                 Batch.objects.create(
                     name=new_batch_name,
                     customer=company,
                     max_tvs=diff_t,
                     max_dispensers=diff_d,
                     max_keypads=diff_k,
                     max_brokers=diff_b,
                     max_leds=diff_l
                 )
                 log_activity(request.user, "Batch Created", f"Auto-created Batch {new_batch_name} for {company.company_name} with NEW licenses (Diff logic)")
            else:
                 log_activity(request.user, "Batch Creation Skipped", f"No new licenses detected for {company.company_name}")

        except Exception as e:
            # Log error but don't fail the response
            logger.error(f"Error creating batch for {company.company_name}: {e}")
            log_activity(request.user, "Batch Creation Failed", f"Failed to create batch for {company.company_name}: {str(e)}")

        return Response({
            "message": "Product authentication saved successfully",
            "data": {
                "company_id": company.company_id,
                "company_name": company.company_name,
                "NumberOfLicence": company.number_of_licence,
                "updated_at": company.updated_at
            }
        })

# --- TEMPLATE VIEWS ---

@login_required
def dashboard(request):
    user = request.user
    
    # Imports inside view to avoid circular dependency in top-level
    from configdetails.models import Device, Mapping
    from licenses.models import License
    from .models import ActivityLog
    
    context = {
        'customers_count': 0,
        'dealers_count': 0,
        'devices_count': 0,
        'licenses_count': 0,
        'alerts': [],
        'recent_activity': []
    }
    
    if user.role == "SUPER_ADMIN":
        context['customers_count'] = Company.objects.filter(company_type='CUSTOMER').count()
        context['dealers_count'] = Company.objects.filter(company_type='DEALER').count()
        context['devices_count'] = Device.objects.count()
        context['licenses_count'] = License.objects.count()
        context['recent_activity'] = ActivityLog.objects.all().order_by('-timestamp')[:5]
    elif user.role == "ADMIN":
        state_filter = user.assigned_state
        if state_filter:
            # state_filter is now a list
            companies_qs = Company.objects.filter(state__in=state_filter)
            context['customers_count'] = companies_qs.filter(company_type='CUSTOMER').count()
            context['dealers_count'] = companies_qs.filter(company_type='DEALER').count()
            context['devices_count'] = Device.objects.filter(company__state__in=state_filter).count()
            # License count might need more complex logic if it's not directly on Company but let's follow existing pattern
            context['licenses_count'] = License.objects.filter(batch__customer__state__in=state_filter).count()
            context['recent_activity'] = ActivityLog.objects.filter(
                Q(user__company_relation__state__in=state_filter) | 
                Q(user__branch_relation__state__in=state_filter)
            ).order_by('-timestamp')[:5]
        else:
            # Fallback for admins without state (same as Super Admin or none?)
            # Assuming if no state, they see everything or nothing. Let's stick to everything for now or none?
            # User said "admin ... is assigned with a state so that he can manage ... in the selected state"
            # If no state, maybe they shouldn't see anything.
            context['customers_count'] = 0
            context['dealers_count'] = 0
            context['devices_count'] = 0
            context['licenses_count'] = 0
            context['recent_activity'] = []
        
        # Pending Batch Requests count
        from licenses.models import BatchRequest
        if user.role == "SUPER_ADMIN":
            context['pending_batch_requests'] = BatchRequest.objects.filter(status=BatchRequest.Status.PENDING).count()
        else:
            context['pending_batch_requests'] = BatchRequest.objects.filter(status=BatchRequest.Status.PENDING, requester__state__in=user.assigned_state).count()
        
        # Alerts Logic
        # 1. License Expiry (Simple check for now, requiring detailed License logic)
        # 2. Unmapped Devices
        unmapped_devices = Device.objects.filter(
            mapped_tvs__isnull=True, 
            mapped_sources__isnull=True,
            mapped_keypads__isnull=True,
            mapped_brokers__isnull=True,
            mapped_leds__isnull=True
        ).count()
        if unmapped_devices > 0:
            context['alerts'].append(f"{unmapped_devices} device(s) are currently unmapped.")
            
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            dealer = user.company_relation
            # Count both DealerCustomer and Company children
            dc_count = DealerCustomer.objects.filter(dealer=dealer).count()
            cc_count = Company.objects.filter(parent_company=dealer).count()
            context['customers_count'] = dc_count + cc_count
            
            # Devices of dealer + devices of all child companies
            context['devices_count'] = Device.objects.filter(
                Q(company=dealer) | Q(company__parent_company=dealer)
            ).count()
            
            # Licenses: Show remaining quota vs total? 
            # Or just total licenses assigned to Dealer by Admin.
            from licenses.models import Batch
            total_dealer_licenses = Batch.objects.filter(customer=dealer).aggregate(
                total=Sum('max_tvs') + Sum('max_dispensers') + Sum('max_keypads') + Sum('max_brokers') + Sum('max_leds')
            )['total'] or 0
            context['licenses_count'] = total_dealer_licenses
            
            # Dealer Alerts
            unmapped_count = Device.objects.filter(
                Q(company=dealer) | Q(company__parent_company=dealer),
                mapped_tvs__isnull=True
            ).count()
            if unmapped_count > 0:
                context['alerts'].append(f"{unmapped_count} device(s) unmapped.")

    elif user.role == "COMPANY_ADMIN":
         # Company View
         if user.company_relation:
            context['devices_count'] = Device.objects.filter(company=user.company_relation).count()
            context['licenses_count'] = 0 # Or total licenses for company if applicable
            # Add more stats as needed
            
    elif user.role == "BRANCH_ADMIN":
        # Branch View
        if user.branch_relation:
            branch = user.branch_relation
            context['devices_count'] = Device.objects.filter(branch=branch).count()
            # Branch-specific alerts
            unmapped_count = Device.objects.filter(branch=branch, mapped_tvs__isnull=True).count()
            if unmapped_count > 0:
                context['alerts'].append(f"{unmapped_count} device(s) unmapped in your branch.")         
            
            # Offline devices logic could be added here if 'last_seen' was available
            # For now, just basic count
            
            # Profiles - if profiles are branch specific, but EmbeddedProfile currently links to Company.
            # Maybe show company profiles available to this branch?
            from configdetails.models import EmbeddedProfile
            context['profiles_count'] = EmbeddedProfile.objects.filter(company=branch.company).count()
            
            # Just to populate the UI "Branches" card with 1 (itself)
            context['branches_count'] = 1
        else:
             pass
    
    elif user.role == "PRODUCTION_ADMIN":
        # Production Admin sees all devices for batch management
        # Device counts per type
        context['token_dispenser_count'] = Device.objects.filter(device_type=Device.DeviceType.TOKEN_DISPENSER).count()
        context['keypad_count'] = Device.objects.filter(device_type=Device.DeviceType.KEYPAD).count()
        context['broker_count'] = Device.objects.filter(device_type=Device.DeviceType.BROKER).count()
        context['led_count'] = Device.objects.filter(device_type=Device.DeviceType.LED).count()
        
        # Count production batches
        from configdetails.models import ProductionBatch
        context['batches_count'] = ProductionBatch.objects.count()
        
        # Hide unrelated stats
        total_devices = Device.objects.count()
        active_devices = Device.objects.filter(is_active=True).count()
        context['devices_count'] = total_devices
        context['active_devices'] = active_devices
        context['offline_devices'] = total_devices - active_devices
        context['branches_count'] = 0
        context['profiles_count'] = 0
        
    elif user.role == "EMPLOYEE":
        # Employee sees companies/users in their assigned states
        if user.assigned_state:
            companies_qs = Company.objects.filter(state__in=user.assigned_state)
            context['customers_count'] = companies_qs.filter(company_type='CUSTOMER').count()
            context['devices_count'] = Device.objects.filter(company__state__in=user.assigned_state).count()
        else:
            context['customers_count'] = 0
            context['devices_count'] = 0
            
    return render(request, 'companydetails/dashboard.html', context)

@login_required
def activity_log_list(request):
    user = request.user
    search_query = request.GET.get('search', '').strip()

    if user.role == "SUPER_ADMIN":
        logs = ActivityLog.objects.all().order_by('-timestamp')
    elif user.role == "ADMIN":
        if user.assigned_state:
            logs = ActivityLog.objects.filter(
                Q(user__company_relation__state__in=user.assigned_state) |
                Q(user__branch_relation__state__in=user.assigned_state)
            ).order_by('-timestamp')
        else:
            logs = ActivityLog.objects.none()
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            logs = ActivityLog.objects.filter(
                Q(user=user) | 
                Q(user__company_relation=user.company_relation) |
                Q(user__company_relation__parent_company=user.company_relation)
            ).distinct().order_by('-timestamp')
        else:
            logs = ActivityLog.objects.none()
    elif user.role == "COMPANY_ADMIN":
        logs = ActivityLog.objects.filter(
            Q(user=user) | Q(user__company_relation=user.company_relation)
        ).distinct().order_by('-timestamp')
    else:
        logs = ActivityLog.objects.none()

    if search_query:
        logs = logs.filter(
            Q(action__icontains=search_query) |
            Q(details__icontains=search_query) |
            Q(user__email__icontains=search_query)
        )

    export_format = request.GET.get('export')
    if export_format:
        filename = f"activity_logs_{date.today()}.{export_format}"
        
        if export_format == 'csv':
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            writer = csv.writer(response)
            writer.writerow(['Timestamp', 'User', 'Role', 'Action', 'Details', 'IP Address'])
            for log in logs:
                writer.writerow([
                    localtime(log.timestamp).strftime("%Y-%m-%d %H:%M:%S"),
                    log.user.email if log.user else "System",
                    log.user.role if log.user else "-",
                    log.action,
                    log.details,
                    getattr(log, 'ip_address', '-')
                ])
            return response

        elif export_format == 'pdf':
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
            elements = []
            
            styles = getSampleStyleSheet()
            elements.append(Paragraph("Activity Logs", styles['Title']))
            
            data = [['Timestamp', 'User', 'Role', 'Action', 'Details', 'IP']]
            for log in logs:
                # Truncate details for PDF to avoid messy layout
                log_details = log.details if log.details else ""
                details = log_details[:50] + "..." if len(log_details) > 50 else log_details
                data.append([
                    localtime(log.timestamp).strftime("%Y-%m-%d %H:%M"),
                    log.user.email if log.user else "System",
                    log.user.role if log.user else "-",
                    log.action,
                    details,
                    getattr(log, 'ip_address', '-')
                ])
            
            table = Table(data, colWidths=[100, 150, 80, 100, 250, 80])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
            ]))
            elements.append(table)
            
            doc.build(elements)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        elif export_format == 'docx':
            doc = Document()
            doc.add_heading('Activity Logs', 0)
            
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Timestamp'
            hdr_cells[1].text = 'User'
            hdr_cells[2].text = 'Role'
            hdr_cells[3].text = 'Action'
            hdr_cells[4].text = 'Details'
            hdr_cells[5].text = 'IP Address'
            
            for log in logs:
                row_cells = table.add_row().cells
                row_cells[0].text = localtime(log.timestamp).strftime("%Y-%m-%d %H:%M:%S")
                row_cells[1].text = log.user.email if log.user else "System"
                row_cells[2].text = log.user.role if log.user else "-"
                row_cells[3].text = log.action
                row_cells[4].text = log.details if log.details else "-"
                row_cells[5].text = getattr(log, 'ip_address', '-')
                
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            doc.save(response)
            return response
    
    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(logs, 8)
    try:
        logs_page = paginator.page(page)
    except PageNotAnInteger:
        logs_page = paginator.page(1)
    except EmptyPage:
        logs_page = paginator.page(paginator.num_pages)
         
    return render(request, 'companydetails/activity_log_list.html', {
        'logs': logs_page,
        'search_query': search_query,
    })

# Ensure dealer_required is used where appropriate or other decorators
@login_required
@user_passes_test(dealer_required)
def customer_list(request):
    user = request.user
    companies = []
    users_list = []
    is_dealer_view = False
    search_query = request.GET.get('search', '').strip()
    
    if user.role == 'SUPER_ADMIN':
        companies_qs = Company.objects.filter(is_dealer_created=False)
    elif user.role == 'ADMIN':
        if user.assigned_state:
            companies_qs = Company.objects.filter(is_dealer_created=False, state__in=user.assigned_state)
        else:
            companies_qs = Company.objects.none()
    elif user.role == 'DEALER_ADMIN':
        is_dealer_view = True
        if user.company_relation:
            # Show users created by the dealer (users in child companies + branch users)
            from userdetails.models import User as UserModel
            # Get all child companies (dealer-created)
            child_companies = Company.objects.filter(parent_company=user.company_relation, is_dealer_created=True)
            users_qs = UserModel.objects.filter(
                Q(company_relation=user.company_relation) |  # Users in dealer's company
                Q(company_relation__in=child_companies) |   # Users in child companies
                Q(branch_relation__company__parent_company=user.company_relation)
            ).exclude(id=user.id).exclude(role='DEALER_ADMIN').distinct()
            if search_query:
                users_qs = users_qs.filter(
                    Q(username__icontains=search_query) |
                    Q(email__icontains=search_query)
                )
            users_list = list(users_qs)
        else:
            users_list = []
        companies_qs = Company.objects.none()
    else:
        companies_qs = Company.objects.none()

    # Apply search filter for companies
    if search_query and not is_dealer_view:
        companies_qs = companies_qs.filter(
            Q(company_name__icontains=search_query) |
            Q(company_email__icontains=search_query) |
            Q(contact_person__icontains=search_query) |
            Q(company_id__icontains=search_query)
        )

    companies = list(companies_qs)

    # Pre-calculate display flags for companies (SUPER_ADMIN/ADMIN)
    for company in companies:
        if isinstance(company, Company):
            if company.product_to_date and company.product_to_date < date.today():
                 company.status_class = 'bg-danger'
                 company.status_text = 'Expired'
            elif company.authentication_status in ['Success', 'Approved', 'Approve']:
                company.status_class = 'bg-success'
                company.status_text = 'Approved'
            else:
                company.status_class = 'bg-warning text-dark'
                if company.authentication_status and 'waiting' in company.authentication_status.lower():
                     company.status_text = 'Pending'
                else:
                     company.status_text = company.authentication_status if company.authentication_status else 'Pending'
            
            from configdetails.models import Device
            company.used_licenses = Device.objects.filter(company=company).count()
            company.total_licenses = company.number_of_licence
        
        company.show_auth_button = False
        if company.status_text != 'Approved':
            if user.role in ['SUPER_ADMIN', 'ADMIN'] or company.status_text == 'Dealer-Created' or (hasattr(company, 'is_dealer_created') and company.is_dealer_created):
                company.show_auth_button = True
    
    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(companies, 8)
    try:
        companies_page = paginator.page(page)
    except PageNotAnInteger:
        companies_page = paginator.page(1)
    except EmptyPage:
        companies_page = paginator.page(paginator.num_pages)
    
    # Paginate users_list for dealer view
    users_paginator = Paginator(users_list, 8)
    try:
        users_page = users_paginator.page(page)
    except PageNotAnInteger:
        users_page = users_paginator.page(1)
    except EmptyPage:
        users_page = users_paginator.page(users_paginator.num_pages)
    
    return render(request, 'companydetails/customer_list.html', {
        'companies': companies_page,
        'users_list': users_page,
        'is_dealer_view': is_dealer_view,
        'search_query': search_query,
    })

@login_required
@user_passes_test(superadmin_required)
def dealer_list(request):
    """List all dealers (companies with company_type='DEALER')."""
    user = request.user
    search_query = request.GET.get('search', '').strip()

    if user.role == 'SUPER_ADMIN':
        dealers = Company.objects.filter(company_type='DEALER')
    elif user.role == 'ADMIN':
        if user.assigned_state:
            dealers = Company.objects.filter(company_type='DEALER', state__in=user.assigned_state)
        else:
            dealers = Company.objects.none()
    else:
        dealers = Company.objects.none()

    if search_query:
        dealers = dealers.filter(
            Q(company_name__icontains=search_query) |
            Q(company_email__icontains=search_query) |
            Q(contact_person__icontains=search_query)
        )

    for dealer in dealers:
        if dealer.product_to_date and dealer.product_to_date < date.today():
            dealer.status_class = 'bg-danger'
            dealer.status_text = 'Expired'
        elif dealer.authentication_status in ['Success', 'Approved', 'Approve']:
            dealer.status_class = 'bg-success'
            dealer.status_text = 'Approved'
        else:
            dealer.status_class = 'bg-warning text-dark'
            dealer.status_text = dealer.authentication_status if dealer.authentication_status else 'Pending'
        
        # Count dealer's customers
        dealer.customer_count = DealerCustomer.objects.filter(dealer=dealer).count()
    
    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(list(dealers), 8)
    try:
        dealers_page = paginator.page(page)
    except PageNotAnInteger:
        dealers_page = paginator.page(1)
    except EmptyPage:
        dealers_page = paginator.page(paginator.num_pages)
    
    return render(request, 'companydetails/dealer_list.html', {
        'dealers': dealers_page,
        'search_query': search_query,
    })

@login_required
@user_passes_test(superadmin_required)
def customer_delete(request, pk):
    """Delete a customer (Company). Only allowed for SUPER_ADMIN and ADMIN."""
    company = get_object_or_404(Company, pk=pk)
    
    # Permission check
    user = request.user
    if user.role == 'ADMIN' and user.assigned_state:
        if company.state not in user.assigned_state:
            messages.error(request, 'You do not have permission to delete this customer.')
            return redirect('customer_list')
    
    if request.method == 'POST':
        company_name = company.company_name
        
        # Delete related records first
        from configdetails.models import Device
        from licenses.models import Batch, License
        
        # Delete devices associated with this company
        Device.objects.filter(company=company).delete()
        
        # Delete batches and licenses
        Batch.objects.filter(customer=company).delete()
        
        # Delete branches
        Branch.objects.filter(company=company).delete()
        
        # Delete the company
        company.delete()
        
        log_activity(request.user, "Customer Deleted", f"Deleted customer {company_name}")
        messages.success(request, f'Customer "{company_name}" has been deleted successfully.')
        return redirect('customer_list')
    
    return redirect('customer_list')

@login_required
def customer_toggle_ads(request, pk):
    """Toggle ads_enabled for a customer. Only allowed for SUPER_ADMIN."""
    if request.user.role != 'SUPER_ADMIN':
        messages.error(request, 'You do not have permission to perform this action.')
        return redirect('customer_list')
    
    company = get_object_or_404(Company, pk=pk)
    
    if request.method == 'POST':
        company.ads_enabled = not company.ads_enabled
        company.save()
        
        status_text = "enabled" if company.ads_enabled else "disabled"
        log_activity(request.user, "Ads Feature Toggled", f"Ads feature {status_text} for customer {company.company_name}")
        messages.success(request, f'Ads feature has been {status_text} for "{company.company_name}".')
    
    return redirect('customer_list')

class CustomLoginView(LoginView):
    template_name = 'login.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['hide_sidebar'] = True
        return context
    
    def form_valid(self, form):
        """Override to check dealer license status before allowing login."""
        # Perform the default login
        response = super().form_valid(form)
        
        user = self.request.user
        
        # Check dealer license validation
        if user.role == 'DEALER_ADMIN' and user.company_relation:
            company = user.company_relation
            
            # Check if authentication status is approved
            approved_statuses = ['Success', 'Approved', 'Approve']
            if company.authentication_status not in approved_statuses:
                logout(self.request)
                messages.error(
                    self.request, 
                    'Your dealer license is not approved. Please contact admin.'
                )
                return redirect('login')
            
            # Check if license has expired
            expiry_date = company.product_to_date
            if expiry_date and expiry_date < date.today():
                logout(self.request)
                messages.error(
                    self.request, 
                    f'Your dealer license expired on {expiry_date.strftime("%d-%m-%Y")}. Please contact admin to renew.'
                )
                return redirect('login')
        
        # Check platform access
        if not user.is_web_user:
            logout(self.request)
            messages.error(
                self.request,
                'Your account does not have access to the web dashboard. Please use the mobile application.'
            )
            return redirect('login')
        
        return response

@login_required
@user_passes_test(dealer_required)
def customer_registration(request):
    if request.method == 'POST':
        # Validate email uniqueness before proceeding (check both tables)
        company_email = request.POST.get('company_email')
        company_name = request.POST.get('company_name')
        gst_number = request.POST.get('gst_number')
        contact_person = request.POST.get('contact_person')
        contact_number = request.POST.get('contact_number')
        address = request.POST.get('address')
        city = request.POST.get('city')
        district = request.POST.get('district')
        state = request.POST.get('state')
        zip_code = request.POST.get('zip_code')
        number_of_licence = request.POST.get('number_of_licence')
        if Company.objects.filter(company_email=company_email).exists() or \
           DealerCustomer.objects.filter(company_email=company_email).exists():
            messages.error(request, f'A customer with email "{company_email}" already exists. Please use a different email address.')
            return render(request, 'companydetails/customer_register.html', {
                'form_data': request.POST  # Preserve form data for user convenience
            })
        
        # Only DEALER_ADMIN creates dealer-created customers (stored in DealerCustomer table)
        # SUPER_ADMIN and ADMIN create regular customers (stored in Company table, need external API validation)
        if request.user.role == 'DEALER_ADMIN':
            dealer_company = request.user.company_relation
            if not dealer_company:
                messages.error(request, "Dealer account not correctly configured.")
                return redirect('customer_list')

            try:
                with transaction.atomic():
                    # Generate customer_id
                    existing_count = DealerCustomer.objects.filter(dealer=dealer_company).count()
                    customer_id = f"{dealer_company.company_id or dealer_company.id}-CST{str(existing_count + 1).zfill(4)}"
                    
                    # 1. Create in DealerCustomer table
                    dealer_customer = DealerCustomer.objects.create(
                        dealer=dealer_company,
                        customer_id=customer_id,
                        company_name=company_name,
                        company_email=company_email,
                        gst_number=gst_number,
                        contact_person=contact_person,
                        contact_number=contact_number,
                        address=address,
                        city=city,
                        district=district,
                        state=state,
                        zip_code=zip_code
                    )

                    # 2. Create in Company table
                    branch_cfg = request.POST.get('branch_configuration', 'SINGLE')
                    company_type = Company.CompanyType.CUSTOMER
                    is_multiple = (branch_cfg == 'MULTIPLE')
                    
                    company = Company.objects.create(
                        company_id=customer_id,
                        company_name=company_name,
                        company_email=company_email,
                        company_type=company_type,
                        parent_company=dealer_company,
                        is_dealer_created=True,
                        branch_configuration=branch_cfg,
                        gst_number=gst_number,
                        contact_person=contact_person,
                        contact_number=contact_number,
                        address=address,
                        city=city,
                        district=district,
                        state=state,
                        zip_code=zip_code,
                        number_of_licence=number_of_licence
                    )

                    # Auto-create branch if SINGLE
                    if not is_multiple:
                        Branch.objects.create(
                            company=company,
                            branch_name="Main Branch",
                            address=address,
                            city=city,
                            district=district,
                            state=state,
                            zip_code=zip_code
                        )

                    log_activity(request.user, "Dealer Customer Registered", f"Registered customer {company_name} (ID: {customer_id})")
                    messages.success(request, f"Customer {company_name} registered successfully.")
                    return redirect('customer_list')

            except Exception as e:
                messages.error(request, f"Error creating customer: {str(e)}")
                return render(request, 'companydetails/customer_register.html', {'form_data': request.POST})
        
        else:
            # SUPER_ADMIN or ADMIN - create regular customer in Company table
            branch_cfg = request.POST.get('branch_configuration', 'SINGLE')
            company_type = request.POST.get('company_type', 'CUSTOMER')
            
            try:
                ads_enabled = request.POST.get('ads_enabled') == '1'
                company = Company.objects.create(
                    company_name=request.POST.get('company_name'),
                    company_type=company_type,
                    company_email=company_email,
                    contact_person=request.POST.get('contact_person'),
                    contact_number=request.POST.get('contact_number'),
                    address=request.POST.get('address'),
                    city=request.POST.get('city'),
                    state=request.POST.get('state'),
                    zip_code=request.POST.get('zip_code'),
                    district=request.POST.get('district'),
                    country=request.POST.get('country', 'India'),
                    gst_number=request.POST.get('gst_number'),
                    number_of_licence=request.POST.get('number_of_licence'),
                    is_dealer_created=False,
                    branch_configuration=branch_cfg,
                    ads_enabled=ads_enabled
                )
                # Auto-create Default Branch if 'SINGLE' configuration is selected
                if branch_cfg == 'SINGLE':
                    Branch.objects.create(
                        company=company,
                        branch_name=company.city if company.city else company.company_name, # Use city as branch name
                        address=company.address,
                        city=company.city,
                        district=company.district,
                        state=company.state,
                        zip_code=company.zip_code
                    )

                log_activity(request.user, f"{company.get_company_type_display()} Created", f"Created {company.company_name} (Requires API validation)")
                messages.success(request, f'{company.get_company_type_display()} "{company.company_name}" created successfully.')
                return redirect('customer_list')
            except Exception as e:
                messages.error(request, f'Error creating customer: {str(e)}')
                return render(request, 'companydetails/customer_register.html', {
                    'form_data': request.POST
                })
    return render(request, 'companydetails/customer_register.html')

@login_required
@user_passes_test(company_required)
def branch_list(request):
    user = request.user
    search_query = request.GET.get('search', '').strip()
    is_dealer_view = False

    if user.role == 'DEALER_ADMIN':
        is_dealer_view = True
        if user.company_relation:
            dealer = user.company_relation
            branches = Branch.objects.filter(
                Q(company=dealer) |
                Q(company__parent_company=dealer)
            )
        else:
            branches = Branch.objects.none()
    else:
        company = request.user.company_relation
        branches = Branch.objects.filter(company=company)

    if search_query:
        branches = branches.filter(
            Q(branch_name__icontains=search_query) |
            Q(city__icontains=search_query) |
            Q(address__icontains=search_query)
        )

    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(list(branches), 8)
    try:
        branches_page = paginator.page(page)
    except PageNotAnInteger:
        branches_page = paginator.page(1)
    except EmptyPage:
        branches_page = paginator.page(paginator.num_pages)

    return render(request, 'companydetails/branch_list.html', {
        'branches': branches_page,
        'is_dealer_view': is_dealer_view,
        'search_query': search_query,
    })

@login_required
@user_passes_test(company_required)
def branch_create(request):
    if request.method == 'POST':
        Branch.objects.create(
            company=request.user.company_relation,
            branch_name=request.POST.get('branch_name'),
            address=request.POST.get('address'),
            city=request.POST.get('city'),
            district=request.POST.get('district'),
            state=request.POST.get('state'),
            zip_code=request.POST.get('zip_code'),
            country=request.POST.get('country', 'India')
        )
        return redirect('branch_list')
    return render(request, 'companydetails/branch_form.html')

@login_required
@user_passes_test(company_required)
def branch_edit(request, pk):
    user = request.user
    if user.role == 'DEALER_ADMIN':
        if user.company_relation:
            dealer = user.company_relation
            # Allow access if branch belongs to dealer or any child company
            branch = get_object_or_404(Branch, pk=pk)
            # Verify permission manually
            is_owned = branch.company == dealer
            is_child = branch.company.parent_company == dealer if branch.company else False
            
            if not (is_owned or is_child):
                return redirect('branch_list') # Or 403
        else:
             return redirect('branch_list')
    else:
        branch = get_object_or_404(Branch, pk=pk, company=user.company_relation)
    if request.method == 'POST':
        branch.branch_name = request.POST.get('branch_name')
        branch.address = request.POST.get('address')
        branch.city = request.POST.get('city')
        branch.district = request.POST.get('district')
        branch.state = request.POST.get('state')
        branch.zip_code = request.POST.get('zip_code')
        branch.country = request.POST.get('country', 'India')
        branch.save()
        return redirect('branch_list')
    return render(request, 'companydetails/branch_form.html', {'branch': branch})

@login_required
@user_passes_test(dealer_required)
def validate_license_view(request, pk):
    if request.user.role not in ['SUPER_ADMIN', 'ADMIN']:
        messages.error(request, 'Permission denied.')
        return redirect('customer_list')
        
    company = get_object_or_404(Company, pk=pk)
    company.authentication_status = 'Approved'
    if not company.product_to_date:
        company.product_from_date = date.today()
        company.product_to_date = date.today() + timedelta(days=365)
    
    company.save()
    
    from callq_core.services import log_activity
    log_activity(request.user, "License Validated", f"Manually approved license for {company.company_name}")
    
    messages.success(request, f'License for {company.company_name} has been validated successfully.')
    return redirect('customer_list')

def sync_company_license_data(company, data, user=None):
    """
    Helper to synchronize external API data with local Company model.
    Matches logic in CompanyViewSet.save_authentication.
    """
    if not data: return

    def parse_date(date_str):
        if not date_str: return None
        try:
            from datetime import datetime
            return datetime.strptime(date_str.split(' ')[0], '%Y-%m-%d').date()
        except (ValueError, IndexError):
            return None

    def parse_int(val, default=0):
        try: return int(val)
        except (ValueError, TypeError): return default

    company.authentication_status = data.get('Authenticationstatus')
    company.product_registration_id = data.get('ProductRegistrationId')
    company.unique_identifier = data.get('UniqueIDentifier')
    company.product_from_date = parse_date(data.get('ProductFromDate'))
    company.product_to_date = parse_date(data.get('ProductToDate'))
    
    company.noof_broker_devices = parse_int(data.get('NoofBrokerdevices'), company.noof_broker_devices)
    company.noof_token_dispensors = parse_int(data.get('NoofTokenDispensors'), company.noof_token_dispensors)
    company.noof_keypad_devices = parse_int(data.get('NoofKeypaddevices'), company.noof_keypad_devices)
    company.noof_television_devices = parse_int(data.get('NoofTelevisiondevices'), company.noof_television_devices)
    company.noof_led_devices = parse_int(data.get('NoofLeddevices'), company.noof_led_devices)
    
    company.number_of_licence = parse_int(data.get('NumberOfLicence'), company.number_of_licence)
    company.save()

    try:
        from licenses.models import Batch
        from django.db.models import Sum

        # API totals from data
        api_tvs = parse_int(data.get('NoofTelevisiondevices'))
        api_dispensers = parse_int(data.get('NoofTokenDispensors'))
        api_keypads = parse_int(data.get('NoofKeypaddevices'))
        api_brokers = parse_int(data.get('NoofBrokerdevices'))
        api_leds = parse_int(data.get('NoofLeddevices'))

        # Sum of all existing batches
        existing_sums = Batch.objects.filter(customer=company).aggregate(
            t=Sum('max_tvs'), d=Sum('max_dispensers'), k=Sum('max_keypads'),
            b=Sum('max_brokers'), l=Sum('max_leds')
        )
        
        # Difference is what goes into the new batch
        diff_t = max(0, api_tvs - (existing_sums['t'] or 0))
        diff_d = max(0, api_dispensers - (existing_sums['d'] or 0))
        diff_k = max(0, api_keypads - (existing_sums['k'] or 0))
        diff_b = max(0, api_brokers - (existing_sums['b'] or 0))
        diff_l = max(0, api_leds - (existing_sums['l'] or 0))

        if any([diff_t, diff_d, diff_k, diff_b, diff_l]) and company.authentication_status in ['Success', 'Approved', 'Approve']:
             existing_batches_count = Batch.objects.filter(customer=company).count()
             new_batch_name = f"B{existing_batches_count + 1}"
             Batch.objects.create(
                 name=new_batch_name,
                 customer=company,
                 max_tvs=diff_t,
                 max_dispensers=diff_d,
                 max_keypads=diff_k,
                 max_brokers=diff_b,
                 max_leds=diff_l
             )
             if user: log_activity(user, "Batch Created", f"Auto-created Batch {new_batch_name} for {company.company_name} with NEW licenses (Diff logic)")
    except Exception as e:
        logger.error(f"Error creating batch in sync for {company.company_name}: {e}")

from django.http import JsonResponse

@login_required
def ajax_authenticate_customer(request, pk):
    logger.info(f"DEBUG: ajax_authenticate_customer hit with pk={pk}")
    try:
        company = Company.objects.get(pk=pk)
        
        # Auto-Registration if ID is missing
        if not company.company_id:
            logger.info(f"Auto-registering {company.company_name}...")
            details = {
                "company_name": company.company_name,
                "company_email": company.company_email,
                "contact_person": company.contact_person,
                "contact_number": company.contact_number,
                "address": company.address,
                "city": company.city,
                "state": company.state,
                "zip_code": company.zip_code,
                "gst_number": company.gst_number,
                "number_of_licence": company.number_of_licence
            }
            
            reg_response = LicenseManagementService.register_product(details)
            logger.info(f"Registration Response for {company.company_name}: {reg_response}")
            
            if reg_response and not reg_response.get('error') and reg_response.get('CustomerId'):
                company.company_id = str(reg_response.get('CustomerId'))
                company.save()
                logger.info(f"Saved new Company ID for {company.company_name}: {company.company_id}")
            else:
                return JsonResponse({'error': f"Registration failed: {reg_response.get('error', 'Unknown error')}"}, status=400)
        
        # Now proceed with Authentication
        if not company.company_id:
             return JsonResponse({'error': 'No Customer ID even after registration attempt.'}, status=400)

        response_data = LicenseManagementService.authenticate_product(company.company_id)
        if response_data and not response_data.get('error'):
            sync_company_license_data(company, response_data, request.user)
            # Refresh company from DB to get updated values
            company.refresh_from_db()
            
        if not response_data:
            return JsonResponse({'error': 'Failed to authenticate with license server.'}, status=503)

        # Calculate used licenses
        from configdetails.models import Device
        used_licenses = Device.objects.filter(company=company).count()
        
        # Determine status text from authentication status
        auth_status = response_data.get('Authenticationstatus', '')
        if auth_status in ['Success', 'Approved', 'Approve']:
            status_text = 'Approved'
        elif 'waiting' in auth_status.lower() or 'pending' in auth_status.lower():
            status_text = 'Pending'
        elif 'expired' in auth_status.lower():
            status_text = 'Expired'
        else:
            status_text = auth_status if auth_status else 'Pending'
        
        # Enhance response with additional data for frontend
        enhanced_response = response_data.copy()
        enhanced_response['used_licenses'] = used_licenses
        enhanced_response['total_licenses'] = company.number_of_licence
        enhanced_response['status_text'] = status_text
        enhanced_response['authentication_status'] = auth_status

        return JsonResponse(enhanced_response)

    except Company.DoesNotExist:
        logger.warning(f"Customer not found for pk={pk}")
        return JsonResponse({'error': 'Customer not found'}, status=404)
    except Exception as e:
        logger.error(f"Error in ajax_authenticate_customer: {e}")
        logger.error(traceback.format_exc())
        return JsonResponse({'error': f"Server Error: {str(e)}"}, status=500)


@login_required
def ajax_get_company_branches(request, company_id):
    """
    Return list of branches for a specific company.
    Used for initializing dropdowns in user creation.
    """
    branches = Branch.objects.filter(company_id=company_id).values('id', 'branch_name')
    return JsonResponse({'branches': list(branches)})


@login_required
def ajax_check_customer_status(request, pk):
    logger.info(f"DEBUG: ajax_check_customer_status hit with pk={pk}")
    try:
        company = Company.objects.get(pk=pk)
        if not company.company_id:
            return JsonResponse({
                'authentication_status': 'No ID',
                'total_licenses': company.number_of_licence
            })
        
        # Optional: We could also call the API here to ensure latest status
        # But for efficiency, maybe just return DB status if authenticate view is called separately.
        # However, user said "call the api ... every 3sec". 
        # So this check status view SHOULD call the API.
        response_data = LicenseManagementService.authenticate_product(company.company_id)
        if response_data and not response_data.get('error'):
            sync_company_license_data(company, response_data, request.user)

        # Calculate used licenses
        from configdetails.models import Device
        used_licenses = Device.objects.filter(company=company).count()

        return JsonResponse({
            'status_text': company.authentication_status, # Current DB status
            'authentication_status': response_data.get('Authenticationstatus') if response_data else 'Unknown', # Latest API status
            'used_licenses': used_licenses,
            'total_licenses': company.number_of_licence
        })

    except Company.DoesNotExist:
        return JsonResponse({'error': 'Customer not found'}, status=404)
    except Exception as e:
        logger.error(f"Error in ajax_check_customer_status: {e}")
        logger.error(traceback.format_exc())
        return JsonResponse({'error': str(e), 'status_text': 'Error', 'authentication_status': 'Error'})

@login_required
def device_report(request):
    """
    Device Report page.
    - Admin: Shows device counts per direct customer and dealer (NOT dealer's customers).
    - Dealer Admin: Shows device counts per dealer's customers.
    """
    from configdetails.models import Device
    from django.db.models import Count
    
    user = request.user
    report_data = []
    
    if user.role in ['SUPER_ADMIN', 'ADMIN']:
        # Admin sees direct customers and dealers (not dealer's child customers)
        if user.role == 'ADMIN' and user.assigned_state:
            companies = Company.objects.filter(
                state__in=user.assigned_state,
                parent_company__isnull=True  # Exclude dealer's child companies
            )
        else:
            companies = Company.objects.filter(parent_company__isnull=True)
        
        # Annotate with device counts
        for company in companies:
            device_count = Device.objects.filter(company=company).count()
            report_data.append({
                'name': company.company_name,
                'type': company.get_company_type_display(),
                'device_count': device_count,
                'is_dealer': company.company_type == 'DEALER'
            })
    
    elif user.role == 'DEALER_ADMIN':
        # Dealer sees their customers' device counts
        if user.company_relation:
            dealer = user.company_relation
            
            # Get dealer's child companies
            child_companies = Company.objects.filter(parent_company=dealer)
            for company in child_companies:
                device_count = Device.objects.filter(company=company).count()
                report_data.append({
                    'name': company.company_name,
                    'type': 'Customer',
                    'device_count': device_count,
                    'is_dealer': False
                })
            
            # Also include DealerCustomers if they have devices
            dealer_customers = DealerCustomer.objects.filter(dealer=dealer)
            for dc in dealer_customers:
                # DealerCustomers might not have devices directly, but check anyway
                device_count = Device.objects.filter(dealer_customer=dc).count() if hasattr(Device, 'dealer_customer') else 0
                report_data.append({
                    'name': dc.company_name,
                    'type': 'Dealer Customer',
                    'device_count': device_count,
                    'is_dealer': False
                })
    else:
        report_data = []
    
    # Calculate totals
    total_devices = sum(item['device_count'] for item in report_data)
    total_customers = len(report_data)
    
    return render(request, 'companydetails/device_report.html', {
        'report_data': report_data,
        'total_devices': total_devices,
        'total_customers': total_customers
    })

@login_required
def customer_export(request, company_id, file_format):
    """Export Customer Devices and Summary"""
    from django.http import HttpResponse, HttpResponseForbidden, FileResponse
    from configdetails.models import Device
    import csv
    import io

    user = request.user
    if user.role not in ['SUPER_ADMIN', 'ADMIN', 'DEALER_ADMIN']:
        return HttpResponseForbidden("Permission Denied")

    company = get_object_or_404(Company, id=company_id)
    
    # Permission Check
    if user.role == 'DEALER_ADMIN':
        # specific check for dealer's customers
        if user.company_relation and (company.parent_company != user.company_relation and getattr(company, 'dealer', None) != user.company_relation):
             # Also allow if company is the dealer itself? Maybe.
             if company != user.company_relation:
                  return HttpResponseForbidden("Permission Denied")

    devices = Device.objects.filter(company=company)
    
    # Calculate Summary
    from django.db.models import Count
    summary = devices.values('device_type').annotate(count=Count('id'))
    summary_dict = {s['device_type']: s['count'] for s in summary}
    total_devices = devices.count()

    if file_format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{company.company_name}_Devices.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Serial Number', 'Device Type', 'MAC Address', 'Model', 'Status', 'Allocated To'])
        for d in devices:
             writer.writerow([
                 d.serial_number, 
                 d.get_device_type_display(), 
                 d.mac_address or '', 
                 d.device_model, 
                 'Active' if d.is_active else 'Inactive',
                 d.dealer_customer.company_name if d.dealer_customer else company.company_name
             ])
        
        writer.writerow([])
        writer.writerow(['--- Summary ---'])
        writer.writerow(['Device Type', 'Count'])
        for dtype, count in summary_dict.items():
            dt_label = dict(Device.DeviceType.choices).get(dtype, dtype)
            writer.writerow([dt_label, count])
        writer.writerow(['Total', total_devices])
             
        return response

    elif file_format == 'pdf':
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, 750, f"Device Report: {company.company_name}")
        p.setFont("Helvetica", 10)
        p.drawString(50, 735, f"Total Devices: {total_devices}")
        
        y = 700
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, y, "Serial Number | Type | MAC | Status")
        y -= 20
        p.setFont("Helvetica", 9)
        
        for d in devices:
            status = 'Active' if d.is_active else 'Inactive'
            text = f"{d.serial_number} | {d.get_device_type_display()} | {d.mac_address or '-'} | {status}"
            p.drawString(50, y, text)
            y -= 15
            if y < 150: # Leave space for summary
                p.showPage()
                y = 750
        
        # Summary Section at bottom (or new page if needed)
        if y < 200:
            p.showPage()
            y = 750
            
        y -= 30
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, "Summary by Type")
        y -= 20
        p.setFont("Helvetica", 10)
        
        for dtype, count in summary_dict.items():
            dt_label = dict(Device.DeviceType.choices).get(dtype, dtype)
            p.drawString(50, y, f"{dt_label}: {count}")
            y -= 15
            
        p.showPage()
        p.save()
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename=f"{company.company_name}_Devices.pdf")

    elif file_format == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(f'Device Report: {company.company_name}', 0)
        document.add_paragraph(f"Total Devices: {total_devices}")
        
        # Device Table
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Serial Number'
        hdr_cells[1].text = 'Device Type'
        hdr_cells[2].text = 'MAC'
        hdr_cells[3].text = 'Status'
        
        for d in devices:
            row_cells = table.add_row().cells
            row_cells[0].text = d.serial_number
            row_cells[1].text = d.get_device_type_display()
            row_cells[2].text = d.mac_address or '-'
            row_cells[3].text = 'Active' if d.is_active else 'Inactive'
            
        document.add_heading('Summary', level=1)
        sum_table = document.add_table(rows=1, cols=2)
        sum_table.style = 'Table Grid'
        sum_cells = sum_table.rows[0].cells
        sum_cells[0].text = 'Device Type'
        sum_cells[1].text = 'Count'
        
        for dtype, count in summary_dict.items():
            dt_label = dict(Device.DeviceType.choices).get(dtype, dtype)
            row = sum_table.add_row().cells
            row[0].text = dt_label
            row[1].text = str(count)
            
        # Total Row
        row = sum_table.add_row().cells
        row[0].text = 'Total'
        row[1].text = str(total_devices)
            
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{company.company_name}_Devices.docx"'
        return response

    return HttpResponse("Invalid format")

def get_states(request):
    country_name = request.GET.get('country', 'India')
    states = State.objects.filter(country__name=country_name).values('id', 'name').order_by('name')
    return JsonResponse({'states': list(states)})

def get_districts(request):
    state_id = request.GET.get('state_id')
    state_name = request.GET.get('state_name')
    
    if state_id:
        districts = District.objects.filter(state_id=state_id).values('id', 'name').order_by('name')
    elif state_name:
        districts = District.objects.filter(state__name=state_name).values('id', 'name').order_by('name')
    else:
        districts = []
    
    return JsonResponse({'districts': list(districts)})

@login_required
@user_passes_test(superadmin_required)
def location_management(request):
    user = request.user
    is_admin = user.role == 'ADMIN'
    
    countries = Country.objects.all().order_by('name')
    
    # For Admin, filter states to only their assigned state(s)
    if is_admin and user.assigned_state:
        states = State.objects.filter(name__in=user.assigned_state).select_related('country').order_by('country__name', 'name')
        # Auto-select the admin's state for the district form if only one state
        admin_state = states.first() if states.count() == 1 else None
    else:
        states = State.objects.all().select_related('country').order_by('country__name', 'name')
        admin_state = None
    
    # Filter districts if state is selected via GET, otherwise show none to avoid heavy load
    selected_state_id = request.GET.get('state_id')
    # For Admin with single state, auto-select their state's districts
    if not selected_state_id and admin_state:
        selected_state_id = str(admin_state.id)
    districts = District.objects.filter(state_id=selected_state_id).order_by('name') if selected_state_id else []

    if request.method == 'POST':
        action = request.POST.get('action')
        
        # Admin users cannot add states
        if action == 'add_state':
            if is_admin:
                messages.error(request, 'Admins cannot add new states.')
                return redirect('location_management')
            name = request.POST.get('name')
            country_id = request.POST.get('country_id')
            if name and country_id:
                try:
                    State.objects.create(name=name, country_id=country_id)
                    messages.success(request, f'State "{name}" added successfully.')
                except Exception as e:
                    messages.error(request, f'Error adding state: {str(e)}')
            return redirect('location_management')

        elif action == 'add_district':
            name = request.POST.get('name')
            state_id = request.POST.get('state_id')
            
            # For Admin, verify they can only add to their assigned state
            if is_admin and user.assigned_state:
                try:
                    state = State.objects.get(id=state_id)
                    if state.name not in user.assigned_state:
                        messages.error(request, 'You can only add districts to your assigned state.')
                        return redirect('location_management')
                except State.DoesNotExist:
                    messages.error(request, 'Invalid state selected.')
                    return redirect('location_management')
            
            if name and state_id:
                try:
                    District.objects.create(name=name, state_id=state_id)
                    messages.success(request, f'District "{name}" added successfully.')
                except Exception as e:
                    messages.error(request, f'Error adding district: {str(e)}')
            return redirect(f"{reverse('location_management')}?state_id={state_id}")

    return render(request, 'companydetails/location_management.html', {
        'countries': countries,
        'states': states,
        'districts': districts,
        'selected_state_id': selected_state_id,
        'is_admin': is_admin,
        'admin_state': admin_state
    })
