from rest_framework import viewsets
from rest_framework.decorators import action
from rest_framework.response import Response
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.http import HttpResponse, HttpResponseForbidden, FileResponse, JsonResponse
from django.contrib import messages
from .models import Device, DeviceConfig, TVConfig, Mapping, Counter, TVCounter, TVAd, LedConfig, EmbeddedProfile, ProductionBatch, ProductionSerialNumber, ButtonMapping, CounterConfig, TVCounterMapping, CounterTokenDispenserMapping, ExternalDeviceCounterLog, GroupMapping, GroupDispenserMapping, GroupCounterButtonMapping, TVDispenserMapping, TVKeypadMapping, get_button_index_char, BUTTON_INDEX_SEQUENCE
from django.db import transaction
from .serializers.config_serializers import DeviceSerializer, DeviceConfigSerializer, TVConfigSerializer, MappingSerializer, EmbeddedProfileSerializer, CounterConfigSerializer, TVCounterMappingSerializer, CounterTokenDispenserMappingSerializer, ExternalDeviceCounterLogSerializer, TVDispenserMappingSerializer, TVKeypadMappingSerializer
from companydetails.models import Company, Branch
from callq_core.services import LicenseManagementService, log_activity
from callq_core.permissions import company_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.conf import settings
import json
from collections import deque
import csv
import io
import os
from pathlib import Path
from datetime import datetime, timedelta

# --- API VIEWS ---
class DeviceViewSet(viewsets.ModelViewSet):
    queryset = Device.objects.all()
    serializer_class = DeviceSerializer

class MappingViewSet(viewsets.ModelViewSet):
    queryset = Mapping.objects.all()
    serializer_class = MappingSerializer

class EmbeddedProfileViewSet(viewsets.ModelViewSet):
    queryset = EmbeddedProfile.objects.all()
    serializer_class = EmbeddedProfileSerializer
    # permission_classes = [IsAuthenticated] # Inherits default or set explicitly

    @action(detail=True, methods=['post'])
    def copy(self, request, pk=None):
        profile = self.get_object()
        new_profile = EmbeddedProfile.objects.create(
            name=f"{profile.name} (Copy)",
            device_type=profile.device_type,
            config_json=profile.config_json,
            is_api_sourced=False
        )
        serializer = self.get_serializer(new_profile)
        return Response(serializer.data)


# Counter-Wise Configuration ViewSets
class CounterConfigViewSet(viewsets.ModelViewSet):
    """
    ViewSet for Counter CRUD operations.
    Counters are scoped to the requesting user's company.
    """
    # Required by DRF router for basename auto-detection; actual filtering is in get_queryset()
    queryset = CounterConfig.objects.none()
    serializer_class = CounterConfigSerializer
    permission_classes = [company_required]

    def _get_user_company(self):
        """Return the Company for the currently logged-in user, or None."""
        user = self.request.user
        if hasattr(user, 'company_relation') and user.company_relation:
            return user.company_relation
        return None

    def get_queryset(self):
        """Return counters belonging to the current user's company only."""
        company = self._get_user_company()
        qs = CounterConfig.objects.all()
        if company:
            qs = qs.filter(company=company)
        return qs.filter(status=True).order_by('counter_name')

    def perform_create(self, serializer):
        """Automatically assign the user's company when creating a counter."""
        company = self._get_user_company()
        serializer.save(company=company)

    def perform_destroy(self, instance):
        """Prevent deletion if counter is mapped to TV or dispenser"""
        if TVCounterMapping.objects.filter(counter=instance).exists():
            raise serializers.ValidationError("Cannot delete counter: It is mapped to one or more TV devices.")
        if CounterTokenDispenserMapping.objects.filter(counter=instance).exists():
            raise serializers.ValidationError("Cannot delete counter: It is mapped to one or more token dispensers.")
        instance.delete()


class TVCounterMappingViewSet(viewsets.ModelViewSet):
    """
    ViewSet for TV-Counter mapping operations.
    """
    queryset = TVCounterMapping.objects.all()
    serializer_class = TVCounterMappingSerializer
    permission_classes = [company_required]

    def get_queryset(self):
        """Filter by TV device if provided"""
        queryset = TVCounterMapping.objects.all()
        tv_id = self.request.query_params.get('tv_id', None)
        if tv_id:
            queryset = queryset.filter(tv_id=tv_id)
        return queryset.select_related('tv', 'counter')


class CounterTokenDispenserMappingViewSet(viewsets.ModelViewSet):
    """
    ViewSet for Counter-Token Dispenser mapping operations.
    """
    queryset = CounterTokenDispenserMapping.objects.all()
    serializer_class = CounterTokenDispenserMappingSerializer
    permission_classes = [company_required]

    def get_queryset(self):
        """Filter by counter or dispenser if provided"""
        queryset = CounterTokenDispenserMapping.objects.all()
        counter_id = self.request.query_params.get('counter_id', None)
        dispenser_id = self.request.query_params.get('dispenser_id', None)
        if counter_id:
            queryset = queryset.filter(counter_id=counter_id)
        if dispenser_id:
            queryset = queryset.filter(dispenser_id=dispenser_id)
        return queryset.select_related('counter', 'dispenser')



from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from django.db.models import Q
from django.db import transaction
import logging
import json

# Setup loggers for API logging
request_logger = logging.getLogger('request_hits')
action_logger = logging.getLogger('actions')

# --- HELPER ---
def log_api_request(api_name, request, payload=None):
    """Helper function to log API requests"""
    try:
        payload_data = payload if payload else request.data if hasattr(request, 'data') else {}
        request_logger.info(f"API REQUEST: {api_name} | Method: {request.method} | Path: {request.path} | Payload: {json.dumps(payload_data)}")
    except Exception as e:
        request_logger.warning(f"API REQUEST: {api_name} | Method: {request.method} | Path: {request.path} | Error logging payload: {str(e)}")

def log_api_response(api_name, status_code, response_data=None, error=None):
    """Helper function to log API responses"""
    try:
        if error:
            request_logger.error(f"API RESPONSE: {api_name} | Status: {status_code} | Error: {error}")
            action_logger.error(f"{api_name} Failed | Status: {status_code} | Error: {error}")
        else:
            # Log only summary for success responses to avoid huge logs
            response_summary = str(response_data)[:200] if response_data else "Success"
            request_logger.info(f"API RESPONSE: {api_name} | Status: {status_code} | Response: {response_summary}")
            action_logger.info(f"{api_name} Success | Status: {status_code}")
    except Exception as e:
        request_logger.warning(f"API RESPONSE: {api_name} | Status: {status_code} | Error logging response: {str(e)}")

# --- HELPER ---
def generate_tv_counters_data(device, tv_config=None):
    if device.device_type != Device.DeviceType.TV:
        return []
    
    # Primary configuration count
    configured_count = getattr(tv_config, 'no_of_counters', 1)
    try:
        configured_count = min(int(configured_count), 8)
    except:
        configured_count = 1

    counters = []

    # 1. Fetch ONLY counters from dispensers explicitly mapped to this TV via TVKeypadMapping.
    #    Chain: TV → TVKeypadMapping.dispenser → GroupCounterButtonMapping → counter
    #    This guarantees we never show counters from unrelated groups or dispensers.
    mapped_counters = {}

    # Collect all dispensers wired to this TV through keypad slots (skip slots with no dispenser)
    tv_keypad_mappings = TVKeypadMapping.objects.filter(tv=device).select_related('dispenser')
    tv_dispenser_ids = [
        kpm.dispenser_id
        for kpm in tv_keypad_mappings
        if kpm.dispenser_id is not None
    ]

    if tv_dispenser_ids:
        from configdetails.models import GroupCounterButtonMapping
        gcbms = (
            GroupCounterButtonMapping.objects
            .filter(dispenser_id__in=tv_dispenser_ids)
            .select_related('dispenser', 'counter')
            .order_by('button_index')
        )
        for gcbm in gcbms:
            try:
                # Convert ASCII button_index (e.g., '1' = 0x31) to 1-based integer
                idx = ord(gcbm.button_index) - 0x31 + 1
                if 1 <= idx <= 8:
                    # First mapping wins; later ones (same slot) are ignored
                    if idx not in mapped_counters:
                        mapped_counters[idx] = gcbm
            except Exception:
                pass

    # 2. Iterate and Fill Slots (Always 8 slots for real-time frontend toggling)
    for i in range(1, 9):
        counter_data = None
        
        if i in mapped_counters:
            gcbm = mapped_counters[i]
            counter_data = {
                'counter_id': f"{gcbm.dispenser.serial_number}-C{i}",
                'default_name': gcbm.counter.counter_name,
                'default_code': gcbm.counter.counter_prefix_code or f"C{i:02d}",
                'source_device': gcbm.dispenser.serial_number,
                'button_index': get_button_index_char(i),
                'name': gcbm.counter.counter_name,
                'code': gcbm.counter.counter_prefix_code or f"C{i:02d}",
                'is_live': False
            }
            
        if not counter_data:
            counter_data = {
                'counter_id': f"Counter-{i}",
                'default_name': f"Counter {i}",
                'default_code': f"C{i:02d}",
                'source_device': 'System',
                'button_index': get_button_index_char(i),
                'name': f"Counter {i}",
                'code': f"C{i:02d}",
                'is_live': False
            }

        # Common attributes
        counter_data.update({
            'row_span': 1,
            'col_span': 1,
            'is_enabled': i <= configured_count, # Default enabled status
            'audio_url': None,
            'audio_name': None,
            'db_obj': None,
            'display_index': i # Helpful for JS hiding
        })
        counters.append(counter_data)

    # 3. Apply TVCounter Database Overrides
    if tv_config:
        db_counters = {c.counter_id: c for c in TVCounter.objects.filter(tv_config=tv_config)}
        for c in counters:
            if c['counter_id'] in db_counters:
                obj = db_counters[c['counter_id']]
                c['db_obj'] = obj
                c['name'] = obj.counter_name or c['default_name']
                c['code'] = obj.counter_code or c['default_code']
                c['row_span'] = obj.row_span
                c['col_span'] = obj.col_span
                c['is_enabled'] = obj.is_enabled
                c['audio_url'] = obj.counter_audio_file.url if obj.counter_audio_file else None
                c['audio_name'] = obj.counter_audio_file.name if obj.counter_audio_file else None
    
    return counters


def get_active_embedded_profile(device, company, target_day, target_time):
    """
    Helper to find the active EmbeddedProfile for a given device based on:
    - Company & Branch
    - Device Type
    - Current Day & Time
    """
    from configdetails.models import EmbeddedProfile
    
    # 1. Check if device has a specifically assigned profile (Direct FK)
    if device.embedded_profile and device.embedded_profile.is_active:
        ep = device.embedded_profile
        # Check constraints for the assigned profile
        # Time Match
        time_match = True
        if ep.start_time and ep.start_time > target_time:
            time_match = False
        if ep.end_time and ep.end_time < target_time:
            time_match = False
            
        # Day Match
        days_list = ep.day or []
        day_match = False
        if not days_list or target_day in days_list or 'ALL' in days_list:
            day_match = True
        elif 'WKDY' in days_list and target_day in ['MON', 'TUE', 'WED', 'THU', 'FRI']:
            day_match = True
        elif 'WKND' in days_list and target_day in ['SAT', 'SUN']:
            day_match = True
            
        if time_match and day_match:
             return {
                'id': ep.id,
                'name': ep.name,
                'start_time': ep.start_time.strftime('%H:%M') if ep.start_time else None,
                'end_time': ep.end_time.strftime('%H:%M') if ep.end_time else None,
                'config': ep.config_json or {}
            }

    # 2. Dynamic Search (Company/Branch/Type) with robust time handling
    # Use Q objects to handle NULL start/end times as "All Day" boundaries
    query = Q(company=company) & Q(device_type=device.device_type) & Q(is_active=True)
    
    # Time logic: (start_time <= target OR start_time is NULL) AND (end_time >= target OR end_time is NULL)
    query &= (Q(start_time__lte=target_time) | Q(start_time__isnull=True))
    query &= (Q(end_time__gte=target_time) | Q(end_time__isnull=True))

    if device.branch:
        query &= Q(branch=device.branch)
    else:
        query &= Q(branch__isnull=True)

    profiles = EmbeddedProfile.objects.filter(query)

    active_profile = None
    for ep in profiles:
        days_list = ep.day or []
        if not days_list or target_day in days_list or 'ALL' in days_list:
            active_profile = ep
            break
        if 'WKDY' in days_list and target_day in ['MON', 'TUE', 'WED', 'THU', 'FRI']:
            active_profile = ep
            break
        if 'WKND' in days_list and target_day in ['SAT', 'SUN']:
            active_profile = ep
            break
            
    if active_profile:
        return {
            'id': active_profile.id,
            'name': active_profile.name,
            'start_time': active_profile.start_time.strftime('%H:%M') if active_profile.start_time else None,
            'end_time': active_profile.end_time.strftime('%H:%M') if active_profile.end_time else None,
            'config': active_profile.config_json or {}
        }
    return None


def get_used_locations(company_ids=None, dealer_customer_ids=None):
    """
    Build a hierarchical location structure (State → District → City) 
    containing only locations that are actually used by customers.
    Queries Company, Branch, and DealerCustomer for unique location combinations.
    """
    from companydetails.models import Company, Branch, DealerCustomer
    from collections import defaultdict
    
    # Collect all unique (state, district, city) tuples from all sources
    locations = set()
    
    # From Company
    company_qs = Company.objects.all()
    if company_ids is not None:
        company_qs = company_qs.filter(id__in=company_ids)
        
    for c in company_qs.values('state', 'district', 'city').distinct():
        if c['state'] and c['city']:
            locations.add((c['state'], c['district'] or '', c['city']))
    
    # From Branch
    branch_qs = Branch.objects.all()
    if company_ids is not None:
        branch_qs = branch_qs.filter(company_id__in=company_ids)
        
    for b in branch_qs.values('state', 'district', 'city').distinct():
        if b['state'] and b['city']:
            locations.add((b['state'], b['district'] or '', b['city']))
    
    # From DealerCustomer
    dc_qs = DealerCustomer.objects.all()
    if dealer_customer_ids is not None:
        dc_qs = dc_qs.filter(id__in=dealer_customer_ids)
    elif company_ids is not None:
        # If we have company_ids but no specific dc_ids, maybe include all DCs for these dealers?
        # Actually, if it's a dealer login, we'll pass company_ids=[dealer_id] and dc_ids=[all_dc_ids]
        dc_qs = dc_qs.filter(dealer_id__in=company_ids)
        
    for dc in dc_qs.values('state', 'district', 'city').distinct():
        if dc['state'] and dc['city']:
            locations.add((dc['state'], dc['district'] or '', dc['city']))
    
    # Build nested structure: State -> District -> Cities
    state_data = defaultdict(lambda: defaultdict(set))
    
    for state, district, city in locations:
        state_data[state][district].add(city)
    
    # Convert to the requested JSON structure
    result = []
    for state_name, districts in sorted(state_data.items()):
        state_entry = {
            "state": state_name,
            "Details": []
        }
        for district_name, cities in sorted(districts.items()):
            district_entry = {
                "district": district_name if district_name else state_name,
                "cities": sorted(list(cities))
            }
            state_entry["Details"].append(district_entry)
        result.append(state_entry)
    
    return result

def get_flattened_used_locations(company_ids=None, dealer_customer_ids=None, states=None):
    """
    Build a HIGHLY flattened location structure (District: City)
    containing only locations that are actually used by customers.
    Returns: [{"state": "...", "district": "...", "city": "..."}, ...]
    
    Params:
        company_ids: Filter to specific companies
        dealer_customer_ids: Filter to specific dealer customers
        states: Filter to specific states (list of state names)
    """
    from companydetails.models import Company, Branch, DealerCustomer
    
    locations = set()
    
    # From Company
    company_qs = Company.objects.all()
    if company_ids is not None:
        company_qs = company_qs.filter(id__in=company_ids)
    if states is not None:
        company_qs = company_qs.filter(state__in=states)
        
    for c in company_qs.values('state', 'district', 'city').distinct():
        if c['city']:
            locations.add((c['state'], c['district'] or '', c['city']))
    
    # From Branch
    branch_qs = Branch.objects.all()
    if company_ids is not None:
        branch_qs = branch_qs.filter(company_id__in=company_ids)
    if states is not None:
        branch_qs = branch_qs.filter(state__in=states)
        
    for b in branch_qs.values('state', 'district', 'city').distinct():
        if b['city']:
            locations.add((b['state'], b['district'] or '', b['city']))
    
    # From DealerCustomer
    dc_qs = DealerCustomer.objects.all()
    if dealer_customer_ids is not None:
        dc_qs = dc_qs.filter(id__in=dealer_customer_ids)
    elif company_ids is not None:
        dc_qs = dc_qs.filter(dealer_id__in=company_ids)
    if states is not None:
        dc_qs = dc_qs.filter(state__in=states)
        
    for dc in dc_qs.values('state', 'district', 'city').distinct():
        if dc['city']:
            locations.add((dc['state'], dc['district'] or '', dc['city']))
    
    # Convert to the requested structure: [{"state": "...", "district": "...", "city": "..."}, ...]
    result = []
    # Sort for consistency
    for state, district, city in sorted(list(locations)):
        result.append({
            "state": state,
            "district": district if district else "",
            "city": city
        })
    
    return result




@api_view(['POST', 'GET'])
@permission_classes([AllowAny])
def get_embedded_config(request):
    """
    API for fetching config for all non-TV embedded devices belonging to a customer.
    Accepts: { "customer_id": "..." }
    Returns: { "devices": [ { "device_id": "...", "config": "$0,...,#" }, ... ] }
    """
    from companydetails.models import DealerCustomer
    
    customer_id = request.data.get('customer_id') or request.GET.get('customer_id')
    
    if not customer_id:
        return JsonResponse({"error": "Missing customer_id"}, status=400)

    # Find Company or DealerCustomer
    company = None
    dealer_customer = None
    
    if str(customer_id).isdigit():
        company = Company.objects.filter(id=customer_id).first()
    
    if not company:
        company = Company.objects.filter(company_id=customer_id).first()
    
    if not company:
        dealer_customer = DealerCustomer.objects.filter(customer_id=customer_id).first()
        if dealer_customer:
            company = dealer_customer.dealer

    if not company and not dealer_customer:
        return JsonResponse({"error": "Invalid customer_id"}, status=404)

    # Get all non-TV devices for this customer
    if dealer_customer:
        devices = Device.objects.filter(dealer_customer=dealer_customer, is_active=True).exclude(device_type='TV')
    else:
        devices = Device.objects.filter(company=company, is_active=True).exclude(device_type='TV')

    result = []
    
    for device in devices:
        # Fetch config
        config = {}
        if hasattr(device, 'config') and device.config:
            config = device.config.config_json or {}
        
        # Find mapping for keypad devices
        mapping = Mapping.objects.filter(
            Q(keypad=device) | 
            Q(keypad_2=device) | 
            Q(keypad_3=device) | 
            Q(keypad_4=device)
        ).first()
        
        config_str = ""
        
        if device.device_type == 'TOKEN_DISPENSER':
            # Token Dispenser format
            random_num = "XXXX"
            serial_number = device.serial_number
            header1 = config.get('header1') or config.get('company_name', 'CallQ')
            header2 = config.get('header2') or config.get('location', 'Token')
            header3 = config.get('header3', 'System')
            footer1 = config.get('footer1', 'Thank You')
            footer2 = config.get('footer2', 'Visit Again')
            day_reset = '2' if config.get('day_wise_reset', '0') == '1' else '1'
            reset_tkn = '1' if config.get('reset_tkn', True) else '0'
            token_cut_val = config.get('token_cut', 'full')
            cutter = '1'
            half_cut = '1' if token_cut_val == 'half' else '0'
            feed = '1' if config.get('feed', True) else '0'
            logo = '1' if config.get('logo_enable', True) else '0'
            button_mode = config.get('button_mode', 'counterwise')
            mode = '1' if button_mode == 'counterwise' else '0'
            label = config.get('token_label', 'Token')
            paper_out = '1' if config.get('paper_out', True) else '0'
            type_val = config.get('type', '0')
            dup_tkn = '1' if config.get('duplicate_print', '0') == '1' else '0'
            
            config_str = f"$0,{random_num},{serial_number},{header1},{header2},{header3},{footer1},{footer2},{day_reset},{reset_tkn},{cutter},{half_cut},{feed},{logo},{mode},{label},{paper_out},{type_val},{dup_tkn},#"
        
        elif device.device_type == 'KEYPAD':
            # Keypad format: $1,SETTINGS,XXXX,ON Screen Text,LogoEnable,Keypad Index,
            #   Single/Multiple,SkipEnable,TransferEnable,VIPEnable,VIPFrom,VIPTo,
            #   KeypadPoolMode,DispenserSlNo,No.of KeypadDev,KPSLNo1,...,KPSLNo5,
            #   MESSAGE1,MESSAGE2,MESSAGE3,XX#
            password = 'XXXX'
            text = config.get('keypad_device_text') or 'Welcome'
            logo_enable = config.get('logo_enable') or '0'
            # Counter Number → use the keypad's button index (slot position) from TVKeypadMapping.
            # This is the ASCII character assigned to this keypad when the group was configured
            # (e.g. '1' for the first keypad, '2' for the second, …).
            # Fall back to the manually configured counter_no, or '1' if neither exists.
            kp_mapping = TVKeypadMapping.objects.filter(keypad=device).first()
            if kp_mapping and kp_mapping.keypad_index:
                counter_num = kp_mapping.keypad_index
            else:
                counter_num = config.get('counter_no') or '1'
            # Convert old values (0/1) to new values (1/2) for backward compatibility
            single_multiple_raw = config.get('single_multiple')
            if single_multiple_raw == '0':
                single_multiple = '1'  # Old single mode → new single mode
            elif single_multiple_raw == '1':
                single_multiple = '2'  # Old multi mode → new multi mode
            else:
                single_multiple = single_multiple_raw or '1'  # Default to single mode (1)
            skip = config.get('skip_enable') or '0'
            transfer = config.get('transfer_enable') or '0'
            vip = config.get('vip_enable') or '0'
            vip_from = config.get('vip_from') or '0'
            vip_to = config.get('vip_to') or '0'
            pool_mode = config.get('keypad_pool_mode') or '0'
            
            dispenser_sn = config.get('dispenser_sl_no', '0000000000000000')
            if not dispenser_sn:
                # Fallback: check mapping for linked dispenser
                if mapping and mapping.token_dispenser:
                    dispenser_sn = mapping.token_dispenser.serial_number
                else:
                    dispenser_sn = '0000000000000000'
            
            # Get keypad serial numbers from mapping
            keypads = []
            if mapping:
                if mapping.keypad: keypads.append(mapping.keypad.serial_number)
                if mapping.keypad_2: keypads.append(mapping.keypad_2.serial_number)
                if mapping.keypad_3: keypads.append(mapping.keypad_3.serial_number)
                if mapping.keypad_4: keypads.append(mapping.keypad_4.serial_number)
            
            keypad_count = int(config.get('no_of_keypad_dev') or '1')
            
            # Fetch keypad serial numbers from config or mapping
            keypads_list = []
            
            # Try getting from config first (UI saved values)
            for i in range(1, 6): # 1 to 5
                k_sl = config.get(f'keypad_sl_no_{i}')
                if k_sl:
                    keypads_list.append(k_sl)
            
            # If no config (or less than count?), fallback or fill logic?
            # The requirement is: "change the UI according to this if the no of keypad is 1 one keypad serial number field"
            # So we trust what is in config.
            
            # Additional logic: If config is empty, maybe fallback to mapped? 
            # But the user wants distinct control. Let's stick to config primarily.
            # However, if config is empty (first run), we might want to populate from mapping?
            # The UI pre-fills standard mapping. So saving from UI will populate config.
            
            # We strictly take 'keypad_count' number of keypads from the collected list
            # But wait, the list is unordered if we just use loop.
            # We should use exactly keypad_sl_no_1, _2 etc.
            
            final_keypads = []
            for i in range(1, keypad_count + 1):
                # UI sends keypad_sl_no_1, keypad_sl_no_2...
                # Note: In previous UI, first one might have been keypad_sl_no. 
                # We will standardize to keypad_sl_no_{i} in UI updates.
                k_val = config.get(f'keypad_sl_no_{i}') or config.get(f'keypad_sl_no' if i==1 else '') or '0'
                final_keypads.append(k_val)

            remaining_bit_flag = config.get('remaining_bit_flag', '1')

            # Button string messages (B, C, D buttons on keypad)
            # Stored in device config as button_b_string_id, button_c_string_id, button_d_string_id
            # Max 16 characters each (embedded device protocol limit)
            msg1 = (config.get('button_b_string_id') or '')[:16]
            msg2 = (config.get('button_c_string_id') or '')[:16]
            msg3 = (config.get('button_d_string_id') or '')[:16]

            # Build string
            parts = ["$1", "SETTINGS", password, text, logo_enable, counter_num,
                      single_multiple, skip, transfer, vip, vip_from, vip_to,
                      pool_mode, dispenser_sn, str(keypad_count)]
            parts.extend(final_keypads)
            parts.extend([msg1, msg2, msg3])

            config_str = ",".join(parts) + f",{remaining_bit_flag}#"
        
        elif device.device_type == 'BROKER':
            # BROKER format: $0,host,port,ssid,password,topic,no_of_keypads,kp1_sn,kp2_sn,...,#
            topic = config.get('topic')
            if not topic:
                # Generate topic: customerid-serialnumber
                customer_id = ""
                if device.dealer_customer:
                    customer_id = device.dealer_customer.customer_id
                elif device.company:
                    customer_id = device.company.company_id or str(device.company.id)
                
                if customer_id and device.serial_number:
                    topic = f"{customer_id}-{device.serial_number}"
                else:
                    topic = "UnknownTopic"
            
            host = config.get('host', '')
            port = config.get('port', '')
            ssid = config.get('ssid', '')
            password = config.get('password', '')

            # Find keypad devices mapped to this broker via GroupMapping (M2M)
            broker_group = GroupMapping.objects.filter(brokers=device).first()
            broker_keypads = []
            if broker_group:
                for kp in broker_group.keypads.all():
                    broker_keypads.append(kp.serial_number)

            no_of_keypads = len(broker_keypads)
            keypad_serials_str = ",".join(broker_keypads)

            config_str = f"$0,{host},{port},{ssid},{password},{topic},#"

            #It is a good to have feature where it can be implimented in future
            # if keypad_serials_str:
            #     config_str = f"$0,{host},{port},{ssid},{password},{topic},{no_of_keypads},{keypad_serials_str},#"
            # else:
            #     config_str = f"$0,{host},{port},{ssid},{password},{topic},{no_of_keypads},#"
        


        elif device.device_type == 'LED':
            # LED format: $0,XXXX,SN,S/B,DT,ND#
            random_num = "XXXX"
            serial_number = device.serial_number
            sound_mode = config.get('sound_mode', '0')
            display_type = config.get('display_type', '1')
            no_of_digits = config.get('no_of_digits', '3')
            
            config_str = f"$0,{random_num},{serial_number},{sound_mode},{display_type},{no_of_digits},#"
        
        else:
            # Fallback
            config_str = f"${device.device_type},{device.serial_number},#"
        
        result.append({
            "device_id": device.serial_number,
            "device_type": device.device_type,
            "config": config_str
        })
    
    return JsonResponse({"devices": result})

def format_tv_serial_number(sn):
    """
    Format the serial number for TV config responses.
    Converts the year part (2026+) to an alphabetical character (A, B, C...).
    e.g., '2026bCAL0K0001' -> 'AbCAL0K0001'
    """
    if not sn:
        return sn
    sn_str = str(sn)
    if len(sn_str) >= 4 and sn_str[:4].isdigit():
        year = int(sn_str[:4])
        if year >= 2026:
            year_char = chr(ord('A') + (year - 2026))
            return f"{year_char}{sn_str[4:]}"
    return sn

@api_view(['POST'])
@permission_classes([AllowAny])
def get_android_tv_config(request):
    """
    API Endpoint for Android TV to fetch configuration.
    Payload: { "mac_address": "XX:XX:XX:XX:XX:XX", "customer_id": "123" }
    
    Returns comprehensive TV configuration including:
    - Device and company details
    - TV configuration settings
    - Mapped embedded devices (token dispensers, keypads, brokers, LEDs)
    - Generated counters from token dispensers
    - Branch counters
    """
    from companydetails.models import DealerCustomer
    from licenses.models import Batch
    from django.db.models import Sum
    
    mac_address = request.data.get('mac_address')
    customer_id = request.data.get('customer_id')
    
    # Check if customer_id starts with 0 and slice it
    if customer_id and str(customer_id).startswith('0'):
        customer_id = str(customer_id)[1:]
    
    # Handle flag/Flag case-insensitively
    flag = request.data.get('Flag') or request.data.get('flag', 'TV')
    flag = flag.upper()

    if not mac_address or not customer_id:
        log_api_response('get_android_tv_config', 400, error='mac_address and customer_id are required')
        return Response({'error': 'mac_address and customer_id are required'}, status=400)

    # Find Company or DealerCustomer
    company = None
    dealer_customer = None
    is_dealer_customer = False
    
    # Try matching Company by ID (if int) or company_id string
    if str(customer_id).isdigit():
        company = Company.objects.filter(id=customer_id).first()
    
    if not company:
        company = Company.objects.filter(company_id=customer_id).first()
    
    # If no Company found, try DealerCustomer
    if not company:
        dealer_customer = DealerCustomer.objects.filter(customer_id=customer_id).first()
        if dealer_customer:
            is_dealer_customer = True
            # For DealerCustomer, use the dealer's company for device association
            company = dealer_customer.dealer
        
    if not company and not dealer_customer:
        log_api_response('get_android_tv_config', 404, error='Invalid customer_id')
        return Response({'error': 'Invalid customer_id'}, status=404)

    # Find Device (TV)
    # Find Device (TV)
    device = Device.objects.select_related('company', 'branch', 'config', 'dealer_customer', 'tv_config').prefetch_related('tv_config__ads', 'tv_config__counters').filter(serial_number=mac_address).first()

    if not device:
        # Check Batch Limits before creating Pending Request
        # 1. Identify owner for filtering batches
        limit_reached = False
        
        if is_dealer_customer:
            # Dealer Customer Batches
            batches = Batch.objects.filter(dealer_customer=dealer_customer)
            current_tvs = Device.objects.filter(dealer_customer=dealer_customer, device_type=Device.DeviceType.TV).count()
        else:
            # Direct Customer Batches
            batches = Batch.objects.filter(customer=company)
            current_tvs = Device.objects.filter(company=company, device_type=Device.DeviceType.TV).exclude(dealer_customer__isnull=False).count() \
                if company else 0 # Be careful with excluding dealer customer devices if company is same
            
            # Actually, simpler logic:
            # Count all devices owned by this entity (Company or DealerCustomer)
            # Batches are also owned by this entity
            pass

        # Robust Logic:
        # Calculate Total Limit from all batches owned by this entity
        if is_dealer_customer:
            total_limit = Batch.objects.filter(dealer_customer=dealer_customer).aggregate(Sum('max_tvs'))['max_tvs__sum'] or 0
            # Count devices assigned to this dealer customer
            used_count = Device.objects.filter(dealer_customer=dealer_customer, device_type=Device.DeviceType.TV).count()
        else:
            total_limit = Batch.objects.filter(customer=company).aggregate(Sum('max_tvs'))['max_tvs__sum'] or 0
            # Count devices belonging to this company AND NOT assigned to a dealer_customer (if that's how it works)
            # Or simply all devices linked to this company context.
            # Usually strict ownership:
            used_count = Device.objects.filter(company=company, dealer_customer__isnull=True, device_type=Device.DeviceType.TV).count()
        
        # If Limit Reached, reject with clear message
        if used_count >= total_limit:
            return Response({
                'error': 'Maximum number of devices reached',
                'message': 'The maximum number of devices has been reached for this customer. Please contact admin to increase the device limit.'
            }, status=403)
            
        # Create Device as Pending Request
        device = Device.objects.create(
            serial_number=mac_address,
            device_type=Device.DeviceType.TV,
            company=company,
            dealer_customer=dealer_customer if is_dealer_customer else None,
            is_active=True,
            licence_status='Pending' # Explicitly Pending
        )

    # Check if the device belongs to the correct company/dealer_customer
    device_matches_customer = False
    
    if is_dealer_customer:
        # For dealer customer, check if device is assigned to this dealer_customer
        # OR if device belongs to the dealer company (parent)
        device_matches_customer = (
            device.dealer_customer == dealer_customer or 
            (device.company == company and device.dealer_customer is None)
        )
    else:
        # For regular company, check if device belongs to this company
        # Accept devices with or without dealer_customer assignment (as long as company matches)
        device_matches_customer = (device.company == company)
    
    # If device exists but belongs to a DIFFERENT customer
    if not device_matches_customer:
        # Check if it's pending for another customer
        if device.licence_status == 'Pending':
            return Response({
                'error': 'Device pending for another customer',
                'message': 'This device is waiting for approval under another customer. Please contact admin for assistance.'
            }, status=403)
        else:
            # Device is active/rejected but belongs to another customer
            return Response({
                'error': 'Device is not mapped to the customer',
                'message': 'The device with the provided MAC address is not registered under this customer. Please contact admin to map the device correctly.'
            }, status=403)
    
    # Device belongs to this customer - check approval status
    if device.licence_status != 'Active':
        log_api_response('get_android_tv_config', 403, error='Device awaiting approval')
        return Response({'status': 'pending', 'message': 'Device awaiting approval'}, status=403)

    # Store FCM token if provided in the request
    fcm_token = request.data.get('fcm_token')
    if fcm_token and device.fcm_token != fcm_token:
        device.fcm_token = fcm_token
        device.save(update_fields=['fcm_token'])
        action_logger.info(f"FCM token updated for device {device.serial_number}")

    # Fetch Device Config
    config_data = {}
    if hasattr(device, 'config') and device.config:
        config_data = device.config.config_json or {}
    
    # TV Config - get or provide defaults
    tv_config = getattr(device, 'tv_config', None)
    
    # Default TV config values
    default_tv_config = {
        'audio_language': 'en',
        'show_ads': False,
        'ad_interval': 5,
        'orientation': 'landscape',
        'layout_type': 'default',
        'save_audio_external': False,
        'enable_counter_announcement': True,
        'enable_token_announcement': True,
        'enable_counter_prifix': False,
        'ad_files': [],
        'display_rows': 3,
        'display_columns': 4,
        'counter_text_color': '#000000',
        'token_text_color': '#000000',
        'scroll_text_color': '#000000',
        'token_font_size': 24,
        'counter_font_size': 24,
        'tokens_per_counter': 5,
        'no_of_counters': 1,
        'ad_placement': 'right',
        
        # Enhanced Token Defaults
        'current_token_color': '#000000',
        'previous_token_color': '#888888',
        'blink_current_token': False,
        'blink_seconds': 1,
        'token_format': 'T1'
    }
    
    tv_config_data = default_tv_config.copy()
    
    # Check if ads feature is enabled for the company
    ads_enabled = False
    if company:
        ads_enabled = company.ads_enabled
    
    if tv_config:
        tv_config_data = {
            'audio_language': tv_config.audio_language or 'en',
            'orientation': tv_config.orientation or 'landscape',
            'layout_type': tv_config.layout_type or 'default',
            'save_audio_external': tv_config.save_audio_external,
            'enable_counter_announcement': tv_config.enable_counter_announcement,
            'enable_token_announcement': tv_config.enable_token_announcement,
            'enable_counter_prifix': getattr(tv_config, 'enable_counter_prifix', False),
            'display_rows': tv_config.display_rows or 3,
            'display_columns': tv_config.display_columns or 4,
            'counter_text_color': tv_config.counter_text_color or '#000000',
            'token_text_color': tv_config.token_text_color or '#000000',
            'scroll_text_color': getattr(tv_config, 'scroll_text_color', '#000000') or '#000000',
            'token_font_size': getattr(tv_config, 'token_font_size', 24) or 24,
            'counter_font_size': getattr(tv_config, 'counter_font_size', 24) or 24,
            'tokens_per_counter': getattr(tv_config, 'tokens_per_counter', 5) or 5,
            'no_of_counters': getattr(tv_config, 'no_of_counters', 1) or 1,
            # Enhanced Token Config
            'current_token_color': getattr(tv_config, 'current_token_color', '#000000'),
            'previous_token_color': getattr(tv_config, 'previous_token_color', '#888888'),
            'blink_current_token': getattr(tv_config, 'blink_current_token', False),
            'blink_seconds': getattr(tv_config, 'blink_seconds', 1),
            'token_format': getattr(tv_config, 'token_format', 'T1'),
            'ad_interval': tv_config.ad_interval or 5,
            'show_ads': 'on' if tv_config.show_ads else 'off',
            'ad_placement': getattr(tv_config, 'ad_placement', 'right')
        }
        
        # Include ad_files only if ads feature is enabled for the company
        if ads_enabled:
            ad_files = []
            for ad in tv_config.ads.all():
                if ad.file:
                    ad_files.append(request.build_absolute_uri(ad.file.url))
                elif ad.ad_url:
                    ad_files.append(ad.ad_url)
            tv_config_data['ad_files'] = ad_files
        else:
            tv_config_data['ad_files'] = []

    # Generate counters from token dispensers
    generated_counters = generate_tv_counters_data(device, tv_config)
    
    # Clean for JSON serialization (remove model objects) and make URLs absolute
    for c in generated_counters:
        c.pop('db_obj', None)
        if c.get('audio_url'):
            c['audio_url'] = request.build_absolute_uri(c['audio_url'])

    # Fetch Embedded Devices (via Mapping)
    embedded_devices = {
        'token_dispensers': [],
        'keypads': [],
        'brokers': [],
        'leds': []
    }
    
    # Helper to safely get device info with config
    def get_device_info(dev):
        if not dev:
            return None
        conf = getattr(dev, 'config', None)
        return {
            'id': dev.id,
            'serial_number': format_tv_serial_number(dev.serial_number),
            'type': dev.device_type,
            'token_type': dev.token_type if dev.device_type == Device.DeviceType.TOKEN_DISPENSER else None,
            'config': conf.config_json if conf else {}
        }

    mappings = Mapping.objects.filter(tv=device)
    
    # Use sets to avoid duplicates if multiple mappings point to same device
    seen_td = set()
    seen_kp = set()
    seen_bk = set()
    seen_led = set()
    
    for m in mappings:
        # Token Dispensers
        if m.token_dispenser and m.token_dispenser.id not in seen_td:
            embedded_devices['token_dispensers'].append(get_device_info(m.token_dispenser))
            seen_td.add(m.token_dispenser.id)
            
        # Primary Keypad
        if m.keypad and m.keypad.id not in seen_kp:
            embedded_devices['keypads'].append(get_device_info(m.keypad))
            seen_kp.add(m.keypad.id)
            
        # Additional Keypads
        for k_attr in ['keypad_2', 'keypad_3', 'keypad_4']:
            kp = getattr(m, k_attr, None)
            if kp and kp.id not in seen_kp:
                embedded_devices['keypads'].append(get_device_info(kp))
                seen_kp.add(kp.id)
        
        # Brokers (fixed duplicate bug)
        if m.broker and m.broker.id not in seen_bk:
            embedded_devices['brokers'].append(get_device_info(m.broker))
            seen_bk.add(m.broker.id)
            
        # LEDs
        if m.led and m.led.id not in seen_led:
            embedded_devices['leds'].append(get_device_info(m.led))
            seen_led.add(m.led.id)
        
    # Build mappings list
    mappings_list = []
    for m in mappings:
        mappings_list.append({
            'id': m.id,
            'token_dispenser': format_tv_serial_number(m.token_dispenser.serial_number) if m.token_dispenser else None,
            'keypad': format_tv_serial_number(m.keypad.serial_number) if m.keypad else None,
            'keypad_2': format_tv_serial_number(m.keypad_2.serial_number) if m.keypad_2 else None,
            'keypad_3': format_tv_serial_number(m.keypad_3.serial_number) if m.keypad_3 else None,
            'keypad_4': format_tv_serial_number(m.keypad_4.serial_number) if m.keypad_4 else None,
            'broker': format_tv_serial_number(m.broker.serial_number) if m.broker else None,
            'led': format_tv_serial_number(m.led.serial_number) if m.led else None
        })

    # Branch information
    branch_info = None
    if device.branch:
        branch_info = {
            'id': device.branch.id,
            'name': device.branch.branch_name,
            'address': device.branch.address,
            'city': device.branch.city,
            'state': device.branch.state
        }
    
    # Get branch counters (from Counter model)
    branch_counters = []
    if device.branch:
        counters = Counter.objects.filter(branch=device.branch)
        for counter in counters:
            branch_counters.append({
                'id': counter.id,
                'counter_name': counter.counter_name,
                'counter_number': counter.counter_number,
                'assigned_device_sn': format_tv_serial_number(counter.assigned_device.serial_number) if counter.assigned_device else None
            })

    # Calculate Shift Details based on current/provided time and date
    from datetime import datetime
    now = datetime.now()
    
    req_time = request.data.get('time')
    req_date = request.data.get('date')
    
    target_time = now.time()
    target_day = now.strftime('%a').upper()
    
    if req_time:
        try:
            # Try formats HH:MM and HH:MM:SS
            try:
                target_time = datetime.strptime(req_time, '%H:%M').time()
            except ValueError:
                target_time = datetime.strptime(req_time, '%H:%M:%S').time()
        except (ValueError, TypeError):
            pass
            
    if req_date:
        try:
            temp_date = datetime.strptime(req_date, '%Y-%m-%d')
            # Use weekday() 0=Mon, 6=Sun
            days_map = {0: 'MON', 1: 'TUE', 2: 'WED', 3: 'THU', 4: 'FRI', 5: 'SAT', 6: 'SUN'}
            target_day = days_map[temp_date.weekday()]
        except (ValueError, TypeError):
            pass

    shift_details = None
    
    # Use the robust helper to find the active profile for the TV itself
    # This handles NULL times and direct 'embedded_profile' assignment correctly
    active_profile_data = get_active_embedded_profile(device, company, target_day, target_time)

    if active_profile_data and active_profile_data.get('config'):
        # Clean up scroll config keys from profile config
        for key in ['scroll_enabled', 'no_of_text_fields', 'scroll_text_lines']:
            active_profile_data['config'].pop(key, None)

    if active_profile_data:
        shift_details = {
            'name': active_profile_data['name'],
            'start_time': active_profile_data['start_time'],
            'end_time': active_profile_data['end_time'],
            'config': active_profile_data['config']
        }
        
        # If a specific profile is found, apply its config
        tv_config_data.update(active_profile_data['config'])

    # If Flag is EMBEDED, return comprehensive branch/device data
    if flag == 'EMBEDED':
        return _get_embedded_branch_config(request, device, company, dealer_customer, is_dealer_customer, mappings_list, generated_counters, branch_counters, shift_details)

    # If Flag is TV, return only TV config and mapped broker
    if flag == 'TV':
        return _get_tv_flag_config(request, device, company, dealer_customer, is_dealer_customer, tv_config_data, active_profile_data, generated_counters, shift_details)

    # Build response for default (fallback)
    response_data = {
        'status': 'success',
        'message': 'Configuration fetched successfully',
        'Route': get_flattened_used_locations(),
        'shift_details': shift_details,
        'device_id': device.id,
        'serial_number': format_tv_serial_number(device.serial_number),
        'company_name': dealer_customer.company_name if is_dealer_customer else (company.company_name if company else None),
        'is_dealer_customer': is_dealer_customer,
        'dealer_customer_id': dealer_customer.customer_id if is_dealer_customer else None,
        'config': config_data,
        'tv_config': tv_config_data,
        'mappings': mappings_list,
        'generated_counters': generated_counters,
        'branch_counters': branch_counters
    }
    
    log_api_response('get_android_tv_config', 200, {'status': 'success', 'device_id': device.id, 'serial_number': format_tv_serial_number(device.serial_number), 'flag': 'default'})
    return Response(response_data)


def _get_embedded_branch_config(request, device, company, dealer_customer, is_dealer_customer, mappings_list, generated_counters, branch_counters, shift_details):
    """Refactored helper for 'EMBEDED' flag in get_android_tv_config"""
    if not device.branch:
        return Response({
            'status': 'error', 
            'message': 'Device is not assigned to any branch. Please contact admin.'
        }, status=400)
        
    branch_devices = Device.objects.filter(company=company, branch=device.branch).select_related('config')
    
    inactive_devices = []
    for dev in branch_devices:
        if dev.is_expired:
            inactive_devices.append({
                'serial_number': format_tv_serial_number(dev.serial_number),
                'device_type': dev.device_type,
                'status': 'Expired',
                'message': f'{dev.get_device_type_display()} ({format_tv_serial_number(dev.serial_number)}) license has expired'
            })
        elif not dev.is_active or dev.licence_status not in ['Active']:
            inactive_devices.append({
                'serial_number': format_tv_serial_number(dev.serial_number),
                'device_type': dev.device_type,
                'status': dev.licence_status,
                'message': f'{dev.get_device_type_display()} ({format_tv_serial_number(dev.serial_number)}) is inactive'
            })
    
    broker_inactive = [d for d in inactive_devices if d['device_type'] == 'BROKER']
    if broker_inactive:
        return Response({
            'status': 'error',
            'message': f'License Status: Your licence is expired. Please contact Admin !!!',
            'inactive_devices': broker_inactive,
            'device_type': 'BROKER',
            'device_serial': broker_inactive[0]['serial_number'],
            'licence_status': broker_inactive[0]['status']
        }, status=403)
    
    def serialize_device_detailed(dev):
        conf = getattr(dev, 'config', None)
        return {
            'id': dev.id,
            'serial_number': format_tv_serial_number(dev.serial_number),
            'type': dev.device_type,
            'name': dev.device_model or dev.get_device_type_display(),
            'token_type': dev.token_type if dev.device_type == Device.DeviceType.TOKEN_DISPENSER else None,
            'config': conf.config_json if conf else {},
            'is_active': dev.is_active,
            'licence_status': dev.licence_status,
            'is_expired': dev.is_expired
        }
        
    button_mappings = ButtonMapping.objects.filter(branch=device.branch)
    button_mappings_data = [
        {
            'source_id': m.source_device.id,
            'source_sn': format_tv_serial_number(m.source_device.serial_number),
            'source_type': m.source_device.device_type,
            'source_button': m.source_button,
            'target_id': m.target_device.id,
            'target_sn': format_tv_serial_number(m.target_device.serial_number),
            'target_type': m.target_device.device_type
        } for m in button_mappings
    ]

    embedded_response = {
        'status': 'success',
        'message': 'Branch configuration fetched successfully',
        'Route': get_flattened_used_locations(),
        'device_id': device.id,
        'serial_number': format_tv_serial_number(device.serial_number),
        'company_name': dealer_customer.company_name if is_dealer_customer else (company.company_name if company else None),
        'branch_id': device.branch.id,
        'branch_name': device.branch.branch_name,
        'shift_details': shift_details,
        'mappings': mappings_list,
        'button_mappings': button_mappings_data,
        'branch_counters': branch_counters,
        'generated_counters': generated_counters,
        'inactive_devices': inactive_devices if inactive_devices else None,
        'token_dispensers': [serialize_device_detailed(d) for d in branch_devices if d.device_type == Device.DeviceType.TOKEN_DISPENSER],
        'keypads': [serialize_device_detailed(d) for d in branch_devices if d.device_type == Device.DeviceType.KEYPAD],
        'brokers': [serialize_device_detailed(d) for d in branch_devices if d.device_type == Device.DeviceType.BROKER],
        'leds': [serialize_device_detailed(d) for d in branch_devices if d.device_type == Device.DeviceType.LED],
        'tvs': [serialize_device_detailed(d) for d in branch_devices if d.device_type == Device.DeviceType.TV],
    }
    log_api_response('get_android_tv_config', 200, {'status': 'success', 'device_id': device.id, 'serial_number': format_tv_serial_number(device.serial_number)})
    return Response(embedded_response)


def _get_tv_flag_config(request, device, company, dealer_customer, is_dealer_customer, tv_config_data, active_profile_data, generated_counters, shift_details):
    """Refactored helper for 'TV' flag in get_android_tv_config"""

    # ---------------------------------------------------------------
    # Build connected_devices — scoped to THIS TV only:
    #   1. Brokers from the GroupMapping the TV belongs to
    #      (brokers are shared at group level, not per-TV)
    #   2. Keypads explicitly mapped to this TV via TVKeypadMapping
    #   3. Dispensers linked to those keypads (TVKeypadMapping.dispenser)
    # Only BROKER devices include their full config_json.
    # ---------------------------------------------------------------
    group = (
        GroupMapping.objects
        .filter(tvs=device)
        .prefetch_related('brokers')
        .first()
    )

    # 1. Brokers — find only the broker(s) specifically mapped to THIS TV
    #    via ButtonMapping (source=BROKER → target=TV).
    #    Fall back to all group brokers only if no explicit mapping exists.
    tv_scoped_devices = []
    broker_ids_via_buttonmap = list(
        ButtonMapping.objects.filter(
            target_device=device,
            source_device__device_type=Device.DeviceType.BROKER,
        ).values_list('source_device_id', flat=True).distinct()
    )
    if broker_ids_via_buttonmap:
        tv_scoped_devices += list(
            Device.objects.filter(id__in=broker_ids_via_buttonmap)
        )
    elif group:
        # No explicit ButtonMapping — fall back to group brokers
        tv_scoped_devices += list(group.brokers.all())

    # 2 & 3. Keypads + their paired dispensers from TVKeypadMapping
    tv_keypad_mappings = (
        TVKeypadMapping.objects
        .filter(tv=device)
        .select_related('keypad', 'dispenser')
    )
    for kpm in tv_keypad_mappings:
        tv_scoped_devices.append(kpm.keypad)
        if kpm.dispenser:
            tv_scoped_devices.append(kpm.dispenser)

    # Deduplicate preserving order (earlier entry wins)
    all_connected_map = {}
    for d in tv_scoped_devices:
        if d.id not in all_connected_map:
            all_connected_map[d.id] = d

    # Use current time/day for embedded-profile lookup
    from datetime import datetime
    now = datetime.now()
    target_time = now.time()
    target_day = now.strftime('%a').upper()

    serialized_devices = []
    for dev in all_connected_map.values():
        dev_data = {
            'id': dev.id,
            'serial_number': format_tv_serial_number(dev.serial_number),
            'device_type': dev.device_type,
            'name': dev.display_name or dev.get_device_type_display(),
            'status': dev.licence_status,
            'is_active': dev.is_active,
        }
        # Only BROKER devices expose their config details
        if dev.device_type == Device.DeviceType.BROKER:
            conf = getattr(dev, 'config', None)
            dev_data['config'] = conf.config_json if conf else {}
        serialized_devices.append(dev_data)

    # ---------------------------------------------------------------
    # Helper: read a device's config_json safely (returns {} on error)
    # ---------------------------------------------------------------
    def _get_device_config_json(dev):
        try:
            cfg = dev.config
            return cfg.config_json if cfg and cfg.config_json else {}
        except Exception:
            return {}

    # ---------------------------------------------------------------
    # Helper: build a list of counter dicts for a given dispenser.
    # Source of truth: GroupCounterButtonMapping scoped to the
    # group that THIS TV belongs to (via GroupDispenserMapping).
    # Filtering by dispenser alone causes duplicates when stale
    # GCBM rows from old group assignments still exist in the DB.
    # ---------------------------------------------------------------
    def _build_counters_for_dispenser(dispenser, keypad_index):
        """
        Build counter dicts for a given dispenser, attaching:
          - keypad_index            : ASCII slot of the keypad on this TV (e.g. '1', '2')
          - dispenser_index         : ASCII slot of the dispenser on this TV (TVDispenserMapping.button_index)
          - button_index            : same as dispenser_index
          - dispenser_button_number : 1-based physical button position on the dispenser
                                     (derived from GCBM order by button_index)
        Source: GroupCounterButtonMapping scoped to the dispenser's current group.
        """
        results = []
        if dispenser is None:
            return results

        # Look up this dispenser's TV-slot index (TVDispenserMapping).
        # Fallback: when no TVDispenserMapping row exists (keypad-first config),
        # use keypad_index as the dispenser slot identifier — each keypad slot
        # maps 1-to-1 with a dispenser slot so the index is semantically identical.
        try:
            tdm = TVDispenserMapping.objects.filter(
                tv=device, dispenser=dispenser
            ).only('button_index').first()
            dispenser_index = tdm.button_index if tdm else keypad_index
        except Exception:
            dispenser_index = keypad_index

        # Resolve the group this dispenser belongs to (via GroupDispenserMapping).
        # This is the only authoritative group for this dispenser — dispensers can
        # only be in ONE group (enforced at view level). Scoping GCBM to this group
        # prevents stale rows from old group assignments causing duplicate counters.
        try:
            gdm = GroupDispenserMapping.objects.filter(
                dispenser=dispenser
            ).select_related('group').first()
            dispenser_group = gdm.group if gdm else None
        except Exception:
            dispenser_group = None

        # Fetch counters from GroupCounterButtonMapping scoped to the correct group
        try:
            gcbm_qs = GroupCounterButtonMapping.objects.filter(
                dispenser=dispenser
            ).select_related('counter').order_by('button_index')

            # Always scope to the group when available to avoid cross-group duplicates
            if dispenser_group:
                gcbm_qs = gcbm_qs.filter(group=dispenser_group)

            gcbms = list(gcbm_qs)
            for btn_num, gcbm in enumerate(gcbms, start=1):
                c = gcbm.counter
                # dispenser_button_number must be the ASCII char (e.g. '1', '2')
                # not a Python int, because the APK uses ASCII-based indexing.
                try:
                    dispenser_button_number_char = get_button_index_char(btn_num)
                except ValueError:
                    dispenser_button_number_char = chr(0x31)
                results.append({
                    'counter_id':              c.counter_name,
                    'default_code':            c.counter_prefix_code,
                    'keypad_index':            keypad_index,
                    'dispenser_index':         dispenser_index,
                    'button_index':            dispenser_index,
                    'dispenser_button_number': dispenser_button_number_char,
                    'name':                    c.counter_display_name,
                    'code':                    c.counter_prefix_code,
                    'row_span':                1,
                    'col_span':                1,
                    'is_enabled':              c.status,
                    'counter_config_id':       c.id,
                    'max_token_number':        c.max_token_number,
                    'dispenser_sn':            dispenser.serial_number,  # used by TV counter filter
                })
        except Exception:
            pass
        return results

    # ----------------------------------------------------------------
    # Build keypads + counters.
    #
    # Dispenser resolution (priority order):
    #   1. kp_mapping.dispenser  (TVKeypadMapping FK — legacy/manual mapping)
    #   2. ButtonMapping where target_device=keypad, source=TOKEN_DISPENSER
    #      (new linear mapping: Dispenser Button N → Keypad)
    #
    # Button strings: returned as {letter: {id, value}} pairs.
    #   id  = "<keypad_db_id>_<letter>"  — stable, unique per keypad+button
    #   value = text stored in keypad's config_json
    # ----------------------------------------------------------------
    mapped_keypads_list = []
    mapped_counters_list = []  # flat aggregate for top-level 'counters' key

    def _resolve_dispenser_for_keypad(kp, kp_mapping_dispenser):
        """
        Return (dispenser_device, source_button) for this keypad.
        Falls back to ButtonMapping when TVKeypadMapping.dispenser is null.
        source_button is the button label on the dispenser (e.g. 'Button 1').
        """
        if kp_mapping_dispenser:
            return kp_mapping_dispenser, None  # legacy direct FK

        bm = (
            ButtonMapping.objects
            .filter(
                target_device=kp,
                source_device__device_type=Device.DeviceType.TOKEN_DISPENSER,
            )
            .select_related('source_device')
            .first()
        )
        if bm:
            return bm.source_device, bm.source_button
        return None, None

    try:
        tv_keypad_mappings = (
            TVKeypadMapping.objects
            .filter(tv=device)
            .select_related('keypad', 'dispenser')
            .order_by('keypad_index')
        )
        for kp_mapping in tv_keypad_mappings:
            kp = kp_mapping.keypad
            kp_cfg = _get_device_config_json(kp)

            # --- Button strings as ordered array ---
            def _btn_entry(letter, str_id, cfg):
                text = cfg.get(f'button_{letter.lower()}_string_id', '') or ''
                return {
                    'id':    str_id,
                    'value': text,
                }

            button_strings = [
                _btn_entry('B', '1', kp_cfg),
                _btn_entry('C', '2', kp_cfg),
                _btn_entry('D', '3', kp_cfg),
            ]

            # --- Resolve dispenser and counters ---
            dispenser, _ = _resolve_dispenser_for_keypad(kp, kp_mapping.dispenser)
            kp_counters = _build_counters_for_dispenser(dispenser, kp_mapping.keypad_index)  # dispenser_index resolved inside
            mapped_counters_list.extend(kp_counters)

            mapped_keypads_list.append({
                'keypad_sn':           format_tv_serial_number(kp.serial_number),
                'keypad_display_name': kp.display_name or format_tv_serial_number(kp.serial_number),
                'keypad_index':        kp_mapping.keypad_index,
                'dispenser_sn':        format_tv_serial_number(dispenser.serial_number) if dispenser else None,
                'button_strings':      button_strings,
                'counters':            kp_counters,
            })
    except Exception:
        mapped_keypads_list = []
        mapped_counters_list = []

    # Fallback: if no TVKeypadMapping rows exist, try TVDispenserMapping directly.
    # In this case keypad_index is None (no keypad slot), dispenser_index is set.
    if not mapped_keypads_list:
        for tdm in TVDispenserMapping.objects.filter(tv=device).select_related('dispenser').order_by('button_index'):
            mapped_counters_list.extend(
                _build_counters_for_dispenser(tdm.dispenser, None)  # no keypad slot; dispenser_index resolved inside
            )

    # ----------------------------------------------------------------
    # Filter counters to only those from dispensers explicitly wired
    # to THIS TV on the config page (via TVKeypadMapping.dispenser or
    # TVDispenserMapping).  Without this, a dispenser that belongs to
    # the same group but is NOT assigned to this TV would still
    # contribute its counters via GCBM, sending the TV data it should
    # not display.
    #
    # Priority: TVKeypadMapping.dispenser FK (one dispenser per keypad
    # slot); fall back to TVDispenserMapping when no keypad mappings
    # exist yet.
    # ----------------------------------------------------------------
    tv_wired_dispenser_sns = set()
    for kpm in TVKeypadMapping.objects.filter(tv=device).select_related('dispenser'):
        if kpm.dispenser:
            tv_wired_dispenser_sns.add(kpm.dispenser.serial_number)
    if not tv_wired_dispenser_sns:
        for tdm in TVDispenserMapping.objects.filter(tv=device).select_related('dispenser'):
            if tdm.dispenser:
                tv_wired_dispenser_sns.add(tdm.dispenser.serial_number)

    if tv_wired_dispenser_sns:
        mapped_counters_list = [
            c for c in mapped_counters_list
            if c.get('dispenser_sn') in tv_wired_dispenser_sns
        ]

    # Scroll Configuration Logic
    # Always use device's default config since profiles don't currently support scrolling text configuration
    # (Existing profiles might have 'off' hardcoded, so we ignore profile config for scroll text)
    device_config_json = device.config.config_json if hasattr(device, 'config') and device.config else {}
    
    scroll_source = device_config_json
    
    # Extract scroll configuration with proper defaults
    scroll_enabled = scroll_source.get('scroll_enabled', 'off') if isinstance(scroll_source, dict) else 'off'
    no_of_text_fields = scroll_source.get('no_of_text_fields', 0) if isinstance(scroll_source, dict) else 0
    scroll_text_lines = scroll_source.get('scroll_text_lines', []) if isinstance(scroll_source, dict) else []
    
    # Ensure scroll_text_lines is a list and validate it
    if not isinstance(scroll_text_lines, list):
        scroll_text_lines = []
    # Filter out empty or invalid lines
    scroll_text_lines = [line for line in scroll_text_lines if line and isinstance(line, str) and len(line.strip()) > 0]

    # Count keypad slots from TVKeypadMapping and update no_of_counters
    # (keypads replace dispensers as the primary TV-side slot count)
    keypad_count = TVKeypadMapping.objects.filter(tv=device).count()
    if keypad_count > 0:
        tv_config_data['no_of_counters'] = keypad_count
    else:
        # Fall back to TVDispenserMapping count if no keypad mappings exist yet
        token_dispenser_count = TVDispenserMapping.objects.filter(tv=device).count()
        tv_config_data['no_of_counters'] = token_dispenser_count

    tv_response = {
        'status': 'success',
        'message': 'TV configuration fetched successfully',
        'device_id': device.id,
        'serial_number': format_tv_serial_number(device.serial_number),
        'company_name': dealer_customer.company_name if is_dealer_customer else (company.company_name if company else None),
        'tv_config': tv_config_data,
        'current_profile': active_profile_data if active_profile_data else None,
        'connected_devices': serialized_devices,
        'counters': mapped_counters_list,
        'keypads': mapped_keypads_list,  # New: keypad slot config with button string IDs
        'shift_details': shift_details,
        'scroll_config': {
            'scroll_enabled': scroll_enabled,
            'no_of_text_fields': no_of_text_fields,
            'scroll_text_lines': scroll_text_lines,
            'scroll_text_color': tv_config_data.get('scroll_text_color', '#000000'),
        }
    }
    return Response(tv_response)


# --- TEMPLATE VIEWS ---


@login_required
def device_config(request, device_id):
    device = get_object_or_404(Device, id=device_id)
    
    config, created = DeviceConfig.objects.get_or_create(device=device)
    
    # Migrate single_multiple from old values (0/1) to new values (1/2) for KEYPAD devices
    # This is a one-time migration: '0' → '1' (single), '1' → '2' (multi)
    if device.device_type == Device.DeviceType.KEYPAD and config.config_json:
        single_multiple = config.config_json.get('single_multiple')
        needs_save = False
        if single_multiple == '0':
            # Old single mode → new single mode (1)
            config.config_json['single_multiple'] = '1'
            needs_save = True
        elif single_multiple == '1':
            # Old multi mode → new multi mode (2)
            # Note: This migrates old '1' (multi) to '2'. 
            # Going forward, new saves will use '1' for single and '2' for multi.
            config.config_json['single_multiple'] = '2'
            needs_save = True
        if needs_save:
            config.save(update_fields=['config_json', 'updated_at'])
    
    tv_config = None
    counters = []
    ads = []
    all_counters = []
    mapped_counter_ids = []
    mapped_counter_by_position = {}
    all_dispensers = []
    mapped_dispenser_by_position = {}
    # Keypad mapping context — defaults for non-TV devices (populated below for TV only)
    all_keypads_for_tv = []
    mapped_keypad_by_position = {}       # slot int (1-8) -> keypad device id
    all_dispensers_for_tv = []           # dispensers in same company (for keypad slot dropdown)
    mapped_dispenser_by_kp_position = {} # slot int (1-8) -> dispenser_id (from kp_mapping.dispenser)
    
    if device.device_type == Device.DeviceType.TV:
        tv_config, _ = TVConfig.objects.get_or_create(tv=device)
        counters = generate_tv_counters_data(device, tv_config)
        ads = tv_config.ads.all()
        
        # Get all active counters for dropdowns (scoped to this device's company)
        try:
            all_counters = CounterConfig.objects.filter(status=True)
            if device.company_id:
                all_counters = all_counters.filter(company_id=device.company_id)
            all_counters = all_counters.order_by('counter_name')
            # Get currently mapped counters in order
            mapped_counters = list(TVCounterMapping.objects.filter(tv=device).select_related('counter').order_by('id'))
            mapped_counter_ids = [mc.counter_id for mc in mapped_counters]
            # Create a dictionary mapping position (1-8) to counter_id for easy template access
            mapped_counter_by_position = {}
            for idx, mc in enumerate(mapped_counters, start=1):
                mapped_counter_by_position[idx] = mc.counter_id
        except Exception:
            # Handle case where tables don't exist yet (migrations not run)
            all_counters = []
            mapped_counter_ids = []
            mapped_counter_by_position = {}
    elif device.device_type == Device.DeviceType.TOKEN_DISPENSER:
        # Get all active counters for token dispenser mapping (scoped to this device's company)
        try:
            all_counters = CounterConfig.objects.filter(status=True)
            if device.company_id:
                all_counters = all_counters.filter(company_id=device.company_id)
            all_counters = all_counters.order_by('counter_name')
            # Get currently mapped counters in order
            mapped_counters = list(CounterTokenDispenserMapping.objects.filter(dispenser=device).select_related('counter').order_by('id'))
            mapped_counter_ids = [mc.counter_id for mc in mapped_counters]
            # Create a dictionary mapping button position (1-4) to counter_id
            mapped_counter_by_button = {}
            for idx, mc in enumerate(mapped_counters, start=1):
                mapped_counter_by_button[idx] = mc.counter_id
        except Exception:
            # Handle case where tables don't exist yet (migrations not run)
            all_counters = []
            mapped_counter_ids = []
            mapped_counter_by_button = {}
    # LED handled via config_json only

    # Fetch mapped devices for Keypad auto-configuration
    mapped_dispenser_sl_no = None
    mapped_keypads = []
    
    if device.device_type == Device.DeviceType.KEYPAD:
        # Find mapping where this device is one of the keypads
        mapping = Mapping.objects.filter(
            Q(keypad=device) | 
            Q(keypad_2=device) | 
            Q(keypad_3=device) | 
            Q(keypad_4=device)
        ).first()
        
        if mapping:
            if mapping.token_dispenser:
                mapped_dispenser_sl_no = mapping.token_dispenser.serial_number
            
            # Collect all mapped keypads in order
            if mapping.keypad: mapped_keypads.append(mapping.keypad.serial_number)
            if mapping.keypad_2: mapped_keypads.append(mapping.keypad_2.serial_number)
            if mapping.keypad_3: mapped_keypads.append(mapping.keypad_3.serial_number)
            if mapping.keypad_4: mapped_keypads.append(mapping.keypad_4.serial_number)

        # Fallback: Check ButtonMapping if dispenser not found in main mapping
        if not mapped_dispenser_sl_no:
            # Look for any button mapping where this keypad is the Target and Source is a Token Dispenser
            btn_mapping = ButtonMapping.objects.filter(
                target_device=device,
                source_device__device_type=Device.DeviceType.TOKEN_DISPENSER
            ).first()
            
            if btn_mapping:
                mapped_dispenser_sl_no = btn_mapping.source_device.serial_number

    # Fetch available keypads for dropdowns
    available_keypads = []
    if device.device_type == Device.DeviceType.KEYPAD:
        if device.dealer_customer:
            available_keypads = Device.objects.filter(
                dealer_customer=device.dealer_customer, 
                device_type=Device.DeviceType.KEYPAD,
                is_active=True
            ).exclude(id=device.id)
        elif device.company:
            available_keypads = Device.objects.filter(
                company=device.company, 
                device_type=Device.DeviceType.KEYPAD,
                is_active=True
            ).exclude(id=device.id)

    available_keypads_serials = [k.serial_number for k in available_keypads]

    # Ensure fixed fields (company_name/location) are always derived from device context
    # Location is taken as branch city (fallback to branch name), else company city.
    def _fixed_company_and_location(dev: Device):
        company_name = dev.company.company_name if dev.company else ''
        if dev.dealer_customer:
            # If device is mapped to a dealer's customer, prefer that customer name/city
            company_name = dev.dealer_customer.company_name or company_name
            location = dev.dealer_customer.city or ''
        else:
            if dev.branch:
                location = dev.branch.city or dev.branch.branch_name or ''
            else:
                # Fallback for company-wide devices
                location = getattr(dev.company, 'city', '') if dev.company else ''
        return company_name, location

    # Remove fixed company/location enforcement for Token Dispenser to allow custom headers
    # if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
    #     fixed_company_name, fixed_location = _fixed_company_and_location(device)
    #     # Persist on GET as well (keeps template values consistent)
    #     if config.config_json.get('company_name') != fixed_company_name or config.config_json.get('location') != fixed_location:
    #         config.config_json = {**(config.config_json or {}), 'company_name': fixed_company_name, 'location': fixed_location}
    #         config.save(update_fields=['config_json', 'updated_at'])

    if request.method == 'POST':
        # DeviceConfig update
        data = request.POST.dict()
        data.pop('csrfmiddlewaretoken', None)
        
        # Remove TV fields from JSON to keep clean
        tv_specific_keys = ['show_ads', 'ad_interval', 'orientation', 'layout_type', 'next']
        # Do not remove new LED config keys: sound_mode, display_type, no_of_digits
        
        # Also remove counter specific keys dynamic keys
        keys_to_remove = [k for k in data.keys() if k.startswith('counter_') or k.startswith('row_') or k.startswith('col_') or k.startswith('enable_c_')]
        
        json_data = data.copy()
        for k in tv_specific_keys + keys_to_remove:
            json_data.pop(k, None)

        # Normalize checkbox/switch values so templates can reliably check against '1'/'0'
        if device.device_type == Device.DeviceType.KEYPAD:
            keypad_checkbox_keys = ['logo_enable', 'skip_enable', 'transfer_enable', 'vip_enable']
            for key in keypad_checkbox_keys:
                json_data[key] = '1' if request.POST.get(key) == 'on' else '0'

            # Enforce max 16 characters for button string messages (embedded device protocol limit)
            for btn_key in ('button_b_string_id', 'button_c_string_id', 'button_d_string_id'):
                if btn_key in json_data and json_data[btn_key]:
                    json_data[btn_key] = json_data[btn_key][:16]

        if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
            dispenser_checkbox_keys = ['day_wise_reset', 'common_pool', 'duplicate_print', 'standalone', 'vip_enable', 'initial_print']
            for key in dispenser_checkbox_keys:
                json_data[key] = '1' if request.POST.get(key) == 'on' else '0'

            # Force fixed fields regardless of POST (prevents tampering)
            # fixed_company_name, fixed_location = _fixed_company_and_location(device)
            # json_data['company_name'] = fixed_company_name
            # json_data['location'] = fixed_location

        if device.device_type == 'BROKER':
            # Automatically generate topic: customerid-serialnumber
            customer_id = ""
            if device.dealer_customer:
                customer_id = device.dealer_customer.customer_id
            elif device.company:
                customer_id = device.company.company_id or str(device.company.id)
            
            if customer_id and device.serial_number:
                json_data['topic'] = f"{customer_id}-{device.serial_number}"
            
        # Handle Scrolling Text for TV (and potentially others if enabled)
        if device.device_type == Device.DeviceType.TV:
            no_of_text_fields = min(int(json_data.get('no_of_text_fields') or 0), 8)
            raw_lines = [
                json_data.get(f'text_content_{i}') 
                for i in range(1, no_of_text_fields + 1)
                if json_data.get(f'text_content_{i}')
            ]
            # Validate: max 150 chars, min 3 non-space chars
            scroll_text_lines = []
            for line in raw_lines:
                line = line[:150]  # enforce max 150 chars
                if len(line.replace(' ', '')) >= 3:  # enforce min 3 non-space chars
                    scroll_text_lines.append(line)
            json_data['scroll_text_lines'] = scroll_text_lines
            json_data['no_of_text_fields'] = len(scroll_text_lines)
            
            # Cleanup individual text fields from json_data
            for i in range(1, 11):
                 json_data.pop(f'text_content_{i}', None)
            
            # Ensure scroll_enabled is 'on' or 'off' (checkbox behavior)
            json_data['scroll_enabled'] = 'on' if request.POST.get('scroll_enabled') == 'on' else 'off'

        config.config_json.update(json_data)
        config.config_json = json_data
        config.save()
        
        # TVConfig update
        if tv_config:
            # Check if ads feature is enabled for the company
            ads_enabled = False
            if hasattr(device, 'company') and device.company:
                ads_enabled = device.company.ads_enabled
            
            # Only allow show_ads and ad_files if ads_enabled is True
            if ads_enabled:
                tv_config.show_ads = request.POST.get('show_ads') == 'on'
            else:
                tv_config.show_ads = False  # Disabled by default if ads feature is not enabled
            
            tv_config.ad_interval = request.POST.get('ad_interval', 5)
            tv_config.orientation = request.POST.get('orientation', 'landscape')
            tv_config.layout_type = request.POST.get('layout_type', 'default')
            tv_config.display_rows = request.POST.get('display_rows', 3)
            tv_config.display_columns = request.POST.get('display_columns', 4)
            tv_config.counter_text_color = request.POST.get('counter_text_color', '#000000')
            tv_config.token_text_color = request.POST.get('token_text_color', '#000000')
            tv_config.scroll_text_color = request.POST.get('scroll_text_color', '#000000')
            tv_config.col_span_single = request.POST.get('col_span_single', 1)
            tv_config.token_font_size = request.POST.get('token_font_size', 24)
            tv_config.counter_font_size = request.POST.get('counter_font_size', 24)
            tv_config.tokens_per_counter = request.POST.get('tokens_per_counter', 5)
            tv_config.no_of_counters = request.POST.get('no_of_counters', request.POST.get('no_of_keypads', 1))
            tv_config.no_of_dispensers = int(request.POST.get('no_of_dispensers', request.POST.get('no_of_keypads', request.POST.get('no_of_counters', 1))))
            tv_config.ad_placement = request.POST.get('ad_placement', 'right')
            
            # Boolean fields
            tv_config.save_audio_external = request.POST.get('save_audio_external') == 'on'
            tv_config.enable_counter_announcement = request.POST.get('enable_counter_announcement') == 'on'
            tv_config.enable_token_announcement = request.POST.get('enable_token_announcement') == 'on'
            tv_config.enable_counter_prifix = request.POST.get('enable_counter_prifix') == 'on'
            
            # Audio Language
            tv_config.audio_language = request.POST.get('audio_language', 'en')

            # Enhanced Token Config Save
            tv_config.current_token_color = request.POST.get('current_token_color', '#000000')
            tv_config.previous_token_color = request.POST.get('previous_token_color', '#888888')
            tv_config.blink_current_token = request.POST.get('blink_current_token') == 'on'
            tv_config.blink_seconds = int(request.POST.get('blink_seconds', 1) or 1)
            tv_config.token_format = request.POST.get('token_format', 'T1')
            
            tv_config.save()
            
            # Save Ads - only if ads_enabled is True
            if ads_enabled:
                files = request.FILES.getlist('ad_files')
                for f in files:
                    TVAd.objects.create(tv_config=tv_config, file=f)
                
                ad_urls = request.POST.getlist('ad_urls[]')
                for url in ad_urls:
                    if url and url.strip():
                        TVAd.objects.create(tv_config=tv_config, ad_url=url.strip())
            
            # Save Counters
            for c in counters:
                cid = c['counter_id']
                # Get existing or new
                tvc, _ = TVCounter.objects.get_or_create(tv_config=tv_config, counter_id=cid)
                tvc.counter_name = request.POST.get(f'counter_name_{cid}')
                tvc.counter_code = request.POST.get(f'counter_code_{cid}')
                tvc.row_span = request.POST.get(f'row_span_{cid}', 1)
                tvc.col_span = request.POST.get(f'col_span_{cid}', 1)
                tvc.is_enabled = request.POST.get(f'enable_c_{cid}') == 'on'
                
                tvc.save()
            
            # ----------------------------------------------------------------
            # Handle Keypad Mapping: Map keypads + their Button-A dispensers to TV
            # Reads: keypad_selection_1…N and dispenser_for_keypad_1…N
            # keypad_index is auto-assigned sequentially via get_button_index_char().
            # Button B/C/D strings are stored in the keypad's own config_json.
            # ----------------------------------------------------------------
            no_of_keypads = int(request.POST.get('no_of_keypads', 0))
            selected_keypad_data = []  # list of (keypad_id, dispenser_id_or_None)

            for i in range(1, no_of_keypads + 1):
                kp_id_raw = request.POST.get(f'keypad_selection_{i}', '').strip()
                if kp_id_raw:
                    try:
                        kp_id_int = int(kp_id_raw)
                        if kp_id_int > 0:
                            disp_id_raw = request.POST.get(f'dispenser_for_keypad_{i}', '').strip()
                            disp_id = int(disp_id_raw) if disp_id_raw else None
                            selected_keypad_data.append((kp_id_int, disp_id))
                    except (ValueError, TypeError):
                        pass

            if selected_keypad_data:
                try:
                    keypad_ids = [row[0] for row in selected_keypad_data]
                    keypads_qs = Device.objects.filter(
                        id__in=keypad_ids,
                        device_type=Device.DeviceType.KEYPAD,
                    )
                    if keypads_qs.count() != len(keypad_ids):
                        messages.error(request, 'One or more keypad IDs are invalid or not KEYPAD devices')
                        return redirect('device_config', device_id=device_id)

                    # Check for keypads already mapped to another TV
                    duplicate_kps = []
                    for kp in keypads_qs:
                        clash = TVKeypadMapping.objects.filter(keypad=kp).exclude(tv=device)
                        if clash.exists():
                            other_tvs = [m.tv.serial_number for m in clash]
                            duplicate_kps.append(f"Keypad '{kp.serial_number}' → TV(s): {', '.join(other_tvs)}")

                    if duplicate_kps:
                        messages.error(request, 'Cannot map keypad(s): ' + ' | '.join(duplicate_kps))
                        return redirect('device_config', device_id=device_id)

                    # Atomically replace keypad mappings for this TV
                    with transaction.atomic():
                        TVKeypadMapping.objects.filter(tv=device).delete()
                        # Clear direct dispenser mappings to prevent stale data
                        TVDispenserMapping.objects.filter(tv=device).delete()
                        for position, (kp_id, disp_id) in enumerate(selected_keypad_data, start=1):
                            try:
                                kp = Device.objects.get(id=kp_id, device_type=Device.DeviceType.KEYPAD)
                                dispenser = Device.objects.get(id=disp_id, device_type=Device.DeviceType.TOKEN_DISPENSER) if disp_id else None
                                # Remove any existing mapping of this keypad to another TV
                                TVKeypadMapping.objects.filter(keypad=kp).exclude(tv=device).delete()
                                ascii_idx = get_button_index_char(position)
                                TVKeypadMapping.objects.create(
                                    tv=device,
                                    keypad=kp,
                                    dispenser=dispenser,
                                    keypad_index=ascii_idx,
                                )
                            except Device.DoesNotExist:
                                pass
                except Exception as e:
                    messages.error(request, f'Error mapping keypads: {e}')
            elif 'no_of_keypads' in request.POST or request.POST.get('clear_keypad_mappings') == '1':
                # Clear all keypad mappings if explicitly submitted with 0 keypads or no valid keypads
                try:
                    TVKeypadMapping.objects.filter(tv=device).delete()
                    TVDispenserMapping.objects.filter(tv=device).delete()
                except Exception:
                    pass


            # Handle Dispenser Mapping: Map dispensers to TV (kept for backward compatibility)
            # Get dispenser selections from dropdowns (dispenser_selection_1, dispenser_selection_2, etc.)
            no_of_dispensers = int(request.POST.get('no_of_dispensers', request.POST.get('no_of_counters', 1)))
            selected_dispenser_ids = []

            for i in range(1, no_of_dispensers + 1):
                dispenser_id = request.POST.get(f'dispenser_selection_{i}') or request.POST.get(f'counter_selection_{i}')
                if dispenser_id:
                    try:
                        dispenser_id_int = int(dispenser_id)
                        if dispenser_id_int > 0:  # Valid dispenser ID
                            selected_dispenser_ids.append(dispenser_id_int)
                    except (ValueError, TypeError):
                        pass  # Skip invalid dispenser IDs


            # Keep old counter mapping logic for backward compatibility, but prioritize dispenser mapping
            selected_counter_ids = []
            if not selected_dispenser_ids:
                # Fallback to counter selection if no dispensers selected
                for i in range(1, no_of_dispensers + 1):
                    counter_id = request.POST.get(f'counter_selection_{i}')
                    if counter_id:
                        try:
                            counter_id_int = int(counter_id)
                            if counter_id_int > 0:
                                selected_counter_ids.append(counter_id_int)
                        except (ValueError, TypeError):
                            pass

            if selected_dispenser_ids:
                # Handle dispenser mapping
                try:
                    # Validate all dispensers exist and are TOKEN_DISPENSER type
                    dispensers = Device.objects.filter(
                        id__in=selected_dispenser_ids,
                        device_type=Device.DeviceType.TOKEN_DISPENSER
                    )
                    if dispensers.count() != len(selected_dispenser_ids):
                        messages.error(request, 'One or more dispenser IDs are invalid or not token dispensers')
                        return redirect('device_config', device_id=device_id)
                    
                    # Validate: Check if any selected dispenser is already mapped to another TV
                    duplicate_dispensers = []
                    for dispenser in dispensers:
                        existing_mappings = TVDispenserMapping.objects.filter(
                            dispenser=dispenser
                        ).exclude(tv=device)
                        
                        if existing_mappings.exists():
                            other_tvs = [m.tv.serial_number for m in existing_mappings]
                            duplicate_dispensers.append({
                                'dispenser': dispenser.serial_number,
                                'other_tvs': other_tvs
                            })
                    
                    if duplicate_dispensers:
                        error_msg = "Cannot map dispensers: One dispenser can only be mapped to one TV at a time. "
                        for dup in duplicate_dispensers:
                            error_msg += f"Dispenser '{dup['dispenser']}' is already mapped to TV(s): {', '.join(dup['other_tvs'])}. "
                        messages.error(request, error_msg)
                        return redirect('device_config', device_id=device_id)
                    
                    # Remove existing mappings for this TV
                    TVDispenserMapping.objects.filter(tv=device).delete()
                    
                    # Remove existing mappings for selected dispensers from other TVs (enforce one-to-one)
                    for dispenser_id in selected_dispenser_ids:
                        try:
                            dispenser = Device.objects.get(id=dispenser_id, device_type=Device.DeviceType.TOKEN_DISPENSER)
                            # Delete any existing mappings for this dispenser to other TVs
                            TVDispenserMapping.objects.filter(dispenser=dispenser).exclude(tv=device).delete()
                        except Device.DoesNotExist:
                            pass
                    
                    # Create new mappings with button_index as ASCII chars (0x31 onwards, continuous per TV)
                    for position, dispenser_id in enumerate(selected_dispenser_ids, start=1):
                        try:
                            dispenser = Device.objects.get(id=dispenser_id, device_type=Device.DeviceType.TOKEN_DISPENSER)
                            TVDispenserMapping.objects.create(
                                tv=device,
                                dispenser=dispenser,
                                button_index=get_button_index_char(position)  # '1','2',... (ASCII 0x31+)
                            )
                        except Device.DoesNotExist:
                            pass  # Skip invalid dispenser IDs
                except Exception as e:
                    messages.error(request, f"Error mapping dispensers: {e}")
            elif selected_counter_ids:
                try:
                    # Validate: Check if any selected counter is already mapped to another TV
                    # Scope check to the same company — counters and TVs are company-specific,
                    # so a mapping in a different company is never a conflict.
                    device_company = device.company
                    duplicate_counters = []
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            existing_qs = TVCounterMapping.objects.filter(
                                counter=counter
                            ).exclude(tv=device)
                            if device_company:
                                existing_qs = existing_qs.filter(tv__company=device_company)
                            
                            if existing_qs.exists():
                                other_tvs = [m.tv.serial_number for m in existing_qs]
                                duplicate_counters.append({
                                    'counter': counter.counter_name,
                                    'other_tvs': other_tvs
                                })
                        except CounterConfig.DoesNotExist:
                            pass  # Skip invalid counter IDs
                    
                    if duplicate_counters:
                        error_msg = "Cannot map counters: One counter can only be mapped to one TV at a time. "
                        for dup in duplicate_counters:
                            error_msg += f"Counter '{dup['counter']}' is already mapped to TV(s): {', '.join(dup['other_tvs'])}. "
                        messages.error(request, error_msg)
                        return redirect('device_config', device_id=device_id)
                    
                    # Remove existing mappings for this TV
                    TVCounterMapping.objects.filter(tv=device).delete()
                    
                    # Remove existing mappings for selected counters from other TVs (enforce one-to-one)
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            # Delete any existing mappings for this counter to other TVs
                            TVCounterMapping.objects.filter(counter=counter).exclude(tv=device).delete()
                        except CounterConfig.DoesNotExist:
                            pass
                    
                    # Create new mappings
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            TVCounterMapping.objects.get_or_create(tv=device, counter=counter)
                        except CounterConfig.DoesNotExist:
                            pass  # Skip invalid counter IDs
                except Exception as e:
                    messages.error(request, f"Error mapping counters: {e}")
            else:
                # If no dispensers and no counters are selected, remove all existing mappings for this TV
                try:
                    TVDispenserMapping.objects.filter(tv=device).delete()
                    TVCounterMapping.objects.filter(tv=device).delete()
                except Exception:
                    pass
        
        # Handle Counter-Dispenser Mapping for Token Dispensers
        if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
            # Get number of buttons from token_type (1_BUTTON = 1, 2_BUTTON = 2, etc.)
            num_buttons = 1  # Default to 1 button
            if device.token_type:
                try:
                    # Extract number from token_type (e.g., '1_BUTTON' -> 1)
                    num_buttons = int(device.token_type.split('_')[0])
                except (ValueError, AttributeError):
                    num_buttons = 1
            
            # Get counter selections from individual dropdowns (dispenser_counter_selection_1, dispenser_counter_selection_2, etc.)
            selected_counter_ids = []
            for i in range(1, num_buttons + 1):
                counter_id = request.POST.get(f'dispenser_counter_selection_{i}')
                if counter_id:
                    try:
                        counter_id_int = int(counter_id)
                        if counter_id_int > 0:  # Valid counter ID
                            selected_counter_ids.append(counter_id_int)
                    except (ValueError, TypeError):
                        pass  # Skip invalid counter IDs
            
            if selected_counter_ids:
                try:
                    # Validate: Check if any selected counter is already mapped to another dispenser.
                    # Scope the check to the same company — counters and dispensers are
                    # company-specific, so a mapping in a different company is never a conflict.
                    device_company = device.company
                    duplicate_counters = []
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            existing_qs = CounterTokenDispenserMapping.objects.filter(
                                counter=counter
                            ).exclude(dispenser=device)
                            if device_company:
                                existing_qs = existing_qs.filter(dispenser__company=device_company)
                            
                            if existing_qs.exists():
                                other_dispensers = [m.dispenser.serial_number for m in existing_qs]
                                duplicate_counters.append({
                                    'counter': counter.counter_name,
                                    'other_dispensers': other_dispensers
                                })
                        except CounterConfig.DoesNotExist:
                            pass  # Skip invalid counter IDs
                    
                    if duplicate_counters:
                        error_msg = "Cannot map counters: One counter can only be mapped to one dispenser at a time. "
                        for dup in duplicate_counters:
                            error_msg += f"Counter '{dup['counter']}' is already mapped to dispenser(s): {', '.join(dup['other_dispensers'])}. "
                        messages.error(request, error_msg)
                        return redirect('device_config', device_id=device_id)
                    
                    # Validate: Check if selected counters exceed dispenser button limit
                    if len(selected_counter_ids) > num_buttons:
                        messages.error(request, f"Cannot map {len(selected_counter_ids)} counter(s): This dispenser is a {num_buttons}-button dispenser and can only have {num_buttons} counter(s) mapped.")
                        return redirect('device_config', device_id=device_id)
                    
                    # Remove existing mappings for this dispenser
                    CounterTokenDispenserMapping.objects.filter(dispenser=device).delete()
                    
                    # Remove existing mappings for selected counters from other dispensers (enforce one-to-one)
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            # Delete any existing mappings for this counter to other dispensers
                            CounterTokenDispenserMapping.objects.filter(counter=counter).exclude(dispenser=device).delete()
                        except CounterConfig.DoesNotExist:
                            pass
                    
                    # Create new mappings in order (button 1 = first counter, button 2 = second counter, etc.)
                    for counter_id in selected_counter_ids:
                        try:
                            counter = CounterConfig.objects.get(id=counter_id, status=True)
                            CounterTokenDispenserMapping.objects.get_or_create(dispenser=device, counter=counter)
                        except CounterConfig.DoesNotExist:
                            pass  # Skip invalid counter IDs

                    # ── Sync GroupCounterButtonMapping (GCBM) for this dispenser only ──────
                    # own_counters reads exclusively from GCBM for grouped dispensers.
                    # Only touch rows for THIS dispenser — sibling dispensers are left alone.
                    import logging as _gcbm_log
                    _gcbm_logger = _gcbm_log.getLogger('actions')
                    try:
                        _gdm_for_device = GroupDispenserMapping.objects.filter(
                            dispenser=device
                        ).select_related('group').first()
                        if _gdm_for_device:
                            _grp = _gdm_for_device.group

                            # Remove only this dispenser's existing GCBM rows.
                            GroupCounterButtonMapping.objects.filter(
                                group=_grp, dispenser=device
                            ).delete()

                            # Compute the correct starting button index for this dispenser.
                            # Count how many GCBM slots are already occupied by dispensers
                            # that come BEFORE this one in the group order.
                            _all_grp_disp = list(
                                _grp.dispenser_slot_mappings
                                    .select_related('dispenser')
                                    .order_by('dispenser_button_index')
                            )
                            _preceding_count = 0
                            for _slot in _all_grp_disp:
                                if _slot.dispenser_id == device.id:
                                    break
                                _preceding_count += GroupCounterButtonMapping.objects.filter(
                                    group=_grp, dispenser=_slot.dispenser
                                ).count()

                            # Insert / update only the new rows for this dispenser.
                            # Use update_or_create (not get_or_create) so button_index
                            # is always written even if the row already existed.
                            _new_ctdm = CounterTokenDispenserMapping.objects.filter(
                                dispenser=device
                            ).select_related('counter').order_by('id')
                            for _local_idx, _cm in enumerate(_new_ctdm, start=1):
                                _abs_pos = _preceding_count + _local_idx
                                try:
                                    _btn_char = get_button_index_char(_abs_pos)
                                except ValueError:
                                    _btn_char = chr(0x31)
                                GroupCounterButtonMapping.objects.update_or_create(
                                    group=_grp,
                                    dispenser=device,
                                    counter=_cm.counter,
                                    defaults={'button_index': _btn_char},
                                )
                            _gcbm_logger.info(
                                f"GCBM synced for dispenser {device.serial_number} "
                                f"in group '{_grp.group_name}': "
                                f"{_new_ctdm.count()} counter(s) written starting at abs_pos {_preceding_count + 1}"
                            )
                        else:
                            _gcbm_logger.info(
                                f"Dispenser {device.serial_number} has no group — GCBM sync skipped (uses CTDM directly)"
                            )
                    except Exception as _gcbm_exc:
                        _gcbm_logger.warning(
                            f"GCBM sync failed for dispenser {device.serial_number}: {_gcbm_exc}"
                        )
                        # GCBM sync is best-effort; CTDM write already succeeded.
                except Exception as e:
                    messages.error(request, f"Error mapping counters to dispenser: {e}")
            else:
                # If no counters are selected, remove all existing mappings for this dispenser
                try:
                    CounterTokenDispenserMapping.objects.filter(dispenser=device).delete()
                    # Also clear GCBM rows for this dispenser if it belongs to a group.
                    try:
                        _gdm_empty = GroupDispenserMapping.objects.filter(
                            dispenser=device
                        ).select_related('group').first()
                        if _gdm_empty:
                            _grp_empty = _gdm_empty.group
                            GroupCounterButtonMapping.objects.filter(
                                group=_grp_empty, dispenser=device
                            ).delete()
                    except Exception:
                        pass
                except Exception:
                    pass

        # LEDConfig update removed (Legacy)

        # Send FCM push notification to TV device if config updated and token available
        if device.device_type == Device.DeviceType.TV:
            import logging as _logging
            action_logger = _logging.getLogger('actions')
            action_logger.info(f"Attempting to send FCM config update to TV device {device.serial_number}")
            try:
                from .fcm_service import send_fcm_notification
                success = send_fcm_notification(
                    device,
                    message="Your TV configuration has been updated. Please reload.",
                    title="Config Update",
                    data={"action": "config_update", "device_id": str(device.id)}
                )
                if success:
                    action_logger.info(f"FCM config update API call succeeded for TV {device.serial_number}")
                else:
                    action_logger.warning(f"FCM config update API call returned False for TV {device.serial_number} (check if token is valid)")
            except Exception as fcm_exc:
                action_logger.error(f"FCM notification crashed after config save: {fcm_exc}")

        next_url = request.POST.get('next')
        if next_url:
            return redirect(next_url)
        messages.success(request, 'Configuration saved successfully.')
        return redirect('device_config', device_id=device_id)

    next_url = request.GET.get('next')
    fixed_company_name, fixed_location = _fixed_company_and_location(device)

    # Ensure mapped_counter_by_button is always defined for the render context.
    # For TOKEN_DISPENSER, it was already populated in the early block (line ~1391).
    # For other device types, default to empty dict.
    if device.device_type != Device.DeviceType.TOKEN_DISPENSER:
        mapped_counter_by_button = {}

    if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
        # Ensure counters are fetched if not already done (edge case: POST path skipped early block)
        if not all_counters:
            try:
                all_counters = CounterConfig.objects.filter(status=True)
                if device.company_id:
                    all_counters = all_counters.filter(company_id=device.company_id)
                all_counters = all_counters.order_by('counter_name')
                # Get currently mapped counters in order
                mapped_counters = list(CounterTokenDispenserMapping.objects.filter(dispenser=device).select_related('counter').order_by('id'))
                mapped_counter_ids = [mc.counter_id for mc in mapped_counters]
                # Create a dictionary mapping button position (1-4) to counter_id
                mapped_counter_by_button = {}
                for idx, mc in enumerate(mapped_counters, start=1):
                    mapped_counter_by_button[idx] = mc.counter_id
            except Exception:
                all_counters = []
                mapped_counter_ids = []
                mapped_counter_by_button = {}
    elif device.device_type == Device.DeviceType.TV:
        # ---------- Keypad mapping context ----------
        all_keypads_for_tv = []
        mapped_keypad_by_position = {}     # slot int (1-8) -> keypad_id
        all_dispensers_for_tv = []         # dispensers in same company for the slot dropdown
        mapped_dispenser_by_kp_position = {}  # slot int (1-8) -> dispenser_id (from kp_mapping.dispenser)
        try:
            all_keypads_for_tv = Device.objects.filter(
                device_type=Device.DeviceType.KEYPAD,
                company=device.company,
            ).order_by('serial_number')

            all_dispensers_for_tv = Device.objects.filter(
                device_type=Device.DeviceType.TOKEN_DISPENSER,
                company=device.company,
            ).order_by('serial_number')

            existing_kp_mappings = list(
                TVKeypadMapping.objects.filter(tv=device)
                    .select_related('keypad', 'dispenser')
                    .order_by('keypad_index')
            )
            for slot_pos, kpm in enumerate(existing_kp_mappings, start=1):
                mapped_keypad_by_position[slot_pos] = kpm.keypad_id
                if kpm.dispenser_id:
                    mapped_dispenser_by_kp_position[slot_pos] = kpm.dispenser_id
        except Exception:
            all_keypads_for_tv = []
            mapped_keypad_by_position = {}
            all_dispensers_for_tv = []
            mapped_dispenser_by_kp_position = {}

        # ---------- Legacy dispenser context (kept for backward compat) ----------
        all_dispensers = []
        mapped_dispenser_by_position = {}
        try:
            all_dispensers = Device.objects.filter(
                device_type=Device.DeviceType.TOKEN_DISPENSER,
                company=device.company,
            ).order_by('serial_number')

            mapped_dispensers = list(
                TVDispenserMapping.objects.filter(tv=device)
                    .select_related('dispenser')
                    .order_by('button_index')
            )
            for mapping in mapped_dispensers:
                mapped_dispenser_by_position[mapping.button_index] = mapping.dispenser_id
        except Exception:
            all_dispensers = []
            mapped_dispenser_by_position = {}



    # Get ads_enabled status from company
    ads_enabled = False
    if hasattr(device, 'company') and device.company:
        ads_enabled = device.company.ads_enabled
    
    return render(request, 'configdetails/device_config.html', {
        'device': device,
        'config': config,
        'tv_config': tv_config,
        'counters': counters,
        'ads': ads,
        'ads_enabled': ads_enabled,
        'all_counters': all_counters,
        # TV keypad mapping context
        'all_keypads_for_tv': all_keypads_for_tv,
        'mapped_keypad_by_position': mapped_keypad_by_position,
        'all_dispensers_for_tv': all_dispensers_for_tv,
        'mapped_dispenser_by_kp_position': mapped_dispenser_by_kp_position,
        # Legacy dispenser mapping context (backward compat)
        'all_dispensers': all_dispensers,
        'mapped_counter_ids': mapped_counter_ids,
        'mapped_counter_by_position': mapped_counter_by_position,
        'mapped_dispenser_by_position': mapped_dispenser_by_position,
        'mapped_counter_by_button': mapped_counter_by_button,
        'next': next_url,
        'mapped_dispenser_sl_no': mapped_dispenser_sl_no,
        'mapped_keypads': mapped_keypads,
        'available_keypads': available_keypads,
        'available_keypads_serials': available_keypads_serials,
        'default_company_name': fixed_company_name,
        'default_location': fixed_location,
    })


@login_required
def delete_ad(request, ad_id):
    """Delete an advertisement file from a TV configuration."""
    # Use filter().first() to avoid 404 if already deleted
    ad = TVAd.objects.filter(id=ad_id).first()
    
    if ad:
        device_id = ad.tv_config.tv.id
        
        # Delete the file from storage
        if ad.file:
            ad.file.delete(save=False)
        
        # Delete the record
        ad.delete()
        
        messages.success(request, 'Advertisement deleted successfully.')
        return redirect('device_config', device_id=device_id)
    else:
        # Ad not found (likely already deleted)
        messages.warning(request, 'Advertisement not found. It may have already been deleted.')
        # Try to return to the previous page (device config)
        referer = request.META.get('HTTP_REFERER')
        if referer:
            return redirect(referer)
        # Fallback if no referer
        return redirect('device_list')



@login_required
def tv_config(request, tv_id):

    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    config, created = TVConfig.objects.get_or_create(tv=tv)
    device_config, _ = DeviceConfig.objects.get_or_create(device=tv)
    
    if request.method == 'POST':
        config.show_ads = False # Hidden from UI - Disabled by default
        config.ad_interval = request.POST.get('ad_interval', 5)
        config.orientation = request.POST.get('orientation', 'landscape')
        config.save()
        
        # Save Scrolling Text to DeviceConfig
        scroll_enabled = request.POST.get('scroll_enabled', 'off')
        
        try:
            no_of_text_fields = int(request.POST.get('no_of_text_fields') or 0)
        except ValueError:
            no_of_text_fields = 0
            
        no_of_text_fields = min(no_of_text_fields, 8)
            
        raw_lines = [
            request.POST.get(f'text_content_{i}') 
            for i in range(1, no_of_text_fields + 1)
            if request.POST.get(f'text_content_{i}')
        ]
        # Validate: max 150 chars, min 3 non-space chars
        scroll_text_lines = []
        for line in raw_lines:
            line = line[:150]  # enforce max 150 chars
            if len(line.replace(' ', '')) >= 3:  # enforce min 3 non-space chars
                scroll_text_lines.append(line)
        
        device_config_data = device_config.config_json or {}
        device_config_data.update({
            'scroll_enabled': scroll_enabled,
            'scroll_text_lines': scroll_text_lines,
            'no_of_text_fields': len(scroll_text_lines)
        })
        device_config.config_json = device_config_data
        device_config.save()

        return redirect('device_list')

    return render(request, 'configdetails/tv_config.html', {'tv': tv, 'config': config, 'device_config': device_config})

@login_required
def device_list(request):
    from django.db.models import Count, Q
    user = request.user
    if user.role == "SUPER_ADMIN":
        devices = Device.objects.all()
        branches = Branch.objects.all()
    elif user.role == "ADMIN":
        if user.assigned_state:
            devices = Device.objects.filter(company__state__in=user.assigned_state)
            branches = Branch.objects.filter(company__state__in=user.assigned_state)
        else:
            devices = Device.objects.none()
            branches = []
    elif user.role == "COMPANY_ADMIN":
        devices = Device.objects.filter(company=user.company_relation)
        branches = Branch.objects.filter(company=user.company_relation)
    elif user.role == "BRANCH_ADMIN":
        devices = Device.objects.filter(branch=user.branch_relation)
        branches = [user.branch_relation]
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            # Dealer sees their own devices AND their child companies' devices
            devices = Device.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
            branches = Branch.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
        else:
            devices = Device.objects.none()
            branches = []
    elif user.role == "PRODUCTION_ADMIN":
        # Production Admin sees all devices for batch management
        devices = Device.objects.all()
        branches = Branch.objects.all()
    else:
        devices = Device.objects.none()
        branches = []
        
    # Add ordering to avoid pagination warning
    devices = devices.order_by('-created_at')

    # Server-side filtering — read params from GET so they survive page navigation
    search_query = request.GET.get('search', '').strip()
    type_filter = request.GET.get('type', '').strip()
    status_filter = request.GET.get('status', '').strip()

    if search_query:
        devices = devices.filter(
            Q(serial_number__icontains=search_query) |
            Q(mac_address__icontains=search_query) |
            Q(display_name__icontains=search_query) |
            Q(device_model__icontains=search_query)
        )

    if type_filter:
        devices = devices.filter(device_type=type_filter)

    if status_filter == 'True':
        devices = devices.filter(is_active=True)
    elif status_filter == 'False':
        devices = devices.filter(is_active=False)
    elif status_filter == 'Expired':
        from django.utils import timezone as tz
        devices = devices.filter(licence_active_to__lt=tz.now().date())

    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(devices, 8)
    try:
        devices_page = paginator.page(page)
    except PageNotAnInteger:
        devices_page = paginator.page(1)
    except EmptyPage:
        devices_page = paginator.page(paginator.num_pages)

    companies = Company.objects.all() if user.role == "SUPER_ADMIN" else []

    return render(request, 'configdetails/device_list.html', {
        'devices': devices_page,
        'branches': branches,
        'companies': companies,
        'search_query': search_query,
        'type_filter': type_filter,
        'status_filter': status_filter,
    })

@login_required
# def device_assign_batch(request, device_id):
#     if request.method == 'POST':
#         device = get_object_or_404(Device, id=device_id)
#         batch_id = request.POST.get('batch_id')
#         
#         # Verify permissions (simple check for now)
#         # TODO: Add stricter permission checks
#         
#         if batch_id:
#             batch = get_object_or_404(Batch, id=batch_id)
#             
#             # Check batch limits
#             device_type_map = {
#                 'TV': 'max_tvs',
#                 'TOKEN_DISPENSER': 'max_dispensers',
#                 'KEYPAD': 'max_keypads',
#                 'BROKER': 'max_brokers',
#                 'LED': 'max_leds'
#             }
#             
#             limit_field = device_type_map.get(device.device_type)
#             if limit_field:
#                 max_limit = getattr(batch, limit_field, 0)
#                 current_count = Device.objects.filter(batch=batch, device_type=device.device_type).count()
#                 
#                 # Check if we are re-assigning to the same batch (count shouldn't increase)
#                 if device.batch != batch:
#                     if current_count >= max_limit:
#                         messages.error(request, f"Cannot assign device. Batch {batch.name} has reached its limit of {max_limit} for {device.get_device_type_display()}s.")
#                         return redirect('device_list')
# 
#             device.batch = batch
#             device.save()
#             messages.success(request, f"Device assigned to batch {batch.name}")
#         else:
#             # Handle unassignment if needed, or error
#             messages.error(request, "No batch selected")
#             
#         return redirect('device_list')
#     return redirect('device_list')

@login_required
def change_device_branch(request, device_id):
    """Change the branch of a device. Company Admin only, within same company."""
    if request.method != 'POST':
        return redirect('device_list')
    
    user = request.user
    if user.role != 'COMPANY_ADMIN':
        messages.error(request, "Only Company Admins can change device branch.")
        return redirect('device_list')
    
    device = get_object_or_404(Device, id=device_id)
    
    # Verify device belongs to user's company
    if device.company != user.company_relation:
        messages.error(request, "You don't have permission to modify this device.")
        return redirect('device_list')
    
    branch_id = request.POST.get('branch_id')
    if not branch_id:
        messages.error(request, "No branch selected.")
        return redirect('device_list')
    
    branch = get_object_or_404(Branch, id=branch_id)
    
    # Verify branch belongs to same company
    if branch.company != user.company_relation:
        messages.error(request, "Invalid branch selection.")
        return redirect('device_list')
    
    device.branch = branch
    device.save()
    messages.success(request, f"Device {device.serial_number} moved to {branch.branch_name}.")
    return redirect('device_list')

@login_required
def change_device_owner(request, device_id):
    """Change the company and branch of a device. Super Admin only.
    
    When the company changes, the device's configuration and all cross-company
    mappings are cleared so the new company starts with a clean slate.
    """
    if request.method != 'POST':
        return redirect('device_list')
    
    user = request.user
    if user.role != 'SUPER_ADMIN':
        messages.error(request, "Only Super Admins can change device ownership.")
        return redirect('device_list')
    
    device = get_object_or_404(Device, id=device_id)
    old_company_id = device.company_id
    company_id = request.POST.get('company_id')
    branch_id = request.POST.get('branch_id')
    
    if company_id:
        from companydetails.models import Company
        company = get_object_or_404(Company, id=company_id)
        device.company = company
    else:
        device.company = None
        
    if branch_id:
        from companydetails.models import Branch
        branch = get_object_or_404(Branch, id=branch_id)
        device.branch = branch
    else:
        device.branch = None
    
    new_company_id = device.company_id

    # If the company actually changed, wipe all company-specific configs and mappings
    # so the new company starts with a clean, uncontaminated device.
    if old_company_id != new_company_id:
        with transaction.atomic():
            # 1. Reset device config (clear stored JSON)
            if hasattr(device, 'config') and device.config:
                device.config.config_json = {}
                device.config.save()
            
            # 2. Detach from any config profile that belonged to the old company
            device.config_profile = None
            device.embedded_profile = None

            # 3. Remove from all groups (GroupDispenserMapping & GroupMapping M2M keypads/leds)
            if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
                # GroupDispenserMapping through-model rows
                GroupDispenserMapping.objects.filter(dispenser=device).delete()
                # GroupCounterButtonMapping rows for this dispenser
                GroupCounterButtonMapping.objects.filter(dispenser=device).delete()
                # Counter-Dispenser mappings
                CounterTokenDispenserMapping.objects.filter(dispenser=device).delete()
                # TV-Dispenser mappings
                TVDispenserMapping.objects.filter(dispenser=device).delete()
                # TVKeypadMapping dispenser references
                TVKeypadMapping.objects.filter(dispenser=device).update(dispenser=None)

            elif device.device_type == Device.DeviceType.KEYPAD:
                device.group_keypads.clear()
                TVKeypadMapping.objects.filter(keypad=device).delete()

            elif device.device_type == Device.DeviceType.TV:
                TVDispenserMapping.objects.filter(tv=device).delete()
                TVKeypadMapping.objects.filter(tv=device).delete()
                TVCounterMapping.objects.filter(tv=device).delete()

            elif device.device_type == Device.DeviceType.LED:
                device.group_leds.clear()

            device.save()
        messages.success(
            request,
            f"Device {device.serial_number} ownership updated. "
            f"Configuration and mappings have been reset for the new company."
        )
    else:
        device.save()
        messages.success(request, f"Device {device.serial_number} ownership updated successfully.")
    
    return redirect('device_list')

@login_required
@user_passes_test(company_required)
def device_register(request):
    user = request.user
    # Restrict dealer-created company admins and dealer customers from registering devices
    if user.role == "DEALER_CUSTOMER" or (user.role == "COMPANY_ADMIN" and user.dealer_customer_relation):
        messages.error(request, "Access denied. Restricted users cannot register devices.")
        return redirect('device_list')
    
    if request.method == 'POST':
        company_id = request.POST.get('company_id')
        branch_id = request.POST.get('branch_id')
        
        # devices_data should be a JSON string from the frontend
        devices_data_str = request.POST.get('devices_data', '[]')
        try:
            devices_data = json.loads(devices_data_str)
        except json.JSONDecodeError:
            devices_data = []

        if not branch_id:
            messages.error(request, "Branch is mandatory.")
            return redirect('device_register')

        if not devices_data:
            # Fallback for single device standard form submission
            mac = request.POST.get('mac_address')
            model = request.POST.get('device_model')
            dtype = request.POST.get('device_type')
            if mac and dtype:
                devices_data = [{
                    'mac_address': mac,
                    'device_type': dtype,
                    'device_model': model or dtype
                }]

        if not devices_data:
            messages.error(request, "No devices provided for registration.")
            return redirect('device_register')

        company = get_object_or_404(Company, id=company_id)
        
        # Sync live device limits from external license API (skip for dealer-created companies)
        if not company.is_dealer_created and company.company_id:
            try:
                from companydetails.views import sync_company_license_data
                api_data = LicenseManagementService.authenticate_product(company.company_id)
                if api_data and not api_data.get('error'):
                    sync_company_license_data(company, api_data, request.user)
                    company.refresh_from_db()
            except Exception as e:
                import logging
                _logger = logging.getLogger(__name__)
                _logger.warning(f"Failed to sync live device limits for {company.company_name}: {e}")
        
        success_count = 0
        fail_count = 0
        errors = []

        for dev in devices_data:
            mac = dev.get('mac_address')          # serial number (production batch SN)
            dtype = dev.get('device_type')
            model = dev.get('device_model') or dtype
            token_type = dev.get('token_type')     # No. of buttons for dispensers
            display_name = dev.get('display_name') # Customer-defined display name
            device_mac = dev.get('device_mac_address') or None  # Physical MAC address

            # 1. Validation against Production Batch
            prod_sn = ProductionSerialNumber.objects.filter(serial_number=mac).first()
            if not prod_sn:
                errors.append(f"Serial Number {mac} not found in any production batch.")
                fail_count += 1
                continue
            
            if prod_sn.device_type != dtype:
                errors.append(f"Serial Number {mac} type mismatch. Expected: {prod_sn.device_type}, Selected: {dtype}")
                fail_count += 1
                continue

            if prod_sn.is_registered:
                errors.append(f"Serial Number {mac} is already registered.")
                fail_count += 1
                continue

            # 1.5 Check Company License Limits (from authentication API)
            limit_field_map = {
                'TV': 'noof_television_devices',
                'TOKEN_DISPENSER': 'noof_token_dispensors',
                'KEYPAD': 'noof_keypad_devices',
                'BROKER': 'noof_broker_devices',
                'LED': 'noof_led_devices'
            }
            limit_field = limit_field_map.get(dtype)
            if limit_field:
                limit = getattr(company, limit_field, 0)
                # Count ALL registered devices for this company (not just active)
                current_count = Device.objects.filter(company=company, device_type=dtype).count()
                
                if current_count >= limit:
                     device_type_display = dtype.replace('_', ' ').title()
                     errors.append(f"Maximum number of {device_type_display} devices reached! Allowed: {limit}, Currently registered: {current_count}.")
                     fail_count += 1
                     continue


            # 2. Register with External API
            DEVICE_TYPE_MAP = {
                'TV': 2,
                'TOKEN_DISPENSER': 3,
                'KEYPAD': 4,
                'BROKER': 5,
                'LED': 6,
                'KEYPAD_4_BTN': 4,
                'KEYPAD_12_BTN': 4,
                'KEYPAD_16_BTN': 4,
            }
            device_type_numeric = DEVICE_TYPE_MAP.get(dtype, 1)
            
            details = {
                'product_registration_id': company.product_registration_id,
                'unique_identifier': company.unique_identifier,
                'device_model': model,
                'mac_address': mac,
                'device_type': device_type_numeric,
                'created_by': request.user.id
            }
            
            try:
                resp = LicenseManagementService.register_device(details)
                ext_id = resp.get('DeviceRegistrationId') if resp else None
                prod_type_id = resp.get('ProductTypeId') if resp else None
                
                # 3. Save locally - Initial status is always 'Pending', will be updated after license check
                new_device, created = Device.objects.update_or_create(
                    serial_number=mac,
                    defaults={
                        'company': company,
                        'display_name': display_name or None,
                        'device_type': dtype,
                        'device_model': model,
                        'token_type': token_type,       # Save token type
                        'mac_address': device_mac,      # Save physical MAC address
                        'device_registration_id': ext_id,
                        'product_type_id': prod_type_id,
                        'licence_status': 'Pending',
                        'is_active': False,
                        'branch_id': branch_id
                    }
                )
                
                # Mark as registered in production batch
                prod_sn.is_registered = True
                prod_sn.save()
                
                # 4. If device was registered successfully, check license status to get actual status
                if ext_id:
                    try:
                        from datetime import datetime, date
                        status_details = {
                            'product_registration_id': company.product_registration_id,
                            'device_registration_id': ext_id,
                            'product_type_id': str(prod_type_id) if prod_type_id else '',
                            'unique_identifier': company.unique_identifier or '',
                            'customer_id': company.company_id or '',
                            'project_name': company.company_name or 'CallQ'
                        }
                        status_resp = LicenseManagementService.check_device_status(status_details)
                        
                        if status_resp and 'error' not in status_resp:
                            # Parse licence_active_to date
                            licence_to_str = status_resp.get('LicenceActiveTo')
                            if licence_to_str:
                                try:
                                    licence_to = datetime.strptime(licence_to_str.split(' ')[0], '%Y-%m-%d').date()
                                    new_device.licence_active_to = licence_to
                                    
                                    # Determine status based on expiry date
                                    days_left = (licence_to - date.today()).days
                                    if days_left < 0:
                                        new_device.is_active = False
                                        new_device.licence_status = 'Expired'
                                    else:
                                        new_device.is_active = True
                                        new_device.licence_status = 'Active'
                                except (ValueError, IndexError):
                                    # Fallback - set as Active if we can't parse date
                                    new_device.is_active = True
                                    new_device.licence_status = 'Active'
                            else:
                                # No expiry date in response - use status from API
                                status_code = status_resp.get('Status', 0)
                                message = status_resp.get('Message', '')
                                if status_code == 1 and message.lower() == 'active':
                                    new_device.licence_status = 'Active'
                                    new_device.is_active = True
                                else:
                                    new_device.licence_status = message or 'Inactive'
                                    new_device.is_active = False
                            
                            # Update APK version if present
                            if status_resp.get('APKVersion'):
                                new_device.apk_version = status_resp.get('APKVersion')
                            
                            new_device.save()
                    except Exception as status_error:
                        # Log but don't fail registration if status check fails
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"License status check failed for new device {mac}: {status_error}")
                        # Keep device as Pending, user can manually refresh status later
                
                success_count += 1
                log_activity(request.user, "Device Registered", f"Device {mac} registered (Ext ID: {ext_id}, Status: {new_device.licence_status})")
            except Exception as e:
                errors.append(f"Error registering {mac}: {str(e)}")
                fail_count += 1

        if success_count > 0:
            messages.success(request, f"Successfully registered {success_count} devices.")
        if fail_count > 0:
            for err in errors[:5]: # Show first 5 errors
                messages.error(request, err)
            if len(errors) > 5:
                messages.error(request, f"...and {len(errors) - 5} more errors.")
        
        return redirect('device_list')

    # GET
    if request.user.role == "SUPER_ADMIN":
        companies = Company.objects.all()
        branches = [] # Load dynamically via JS
    elif request.user.role == "ADMIN":
        if request.user.assigned_state:
            companies = Company.objects.filter(state__in=request.user.assigned_state)
            branches = [] # Load dynamically via JS
        else:
            companies = []
            branches = []
    elif request.user.role == "COMPANY_ADMIN":
        companies = [request.user.company_relation]
        branches = Branch.objects.filter(company=request.user.company_relation)
    elif request.user.role == "DEALER_ADMIN":
        if request.user.company_relation:
            # Dealer can register for themselves or their customers
            companies = Company.objects.filter(
                Q(id=request.user.company_relation.id) | 
                Q(parent_company=request.user.company_relation)
            )
            branches = [] # Load dynamically via JS
        else:
            companies = []
            branches = []
    elif request.user.role == "PRODUCTION_ADMIN":
        # Production Admin can register devices for any company
        companies = Company.objects.all()
        branches = [] # Load dynamically via JS
    else:
        companies = []
        branches = []
    
    return render(request, 'configdetails/device_register.html', {
        'companies': companies,
        'branches': branches
    })

@login_required
def assign_device_branch(request, device_id):
    if request.method == 'POST':
        device = get_object_or_404(Device, id=device_id)
        branch_id = request.POST.get('branch_id')
        
        if branch_id:
            branch = get_object_or_404(Branch, id=branch_id)
            device.branch = branch
        else:
            device.branch = None
            
        device.save()
        messages.success(request, f"Device {device.serial_number} assigned to {branch.branch_name if branch_id else 'No Branch'}")
        
    return redirect('device_list')

def _create_tv_slot_mappings(group):
    """
    Create (or re-create) TVDispenserMapping and TVKeypadMapping rows for every
    TV in the group, assigning ASCII slot indices starting at chr(0x31)='1'.

    This is ALWAYS called when a group is created — regardless of whether the
    wizard or the auto-generation path was used — so that dispenser_button_index
    and keypad_index are never null.

    Rules:
      - Each group starts its own slot numbering from position 1.
      - Dispensers are distributed across TVs round-robin (one dispenser per TV slot).
      - Keypads are distributed the same way, linked to the corresponding dispenser.
    """
    dispensers_list = list(group.dispensers.all().order_by('id'))
    keypads_list    = list(group.keypads.all().order_by('id'))
    tvs_list        = list(group.tvs.all().order_by('id'))

    if not tvs_list:
        return  # Nothing to do without at least one TV

    # ---------------------------------------------------------------
    # Step A: TVDispenserMapping — dispenser slot on each TV
    # ---------------------------------------------------------------
    if dispensers_list:
        # Fresh per-group counter — always starts at position 1 (chr 0x31 = '1')
        tv_positions = {tv.id: 1 for tv in tvs_list}
        for idx, dispenser in enumerate(dispensers_list):
            tv = tvs_list[idx % len(tvs_list)]

            # Remove any pre-existing mapping for this dispenser (one-to-one constraint)
            TVDispenserMapping.objects.filter(dispenser=dispenser).delete()

            try:
                button_char = get_button_index_char(tv_positions[tv.id])
            except ValueError:
                button_char = chr(0x31)  # Fallback to '1'

            TVDispenserMapping.objects.create(
                tv=tv,
                dispenser=dispenser,
                button_index=button_char,
            )
            tv_positions[tv.id] += 1

        # Write / refresh GroupDispenserMapping rows so each dispenser in the group
        # has a stored sequential position (one row per dispenser, not per button).
        GroupDispenserMapping.objects.filter(group=group).delete()
        for pos, dispenser in enumerate(dispensers_list, start=1):
            try:
                grp_btn_char = get_button_index_char(pos)
            except ValueError:
                grp_btn_char = chr(0x31)  # Fallback
            GroupDispenserMapping.objects.create(
                group=group,
                dispenser=dispenser,
                dispenser_button_index=grp_btn_char,
            )

        # ── Build GroupCounterButtonMapping ──────────────────────────────────
        # One row per (group, dispenser, counter) with a group-wide sequential
        # ASCII button index.  The counter sequence NEVER resets between
        # dispensers — a 4-button dispenser claims 4 consecutive slots.
        GroupCounterButtonMapping.objects.filter(group=group).delete()
        grp_btn_pos = 0  # Group-wide counter, increments per counter (not per dispenser)
        for dispenser in dispensers_list:
            counter_qs = CounterTokenDispenserMapping.objects.filter(
                dispenser=dispenser
            ).select_related('counter').order_by('id')
            for cm in counter_qs:
                grp_btn_pos += 1
                try:
                    btn_char = get_button_index_char(grp_btn_pos)
                except ValueError:
                    btn_char = chr(0x31)
                GroupCounterButtonMapping.objects.get_or_create(
                    group=group,
                    dispenser=dispenser,
                    counter=cm.counter,
                    defaults={'button_index': btn_char},
                )

    # ---------------------------------------------------------------
    # Step B: TVKeypadMapping — keypad slot on each TV, linked to dispenser
    # ---------------------------------------------------------------
    if keypads_list:
        # Collect per-TV slot counters, seeded from any already-existing explicit mappings
        # so new round-robin entries don't collide with manually assigned slot indices.
        tv_kp_positions = {tv.id: 1 for tv in tvs_list}

        # Pre-seed positions from existing explicit mappings so we don't reuse taken slots
        for tv in tvs_list:
            existing_count = TVKeypadMapping.objects.filter(tv=tv).count()
            if existing_count > 0:
                tv_kp_positions[tv.id] = existing_count + 1

        for idx, keypad in enumerate(keypads_list):
            # If an explicit TVKeypadMapping already exists for this keypad (set via the
            # Group Editor UI), PRESERVE it — do NOT overwrite with round-robin logic.
            if TVKeypadMapping.objects.filter(keypad=keypad).exists():
                continue

            # No explicit assignment — apply round-robin default
            tv = tvs_list[idx % len(tvs_list)]

            # Link the corresponding dispenser (round-robin if fewer dispensers than keypads)
            dispenser = dispensers_list[idx % len(dispensers_list)] if dispensers_list else None

            try:
                keypad_char = get_button_index_char(tv_kp_positions[tv.id])
            except ValueError:
                keypad_char = chr(0x31)  # Fallback to '1'

            TVKeypadMapping.objects.create(
                tv=tv,
                keypad=keypad,
                dispenser=dispenser,
                keypad_index=keypad_char,
            )
            tv_kp_positions[tv.id] += 1


def create_group_button_mappings(group, company, branch, dealer_customer):
    """
    Linear mapping hierarchy (ButtonMapping rows only — steps 1-4):
    1. Token Dispenser Button N -> Keypad[N] (sequential across ALL dispensers)
       e.g. 1 dispenser with 4 buttons -> 4 keypads (one per button)
    2. Keypad -> Broker ("Main" — linear, one broker per keypad)
    3. Keypad -> LED   ("LED"  — linear, one LED per keypad)
    4. Broker -> TV    ("Main" — linear, one TV per broker)

    TV slot index mappings (steps 5 & 6) are handled by _create_tv_slot_mappings
    which is always called unconditionally from the group creation view.
    """
    dispensers_list = list(group.dispensers.all().order_by('id'))
    keypads_list    = list(group.keypads.all().order_by('id'))
    leds_list       = list(group.leds.all().order_by('id'))
    brokers_list    = list(group.brokers.all().order_by('id'))
    tvs_list        = list(group.tvs.all().order_by('id'))

    # 1. Dispenser Button N -> Keypad[sequential]
    # Each button of each dispenser claims the NEXT keypad slot in order.
    keypad_idx = 0
    for dispenser in dispensers_list:
        button_count = 1
        if dispenser.token_type:
            try:
                button_count = int(dispenser.token_type.split('_')[0])
            except (ValueError, IndexError):
                button_count = 1
        for btn_num in range(1, button_count + 1):
            if keypad_idx < len(keypads_list):
                keypad = keypads_list[keypad_idx]
                ButtonMapping.objects.get_or_create(
                    company=company, branch=branch,
                    dealer_customer=dealer_customer,
                    source_device=dispenser,
                    source_button=f"Button {btn_num}",
                    defaults={'target_device': keypad}
                )
            keypad_idx += 1

    # 2. Keypad -> Broker ("Main", linear — modulo if fewer brokers than keypads)
    if brokers_list:
        for idx, keypad in enumerate(keypads_list):
            broker = brokers_list[idx % len(brokers_list)]
            ButtonMapping.objects.get_or_create(
                company=company, branch=branch,
                dealer_customer=dealer_customer,
                source_device=keypad,
                source_button="Main",
                defaults={'target_device': broker}
            )

    # 3. Keypad -> LED ("LED", linear — modulo if fewer LEDs than keypads)
    if leds_list:
        for idx, keypad in enumerate(keypads_list):
            led = leds_list[idx % len(leds_list)]
            ButtonMapping.objects.get_or_create(
                company=company, branch=branch,
                dealer_customer=dealer_customer,
                source_device=keypad,
                source_button="LED",
                defaults={'target_device': led}
            )

    # 4. Broker -> TV ("Main", linear — modulo if fewer TVs than brokers)
    if tvs_list:
        for idx, broker in enumerate(brokers_list):
            tv = tvs_list[idx % len(tvs_list)]
            ButtonMapping.objects.get_or_create(
                company=company, branch=branch,
                dealer_customer=dealer_customer,
                source_device=broker,
                source_button="Main",
                defaults={'target_device': tv}
            )

    # NOTE: TV slot index mappings (TVDispenserMapping + TVKeypadMapping) are
    # created by _create_tv_slot_mappings(), called unconditionally by the
    # group creation view — so they are always set even when the wizard is used.


@api_view(['GET'])
@login_required
def get_branch_group_mapping_api(request, branch_id):
    """
    Return the most-recent GroupMapping for a branch, including:
    - device lists per type
    - existing ButtonMapping rows for those devices
    Used by the mapping wizard to auto-fetch existing configuration.
    """
    branch = get_object_or_404(Branch, id=branch_id)
    group = GroupMapping.objects.filter(branch=branch).order_by('-created_at').first()

    if not group:
        return Response({'exists': False})

    def _device_dict(d):
        return {
            'id': d.id,
            'serial_number': d.serial_number,
            'display_name': d.display_name,
            'get_display_identifier': d.get_display_identifier,
            'device_model': d.device_model,
            'device_type': d.device_type,
            'token_type': d.token_type,
        }

    devices_by_type = {}
    for dtype, qs in [
        ('TOKEN_DISPENSER', group.dispensers.all()),
        ('KEYPAD',          group.keypads.all()),
        ('BROKER',          group.brokers.all()),
        ('TV',              group.tvs.all()),
        ('LED',             group.leds.all()),
    ]:
        devs = [_device_dict(d) for d in qs]
        if devs:
            devices_by_type[dtype] = devs

    all_group_devices = (
        list(group.dispensers.all()) + list(group.keypads.all()) +
        list(group.brokers.all())    + list(group.tvs.all()) +
        list(group.leds.all())
    )
    bm_qs = ButtonMapping.objects.filter(
        branch=branch, source_device__in=all_group_devices
    ).select_related('source_device', 'target_device')

    button_mappings = [
        {
            'id': m.id,
            'source_device_id': m.source_device.id,
            'source_button': m.source_button,
            'target_device_id': m.target_device.id,
            'target_device_type': m.target_device.device_type,
        }
        for m in bm_qs
    ]

    return Response({
        'exists': True,
        'group_id': group.id,
        'group_name': group.group_name,
        'no_of_dispensers': group.no_of_dispensers,
        'no_of_keypads':    group.no_of_keypads,
        'no_of_tvs':        group.no_of_tvs,
        'no_of_brokers':    group.no_of_brokers,
        'no_of_leds':       group.no_of_leds,
        'devices_by_type':  devices_by_type,
        'button_mappings':  button_mappings,
    })

@login_required
def mapping_view(request):
    user = request.user
    company = None
    if user.role == "SUPER_ADMIN":
        company = None # Handle superadmin selection if needed
        mappings = GroupMapping.objects.all()
        branches = Branch.objects.all()
    elif user.role == "ADMIN":
        if user.assigned_state:
            company = None # ADMIN might manage multiple companies in a state
            mappings = GroupMapping.objects.filter(company__state__in=user.assigned_state)
            branches = Branch.objects.filter(company__state__in=user.assigned_state)
        else:
            mappings = GroupMapping.objects.none()
            branches = []
    elif user.role == "COMPANY_ADMIN":
        company = user.company_relation
        mappings = GroupMapping.objects.filter(company=company)
        branches = Branch.objects.filter(company=company)
    elif user.role == "BRANCH_ADMIN":
        company = user.branch_relation.company
        mappings = GroupMapping.objects.filter(branch=user.branch_relation)
        branches = [user.branch_relation]
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            mappings = GroupMapping.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
            branches = Branch.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
        else:
            mappings = GroupMapping.objects.none()
            branches = []
    else:
        mappings = GroupMapping.objects.none()
        branches = []

    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'create':
            # Handle Dealer Customer logic
            dealer_customer_id = request.POST.get('dealer_customer_id')
            branch_id = request.POST.get('branch_id')
            group_name = request.POST.get('group_name')
            
            branch = None
            dealer_customer = None
            mapping_company = None

            if branch_id:
                branch = get_object_or_404(Branch, id=branch_id)
                mapping_company = branch.company
            elif dealer_customer_id:
                from companydetails.models import DealerCustomer
                dealer_customer = get_object_or_404(DealerCustomer, id=dealer_customer_id)
                mapping_company = dealer_customer.dealer
            
            if not mapping_company or not group_name:
                messages.error(request, "Company and Group Name are required.")
                return redirect('mapping_view')

            # Get device lists from form (for group mapping)
            dispensers = request.POST.getlist('dispensers[]')
            keypads = request.POST.getlist('keypads[]')
            tvs = request.POST.getlist('tvs[]')
            brokers = request.POST.getlist('brokers[]')
            leds = request.POST.getlist('leds[]')
            
            # Parse quantities
            no_of_dispensers = int(request.POST.get('qty_dispenser', 1) or 1)
            no_of_keypads = int(request.POST.get('qty_keypad', 1) or 1)
            no_of_tvs = int(request.POST.get('qty_tv', 1) or 1)
            no_of_brokers = int(request.POST.get('qty_broker', 0) or 0)
            no_of_leds = int(request.POST.get('qty_led', 0) or 0)

            # Validation: Check device counts match quantities
            error_msgs = []
            
            # Filter out empty strings from device lists
            dispensers = [d for d in dispensers if d]
            keypads = [k for k in keypads if k]
            tvs = [t for t in tvs if t]
            brokers = [b for b in brokers if b]
            leds = [l for l in leds if l]
            
            # Validate device counts
            if len(dispensers) != no_of_dispensers:
                error_msgs.append(f"Token Dispenser count mismatch: Expected {no_of_dispensers}, but {len(dispensers)} device(s) selected.")
            
            if len(keypads) != no_of_keypads:
                error_msgs.append(f"Keypad count mismatch: Expected {no_of_keypads}, but {len(keypads)} device(s) selected.")
            
            if no_of_leds > 0 and len(leds) != no_of_leds:
                error_msgs.append(f"LED count mismatch: Expected {no_of_leds}, but {len(leds)} device(s) selected.")
            elif no_of_leds == 0 and len(leds) > 0:
                error_msgs.append(f"LED count mismatch: Expected 0, but {len(leds)} device(s) selected.")
            
            if no_of_brokers > 0 and len(brokers) != no_of_brokers:
                error_msgs.append(f"Broker count mismatch: Expected {no_of_brokers}, but {len(brokers)} device(s) selected.")
            elif no_of_brokers == 0 and len(brokers) > 0:
                error_msgs.append(f"Broker count mismatch: Expected 0, but {len(brokers)} device(s) selected.")
            
            if len(tvs) != no_of_tvs:
                error_msgs.append(f"TV count mismatch: Expected {no_of_tvs}, but {len(tvs)} device(s) selected.")

            def check_device_in_group(device_ids, device_type_name):
                """Check if devices are already in another group"""
                in_use = []
                for d_id in device_ids:
                    if d_id:
                        d = Device.objects.filter(id=d_id).first()
                        if d:
                            # Check if device is already in a GroupMapping
                            if device_type_name == 'Token Dispenser':
                                groups = d.group_dispensers.all()
                            elif device_type_name == 'Keypad':
                                groups = d.group_keypads.all()
                            elif device_type_name == 'LED':
                                groups = d.group_leds.all()
                            else:
                                continue  # Broker and TV can be in multiple groups
                            
                            if groups.exists():
                                in_use.append(d.get_display_identifier())
                if in_use:
                    error_msgs.append(f"The following {device_type_name}s are already mapped to another group: {', '.join(in_use)}.")

            check_device_in_group(dispensers, 'Token Dispenser')
            check_device_in_group(keypads, 'Keypad')
            check_device_in_group(leds, 'LED')

            if error_msgs:
                for msg in error_msgs:
                    messages.error(request, msg)
                return redirect('mapping_view')

            # Validation: Check if Token Dispenser, Keypad, or LED are already in another group
            
            # Create GroupMapping
            group = GroupMapping.objects.create(
                group_name=group_name,
                company=mapping_company,
                branch=branch,
                dealer_customer=dealer_customer,
                no_of_dispensers=no_of_dispensers,
                no_of_keypads=no_of_keypads,
                no_of_tvs=no_of_tvs,
                no_of_brokers=no_of_brokers,
                no_of_leds=no_of_leds
            )

            # Helper to assign dispensers using the through model (stores button index)
            def assign_dispensers(device_ids, group_obj):
                pos = 0
                for d_id in device_ids:
                    if d_id:
                        d = Device.objects.filter(id=d_id).first()
                        if d:
                            pos += 1
                            try:
                                btn_char = get_button_index_char(pos)
                            except ValueError:
                                btn_char = chr(0x31)
                            GroupDispenserMapping.objects.get_or_create(
                                group=group_obj,
                                dispenser=d,
                                defaults={'dispenser_button_index': btn_char},
                            )

            # Helper to assign non-dispenser devices to group
            def assign_devices(device_ids, m2m_field):
                for d_id in device_ids:
                    if d_id:
                        d = Device.objects.filter(id=d_id).first()
                        if d:
                            m2m_field.add(d)

            assign_dispensers(dispensers, group)
            assign_devices(keypads, group.keypads)
            assign_devices(tvs, group.tvs)
            assign_devices(brokers, group.brokers)
            assign_devices(leds, group.leds)

            # session_mapping_ids: IDs of ButtonMapping rows already saved by the
            # wizard during Steps 2-5.  When present, the user has explicitly
            # configured every button — skip auto-generation to avoid duplicates.
            session_mapping_ids = [
                sid for sid in request.POST.getlist('session_mapping_ids[]') if sid
            ]

            if session_mapping_ids:
                # The wizard already persisted the user's button choices.
                # We only need to ensure those ButtonMapping rows are associated
                # with the correct company / branch (they already are since
                # save_button_mapping_api used the same context).
                log_activity(
                    request.user, "Group Created",
                    f"New group '{group_name}' created using {len(session_mapping_ids)} "
                    f"wizard-configured button mapping(s)."
                )
                messages.success(
                    request,
                    f"Group '{group_name}' created successfully with your custom button mappings."
                )
            else:
                # No wizard mappings were saved — fall back to auto-generation.
                create_group_button_mappings(group, mapping_company, branch, dealer_customer)
                log_activity(
                    request.user, "Group Created",
                    f"New group '{group_name}' created with automatic button mappings."
                )
                messages.success(
                    request,
                    f"Group '{group_name}' created successfully with automatic button mappings."
                )

            # Always create TV slot index mappings (TVDispenserMapping + TVKeypadMapping)
            # regardless of wizard or auto-generation path, so dispenser_button_index
            # and keypad_index are never null after group creation.
            _create_tv_slot_mappings(group)
            
        elif action == 'unmap':
            mapping_id = request.POST.get('mapping_id')
            GroupMapping.objects.filter(id=mapping_id).delete()
            log_activity(request.user, "Group Deleted", f"Group Mapping {mapping_id} deleted")
            messages.success(request, "Group deleted successfully.")
            
        return redirect('mapping_view')

    # Get available devices for the company
    if user.role == "SUPER_ADMIN":
        devices = Device.objects.all()
        companies = Company.objects.all()
    elif user.role == "ADMIN":
        if user.assigned_state:
            devices = Device.objects.filter(company__state__in=user.assigned_state)
            companies = Company.objects.filter(state__in=user.assigned_state)
        else:
            devices = Device.objects.none()
            companies = Company.objects.none()
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            devices = Device.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
            companies = Company.objects.filter(
                Q(id=user.company_relation.id) |
                Q(parent_company=user.company_relation)
            )
        else:
            devices = Device.objects.none()
            companies = Company.objects.none()
    elif user.role == "COMPANY_ADMIN":
        devices = Device.objects.filter(company=user.company_relation)
        companies = [user.company_relation] if user.company_relation else []
    elif user.role == "BRANCH_ADMIN":
        devices = Device.objects.filter(branch=user.branch_relation)
        companies = [user.branch_relation.company] if user.branch_relation else []
    else:
        devices = Device.objects.none()
        companies = Company.objects.none()
    
    # Determine if this is a dealer view (for showing customer mapping features)
    is_dealer_view = user.role == "DEALER_ADMIN"
    
    # Calculate statistics based on role
    total_devices = devices.count()
    
    if is_dealer_view:
        # For dealers: count devices mapped to dealer customers
        mapped_count = devices.filter(dealer_customer__isnull=False).count()
    else:
        # For company admins: count devices that are in a GroupMapping
        mapped_count = devices.filter(
            Q(group_dispensers__isnull=False) | 
            Q(group_keypads__isnull=False) | 
            Q(group_leds__isnull=False) |
            Q(group_brokers__isnull=False) |
            Q(group_tvs__isnull=False)
        ).distinct().count()
    
    unmapped_count = total_devices - mapped_count
    coverage_percent = int((mapped_count / total_devices * 100)) if total_devices > 0 else 0
    
    # Get dealer customers for the dropdown (only for DEALER_ADMIN)
    from companydetails.models import DealerCustomer
    dealer_customers = []
    if is_dealer_view and user.company_relation:
        dealer_customers = DealerCustomer.objects.filter(dealer=user.company_relation, is_active=True)
    
    # Get devices by type for the mapping form (for non-dealer users)
    tvs = devices.filter(device_type=Device.DeviceType.TV)
    token_dispensers = devices.filter(device_type=Device.DeviceType.TOKEN_DISPENSER)
    keypads = devices.filter(device_type=Device.DeviceType.KEYPAD)
    brokers = devices.filter(device_type=Device.DeviceType.BROKER)
    leds = devices.filter(device_type=Device.DeviceType.LED)
    
    context = {
        'devices': devices,
        'mappings': mappings,
        'branches': branches,
        'companies': companies,
        'dealer_customers': dealer_customers,
        'is_dealer_view': is_dealer_view,  # Flag to show/hide dealer-specific features
        'total_devices': total_devices,
        'mapped_count': mapped_count,
        'unmapped_count': unmapped_count,
        'coverage_percent': coverage_percent,
        # Determine available batches for context (assuming company is known)
        'batches': Batch.objects.filter(customer=company) if company else [],
        # Device lists by type for mapping form
        'tvs': tvs,
        'token_dispensers': token_dispensers,
        'keypads': keypads,
        'brokers': brokers,
        'leds': leds,
    }

    # Pagination - 8 rows per page
    page = request.GET.get('page', 1)
    paginator = Paginator(list(devices), 8)
    try:
        devices_page = paginator.page(page)
    except PageNotAnInteger:
        devices_page = paginator.page(1)
    except EmptyPage:
        devices_page = paginator.page(paginator.num_pages)
    
    context['devices'] = devices_page

    return render(request, 'configdetails/mapping.html', context)


@login_required
def map_device_to_customer(request):
    """View to bulk assign devices to a dealer customer."""
    user = request.user
    
    # Permission check - DEAlER_ADMIN only
    if user.role != "DEALER_ADMIN":
        messages.error(request, "Permission denied. Only Dealer Admins can access this page.")
        return redirect('device_list')
    
    company = user.company_relation
    if not company:
         messages.error(request, "Dealer context associated with user not found.")
         return redirect('device_list')

    from companydetails.models import DealerCustomer

    if request.method == 'POST':
        dealer_customer_id = request.POST.get('dealer_customer_id')
        device_ids = request.POST.getlist('device_ids')
        
        if dealer_customer_id and device_ids:
            dealer_customer = get_object_or_404(DealerCustomer, id=dealer_customer_id)
            
            # Verify ownership
            if dealer_customer.dealer != company:
                messages.error(request, "Invalid Customer.")
                return redirect('map_device_to_customer')
                
            # Update devices
            # Ensure we only update devices owned by this dealer
            count = Device.objects.filter(
                id__in=device_ids, 
                company=company
            ).update(dealer_customer=dealer_customer, branch=None) # Reset branch when assigning to customer
            
            messages.success(request, f"{count} devices assigned to {dealer_customer.company_name} successfully.")
        else:
            messages.error(request, "Please select a customer and at least one device.")
            
        return redirect('map_device_to_customer')
    
    # GET: Prepare configuration form
    dealer_customers = DealerCustomer.objects.filter(dealer=company, is_active=True)
    devices = Device.objects.filter(
        Q(company=company) | Q(company__parent_company=company)
    ).order_by('dealer_customer', 'serial_number')
        
    context = {
        'devices': devices,
        'dealer_customers': dealer_customers
    }
    return render(request, 'configdetails/map_device_to_customer.html', context)


@login_required
def unmap_device(request):
    """Remove device-to-customer mapping."""
    if request.method == 'POST':
        device_id = request.POST.get('device_id')
        
        if device_id:
            device = get_object_or_404(Device, id=device_id)
            
            # Verify permission
            if request.user.role == "DEALER_ADMIN" and request.user.company_relation:
                if device.company == request.user.company_relation:
                    customer_name = device.dealer_customer.company_name if device.dealer_customer else "None"
                    device.dealer_customer = None
                    device.save()
                    messages.success(request, f"Device {device.serial_number} unmapped from {customer_name}")
                    log_activity(request.user, "Device Unmapped", f"Device {device.serial_number} unmapped")
                else:
                    messages.error(request, "Permission denied.")
            else:
                messages.error(request, "Only dealers can unmap devices.")
    
    return redirect('mapping_view')


@login_required
def device_delete(request, device_id):
    """Unassign (soft delete) a device."""
    user = request.user
    # Restrict dealer-created company admins from deleting devices
    if user.role == "DEALER_CUSTOMER" or (user.role == "COMPANY_ADMIN" and user.dealer_customer_relation):
        messages.error(request, "Access denied. Restricted users cannot delete devices.")
        return redirect('device_list')

    if request.method == 'POST':
        device = get_object_or_404(Device, id=device_id)
        
        # Check permissions
        can_delete = False
        if user.role == "SUPER_ADMIN":
            can_delete = True
        elif user.role == "ADMIN" and user.assigned_state:
            if device.company and device.company.state in user.assigned_state:
                can_delete = True
        elif user.role == "COMPANY_ADMIN":
            if device.company == user.company_relation:
                can_delete = True
        elif user.role == "DEALER_ADMIN":
            # Dealer can delete if it's their own device OR if it belongs to a company they created
            if device.company == user.company_relation:
                can_delete = True
            elif device.company and device.company.is_dealer_created and device.company.parent_company == user.company_relation:
                can_delete = True
        
        if can_delete:
            serial = device.serial_number
            
            # Reset the production serial number status
            from .models import ProductionSerialNumber
            prod_sn = ProductionSerialNumber.objects.filter(serial_number=serial).first()
            if prod_sn:
                prod_sn.is_registered = False
                prod_sn.save()

            # Clear configs and mappings for the next user
            DeviceConfig.objects.filter(device=device).delete()
            TVConfig.objects.filter(tv=device).delete()
            Mapping.objects.filter(Q(token_dispenser=device) | Q(tv=device) | Q(keypad=device) | Q(broker=device) | Q(led=device) | Q(keypad_2=device) | Q(keypad_3=device) | Q(keypad_4=device)).delete()
            ButtonMapping.objects.filter(Q(source_device=device) | Q(target_device=device)).delete()
            
            # Unassign the device
            device.company = None
            device.branch = None
            device.dealer_customer = None
            device.device_registration_id = None
            device.licence_status = 'Pending'
            device.is_active = False
            device.save()
            
            messages.success(request, f"Device {serial} unassigned/deleted successfully.")
            log_activity(request.user, "Device Deleted (Unassigned)", f"Device {serial} unassigned")
        else:
            messages.error(request, "Permission denied.")
    
    return redirect('device_list')


@api_view(['POST'])
@login_required
def device_authenticate_api(request, device_id):
    """
    Register device with external license server (similar to customer authentication).
    """
    device = get_object_or_404(Device, id=device_id)
    
    if not device.company:
        return Response({'status': 'error', 'message': 'Device not assigned to a company'}, status=400)
    
    company = device.company
    
    if not company.product_registration_id:
        return Response({'status': 'error', 'message': 'Company not registered with license server. Please authenticate the company first.'}, status=400)
    
    # Map device type string to numeric value for the external API
    # These must match the DEVICE_MODEL_MAP in core/services.py
    DEVICE_TYPE_MAP = {
        'TV': 2,                # Android TV
        'TOKEN_DISPENSER': 3,   # Token Dispenser
        'KEYPAD': 4,            # Keypad / Embedded
        'BROKER': 5,            # Broker
        'LED': 6,               # LED Device
    }
    
    # Get the numeric device type for the API (default to 1 for PC/unknown)
    device_type_numeric = DEVICE_TYPE_MAP.get(device.device_type, 1)
    
    # Register device with external API - include company details
    details = {
        'product_registration_id': company.product_registration_id,
        'unique_identifier': company.unique_identifier,
        'device_model': device.device_model or device.get_device_type_display(),  # Use device type display name if no model set
        'mac_address': device.serial_number,
        'device_type': device_type_numeric,  # Correctly mapped device type
        'created_by': request.user.id,
        # Include company details in the payload
        'customer_name': company.company_name or '',
        'customer_contact_person': company.contact_person or '',
        'customer_address': company.address or '',
        'customer_city': company.city or '',
        'customer_state': company.state or '',
        'customer_zip': company.zip_code or '',
        'customer_contact': company.contact_number or '',
        'customer_email': company.company_email or '',
    }
    
    try:
        resp = LicenseManagementService.register_device(details)
        if resp and resp.get('DeviceRegistrationId'):
            device.device_registration_id = resp.get('DeviceRegistrationId')
            device.product_type_id = resp.get('ProductTypeId')
            device.licence_status = 'Active' # Activate device upon successful authentication
            device.is_active = True
            device.save()
            
            log_activity(request.user, "Device Authenticated", f"Device {device.serial_number} registered with license server (ID: {device.device_registration_id})")
            
            return Response({
                'status': 'success',
                'message': f'Device registered successfully (ID: {device.device_registration_id})',
                'device_registration_id': device.device_registration_id
            })
        else:
            return Response({'status': 'error', 'message': 'Registration failed - no ID returned'}, status=400)
    except Exception as e:
        return Response({'status': 'error', 'message': f'Registration failed: {str(e)}'}, status=500)


@api_view(['POST'])
@login_required
def check_device_status_api(request, device_id):
    """
    Refresh device status with external API.
    1. Calls DeviceRegistration API to update/refresh registration
    2. Calls CheckDeviceStatus API to get current license status
    Updates licence_active_to, is_active based on API responses.
    """
    from datetime import datetime, date
    import logging
    logger = logging.getLogger(__name__)
    
    device = get_object_or_404(Device, id=device_id)
    company = device.company
    
    if not company:
        return Response({'status': 'error', 'message': 'Device not assigned to a company'}, status=400)
    
    if not company.product_registration_id:
        return Response({'status': 'error', 'message': 'Company not registered with license server. Please authenticate the company first.'}, status=400)
    
    # Map device type string to numeric value for the external API
    DEVICE_TYPE_MAP = {
        'TV': 2,
        'TOKEN_DISPENSER': 3,
        'KEYPAD': 4,
        'BROKER': 5,
        'LED': 6,
    }
    device_type_numeric = DEVICE_TYPE_MAP.get(device.device_type, 1)
    
    # Step 1: Call DeviceRegistration API to update/refresh registration
    reg_details = {
        'product_registration_id': company.product_registration_id,
        'unique_identifier': company.unique_identifier,
        'device_model': device.device_model or device.get_device_type_display(),
        'mac_address': device.serial_number,
        'device_type': device_type_numeric,
        'created_by': request.user.id,
        'customer_name': company.company_name or '',
        'customer_contact_person': company.contact_person or '',
        'customer_address': company.address or '',
        'customer_city': company.city or '',
        'customer_state': company.state or '',
        'customer_zip': company.zip_code or '',
        'customer_contact': company.contact_number or '',
        'customer_email': company.company_email or '',
    }
    
    try:
        reg_resp = LicenseManagementService.register_device(reg_details)
        logger.info(f"DeviceRegistration API Response for device {device_id}: {reg_resp}")
        
        if reg_resp and reg_resp.get('DeviceRegistrationId'):
            device.device_registration_id = reg_resp.get('DeviceRegistrationId')
            device.product_type_id = reg_resp.get('ProductTypeId')
            device.save()
    except Exception as e:
        logger.warning(f"DeviceRegistration API failed for device {device_id}: {e}")
        # Continue to check status even if registration fails
    
    # Step 2: Call CheckDeviceStatus API
    if not device.device_registration_id:
        return Response({'status': 'error', 'message': 'Device registration failed. No registration ID.'}, status=400)
    
    status_details = {
        'product_registration_id': company.product_registration_id,
        'device_registration_id': device.device_registration_id,
        'product_type_id': str(device.product_type_id) if device.product_type_id else '',
        'unique_identifier': company.unique_identifier or '',
        'customer_id': company.company_id or '',  # Use external company_id from license portal
        'project_name': device.project_name or company.company_name or ''
    }
    
    try:
        resp = LicenseManagementService.check_device_status(status_details)
        logger.info(f"CheckDeviceStatus API Response for device {device_id}: {resp}")
        
        if resp and 'error' not in resp:
            # Update device with response data
            status_code = resp.get('Status', 0)
            message = resp.get('Message', '')
            
            # Parse licence_active_to date
            licence_to_str = resp.get('LicenceActiveTo')
            if licence_to_str:
                try:
                    licence_to = datetime.strptime(licence_to_str.split(' ')[0], '%Y-%m-%d').date()
                    device.licence_active_to = licence_to
                except (ValueError, IndexError):
                    pass
            
            # Update other fields if present
            if resp.get('APKVersion'):
                device.apk_version = resp.get('APKVersion')
            
            # Determine status - prioritize actual date over API message
            days_left = None
            expiry_warning = None
            
            if device.licence_active_to:
                days_left = (device.licence_active_to - date.today()).days
                
                if days_left < 0:
                    device.is_active = False
                    device.licence_status = 'Expired'
                    expiry_warning = f'License expired on {device.licence_active_to}!'
                elif days_left <= 10:
                    device.is_active = True
                    device.licence_status = 'Active'
                    expiry_warning = f'License expiring in {days_left} days!'
                else:
                    device.is_active = True
                    device.licence_status = 'Active'
            else:
                # No expiry date - use API status/message
                if status_code == 1 and message.lower() == 'active':
                    device.licence_status = 'Active'
                    device.is_active = True
                else:
                    device.licence_status = message or 'Inactive'
                    device.is_active = False
            
            device.save()
            
            log_activity(request.user, "Device Refreshed", f"Device {device.serial_number} - Status: {device.licence_status}")
            
            return Response({
                'status': 'success',
                'message': f'Device status: {device.licence_status}',
                'licence_status': device.licence_status,
                'licence_active_to': str(device.licence_active_to) if device.licence_active_to else None,
                'is_active': device.is_active,
                'days_until_expiry': days_left,
                'expiry_warning': expiry_warning
            })
        else:
            error_msg = resp.get('error', 'Unknown error') if resp else 'No response from server'
            return Response({'status': 'error', 'message': f'Status check failed: {error_msg}'}, status=400)
            
    except Exception as e:
        return Response({'status': 'error', 'message': f'Status check failed: {str(e)}'}, status=500)

# API for dynamic loading
@api_view(['POST'])
@login_required
def get_available_devices_api(request):
    branch_id = request.data.get('branch_id')
    batch_id = request.data.get('batch_id')
    
    if not branch_id:
        return Response({'error': 'Branch ID required'}, status=400)
        
    branch = get_object_or_404(Branch, id=branch_id)
    company = branch.company
    
    # Base query: Devices for this company/branch
    # Note: Devices might be assigned to branch or just company. 
    # If assigned to branch, filter by branch. If just company, filter by company but unassigned?
    # Logic: "Available for mapping in this branch". 
    # If device is already assigned to THIS branch, show it.
    # If device is pending/unassigned branch, maybe allow mapping?
    # Usually devices are assigned to branch first (Device Register or Assign Branch).
    # So we filter matching branch_id or NULL branch (if company matches)?
    # To keep it strict: Devices must be assigned to the selected branch logic?
    # Or strict: filter(branch=branch).
    
    devices = Device.objects.filter(company=company)
    
    # Filter by Batch if provided
    if batch_id:
        devices = devices.filter(batch_id=batch_id)
    else:
        # If no batch selected? Maybe require batch.
        pass
        
    # Exclude already mapped devices
    # Logic: device.mapped_tvs.exists() etc.
    # We need to exclude devices that are present in ANY Mapping object.
    
    # Get IDs of devices in Mappings
    mapped_tvs = Mapping.objects.values_list('tv', flat=True)
    mapped_tds = Mapping.objects.values_list('token_dispenser', flat=True)
    mapped_kps = Mapping.objects.values_list('keypad', flat=True) # Check 1-4?
    # Complex to check all keypad columns efficiently, but doable.
    # Using exclude(mapping__isnull=False) might work if reverse relation exists.
    # 'mapping' related name standard? Let's check models. 
    # Usually Django adds related_name 'mapping_set' or similar if not defined.
    # Step 22 `Mapping` model:
    # tv = ForeignKey(..., related_name='mapped_tvs')
    # token_dispenser = ForeignKey(..., related_name='mapped_dispensers') (Wait, checked model?)
    # Models snippet in Step 22 didn't show related_names. I should assume default or check.
    # Step 89 used `mapped_tvs__isnull=True`. So related_name='mapped_tvs' exists.
    # Let's assume standard names or specific ones I saw.
    
    tvs = devices.filter(device_type='TV', mapped_tvs__isnull=True).values('id', 'serial_number')
    
    # Token Dispensers
    # Assuming related_name='mapped_token_dispensers' or 'mapping_set'?
    # View in Step 89 used `mapped_tvs`.
    # Let's check `configdetails/models.py` quickly if needed, or deduce.
    # Previous view code (Step 89 Update):
    # unmapped_devices = Device.objects.filter(mapped_tvs__isnull=True, mapped_sources__isnull=True ... )
    # mapped_sources? `Mapping` has `source`? Step 22 didn't show `source` field.
    # Maybe `mapped_tvs` is correct for TV.
    # For others, I'll use exclusion by ID list to be safe if names unverified.
    
    # Recalculate IDs
    qs_maps = Mapping.objects.all()
    used_ids = set()
    for m in qs_maps:
        if m.tv_id: used_ids.add(m.tv_id)
        if m.token_dispenser_id: used_ids.add(m.token_dispenser_id)
        if m.broker_id: used_ids.add(m.broker_id)
        if m.led_id: used_ids.add(m.led_id)
        if m.keypad_id: used_ids.add(m.keypad_id)
        if m.keypad_2_id: used_ids.add(m.keypad_2_id)
        if m.keypad_3_id: used_ids.add(m.keypad_3_id)
        if m.keypad_4_id: used_ids.add(m.keypad_4_id)
        
    def serialize(qs):
        return [{'serial_number': d.serial_number, 'id': d.id} for d in qs if d.id not in used_ids]
        
    return Response({
        'tvs': serialize(devices.filter(device_type='TV')),
        'token_dispensers': serialize(devices.filter(device_type='TOKEN_DISPENSER')),
        'keypads': serialize(devices.filter(device_type='KEYPAD')),
        'brokers': serialize(devices.filter(device_type='BROKER')),
        'leds': serialize(devices.filter(device_type='LED')),
    })

# Import branch_required here if it wasn't available at module level (standard import preferred at top)
from callq_core.permissions import branch_required
from django.contrib.auth.decorators import user_passes_test
from licenses.models import Batch # Import Batch for view context



# --- CONFIG LIST AND PROFILE VIEWS ---



@login_required
def device_approval_list(request):
    """
    List devices with filtering by licence_status (Pending, Approved, Rejected).
    """
    user = request.user
    status_filter = request.GET.get('status', 'Pending')  # Default to Pending
    
    # Validate status filter
    valid_statuses = ['Pending', 'Active', 'Rejected', 'All']
    if status_filter not in valid_statuses:
        status_filter = 'Pending'
    
    # Base queryset based on role
    if user.role == "SUPER_ADMIN":
        devices = Device.objects.all()
    elif user.role == "ADMIN":
        if user.assigned_state:
            devices = Device.objects.filter(company__state__in=user.assigned_state)
        else:
            devices = Device.objects.none()
    elif user.role == "COMPANY_ADMIN":
        devices = Device.objects.filter(company=user.company_relation)
    elif user.role == "DEALER_ADMIN":
        if user.company_relation:
            # Import DealerCustomer
            from companydetails.models import DealerCustomer
            
            # Get all dealer customers of this dealer
            dealer_customer_ids = DealerCustomer.objects.filter(
                dealer=user.company_relation
            ).values_list('id', flat=True)
            
            # Include devices belonging to the company, child companies, OR assigned to dealer customers
            devices = Device.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation) |
                Q(dealer_customer_id__in=dealer_customer_ids)
            )
        else:
            devices = Device.objects.none()
    elif user.role == "BRANCH_ADMIN":
        if user.branch_relation:
            devices = Device.objects.filter(branch=user.branch_relation)
        else:
            devices = Device.objects.none()
    else:
        devices = Device.objects.none()
    
    # Apply status filter
    if status_filter != 'All':
        devices = devices.filter(licence_status=status_filter)
    
    context = {
        'devices': devices,
        'current_status': status_filter
    }

    # If Super Admin, provide all companies for the assignment UI during approval
    if user.role == "SUPER_ADMIN":
        from companydetails.models import Company
        context['companies'] = Company.objects.all()

    return render(request, 'configdetails/device_approval_list.html', context)

@login_required
def assign_devices_to_branch(request):
    """View to bulk assign devices to a branch."""
    user = request.user
    
    # Permission check - mainly for Company Admin
    if user.role not in ["COMPANY_ADMIN", "SUPER_ADMIN", "ADMIN"]:
        messages.error(request, "Permission denied.")
        return redirect('device_list')
    
    company = None
    if user.role == "COMPANY_ADMIN":
        company = user.company_relation
    elif user.role == "ADMIN" and user.assigned_state:
        # Complex for Admin with states - maybe allow selecting company?
        # For simplicity, if Admin manages multiple companies, this view might need a company selector first.
        # But assuming context is single company or simplified.
        pass 
    
    if not company and user.role != "SUPER_ADMIN":
         messages.error(request, "Company context required.")
         return redirect('device_list')

    if request.method == 'POST':
        branch_id = request.POST.get('branch_id')
        device_ids = request.POST.getlist('device_ids')
        
        if branch_id and device_ids:
            branch = get_object_or_404(Branch, id=branch_id)
            
            # Verify branch belongs to company
            if company and branch.company != company:
                messages.error(request, "Invalid Branch.")
                return redirect('assign_devices_to_branch')
                
            count = Device.objects.filter(id__in=device_ids).update(branch=branch)
            messages.success(request, f"{count} devices assigned to {branch.branch_name} successfully.")
        else:
            messages.error(request, "Please select a branch and at least one device.")
            
        return redirect('assign_devices_to_branch')
    
    # Get Devices
    if user.role == "SUPER_ADMIN":
        devices = Device.objects.all()
        branches = Branch.objects.all()
    elif user.role == "COMPANY_ADMIN":
        devices = Device.objects.filter(company=user.company_relation)
        branches = Branch.objects.filter(company=user.company_relation)
    else:
        devices = Device.objects.none()
        branches = []
        
    context = {
        'devices': devices,
        'branches': branches
    }
    return render(request, 'configdetails/assign_devices_to_branch.html', context)

# ============================================
# Embedded Profile Template Views
# ============================================

@login_required
def embedded_profile_list(request):
    """List all reusable embedded profiles with device type filtering"""
    user = request.user
    
    # Get all profiles (they are company-agnostic for now)
    # Update for isolation: Filter by Company/Role
    if user.role == "SUPER_ADMIN":
        profiles = EmbeddedProfile.objects.all()
    elif user.role == "ADMIN":
        if user.assigned_state:
            profiles = EmbeddedProfile.objects.filter(company__state__in=user.assigned_state)
        else:
            profiles = EmbeddedProfile.objects.none()
    elif user.role == "COMPANY_ADMIN":
        profiles = EmbeddedProfile.objects.filter(company=user.company_relation)
    elif user.role == "BRANCH_ADMIN":
        # Request: profile section only show profiles created by branch only
        profiles = EmbeddedProfile.objects.filter(branch=user.branch_relation)
    elif user.role == "DEALER_ADMIN":
         if user.company_relation:
            profiles = EmbeddedProfile.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
         else:
            profiles = EmbeddedProfile.objects.none()
    else:
        # Branch Admin etc - maybe see company profiles?
        if user.role == "BRANCH_ADMIN" and user.branch_relation:
            # Request: profile section only show profiles created by branch only
            profiles = EmbeddedProfile.objects.filter(branch=user.branch_relation)
        else:
            profiles = EmbeddedProfile.objects.none()
    
    # Device type filter
    device_type_filter = request.GET.get('device_type', '')
    if device_type_filter:
        profiles = profiles.filter(device_type=device_type_filter)
    
    # Search filter
    search_query = request.GET.get('search', '')
    if search_query:
        profiles = profiles.filter(name__icontains=search_query)
        
    # Order by update needed?
    profiles = profiles.select_related('company')
    
    # Annotate with device count
    from django.db.models import Count
    profiles = profiles.annotate(device_count=Count('devices'))
    
    # Filter out device types that don't need profiles (Broker, LED)
    excluded_types = [Device.DeviceType.BROKER, Device.DeviceType.LED]
    allowed_device_types = [
        choice for choice in Device.DeviceType.choices 
        if choice[0] not in excluded_types
    ]

    context = {
        'profiles': profiles,
        'device_types': allowed_device_types,
        'device_type_filter': device_type_filter,
        'search_query': search_query,
    }
    return render(request, 'configdetails/embedded_profile_list.html', context)


@login_required
def embedded_profile_create(request):
    """Create a new embedded profile template"""
    device_type = request.GET.get('device_type', 'TV')
    
    if request.method == 'POST':
        user = request.user
        name = request.POST.get('name')
        device_type = request.POST.get('device_type')
        
        # Scheduling Fields
        day = request.POST.getlist('day')
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        is_active = request.POST.get('is_active') == 'on'
        
        # Determine company based on user
        company = None
        branch = None
        if user.role == "COMPANY_ADMIN":
            company = user.company_relation
        elif user.role == "DEALER_ADMIN":
            company = user.company_relation
        elif user.role == "BRANCH_ADMIN" and user.branch_relation:
            company = user.branch_relation.company
            # Branch Context
            branch = user.branch_relation
            
        if not company and user.role != "SUPER_ADMIN":
            # Fallback or error if not super admin
            messages.error(request, "Company context missing.")
            return redirect('embedded_profile_list')

        # Build config_json based on device type
        config_json = {}
        if device_type == 'TV':
            config_json = {
                'orientation': request.POST.get('orientation', 'landscape'),
                'layout_type': request.POST.get('layout_type', 'default'),
                'ad_interval': request.POST.get('ad_interval', 5),
                'audio_language': request.POST.get('audio_language', 'en'),
                'show_ads': request.POST.get('show_ads', 'off'),
                'enable_counter_announcement': request.POST.get('enable_counter_announcement', 'off'),
                'enable_token_announcement': request.POST.get('enable_token_announcement', 'off'),
            }
        elif device_type == 'TOKEN_DISPENSER':
            # Fixed fields
            fixed_company_name = company.company_name if company else ''
            fixed_location = ''
            if branch:
                fixed_location = branch.city or branch.branch_name or ''
            elif company:
                fixed_location = getattr(company, 'city', '')

            config_json = {
                # Dispenser settings
                'day_wise_reset': '1' if request.POST.get('day_wise_reset') == 'on' else '0',
                'reset_token_number': request.POST.get('reset_token_number', ''),
                'common_pool': '1' if request.POST.get('common_pool') == 'on' else '0',
                'duplicate_print': '1' if request.POST.get('duplicate_print') == 'on' else '0',
                'standalone': '1' if request.POST.get('standalone') == 'on' else '0',
                'initial_print': '1' if request.POST.get('initial_print') == 'on' else '0',
                'token_cut': request.POST.get('token_cut', 'full'),
                # Existing fields
                'company_name': fixed_company_name,
                'location': fixed_location,
                'footer_text': request.POST.get('footer_text', 'Thank you, visit again'),
                # Button / counters
                'button_mode': request.POST.get('button_mode', 'counterwise'),
                'token_label': request.POST.get('token_label', ''),
                'start_counter': request.POST.get('start_counter', ''),
                'end_counter': request.POST.get('end_counter', ''),
                # VIP
                'vip_enable': '1' if request.POST.get('vip_enable') == 'on' else '0',
                'vip_count_from': request.POST.get('vip_count_from', ''),
                'vip_count_to': request.POST.get('vip_count_to', ''),
            }
        elif device_type == 'KEYPAD':
            config_json = {
                'keypad_device_text': request.POST.get('keypad_device_text', ''),
                'counter_no': request.POST.get('counter_no', ''),
            }
        elif device_type == 'BROKER':
            config_json = {
                'ssid': request.POST.get('ssid', ''),
                'password': request.POST.get('password', ''),
                'host': request.POST.get('host', ''),
                'port': request.POST.get('port', '8000'),
            }
        elif device_type == 'LED':
            config_json = {
                'led_identifier_name': request.POST.get('led_identifier_name', ''),
                'voice_announcement': request.POST.get('voice_announcement', 'off'),
                'counter_number': request.POST.get('counter_number', 1),
                'token_calling': request.POST.get('token_calling', 'Single'),
                'counter_voice': request.POST.get('counter_voice', 'off'),
                'token_voice': request.POST.get('token_voice', 'First'),
            }
        
        # Handle time input if empty
        if not start_time: start_time = None
        if not end_time: end_time = None

        profile = EmbeddedProfile.objects.create(
            name=name,
            device_type=device_type,
            company=company,
            day=day,
            start_time=start_time,
            end_time=end_time,
            is_active=is_active,
            config_json=config_json,
            branch=branch if user.role == "BRANCH_ADMIN" else None
        )
        messages.success(request, f'Profile "{name}" created successfully!')
        return redirect('embedded_profile_list')
    
    # Filter out device types that don't need profiles (Broker, LED)
    excluded_types = [Device.DeviceType.BROKER, Device.DeviceType.LED]
    allowed_device_types = [
        choice for choice in Device.DeviceType.choices 
        if choice[0] not in excluded_types
    ]

    context = {
        'device_type': device_type,
        'device_types': allowed_device_types,
    }
    return render(request, 'configdetails/embedded_profile_form.html', context)


@login_required
def embedded_profile_edit(request, pk):
    """Edit an existing embedded profile template"""
    profile = get_object_or_404(EmbeddedProfile, pk=pk)
    
    if request.method == 'POST':
        user = request.user
        profile.name = request.POST.get('name')
        profile.day = request.POST.getlist('day')
        
        start_time = request.POST.get('start_time')
        end_time = request.POST.get('end_time')
        
        if start_time: profile.start_time = start_time
        if end_time: profile.end_time = end_time
        
        profile.is_active = request.POST.get('is_active') == 'on'
        
        # Update config_json based on device type
        if profile.device_type == 'TV':
            profile.config_json = {
                'orientation': request.POST.get('orientation', 'landscape'),
                'layout_type': request.POST.get('layout_type', 'default'),
                'ad_interval': request.POST.get('ad_interval', 5),
                'audio_language': request.POST.get('audio_language', 'en'),
                'show_ads': request.POST.get('show_ads', 'off'),
                'enable_counter_announcement': request.POST.get('enable_counter_announcement', 'off'),
                'enable_token_announcement': request.POST.get('enable_token_announcement', 'off'),
            }
        elif profile.device_type == 'TOKEN_DISPENSER':
            # Fixed fields (derive from profile context if available; else user context)
            ctx_company = profile.company
            ctx_branch = profile.branch
            if not ctx_company:
                if user.role in ("COMPANY_ADMIN", "DEALER_ADMIN"):
                    ctx_company = user.company_relation
                elif user.role == "BRANCH_ADMIN" and user.branch_relation:
                    ctx_company = user.branch_relation.company
                    ctx_branch = user.branch_relation

            fixed_company_name = ctx_company.company_name if ctx_company else ''
            fixed_location = ''
            if ctx_branch:
                fixed_location = ctx_branch.city or ctx_branch.branch_name or ''
            elif ctx_company:
                fixed_location = getattr(ctx_company, 'city', '')

            profile.config_json = {
                # Dispenser settings
                'day_wise_reset': '1' if request.POST.get('day_wise_reset') == 'on' else '0',
                'reset_token_number': request.POST.get('reset_token_number', ''),
                'common_pool': '1' if request.POST.get('common_pool') == 'on' else '0',
                'duplicate_print': '1' if request.POST.get('duplicate_print') == 'on' else '0',
                'standalone': '1' if request.POST.get('standalone') == 'on' else '0',
                'initial_print': '1' if request.POST.get('initial_print') == 'on' else '0',
                'token_cut': request.POST.get('token_cut', 'full'),
                # Existing fields
                'company_name': fixed_company_name,
                'location': fixed_location,
                'footer_text': request.POST.get('footer_text', 'Thank you, visit again'),
                # Button / counters
                'button_mode': request.POST.get('button_mode', 'counterwise'),
                'token_label': request.POST.get('token_label', ''),
                'start_counter': request.POST.get('start_counter', ''),
                'end_counter': request.POST.get('end_counter', ''),
                # VIP
                'vip_enable': '1' if request.POST.get('vip_enable') == 'on' else '0',
                'vip_count_from': request.POST.get('vip_count_from', ''),
                'vip_count_to': request.POST.get('vip_count_to', ''),
            }
        elif profile.device_type == 'KEYPAD':
            profile.config_json = {
                'keypad_device_text': request.POST.get('keypad_device_text', ''),
                'counter_no': request.POST.get('counter_no', ''),
            }
        elif profile.device_type == 'BROKER':
            profile.config_json = {
                'ssid': request.POST.get('ssid', ''),
                'password': request.POST.get('password', ''),
                'host': request.POST.get('host', ''),
                'port': request.POST.get('port', '8000'),
            }
        elif profile.device_type == 'LED':
            profile.config_json = {
                'led_identifier_name': request.POST.get('led_identifier_name', ''),
                'voice_announcement': request.POST.get('voice_announcement', 'off'),
                'counter_number': request.POST.get('counter_number', 1),
                'token_calling': request.POST.get('token_calling', 'Single'),
                'counter_voice': request.POST.get('counter_voice', 'off'),
                'token_voice': request.POST.get('token_voice', 'First'),
            }
        
        profile.save()
        messages.success(request, f'Profile "{profile.name}" updated successfully!')
        return redirect('embedded_profile_list')
    
    # Filter out device types that don't need profiles (Broker, LED)
    excluded_types = [Device.DeviceType.BROKER, Device.DeviceType.LED]
    allowed_device_types = [
        choice for choice in Device.DeviceType.choices 
        if choice[0] not in excluded_types
    ]

    context = {
        'profile': profile,
        'device_type': profile.device_type,
        'device_types': allowed_device_types,
    }
    return render(request, 'configdetails/embedded_profile_form.html', context)


@login_required
def embedded_profile_delete(request, pk):
    """Delete an embedded profile template"""
    profile = get_object_or_404(EmbeddedProfile, pk=pk)
    device_count = profile.devices.count()
    
    if request.method == 'POST':
        name = profile.name
        # Remove profile from all devices first
        profile.devices.update(embedded_profile=None)
        profile.delete()
        messages.success(request, f'Profile "{name}" deleted successfully!')
        return redirect('embedded_profile_list')
    
    context = {
        'profile': profile,
        'device_count': device_count,
    }
    return render(request, 'configdetails/embedded_profile_confirm_delete.html', context)


@login_required
def embedded_profile_allocate(request, pk):
    """Allocate an embedded profile to devices"""
    profile = get_object_or_404(EmbeddedProfile, pk=pk)
    user = request.user
    
    # Get devices of matching type based on user role
    if user.role == "SUPER_ADMIN":
        devices = Device.objects.filter(device_type=profile.device_type)
    elif user.role == "ADMIN":
        if user.assigned_state:
            devices = Device.objects.filter(device_type=profile.device_type, company__state__in=user.assigned_state)
        else:
            devices = Device.objects.none()
    elif user.role == "COMPANY_ADMIN":
        devices = Device.objects.filter(device_type=profile.device_type, company=user.company_relation)
    elif user.role == "DEALER_ADMIN":
        devices = Device.objects.filter(device_type=profile.device_type, company=user.company_relation)
    else:
        devices = Device.objects.none()
    
    if request.method == 'POST':
        device_ids = request.POST.getlist('device_ids')
        action = request.POST.get('action', 'allocate')
        
        if action == 'allocate':
            Device.objects.filter(id__in=device_ids).update(embedded_profile=profile)
            messages.success(request, f'Profile allocated to {len(device_ids)} device(s)!')
        elif action == 'deallocate':
            Device.objects.filter(id__in=device_ids).update(embedded_profile=None)
            messages.success(request, f'Profile removed from {len(device_ids)} device(s)!')
        
        return redirect('embedded_profile_allocate', pk=pk)
    
    context = {
        'profile': profile,
        'devices': devices,
    }
    return render(request, 'configdetails/embedded_profile_allocate.html', context)


# ============================================
# Device Config Profile Views (No Scheduling)
# ============================================

@login_required
def config_profile_list(request):
    """List all Device Config Profiles for the current user's scope."""
    from .models import DeviceConfigProfile
    user = request.user

    if user.role == 'SUPER_ADMIN':
        profiles = DeviceConfigProfile.objects.all()
    elif user.role == 'ADMIN':
        profiles = DeviceConfigProfile.objects.filter(company__state__in=user.assigned_state) if user.assigned_state else DeviceConfigProfile.objects.none()
    elif user.role == 'COMPANY_ADMIN':
        profiles = DeviceConfigProfile.objects.filter(company=user.company_relation)
    elif user.role == 'DEALER_ADMIN':
        profiles = DeviceConfigProfile.objects.filter(
            Q(company=user.company_relation) | Q(company__parent_company=user.company_relation)
        ) if user.company_relation else DeviceConfigProfile.objects.none()
    elif user.role == 'BRANCH_ADMIN':
        profiles = DeviceConfigProfile.objects.filter(branch=user.branch_relation) if user.branch_relation else DeviceConfigProfile.objects.none()
    else:
        profiles = DeviceConfigProfile.objects.none()

    device_type_filter = request.GET.get('device_type', '')
    if device_type_filter:
        profiles = profiles.filter(device_type=device_type_filter)

    search_query = request.GET.get('search', '')
    if search_query:
        profiles = profiles.filter(name__icontains=search_query)

    from django.db.models import Count
    profiles = profiles.annotate(device_count=Count('assigned_devices'))

    context = {
        'profiles': profiles,
        'device_types': Device.DeviceType.choices,
        'device_type_filter': device_type_filter,
        'search_query': search_query,
    }
    return render(request, 'configdetails/config_profile_list.html', context)


@login_required
def config_profile_create(request):
    """Create a new Device Config Profile."""
    from .models import DeviceConfigProfile
    user = request.user
    device_type = request.GET.get('device_type', 'TV')

    if request.method == 'POST':
        name = request.POST.get('name')
        device_type = request.POST.get('device_type')
        is_active = request.POST.get('is_active') == 'on'

        company = None
        branch = None
        if user.role == 'COMPANY_ADMIN':
            company = user.company_relation
        elif user.role in ('DEALER_ADMIN',):
            company = user.company_relation
        elif user.role == 'BRANCH_ADMIN' and user.branch_relation:
            company = user.branch_relation.company
            branch = user.branch_relation

        config_json = _build_config_json(device_type, request.POST, company, branch)

        DeviceConfigProfile.objects.create(
            name=name,
            device_type=device_type,
            company=company,
            branch=branch,
            config_json=config_json,
            is_active=is_active,
        )
        messages.success(request, f'Config Profile "{name}" created successfully!')
        return redirect('config_profile_list')

    context = {
        'device_type': device_type,
        'device_types': Device.DeviceType.choices,
    }
    return render(request, 'configdetails/config_profile_form.html', context)


@login_required
def config_profile_edit(request, pk):
    """Edit an existing Device Config Profile."""
    from .models import DeviceConfigProfile
    profile = get_object_or_404(DeviceConfigProfile, pk=pk)
    user = request.user

    if request.method == 'POST':
        profile.name = request.POST.get('name')
        profile.is_active = request.POST.get('is_active') == 'on'

        ctx_company = profile.company
        ctx_branch = profile.branch
        if not ctx_company:
            if user.role in ('COMPANY_ADMIN', 'DEALER_ADMIN'):
                ctx_company = user.company_relation
            elif user.role == 'BRANCH_ADMIN' and user.branch_relation:
                ctx_company = user.branch_relation.company
                ctx_branch = user.branch_relation

        profile.config_json = _build_config_json(profile.device_type, request.POST, ctx_company, ctx_branch)
        profile.save()
        messages.success(request, f'Config Profile "{profile.name}" updated successfully!')
        return redirect('config_profile_list')

    context = {
        'profile': profile,
        'device_type': profile.device_type,
        'device_types': Device.DeviceType.choices,
    }
    return render(request, 'configdetails/config_profile_form.html', context)


@login_required
def config_profile_delete(request, pk):
    """Delete a Device Config Profile."""
    from .models import DeviceConfigProfile
    profile = get_object_or_404(DeviceConfigProfile, pk=pk)
    device_count = profile.assigned_devices.count()

    if request.method == 'POST':
        name = profile.name
        profile.assigned_devices.update(config_profile=None)
        profile.delete()
        messages.success(request, f'Config Profile "{name}" deleted successfully!')
        return redirect('config_profile_list')

    context = {
        'profile': profile,
        'device_count': device_count,
    }
    return render(request, 'configdetails/config_profile_confirm_delete.html', context)


@login_required
def config_profile_allocate(request, pk):
    """Allocate / deallocate a Device Config Profile to devices."""
    from .models import DeviceConfigProfile
    profile = get_object_or_404(DeviceConfigProfile, pk=pk)
    user = request.user

    if user.role == 'SUPER_ADMIN':
        devices = Device.objects.filter(device_type=profile.device_type)
    elif user.role == 'ADMIN':
        devices = Device.objects.filter(device_type=profile.device_type, company__state__in=user.assigned_state) if user.assigned_state else Device.objects.none()
    elif user.role == 'COMPANY_ADMIN':
        devices = Device.objects.filter(device_type=profile.device_type, company=user.company_relation)
    elif user.role == 'DEALER_ADMIN':
        devices = Device.objects.filter(device_type=profile.device_type, company=user.company_relation) if user.company_relation else Device.objects.none()
    elif user.role == 'BRANCH_ADMIN':
        devices = Device.objects.filter(device_type=profile.device_type, branch=user.branch_relation) if user.branch_relation else Device.objects.none()
    else:
        devices = Device.objects.none()

    if request.method == 'POST':
        device_ids = request.POST.getlist('device_ids')
        action = request.POST.get('action', 'allocate')
        if action == 'allocate':
            Device.objects.filter(id__in=device_ids).update(config_profile=profile)
            messages.success(request, f'Config Profile allocated to {len(device_ids)} device(s)!')
        elif action == 'deallocate':
            Device.objects.filter(id__in=device_ids).update(config_profile=None)
            messages.success(request, f'Config Profile removed from {len(device_ids)} device(s)!')
        return redirect('config_profile_allocate', pk=pk)

    context = {
        'profile': profile,
        'devices': devices,
    }
    return render(request, 'configdetails/config_profile_allocate.html', context)


def _build_config_json(device_type, post, company=None, branch=None):
    """Helper: build config_json dict from POST data for any device type."""
    if device_type == 'TV':
        return {
            'orientation': post.get('orientation', 'landscape'),
            'layout_type': post.get('layout_type', 'default'),
            'audio_language': post.get('audio_language', 'en'),
            'enable_counter_announcement': post.get('enable_counter_announcement', 'off'),
            'enable_token_announcement': post.get('enable_token_announcement', 'off'),
        }
    elif device_type == 'TOKEN_DISPENSER':
        fixed_company_name = company.company_name if company else ''
        fixed_location = ''
        if branch:
            fixed_location = branch.city or branch.branch_name or ''
        elif company:
            fixed_location = getattr(company, 'city', '')
        return {
            'day_wise_reset': '1' if post.get('day_wise_reset') == 'on' else '0',
            'reset_token_number': post.get('reset_token_number', ''),
            'common_pool': '1' if post.get('common_pool') == 'on' else '0',
            'duplicate_print': '1' if post.get('duplicate_print') == 'on' else '0',
            'standalone': '1' if post.get('standalone') == 'on' else '0',
            'initial_print': '1' if post.get('initial_print') == 'on' else '0',
            'token_cut': post.get('token_cut', 'full'),
            'company_name': fixed_company_name,
            'location': fixed_location,
            'footer_text': post.get('footer_text', 'Thank you, visit again'),
            'button_mode': post.get('button_mode', 'counterwise'),
            'token_label': post.get('token_label', ''),
            'start_counter': post.get('start_counter', ''),
            'end_counter': post.get('end_counter', ''),
            'vip_enable': '1' if post.get('vip_enable') == 'on' else '0',
            'vip_count_from': post.get('vip_count_from', ''),
            'vip_count_to': post.get('vip_count_to', ''),
        }
    elif device_type == 'KEYPAD':
        keypad_data = {
            'keypad_device_text': post.get('keypad_device_text', ''),
            'counter_no': post.get('counter_no', ''),
            'logo_enable': '1' if post.get('logo_enable') == 'on' else '0',
            'single_multiple': post.get('single_multiple', '1'),  # Default to '1' (single mode)
            'skip_enable': '1' if post.get('skip_enable') == 'on' else '0',
            'transfer_enable': '1' if post.get('transfer_enable') == 'on' else '0',
            'vip_enable': '1' if post.get('vip_enable') == 'on' else '0',
            'vip_from': post.get('vip_from', '0'),
            'vip_to': post.get('vip_to', '0'),
            'keypad_pool_mode': post.get('keypad_pool_mode', '0'),
            'dispenser_sl_no': post.get('dispenser_sl_no', ''),
            'no_of_keypad_dev': post.get('no_of_keypad_dev', '1'),
            'remaining_bit_flag': post.get('remaining_bit_flag', '1'),
            # Button string IDs for buttons B, C, D (passed to Android TV via keypad config)
            'button_b_string_id': post.get('button_b_string_id', ''),
            'button_c_string_id': post.get('button_c_string_id', ''),
            'button_d_string_id': post.get('button_d_string_id', ''),
        }
        # Save keypad serial numbers (up to 5)
        for i in range(1, 6):
            keypad_data[f'keypad_sl_no_{i}'] = post.get(f'keypad_sl_no_{i}', '')
        return keypad_data

    elif device_type == 'BROKER':
        return {
            'ssid': post.get('ssid', ''),
            'password': post.get('password', ''),
            'host': post.get('host', ''),
            'port': post.get('port', '8000'),
        }
    elif device_type == 'LED':
        return {
            'led_identifier_name': post.get('led_identifier_name', ''),
            'voice_announcement': post.get('voice_announcement', 'off'),
            'counter_number': post.get('counter_number', 1),
            'token_calling': post.get('token_calling', 'Single'),
            'counter_voice': post.get('counter_voice', 'off'),
            'token_voice': post.get('token_voice', 'First'),
        }
    return {}

@login_required
def approve_device_request(request, device_id):
    if request.method == 'POST':
        device = get_object_or_404(Device, id=device_id)
        # Add permission check here if needed

        # Save Display Name
        display_name = request.POST.get('display_name')
        if display_name:
            device.display_name = display_name
        
        # Assign Company if selected (for Super Admin assignment flow)
        company_id = request.POST.get('company_id')
        if company_id:
            from companydetails.models import Company
            company = get_object_or_404(Company, id=company_id)
            device.company = company

        # Assign Branch if selected
        branch_id = request.POST.get('branch_id')
        if branch_id:
            branch = get_object_or_404(Branch, id=branch_id)
            device.branch = branch
        
        # Final validation: device must have a company before approval
        if not device.company:
            messages.error(request, "Device must be assigned to a company before approval.")
            return redirect('device_approval_list')

        device.licence_status = 'Active'
        device.is_active = True
        device.save()
        messages.success(request, f"Device {device.serial_number} approved successfully (Status: Active).")
        return redirect('device_approval_list')
    return redirect('device_approval_list')

@login_required
def reject_device_request(request, device_id):
    if request.method == 'POST':
        device = get_object_or_404(Device, id=device_id)
        # Add permission check here if needed
        
        serial = device.serial_number
        # device.delete()
        device.licence_status = 'Rejected'
        device.is_active = False
        device.save()
        messages.warning(request, f"Device request for {serial} rejected.")
        return redirect('device_approval_list')
    return redirect('device_approval_list')

@login_required
def production_batch_upload(request):
    """
    View to upload production batches (serial numbers).
    Accessible by Super Admin, Admin, and Production Admin.
    """
    if request.user.role not in ['SUPER_ADMIN', 'ADMIN', 'PRODUCTION_ADMIN']:
        messages.error(request, "Permission denied.")
        return redirect('dashboard')

    if request.method == 'POST':
        batch_id = request.POST.get('batch_id')
        serial_numbers_text = request.POST.get('serial_numbers', '')
        excel_file = request.FILES.get('excel_upload')

        if not batch_id:
            messages.error(request, "Batch ID is required.")
            return redirect('production_batch_upload')

        batch, created = ProductionBatch.objects.get_or_create(
            batch_id=batch_id
        )

        # List to store tuples of (serial_number, device_type, mac_address)
        serial_data = []
        
        # Handle manual entry (assumes default device type TV for manual entry, no MAC)
        if serial_numbers_text:
            for s in serial_numbers_text.split('\n'):
                if s.strip():
                    serial_data.append((s.strip(), 'TV', None))

        # Handle File upload (Excel or CSV)
        if excel_file:
            import os
            try:
                filename = excel_file.name
                ext = os.path.splitext(filename)[1].lower()
                
                if ext == '.csv':
                    import csv
                    import io
                    # Read CSV file
                    file_data = excel_file.read().decode('utf-8')
                    csv_reader = csv.DictReader(io.StringIO(file_data))
                    
                    # Check if required columns exist
                    fieldnames = csv_reader.fieldnames
                    if not fieldnames:
                         raise ValueError("Empty CSV file")
                    
                    # Find column names (case-insensitive)
                    fieldnames_lower = {f.lower().strip(): f for f in fieldnames}
                    sn_col = fieldnames_lower.get('serial number') or fieldnames[0]
                    dt_col = fieldnames_lower.get('device type')
                    mac_col = fieldnames_lower.get('mac address') or fieldnames_lower.get('mac_address')
                    
                    for row in csv_reader:
                        sn_value = row.get(sn_col, '').strip()
                        if sn_value:
                            # Get device type from row or default to TV
                            device_type = 'TV'
                            if dt_col and row.get(dt_col):
                                dt_value = row.get(dt_col, '').strip().upper().replace(' ', '_')
                                # Map common variations
                                dt_mapping = {
                                    'TV': 'TV',
                                    'TOKEN_DISPENSER': 'TOKEN_DISPENSER',
                                    'TOKEN DISPENSER': 'TOKEN_DISPENSER',
                                    'TOKENDISPENSER': 'TOKEN_DISPENSER',
                                    'KEYPAD': 'KEYPAD',
                                    'BROKER': 'BROKER',
                                    'LED': 'LED',
                                }
                                device_type = dt_mapping.get(dt_value, 'TV')
                            # Get optional MAC address
                            mac_value = row.get(mac_col, '').strip() if mac_col else ''
                            serial_data.append((sn_value, device_type, mac_value or None))
                            
                elif ext in ['.xlsx', '.xls']:
                    from openpyxl import load_workbook
                    wb = load_workbook(excel_file)
                    ws = wb.active
                    
                    headers = [str(cell.value).lower().strip() if cell.value else '' for cell in ws[1]]
                    original_headers = [cell.value for cell in ws[1]]
                    
                    # Find column indices
                    sn_idx = 0
                    dt_idx = None
                    mac_idx = None
                    
                    for i, h in enumerate(headers):
                        if 'serial' in h and 'number' in h:
                            sn_idx = i
                        elif 'device' in h and 'type' in h:
                            dt_idx = i
                        elif 'mac' in h:
                            mac_idx = i
                    
                    # Iterate from second row
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        if row[sn_idx]:
                            sn_value = str(row[sn_idx]).strip()
                            device_type = 'TV'
                            if dt_idx is not None and row[dt_idx]:
                                dt_value = str(row[dt_idx]).strip().upper().replace(' ', '_')
                                dt_mapping = {
                                    'TV': 'TV',
                                    'TOKEN_DISPENSER': 'TOKEN_DISPENSER',
                                    'TOKEN DISPENSER': 'TOKEN_DISPENSER',
                                    'TOKENDISPENSER': 'TOKEN_DISPENSER',
                                    'KEYPAD': 'KEYPAD',
                                    'BROKER': 'BROKER',
                                    'LED': 'LED',
                                }
                                device_type = dt_mapping.get(dt_value, 'TV')
                            # Get optional MAC address
                            mac_value = str(row[mac_idx]).strip() if mac_idx is not None and row[mac_idx] else None
                            serial_data.append((sn_value, device_type, mac_value))
                    
                else:
                    raise ValueError("Unsupported file format. Please upload .csv or .xlsx")

            except Exception as e:
                messages.error(request, f"Error processing file: {str(e)}")
                return redirect('production_batch_upload')

        # Save serial numbers with device types and optional MAC addresses
        created_count = 0
        duplicate_count = 0
        for entry in serial_data:
            sn, device_type = entry[0], entry[1]
            mac = entry[2] if len(entry) > 2 else None
            if not ProductionSerialNumber.objects.filter(serial_number=sn).exists():
                ProductionSerialNumber.objects.create(
                    batch=batch,
                    serial_number=sn,
                    device_type=device_type,
                    mac_address=mac or None,
                )
                created_count += 1
            else:
                duplicate_count += 1

        messages.success(request, f"Batch {batch_id} processed. {created_count} serial numbers added. {duplicate_count} duplicates skipped.")
        return redirect('production_batch_upload')

    batches_qs = ProductionBatch.objects.all().order_by('-created_at')
    batches = []
    for b in batches_qs:
        b.type_display = b.get_device_type_display()
        batches.append(b)
    
    # Exclude TV from the choices
    device_choices = [c for c in Device.DeviceType.choices if c[0] != 'TV']
    
    return render(request, 'configdetails/production_batch_upload.html', {
        'batches': batches,
        # 'device_types': device_choices # Not needed for UI selection anymore
    })

@login_required
def batch_download(request, batch_id, file_format):
    """Download batch details in CSV, PDF, or DOCX"""
    if request.user.role not in ['SUPER_ADMIN', 'ADMIN', 'PRODUCTION_ADMIN']:
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("Permission Denied")
        
    batch = get_object_or_404(ProductionBatch, id=batch_id)
    items = batch.serial_numbers.all()
    
    if file_format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="Batch_{batch.batch_id}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Serial Number', 'MAC Address', 'Device Type', 'Status'])
        for item in items:
            writer.writerow([
                item.serial_number,
                item.mac_address or '',
                item.get_device_type_display(),
                'Registered' if item.is_registered else 'Pending',
            ])
        return response

    elif file_format == 'pdf':
        buffer = io.BytesIO()
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from django.http import FileResponse
        
        p = canvas.Canvas(buffer, pagesize=letter)
        p.drawString(100, 750, f"Batch Report: {batch.batch_id}")
        y = 730
        p.drawString(100, y, "Serial Number | Device Type | Status")
        y -= 20
        
        for item in items:
            text = f"{item.serial_number} | {item.get_device_type_display()} | {'Registered' if item.is_registered else 'Pending'}"
            p.drawString(100, y, text)
            y -= 15
            if y < 50:
                p.showPage()
                y = 750
        
        p.showPage()
        p.save()
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename=f"Batch_{batch.batch_id}.pdf")

    elif file_format == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(f'Batch Report: {batch.batch_id}', 0)
        
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Serial Number'
        hdr_cells[1].text = 'MAC Address'
        hdr_cells[2].text = 'Device Type'
        hdr_cells[3].text = 'Status'
        
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.serial_number
            row_cells[1].text = item.mac_address or ''
            row_cells[2].text = item.get_device_type_display()
            row_cells[3].text = 'Registered' if item.is_registered else 'Pending'
            
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="Batch_{batch.batch_id}.docx"'
        return response
    
    return HttpResponse("Invalid format")

@login_required
def production_report_view(request):
    """
    View for Production Admin to see past batch uploads.
    """
    if request.user.role not in ['SUPER_ADMIN', 'ADMIN', 'PRODUCTION_ADMIN']:
        messages.error(request, "Permission denied.")
        return redirect('dashboard')
        
    batches_qs = ProductionBatch.objects.all().order_by('-created_at')
    
    # Same aggregation logic
    batches = []
    from django.db.models import Count
    
    for b in batches_qs:
        stats = b.serial_numbers.values('device_type').annotate(count=Count('id'))
        summary_parts = []
        for s in stats:
            d_label = dict(Device.DeviceType.choices).get(s['device_type'], s['device_type'])
            summary_parts.append(f"{d_label}: {s['count']}")
        b.summary = ", ".join(summary_parts) if summary_parts else "No devices"
        batches.append(b)

    return render(request, 'configdetails/production_report.html', {'batches': batches})

@api_view(['GET'])
@login_required
def get_branch_devices_api(request, branch_id):
    """Fetch devices for a branch filtered by type."""
    branch = get_object_or_404(Branch, id=branch_id)
    dtype = request.GET.get('type')
    
    devices = Device.objects.filter(branch=branch)
    if dtype:
        devices = devices.filter(device_type=dtype)
        
    data = [
        {
            'id': d.id,
            'serial_number': d.serial_number,
            'display_name': d.display_name,
            'get_display_identifier': d.get_display_identifier,
            'device_model': d.device_model,
            'device_type': d.device_type,
            'token_type': d.token_type
        } for d in devices
    ]
    return Response({'devices': data})

@api_view(['POST'])
@login_required
def save_button_mapping_api(request):
    """Save or update button-level mapping."""
    branch_id = request.data.get('branch_id')
    dealer_customer_id = request.data.get('dealer_customer_id')

    source_device_id = request.data.get('source_device_id')
    source_button = request.data.get('source_button')
    target_device_id = request.data.get('target_device_id')
    
    branch = None
    dealer_customer = None
    company = None

    if branch_id:
        branch = get_object_or_404(Branch, id=branch_id)
        company = branch.company
    elif dealer_customer_id:
        from companydetails.models import DealerCustomer
        dealer_customer = get_object_or_404(DealerCustomer, id=dealer_customer_id)
        company = dealer_customer.dealer
    else:
         return Response({'status': 'error', 'message': 'Branch or Dealer Customer required'}, status=400)

    source_device = get_object_or_404(Device, id=source_device_id)
    target_device = get_object_or_404(Device, id=target_device_id)
    
    # Update or create mapping
    mapping, created = ButtonMapping.objects.update_or_create(
        company=company,
        source_device=source_device,
        source_button=source_button,
        defaults={
            'target_device': target_device,
            'branch': branch,
            'dealer_customer': dealer_customer
        }
    )
    
    return Response({
        'status': 'success',
        'mapping_id': mapping.id,
        'created': created
    })

@api_view(['GET'])
@login_required
def get_branch_mappings_api(request, branch_id):
    """Fetch all button mappings for a branch."""
    mappings = ButtonMapping.objects.filter(branch_id=branch_id)
    data = [
        {
            'id': m.id,
            'source_device_id': m.source_device.id,
            'source_device_sn': m.source_device.serial_number,
            'source_device_display_name': m.source_device.get_display_identifier,
            'source_button': m.source_button,
            'target_device_id': m.target_device.id,
            'target_device_sn': m.target_device.serial_number,
            'target_device_display_name': m.target_device.get_display_identifier,
            'target_device_type': m.target_device.device_type
        } for m in mappings
    ]
    return Response({'mappings': data})

@api_view(['POST'])
@login_required
def delete_button_mapping_api(request, mapping_id):
    """Delete a button mapping."""
    mapping = get_object_or_404(ButtonMapping, id=mapping_id)
    mapping.delete()
    return Response({'status': 'success'})

@api_view(['POST'])
@login_required
def delete_family_mapping_api(request, family_id):
    """Delete a group mapping and all associated button mappings."""
    from django.db import transaction
    
    group = get_object_or_404(GroupMapping, id=family_id)
    group_name = group.group_name
    
    # Check permissions
    user = request.user
    if user.role == 'BRANCH_ADMIN' and group.branch != user.branch_relation:
        return Response({'status': 'error', 'message': 'You do not have permission to delete this group.'}, status=403)
    elif user.role == 'COMPANY_ADMIN' and group.company != user.company_relation:
        return Response({'status': 'error', 'message': 'You do not have permission to delete this group.'}, status=403)
    
    with transaction.atomic():
        # Get all devices in this group
        all_devices = list(group.dispensers.all()) + list(group.keypads.all()) + \
                     list(group.leds.all()) + list(group.brokers.all()) + list(group.tvs.all())
        
        # Delete all button mappings associated with devices in this group
        ButtonMapping.objects.filter(
            Q(source_device__in=all_devices) | Q(target_device__in=all_devices)
        ).delete()
        
        # Delete the group mapping
        group.delete()
        
        # Log the activity
        log_activity(request.user, "Group Deleted", f"Group '{group_name}' (ID: {family_id}) and all associated button mappings deleted")
    
    return Response({'status': 'success', 'message': f'Group "{group_name}" deleted successfully.'})

@api_view(['GET'])
def get_available_serial_numbers_api(request):
    """
    API to fetch available (unregistered) serial numbers from production batches.
    Filters by device_type if provided.
    Supports search query for real-time filtering.
    """
    device_type = request.GET.get('device_type', '')
    search = request.GET.get('search', '')
    
    # Get all unregistered serial numbers
    queryset = ProductionSerialNumber.objects.filter(is_registered=False)
    
    # Filter by device type if provided
    if device_type:
        queryset = queryset.filter(device_type=device_type)
    
    # Filter by search query
    if search:
        queryset = queryset.filter(serial_number__icontains=search)
    
    # Limit to 50 results for performance
    serial_numbers = queryset.select_related('batch')[:50]
    
    data = [
        {
            'serial_number': sn.serial_number,
            'mac_address': sn.mac_address or '',
            'batch_id': sn.batch.batch_id,
            'device_type': sn.device_type
        }
        for sn in serial_numbers
    ]
    
    return Response({'serial_numbers': data})
def get_device_order_priority(device):
    """
    Get hierarchical order priority for devices.
    Order: Token Dispenser (1) -> Keypad (2) -> Broker (3) -> TV/LED (4)
    Lower number = appears first
    """
    if device.device_type == Device.DeviceType.TOKEN_DISPENSER:
        return 1
    elif device.device_type == Device.DeviceType.KEYPAD:
        return 2
    elif device.device_type == Device.DeviceType.BROKER:
        return 3
    elif device.device_type == Device.DeviceType.TV:
        return 4
    elif device.device_type == Device.DeviceType.LED:
        return 4
    else:
        return 99  # Unknown types go last

def detect_device_families(button_mappings):
    """
    Detect device families (chains) from ButtonMappings.
    A family is a connected chain: TD → Keypad → Broker → LED
    TV devices are excluded from families - if same TV is mapped to different brokers, they are separate families.
    Returns: { family_id: { 'device_ids': set(), 'mapping_ids': set() } }
    
    This function properly separates independent families by finding connected components
    in the device graph. Each family is a separate connected component.
    TV devices do not connect families together - if same TV is mapped to two brokers, they are separate families.
    """
    if not button_mappings:
        return {}
    
    # Get all devices from mappings to check their types
    all_devices = {}
    for m in button_mappings:
        all_devices[m.source_device.id] = m.source_device
        all_devices[m.target_device.id] = m.target_device
    
    # Build a bidirectional graph of device connections (excluding TV from connections)
    # TV devices should not connect families together
    device_connections = {}  # { device_id: set of connected device_ids }
    device_to_mappings = {}  # { device_id: [mappings involving this device] }
    
    for m in button_mappings:
        source_id = m.source_device.id
        target_id = m.target_device.id
        source_device = m.source_device
        target_device = m.target_device
        
        # Only connect devices if neither is a TV
        # TV devices should not be part of family connections
        # If source is TV, don't add connection (TV doesn't connect families)
        # If target is TV, don't add connection (TV doesn't connect families)
        # This ensures that if same TV is mapped to different brokers, they are separate families
        
        # Build bidirectional connection graph (undirected graph for family detection)
        # But exclude TV from connections - TV should not connect families
        if source_device.device_type != Device.DeviceType.TV and target_device.device_type != Device.DeviceType.TV:
            # Both are non-TV devices, create connection
            if source_id not in device_connections:
                device_connections[source_id] = set()
            device_connections[source_id].add(target_id)
            
            if target_id not in device_connections:
                device_connections[target_id] = set()
            device_connections[target_id].add(source_id)
        
        # Track mappings per device (for both source and target) - but only for family devices
        # Family devices are: TOKEN_DISPENSER, KEYPAD, BROKER, LED (not TV)
        if source_device.device_type != Device.DeviceType.TV:
            if source_id not in device_to_mappings:
                device_to_mappings[source_id] = []
            device_to_mappings[source_id].append(m)
        
        if target_device.device_type != Device.DeviceType.TV:
            if target_id not in device_to_mappings:
                device_to_mappings[target_id] = []
            device_to_mappings[target_id].append(m)
    
    # Find connected components (families) using DFS
    # Each connected component is a separate family
    # Only include non-TV devices in families
    visited = set()
    families = {}
    family_id = 0
    
    def dfs(device_id, current_family):
        """Depth-first search to find all devices in the same connected component"""
        if device_id in visited:
            return
        visited.add(device_id)
        
        # Only add non-TV devices to family
        if device_id in all_devices and all_devices[device_id].device_type != Device.DeviceType.TV:
            current_family['device_ids'].add(device_id)
        
        # Add all mappings involving this device to the family
        if device_id in device_to_mappings:
            for mapping in device_to_mappings[device_id]:
                if mapping.id not in current_family['mapping_ids']:
                    current_family['mapping_ids'].add(mapping.id)
        
        # Visit all connected devices (recursively)
        if device_id in device_connections:
            for connected_id in device_connections[device_id]:
                if connected_id not in visited:
                    dfs(connected_id, current_family)
    
    # Find all families by traversing each unvisited device
    # Only start DFS from non-TV devices
    all_device_ids = set(device_connections.keys())
    for device_id in all_device_ids:
        if device_id not in visited:
            # Only start family detection from non-TV devices
            if device_id in all_devices and all_devices[device_id].device_type != Device.DeviceType.TV:
                family_id += 1
                current_family = {
                    'device_ids': set(),
                    'mapping_ids': set()
                }
                dfs(device_id, current_family)
                
                # Only create family if it has multiple devices (a meaningful chain)
                # Single device mappings are considered standalone, not families
                if len(current_family['device_ids']) > 1:
                    families[family_id] = current_family
    
    return families

@login_required
def mapping_list_view(request):
    """View to list all GroupMapping objects grouped by Branch."""
    user = request.user
    
    # Filter GroupMapping based on role
    if user.role in ['SUPER_ADMIN', 'ADMIN']:
        qs = GroupMapping.objects.all()
    elif user.role == 'DEALER_ADMIN':
        if user.company_relation:
            qs = GroupMapping.objects.filter(
                Q(company=user.company_relation) | 
                Q(company__parent_company=user.company_relation)
            )
        else:
            qs = GroupMapping.objects.none() 
    elif user.role == 'COMPANY_ADMIN':
        qs = GroupMapping.objects.filter(company=user.company_relation)
    elif user.role == 'BRANCH_ADMIN':
        qs = GroupMapping.objects.filter(branch=user.branch_relation)
    else:
        qs = GroupMapping.objects.none()

    # Prefetch related fields for performance
    qs = qs.select_related('branch', 'company').prefetch_related(
        'dispensers', 'keypads', 'tvs', 'brokers', 'leds'
    ).order_by('branch__branch_name', 'group_name')

    # Group by Branch and structure data for template
    branch_mappings = {}
    for fm in qs:
        # If branch is None, map it under the company
        key = fm.branch if fm.branch else fm.company
        if key not in branch_mappings:
            branch_mappings[key] = {'groups': [], 'standalone_by_type': {}}
        
        # Get all devices in this group
        all_group_devices = list(fm.dispensers.all()) + list(fm.keypads.all()) + \
                            list(fm.leds.all()) + list(fm.brokers.all()) + list(fm.tvs.all())
        
        # Get all button mappings for devices in this group
        button_mappings = ButtonMapping.objects.filter(
            Q(source_device__in=all_group_devices) | Q(target_device__in=all_group_devices)
        ).select_related('source_device', 'target_device')
        
        # Group devices by type
        devices_by_type = {
            'TOKEN_DISPENSER': list(fm.dispensers.all()),
            'KEYPAD': list(fm.keypads.all()),
            'LED': list(fm.leds.all()),
            'BROKER': list(fm.brokers.all()),
            'TV': list(fm.tvs.all())
        }
        
        # Organize button mappings by source device
        button_ordered_mappings = {}
        device_mappings = {}
        
        for bm in button_mappings:
            source = bm.source_device
            if source not in device_mappings:
                device_mappings[source] = []
            
            device_mappings[source].append(bm)
            
            # For keypads, separate broker and LED mappings
            if source.device_type == Device.DeviceType.KEYPAD:
                if source not in button_ordered_mappings:
                    button_ordered_mappings[source] = {'broker': [], 'led': [], 'broker_count': 0, 'led_count': 0}
                
                target = bm.target_device
                if target.device_type == Device.DeviceType.BROKER:
                    button_ordered_mappings[source]['broker'].append({
                        'mappings': [bm],
                        'target_device': target
                    })
                    button_ordered_mappings[source]['broker_count'] += 1
                elif target.device_type == Device.DeviceType.LED:
                    button_ordered_mappings[source]['led'].append({
                        'mappings': [bm],
                        'target_device': target
                    })
                    button_ordered_mappings[source]['led_count'] += 1
            else:
                if source not in button_ordered_mappings:
                    button_ordered_mappings[source] = []
                button_ordered_mappings[source].append(bm)
        
        # TV devices (as targets)
        tv_devices = {}
        for tv in fm.tvs.all():
            tv_mappings = ButtonMapping.objects.filter(target_device=tv)
            if tv_mappings.exists():
                tv_devices[tv] = list(tv_mappings)
        
        # Create group data structure
        group_data = {
            'group_mapping': fm,  # The actual GroupMapping object
            'group_name': fm.group_name,  # Group name for display
            'devices_by_type': devices_by_type,
            'button_ordered_mappings': button_ordered_mappings,
            'device_mappings': device_mappings,
            'tv_devices': tv_devices,
            'devices': {d: [] for d in all_group_devices}  # For compatibility with template
        }
        
        branch_mappings[key]['groups'].append(group_data)
    
    # Structure: { branch: {'families': [...], 'standalone_by_type': {...}} }
    
    return render(request, 'configdetails/mapping_list.html', {
        'grouped_mappings': branch_mappings
    })

def link_broker_devices(broker_device):
    """
    Auto-link all devices connected to the same Broker.
    When a user maps TokenDispenser -> Broker, and Broker -> TV, 
    this function creates ButtonMapping(TokenDispenser -> TV).
    """
    if not broker_device or broker_device.device_type != Device.DeviceType.BROKER:
        return

    # Find all mappings involving this broker
    mappings = Mapping.objects.filter(broker=broker_device)
    
    # 1. Identify Inputs (Sources)
    input_devices = []
    
    # Track Keypads to find devices connected to them (Multi-hop: TD -> Keypad -> Broker)
    connected_keypads = []

    for m in mappings:
        # Token Dispensers directly connected
        if m.token_dispenser:
            input_devices.append(m.token_dispenser)
        
        # Keypads directly connected
        if m.keypad: 
            input_devices.append(m.keypad)
            connected_keypads.append(m.keypad)
        if m.keypad_2: 
            input_devices.append(m.keypad_2)
            connected_keypads.append(m.keypad_2)
        if m.keypad_3: 
            input_devices.append(m.keypad_3)
            connected_keypads.append(m.keypad_3)
        if m.keypad_4: 
            input_devices.append(m.keypad_4)
            connected_keypads.append(m.keypad_4)
            
    # Find Token Dispensers connected to these Keypads
    for kp in connected_keypads:
        # Find mappings where this keypad is present and a TD is present
        # Note: In a mapping TD->Keypad, the Keypad field is populated.
        kp_mappings = Mapping.objects.filter(keypad=kp, token_dispenser__isnull=False)
        for kpm in kp_mappings:
            if kpm.token_dispenser:
                input_devices.append(kpm.token_dispenser)
        
        # Also check keypad_2..4 slots if TD is somehow mapped there (unlikely but safe)
        # Actually, usually TD is main input.
    
    # Unique Inputs
    input_devices = list(set(input_devices))

    # 2. Identify Outputs (Targets)
    output_devices = []
    for m in mappings:
        # TV
        if m.tv: output_devices.append(m.tv)
        # LED
        if m.led: output_devices.append(m.led)
    
    # Unique Outputs
    output_devices = list(set(output_devices))

    # 3. Create Button Mappings
    for source in input_devices:
        for target in output_devices:
            # Determine buttons based on source type
            buttons = []
            if source.device_type == Device.DeviceType.TOKEN_DISPENSER:
                match = source.token_type
                count = 4 # Default
                if match:
                    try:
                         count = int(match.split('_')[0])
                    except:
                        pass
                for i in range(1, count + 1):
                    buttons.append(f"Button {i}")
            
            elif source.device_type == Device.DeviceType.KEYPAD:
                # Assuming Keypad has buttons like "Call", "Recall" etc or numeric
                # For now, map generic buttons or standard set matching TD
                # Let's map Button 1 to 4 for consistency
                for i in range(1, 5):
                     buttons.append(f"Button {i}")

            # Create Mapping
            for btn in buttons:
                ButtonMapping.objects.get_or_create(
                    company=source.company,
                    branch=source.branch,
                    source_device=source,
                    source_button=btn,
                    target_device=target
                )
@api_view(['GET'])
@login_required
def get_dealer_customer_devices_api(request, customer_id):
    """Fetch all devices for a dealer customer."""
    devices = Device.objects.filter(dealer_customer_id=customer_id, is_active=True)
    
    data = []
    for d in devices:
        data.append({
            'id': d.id,
            'serial_number': d.serial_number,
            'display_name': d.display_name,
            'get_display_identifier': d.get_display_identifier,
            'device_type': d.device_type,
            'device_model': d.device_model,
            'token_type': d.token_type if d.device_type == Device.DeviceType.TOKEN_DISPENSER else None
        })
    return Response({'devices': data})

@api_view(['GET'])
@login_required
def get_dealer_customer_mappings_api(request, customer_id):
    """Fetch all button mappings for a dealer customer."""
    mappings = ButtonMapping.objects.filter(dealer_customer_id=customer_id)
    data = [
        {
            'id': m.id,
            'source_device_id': m.source_device.id,
            'source_device_sn': m.source_device.serial_number,
            'source_device_display_name': m.source_device.get_display_identifier,
            'source_button': m.source_button,
            'target_device_id': m.target_device.id,
            'target_device_sn': m.target_device.serial_number,
            'target_device_display_name': m.target_device.get_display_identifier,
            'target_device_type': m.target_device.device_type
        } for m in mappings
    ]
    return Response({'mappings': data})


# ============================================================================
# Counter-Wise Configuration API Endpoints
# ============================================================================

@api_view(['GET'])
@login_required
def get_counters_api(request):
    """Get list of active counters scoped to the current user's company."""
    log_api_request('get_counters_api', request)
    company = getattr(request.user, 'company_relation', None)
    qs = CounterConfig.objects.filter(status=True)
    if company:
        qs = qs.filter(company=company)
    qs = qs.order_by('counter_name')
    serializer = CounterConfigSerializer(qs, many=True)
    log_api_response('get_counters_api', 200, {'counters_count': len(serializer.data)})
    return Response({'counters': serializer.data})


@api_view(['POST'])
@login_required
def create_counter_api(request):
    """Create a new counter scoped to the current user's company."""
    log_api_request('create_counter_api', request)
    company = getattr(request.user, 'company_relation', None)
    serializer = CounterConfigSerializer(data=request.data, context={'request': request})
    if serializer.is_valid():
        counter = serializer.save(company=company)
        log_api_response('create_counter_api', 201, {'counter_id': counter.id, 'counter_name': counter.counter_name})
        action_logger.info(f"Counter Created | ID: {counter.id} | Name: {counter.counter_name}")
        return Response({
            'status': 'success',
            'message': 'Counter created successfully',
            'data': serializer.data
        }, status=201)
    return Response({
        'status': 'error',
        'message': 'Validation failed',
        'errors': serializer.errors
    }, status=400)


@api_view(['PUT', 'PATCH'])
@login_required
def update_counter_api(request, counter_id):
    """Update an existing counter"""
    log_api_request('update_counter_api', request, {'counter_id': counter_id})
    counter = get_object_or_404(CounterConfig, id=counter_id)
    serializer = CounterConfigSerializer(counter, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        log_api_response('update_counter_api', 200, {'counter_id': counter_id})
        action_logger.info(f"Counter Updated | ID: {counter_id}")
        return Response({
            'status': 'success',
            'message': 'Counter updated successfully',
            'data': serializer.data
        })
    log_api_response('update_counter_api', 400, error='Validation failed')
    return Response({
        'status': 'error',
        'message': 'Validation failed',
        'errors': serializer.errors
    }, status=400)


@api_view(['DELETE'])
@login_required
def delete_counter_api(request, counter_id):
    """Delete a counter (with validation)"""
    log_api_request('delete_counter_api', request, {'counter_id': counter_id})
    counter = get_object_or_404(CounterConfig, id=counter_id)
    
    # Check if counter is mapped
    if TVCounterMapping.objects.filter(counter=counter).exists():
        log_api_response('delete_counter_api', 400, error='Counter is mapped to TV devices')
        return Response({
            'status': 'error',
            'message': 'Cannot delete counter: It is mapped to one or more TV devices.'
        }, status=400)
    
    if CounterTokenDispenserMapping.objects.filter(counter=counter).exists():
        log_api_response('delete_counter_api', 400, error='Counter is mapped to token dispensers')
        return Response({
            'status': 'error',
            'message': 'Cannot delete counter: It is mapped to one or more token dispensers.'
        }, status=400)
    
    counter_name = counter.counter_name
    counter.delete()
    log_api_response('delete_counter_api', 200, {'counter_id': counter_id, 'counter_name': counter_name})
    action_logger.info(f"Counter Deleted | ID: {counter_id} | Name: {counter_name}")
    return Response({
        'status': 'success',
        'message': 'Counter deleted successfully'
    })


@api_view(['GET'])
@login_required
def get_tv_counter_mappings_api(request, tv_id):
    """Get all counters mapped to a TV device"""
    log_api_request('get_tv_counter_mappings_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    mappings = TVCounterMapping.objects.filter(tv=tv).select_related('counter')
    serializer = TVCounterMappingSerializer(mappings, many=True)
    log_api_response('get_tv_counter_mappings_api', 200, {'tv_id': tv_id, 'mappings_count': len(serializer.data)})
    return Response({'mappings': serializer.data})


@api_view(['GET'])
@login_required
def get_tv_dispenser_mappings_api(request, tv_id):
    """Get all dispensers mapped to a TV device with their button_index and associated counters"""
    log_api_request('get_tv_dispenser_mappings_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    mappings = TVDispenserMapping.objects.filter(tv=tv).select_related('dispenser').order_by('button_index')
    serializer = TVDispenserMappingSerializer(mappings, many=True)
    log_api_response('get_tv_dispenser_mappings_api', 200, {'tv_id': tv_id, 'mappings_count': len(serializer.data)})
    return Response({'mappings': serializer.data})


@api_view(['POST'])
@login_required
def map_tv_counters_api(request, tv_id):
    """Map one or more counters to a TV device (one counter = one TV rule enforced)"""
    log_api_request('map_tv_counters_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    counter_ids = request.data.get('counter_ids', [])
    
    if not counter_ids:
        log_api_response('map_tv_counters_api', 400, error='counter_ids is required and must be a list')
        return Response({
            'status': 'error',
            'message': 'counter_ids is required and must be a list'
        }, status=400)
    
    # Validate all counters exist
    counters = CounterConfig.objects.filter(id__in=counter_ids, status=True)
    if counters.count() != len(counter_ids):
        log_api_response('map_tv_counters_api', 400, error='One or more counter IDs are invalid')
        return Response({
            'status': 'error',
            'message': 'One or more counter IDs are invalid'
        }, status=400)
    
    # Validate: Check if any selected counter is already mapped to another TV
    duplicate_counters = []
    for counter in counters:
        existing_mappings = TVCounterMapping.objects.filter(
            counter=counter
        ).exclude(tv=tv)
        
        if existing_mappings.exists():
            other_tvs = [m.tv.serial_number for m in existing_mappings]
            duplicate_counters.append({
                'counter_id': counter.id,
                'counter_name': counter.counter_name,
                'other_tvs': other_tvs
            })
    
    if duplicate_counters:
        error_details = []
        for dup in duplicate_counters:
            error_details.append(f"Counter '{dup['counter_name']}' (ID: {dup['counter_id']}) is already mapped to TV(s): {', '.join(dup['other_tvs'])}")
        
        return Response({
            'status': 'error',
            'message': 'One counter can only be mapped to one TV at a time.',
            'details': error_details
        }, status=400)
    
    # Remove existing mappings for this TV
    TVCounterMapping.objects.filter(tv=tv).delete()
    
    # Remove existing mappings for selected counters from other TVs (enforce one-to-one)
    for counter in counters:
        TVCounterMapping.objects.filter(counter=counter).exclude(tv=tv).delete()
    
    # Create new mappings
    created_mappings = []
    for counter in counters:
        mapping, created = TVCounterMapping.objects.get_or_create(tv=tv, counter=counter)
        created_mappings.append(mapping)
    
    serializer = TVCounterMappingSerializer(created_mappings, many=True)
    log_api_response('map_tv_counters_api', 200, {'tv_id': tv_id, 'counters_mapped': len(created_mappings)})
    action_logger.info(f"TV Counters Mapped | TV ID: {tv_id} | Counters: {len(created_mappings)}")
    return Response({
        'status': 'success',
        'message': f'{len(created_mappings)} counter(s) mapped successfully',
        'data': serializer.data
    })


@api_view(['POST'])
@login_required
def map_tv_dispensers_api(request, tv_id):
    """Map one or more dispensers to a TV device (one dispenser = one TV rule enforced)"""
    log_api_request('map_tv_dispensers_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    dispenser_ids = request.data.get('dispenser_ids', [])
    
    if not dispenser_ids:
        log_api_response('map_tv_dispensers_api', 400, error='dispenser_ids is required and must be a list')
        return Response({
            'status': 'error',
            'message': 'dispenser_ids is required and must be a list'
        }, status=400)
    
    # Validate all dispensers exist and are TOKEN_DISPENSER type
    dispensers = Device.objects.filter(
        id__in=dispenser_ids,
        device_type=Device.DeviceType.TOKEN_DISPENSER
    )
    if dispensers.count() != len(dispenser_ids):
        log_api_response('map_tv_dispensers_api', 400, error='One or more dispenser IDs are invalid or not token dispensers')
        return Response({
            'status': 'error',
            'message': 'One or more dispenser IDs are invalid or not token dispensers'
        }, status=400)
    
    # Validate: Check if any selected dispenser is already mapped to another TV
    duplicate_dispensers = []
    for dispenser in dispensers:
        existing_mappings = TVDispenserMapping.objects.filter(
            dispenser=dispenser
        ).exclude(tv=tv)
        
        if existing_mappings.exists():
            other_tvs = [m.tv.serial_number for m in existing_mappings]
            duplicate_dispensers.append({
                'dispenser_id': dispenser.id,
                'dispenser_serial_number': dispenser.serial_number,
                'other_tvs': other_tvs
            })
    
    if duplicate_dispensers:
        error_details = []
        for dup in duplicate_dispensers:
            error_details.append(f"Dispenser '{dup['dispenser_serial_number']}' (ID: {dup['dispenser_id']}) is already mapped to TV(s): {', '.join(dup['other_tvs'])}")
        
        return Response({
            'status': 'error',
            'message': 'One dispenser can only be mapped to one TV at a time.',
            'details': error_details
        }, status=400)
    
    # Remove existing mappings for this TV
    TVDispenserMapping.objects.filter(tv=tv).delete()
    
    # Remove existing mappings for selected dispensers from other TVs (enforce one-to-one)
    for dispenser in dispensers:
        TVDispenserMapping.objects.filter(dispenser=dispenser).exclude(tv=tv).delete()
    
    # Create new mappings with button_index as ASCII chars (0x31 onwards, continuous per TV)
    created_mappings = []
    for position, dispenser in enumerate(dispensers, start=1):
        mapping = TVDispenserMapping.objects.create(
            tv=tv,
            dispenser=dispenser,
            button_index=get_button_index_char(position)  # '1','2',... (ASCII 0x31+)
        )
        created_mappings.append(mapping)
    
    serializer = TVDispenserMappingSerializer(created_mappings, many=True)
    log_api_response('map_tv_dispensers_api', 200, {'tv_id': tv_id, 'dispensers_mapped': len(created_mappings)})
    action_logger.info(f"TV Dispensers Mapped | TV ID: {tv_id} | Dispensers: {len(created_mappings)}")
    return Response({
        'status': 'success',
        'message': f'{len(created_mappings)} dispenser(s) mapped successfully',
        'data': serializer.data
    })


@api_view(['GET'])
@login_required
def get_tv_keypad_mappings_api(request, tv_id):
    """Get all keypads mapped to a TV device with their ASCII indices and button string IDs"""
    log_api_request('get_tv_keypad_mappings_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    mappings = TVKeypadMapping.objects.filter(tv=tv).select_related('keypad').order_by('keypad_index')
    serializer = TVKeypadMappingSerializer(mappings, many=True)
    log_api_response('get_tv_keypad_mappings_api', 200, {'tv_id': tv_id, 'mappings_count': len(serializer.data)})
    return Response({'mappings': serializer.data})


@api_view(['POST'])
@login_required
def map_tv_keypads_api(request, tv_id):
    """
    Map one or more keypads to a TV device.
    Enforces one-keypad-one-TV constraint.

    Expected body:
    {
        "keypad_entries": [
            {
                "keypad_id": 5,
                "button_b_string_id": "some_string",
                "button_c_string_id": "another_string",
                "button_d_string_id": "yet_another"
            },
            ...
        ]
    }
    Both keypad_index and dispenser_index are auto-assigned sequentially
    using get_button_index_char() (same ASCII series chr(0x31)…).
    """
    log_api_request('map_tv_keypads_api', request, {'tv_id': tv_id})
    tv = get_object_or_404(Device, id=tv_id, device_type=Device.DeviceType.TV)
    keypad_entries = request.data.get('keypad_entries', [])

    if not keypad_entries:
        log_api_response('map_tv_keypads_api', 400, error='keypad_entries is required and must be a non-empty list')
        return Response({
            'status': 'error',
            'message': 'keypad_entries is required and must be a non-empty list'
        }, status=400)

    # Validate all keypad IDs
    keypad_ids = [int(e['keypad_id']) for e in keypad_entries if e.get('keypad_id')]
    keypads_qs = Device.objects.filter(id__in=keypad_ids, device_type=Device.DeviceType.KEYPAD)
    if keypads_qs.count() != len(keypad_ids):
        log_api_response('map_tv_keypads_api', 400, error='One or more keypad IDs are invalid or not KEYPAD devices')
        return Response({
            'status': 'error',
            'message': 'One or more keypad IDs are invalid or not KEYPAD devices'
        }, status=400)

    # Check for keypads already mapped to another TV
    duplicate_kps = []
    for kp in keypads_qs:
        clash = TVKeypadMapping.objects.filter(keypad=kp).exclude(tv=tv)
        if clash.exists():
            other_tvs = [m.tv.serial_number for m in clash]
            duplicate_kps.append({
                'keypad_id': kp.id,
                'keypad_serial_number': kp.serial_number,
                'other_tvs': other_tvs
            })

    if duplicate_kps:
        error_details = [
            f"Keypad '{d['keypad_serial_number']}' (ID: {d['keypad_id']}) is already mapped to TV(s): {', '.join(d['other_tvs'])}"
            for d in duplicate_kps
        ]
        return Response({
            'status': 'error',
            'message': 'One keypad can only be mapped to one TV at a time.',
            'details': error_details
        }, status=400)

    # Atomically replace all keypad mappings for this TV
    created_mappings = []
    with transaction.atomic():
        TVKeypadMapping.objects.filter(tv=tv).delete()
        for position, entry in enumerate(keypad_entries, start=1):
            kp_id = int(entry.get('keypad_id', 0))
            if not kp_id:
                continue
            try:
                kp = Device.objects.get(id=kp_id, device_type=Device.DeviceType.KEYPAD)
                # Remove any existing mapping of this keypad to another TV
                TVKeypadMapping.objects.filter(keypad=kp).exclude(tv=tv).delete()
                ascii_idx = get_button_index_char(position)  # same fn for both indices
                mapping = TVKeypadMapping.objects.create(
                    tv=tv,
                    keypad=kp,
                    keypad_index=ascii_idx,
                    dispenser_index=ascii_idx,
                    button_b_string_id=entry.get('button_b_string_id', '') or '',
                    button_c_string_id=entry.get('button_c_string_id', '') or '',
                    button_d_string_id=entry.get('button_d_string_id', '') or '',
                )
                created_mappings.append(mapping)
            except Device.DoesNotExist:
                pass

    serializer = TVKeypadMappingSerializer(created_mappings, many=True)
    log_api_response('map_tv_keypads_api', 200, {'tv_id': tv_id, 'keypads_mapped': len(created_mappings)})
    action_logger.info(f"TV Keypads Mapped | TV ID: {tv_id} | Keypads: {len(created_mappings)}")
    return Response({
        'status': 'success',
        'message': f'{len(created_mappings)} keypad(s) mapped successfully',
        'data': serializer.data
    })


@api_view(['GET'])
@login_required
def get_counter_dispenser_mappings_api(request, counter_id):

    """Get token dispensers mapped to a counter"""
    log_api_request('get_counter_dispenser_mappings_api', request, {'counter_id': counter_id})
    counter = get_object_or_404(CounterConfig, id=counter_id)
    mappings = CounterTokenDispenserMapping.objects.filter(counter=counter).select_related('dispenser')
    serializer = CounterTokenDispenserMappingSerializer(mappings, many=True)
    log_api_response('get_counter_dispenser_mappings_api', 200, {'counter_id': counter_id, 'mappings_count': len(serializer.data)})
    return Response({'mappings': serializer.data})


@api_view(['POST'])
@login_required
def map_counter_dispenser_api(request, counter_id):
    """Map a counter to a token dispenser (one counter = one dispenser rule enforced)"""
    log_api_request('map_counter_dispenser_api', request, {'counter_id': counter_id})
    counter = get_object_or_404(CounterConfig, id=counter_id)
    dispenser_id = request.data.get('dispenser_id')
    
    if not dispenser_id:
        log_api_response('map_counter_dispenser_api', 400, error='dispenser_id is required')
        return Response({
            'status': 'error',
            'message': 'dispenser_id is required'
        }, status=400)
    
    dispenser = get_object_or_404(Device, id=dispenser_id, device_type=Device.DeviceType.TOKEN_DISPENSER)
    
    # Validate: Check if counter is already mapped to another dispenser within the same company.
    # Counters and dispensers are company-scoped, so cross-company mappings are not a conflict.
    existing_qs = CounterTokenDispenserMapping.objects.filter(
        counter=counter
    ).exclude(dispenser=dispenser)
    if dispenser.company:
        existing_qs = existing_qs.filter(dispenser__company=dispenser.company)
    
    if existing_qs.exists():
        other_dispensers = [m.dispenser.serial_number for m in existing_qs]
        return Response({
            'status': 'error',
            'message': f'Counter "{counter.counter_name}" (ID: {counter.id}) is already mapped to dispenser(s): {", ".join(other_dispensers)}. One counter can only be mapped to one dispenser at a time.',
            'details': {
                'counter_id': counter.id,
                'counter_name': counter.counter_name,
                'existing_dispensers': other_dispensers
            }
        }, status=400)
    
    # Validate: Check if dispenser has reached its button limit
    num_buttons = 1  # Default
    if dispenser.token_type:
        try:
            num_buttons = int(dispenser.token_type.split('_')[0])
        except (ValueError, AttributeError):
            num_buttons = 1
    
    current_mappings_count = CounterTokenDispenserMapping.objects.filter(
        dispenser=dispenser
    ).exclude(counter=counter).count()
    
    if current_mappings_count >= num_buttons:
        return Response({
            'status': 'error',
            'message': f'Cannot map counter: Dispenser "{dispenser.serial_number}" is a {num_buttons}-button dispenser and already has {current_mappings_count} counter(s) mapped. Maximum {num_buttons} counter(s) allowed.',
            'details': {
                'dispenser_serial_number': dispenser.serial_number,
                'max_counters': num_buttons,
                'current_count': current_mappings_count
            }
        }, status=400)
    
    # Remove any existing mappings for this counter to other dispensers (enforce one-to-one)
    CounterTokenDispenserMapping.objects.filter(counter=counter).exclude(dispenser=dispenser).delete()
    
    mapping, created = CounterTokenDispenserMapping.objects.get_or_create(
        counter=counter,
        dispenser=dispenser
    )
    
    serializer = CounterTokenDispenserMappingSerializer(mapping)
    log_api_response('map_counter_dispenser_api', 201 if created else 200, {
        'counter_id': counter_id,
        'dispenser_id': dispenser_id,
        'created': created
    })
    action_logger.info(f"Counter-Dispenser Mapped | Counter ID: {counter_id} | Dispenser ID: {dispenser_id}")
    return Response({
        'status': 'success',
        'message': 'Counter mapped to token dispenser successfully',
        'data': serializer.data
    }, status=201 if created else 200)


@api_view(['DELETE'])
@login_required
def unmap_counter_dispenser_api(request, counter_id, dispenser_id):
    """Unmap a counter from a token dispenser"""
    mapping = get_object_or_404(
        CounterTokenDispenserMapping,
        counter_id=counter_id,
        dispenser_id=dispenser_id
    )
    mapping.delete()
    return Response({
        'status': 'success',
        'message': 'Counter unmapped from token dispenser successfully'
    })


@api_view(['POST'])
@permission_classes([AllowAny])
def swap_counters_api(request):
    """
    Update counter-to-button mappings on a single token dispenser.

    The entire 'buttons' array is the DESIRED final state.
    Changes are applied atomically in one transaction.

    Request body:
    {
        "serial_number": "DISP-001",
        "buttons": [
            {"button_index": "1", "counter_id": 5},
            {"button_index": "2", "counter_id": 7},
            {"button_index": "3", "counter_id": null}
        ]
    }
    button_index = GroupCounterButtonMapping.button_index (fixed, group-wide)
    counter_id   = null means unmap that slot
    """
    try:
        serial_number = request.data.get('serial_number')
        buttons       = request.data.get('buttons')

        if not serial_number:
            return Response({'status': 'error', 'message': 'serial_number is required'}, status=400)

        if not isinstance(buttons, list) or len(buttons) == 0:
            return Response({
                'status': 'error',
                'message': 'buttons must be a non-empty array of {button_index, counter_id} objects'
            }, status=400)

        # resolve dispenser
        try:
            dispenser = Device.objects.get(
                serial_number=serial_number,
                device_type=Device.DeviceType.TOKEN_DISPENSER,
            )
        except Device.DoesNotExist:
            return Response({
                'status': 'error',
                'message': f'Token dispenser with serial number "{serial_number}" not found'
            }, status=404)

        # group (may be None)
        gdm_row = (
            GroupDispenserMapping.objects
            .filter(dispenser=dispenser)
            .select_related('group')
            .first()
        )
        group = gdm_row.group if gdm_row else None

        # parse & validate desired state
        # desired_map: {button_index (str) -> CounterConfig | None}
        desired_map = {}
        for entry in buttons:
            btn_idx    = entry.get('button_index')
            counter_id = entry.get('counter_id')

            if not btn_idx:
                return Response({
                    'status': 'error',
                    'message': f'Each button entry must have a "button_index". Got: {entry}'
                }, status=400)

            if counter_id is not None:
                try:
                    desired_map[btn_idx] = CounterConfig.objects.get(id=counter_id)
                except CounterConfig.DoesNotExist:
                    return Response({
                        'status': 'error',
                        'message': f'Counter with id {counter_id} does not exist'
                    }, status=404)
            else:
                desired_map[btn_idx] = None  # explicit unmap

        # read current state: {button_index -> CounterConfig}
        if group:
            current_gcbm_rows = list(
                GroupCounterButtonMapping.objects
                .filter(group=group, dispenser=dispenser)
                .select_related('counter')
            )
            current_map = {row.button_index: row.counter for row in current_gcbm_rows}
        else:
            ctdm_rows = list(
                CounterTokenDispenserMapping.objects
                .filter(dispenser=dispenser)
                .select_related('counter')
                .order_by('id')
            )
            current_map = {
                BUTTON_INDEX_SEQUENCE[i]: row.counter
                for i, row in enumerate(ctdm_rows)
                if i < len(BUTTON_INDEX_SEQUENCE)
            }


        # Collect displaced dispenser IDs BEFORE the transaction so we can
        # rebuild their GCBM after the payload dispenser's new state is committed.
        # A displaced dispenser is any group dispenser (other than the payload one)
        # that currently holds a GCBM row for a counter being moved into this
        # dispenser's desired state.
        displaced_dispenser_ids = set()
        if group:
            for incoming_ctr in desired_map.values():
                if incoming_ctr is not None:
                    for old_gcbm in GroupCounterButtonMapping.objects.filter(
                        group=group, counter=incoming_ctr
                    ).exclude(dispenser=dispenser):
                        displaced_dispenser_ids.add(old_gcbm.dispenser_id)

        # apply atomically
        results = []
        with transaction.atomic():

            if group:
                # ----------------------------------------------------------------
                # "buttons" is the COMPLETE desired final state for this dispenser.
                # ----------------------------------------------------------------

                # Snapshot EVICTED counters BEFORE wiping CTDM.
                # Evicted = currently on this dispenser's CTDM but NOT in the payload.
                # These will be handed to the displaced dispenser to complete the swap.
                # (e.g. D2 has Counter B; payload asks D2 to take Counter A instead;
                #  Counter B is evicted → must go to D1 which lost Counter A)
                incoming_counter_ids = {c.id for c in desired_map.values() if c is not None}
                evicted_counter_objs = [
                    row.counter
                    for row in CounterTokenDispenserMapping.objects
                        .filter(dispenser=dispenser)
                        .exclude(counter_id__in=incoming_counter_ids)
                        .select_related('counter')
                ]

                # Wipe every existing GCBM row for this dispenser.
                GroupCounterButtonMapping.objects.filter(
                    group=group,
                    dispenser=dispenser,
                ).delete()

                # Wipe ALL CTDM rows for this dispenser atomically.
                CounterTokenDispenserMapping.objects.filter(dispenser=dispenser).delete()

                # Now apply the desired state.
                for btn_idx, new_ctr in desired_map.items():
                    prev_ctr = current_map.get(btn_idx)

                    if new_ctr is None:
                        # Explicit unmap of this slot (counter_id: null in payload).
                        results.append({
                            'button_index': btn_idx,
                            'operation': 'unmap',
                            'unmapped_counter_id': prev_ctr.id if prev_ctr else None,
                        })
                    else:
                        # Remove this counter from any other dispenser's GCBM slot.
                        GroupCounterButtonMapping.objects.filter(
                            group=group, counter=new_ctr
                        ).delete()

                        # Ensure CTDM: new counter -> this dispenser.
                        CounterTokenDispenserMapping.objects.filter(
                            counter=new_ctr
                        ).exclude(dispenser=dispenser).delete()
                        CounterTokenDispenserMapping.objects.get_or_create(
                            counter=new_ctr, dispenser=dispenser
                        )

                        # Create the new GCBM row.
                        GroupCounterButtonMapping.objects.create(
                            group=group,
                            dispenser=dispenser,
                            counter=new_ctr,
                            button_index=btn_idx,
                        )

                        # Label operation for response.
                        if prev_ctr is None:
                            op_name = 'remap'
                        elif prev_ctr.id == new_ctr.id:
                            op_name = 'no_change'
                        else:
                            current_counter_ids = {c.id for c in current_map.values() if c}
                            op_name = 'swap' if new_ctr.id in current_counter_ids else 'remap'

                        results.append({
                            'button_index': btn_idx,
                            'operation': op_name,
                            'new_counter_id': new_ctr.id,
                            'new_counter_name': new_ctr.counter_name,
                            'previous_counter_id': prev_ctr.id if prev_ctr else None,
                        })

                # ----------------------------------------------------------------
                # Rebuild GCBM for displaced dispensers and complete the true swap.
                #
                # Root cause of the original bug:
                #   - D2's CTDM was wiped (Counter B deleted from D2).
                #   - Counter A's CTDM was removed from D1 (line above).
                #   - D1's CTDM is now empty → rebuild gives D1 zero counters.
                #
                # Fix: Give D1 the evicted counter(s) from D2 (Counter B), then
                # rebuild D1's GCBM so the swap is truly bidirectional.
                # ----------------------------------------------------------------
                for d_id in displaced_dispenser_ids:
                    try:
                        disp_obj = Device.objects.get(id=d_id)

                        # Step 1: Route evicted counters to the displaced dispenser.
                        for ev_ctr in evicted_counter_objs:
                            CounterTokenDispenserMapping.objects.filter(
                                counter=ev_ctr
                            ).exclude(dispenser=disp_obj).delete()
                            CounterTokenDispenserMapping.objects.get_or_create(
                                counter=ev_ctr, dispenser=disp_obj
                            )

                        # Step 2: Clear stale GCBM rows for displaced dispenser.
                        GroupCounterButtonMapping.objects.filter(
                            group=group, dispenser=disp_obj
                        ).delete()

                        # Step 3: Rebuild GCBM from post-swap CTDM state
                        # (includes evicted counters added in step 1).
                        remaining = list(
                            CounterTokenDispenserMapping.objects
                            .filter(dispenser=disp_obj)
                            .select_related('counter')
                            .order_by('id')
                        )
                        for pos, ctdm_row in enumerate(remaining, start=1):
                            try:
                                btn_char = get_button_index_char(pos)
                            except ValueError:
                                btn_char = chr(0x31)
                            GroupCounterButtonMapping.objects.get_or_create(
                                group=group,
                                dispenser=disp_obj,
                                counter=ctdm_row.counter,
                                defaults={'button_index': btn_char},
                            )
                    except Exception:
                        pass  # Never fail the whole swap due to displaced rebuild


            else:
                # No group: CTDM insertion-order is the source of truth.
                # Merge current + desired, then recreate CTDM rows in index order.
                merged = dict(current_map)
                merged.update(desired_map)

                CounterTokenDispenserMapping.objects.filter(dispenser=dispenser).delete()

                for btn_idx in sorted(
                    merged.keys(),
                    key=lambda x: BUTTON_INDEX_SEQUENCE.index(x) if x in BUTTON_INDEX_SEQUENCE else 999
                ):
                    ctr = merged[btn_idx]
                    if ctr is not None:
                        CounterTokenDispenserMapping.objects.filter(
                            counter=ctr
                        ).exclude(dispenser=dispenser).delete()
                        CounterTokenDispenserMapping.objects.get_or_create(
                            counter=ctr, dispenser=dispenser
                        )

                for btn_idx, new_ctr in desired_map.items():
                    prev_ctr = current_map.get(btn_idx)
                    if new_ctr is None:
                        results.append({'button_index': btn_idx, 'operation': 'unmap',
                                        'unmapped_counter_id': prev_ctr.id if prev_ctr else None})
                    elif prev_ctr and prev_ctr.id == new_ctr.id:
                        results.append({'button_index': btn_idx, 'operation': 'no_change',
                                        'counter_id': new_ctr.id})
                    else:
                        results.append({
                            'button_index': btn_idx,
                            'operation': 'updated',
                            'new_counter_id': new_ctr.id,
                            'new_counter_name': new_ctr.counter_name,
                            'previous_counter_id': prev_ctr.id if prev_ctr else None,
                        })

        # FCM to affected TVs
        try:
            from .fcm_service import send_fcm_notifications
            affected_tv_ids = TVDispenserMapping.objects.filter(
                dispenser=dispenser
            ).values_list('tv_id', flat=True).distinct()
            affected_tvs = list(Device.objects.filter(
                id__in=affected_tv_ids, device_type=Device.DeviceType.TV
            ))
            if affected_tvs:
                send_fcm_notifications(
                    affected_tvs,
                    message="Counter mapping updated. Please reload your configuration.",
                    title="Counter Mapping Update",
                    data={"action": "counter_swap"},
                )
        except Exception as fcm_exc:
            action_logger.error(f"FCM notification error during swap: {fcm_exc}")

        return Response({
            'status': 'success',
            'message': f'Dispenser "{serial_number}" button mappings updated successfully.',
            'dispenser_serial_number': serial_number,
            'results': results,
        }, status=200)

    except Exception as e:
        log_api_response('swap_counters_api', 500, error=str(e))
        action_logger.error(f"Counter Swap Failed | Error: {str(e)}")
        return Response({
            'status': 'error',
            'message': f'An error occurred while updating counter mappings: {str(e)}'
        }, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
def get_token_dispenser_config_api(request):
    """
    API for fetching token dispenser configuration in JSON format with mapped counter details.
    Accepts: { "serial_number": "..." } or { "dispenser_id": ... }
    Returns: JSON response with device info, config, and mapped counters with button_index
    """
    serial_number = request.data.get('serial_number')
    dispenser_id = request.data.get('dispenser_id')
    
    if not serial_number and not dispenser_id:
        log_api_response('get_token_dispenser_config_api', 400, error='Either serial_number or dispenser_id is required')
        return Response({
            'status': 'error',
            'message': 'Either serial_number or dispenser_id is required'
        }, status=400)
    
    # Fetch token dispenser device
    try:
        if dispenser_id:
            # Check if dispenser_id is a numeric ID or a serial number
            try:
                # Try to convert to integer - if successful, it's an ID
                dispenser_id_int = int(dispenser_id)
                dispenser = Device.objects.get(id=dispenser_id_int, device_type=Device.DeviceType.TOKEN_DISPENSER)
            except (ValueError, TypeError):
                # If not numeric, treat it as a serial number
                dispenser = Device.objects.get(serial_number=dispenser_id, device_type=Device.DeviceType.TOKEN_DISPENSER)
        else:
            dispenser = Device.objects.get(serial_number=serial_number, device_type=Device.DeviceType.TOKEN_DISPENSER)
    except Device.DoesNotExist:
        log_api_response('get_token_dispenser_config_api', 404, error='Token dispenser not found')
        return Response({
            'status': 'error',
            'message': 'Token dispenser not found'
        }, status=404)
    
    # Get device configuration
    config_data = {}
    if hasattr(dispenser, 'config') and dispenser.config:
        config_data = dispenser.config.config_json or {}
    
    # Build configuration object in JSON format (not string)
    config_json = {
        'header1': config_data.get('header1') or config_data.get('company_name', 'CallQ'),
        'header2': config_data.get('header2') or config_data.get('location', 'Token'),
        'header3': config_data.get('header3', 'System'),
        'footer1': config_data.get('footer1', 'Thank You'),
        'footer2': config_data.get('footer2', 'Visit Again'),
        'day_wise_reset': '2' if config_data.get('day_wise_reset', '0') == '1' else '1',
        'reset_tkn': '1' if config_data.get('reset_tkn', True) else '0',
        'token_cut': config_data.get('token_cut', 'full'),
        'feed': '1' if config_data.get('feed', True) else '0',
        'logo_enable': '1' if config_data.get('logo_enable', True) else '0',
        'button_mode': config_data.get('button_mode', 'counterwise'),
        'token_label': config_data.get('token_label', 'Token'),
        'paper_out': '1' if config_data.get('paper_out', True) else '0',
        'type': config_data.get('type', '0'),
        'duplicate_print': '1' if config_data.get('duplicate_print', '0') == '1' else '0',
        'common_pool': '1' if config_data.get('common_pool', '0') == '1' else '0',
        'standalone': '1' if config_data.get('standalone', '0') == '1' else '0',
        'initial_print': '1' if config_data.get('initial_print', '0') == '1' else '0',
    }
    
    # Get all counters mapped to the TV connected to this token dispenser
    # First find TVs via TVDispenserMapping (new system), then fallback to old Mapping/ButtonMapping methods
    # NOTE: Build tv_counters FIRST so we can use their button_index for mapped_counters
    tv_counters = []
    connected_tvs = set()
    
    try:
        # Method 1: Find TVs via TVDispenserMapping (new system - primary method)
        # Find all TVs that have this dispenser mapped
        tv_dispenser_mappings_for_this_dispenser = TVDispenserMapping.objects.filter(
            dispenser=dispenser
        ).select_related('tv')
        
        for tv_mapping in tv_dispenser_mappings_for_this_dispenser:
            if tv_mapping.tv:
                connected_tvs.add(tv_mapping.tv)
                
        # Method 1.5: Find TVs via TVKeypadMapping
        tv_keypad_mappings_for_this_dispenser = TVKeypadMapping.objects.filter(
            dispenser=dispenser
        ).select_related('tv')
        
        for tv_mapping in tv_keypad_mappings_for_this_dispenser:
            if tv_mapping.tv:
                connected_tvs.add(tv_mapping.tv)
        
        # Method 2: Fallback - Direct Mapping (TD -> TV) - old system
        if not connected_tvs:
            mapping = Mapping.objects.filter(token_dispenser=dispenser).select_related('tv').first()
            if mapping and mapping.tv:
                connected_tvs.add(mapping.tv)
        
        # Method 3: Fallback - Indirect connection via ButtonMapping chain (TD -> Keypad -> Broker -> TV)
        if not connected_tvs:
            button_mappings = ButtonMapping.objects.filter(source_device=dispenser).select_related('target_device')
            
            for bm in button_mappings:
                target = bm.target_device
                
                # If target is a TV, add it
                if target.device_type == Device.DeviceType.TV:
                    connected_tvs.add(target)
                
                # If target is a Keypad, check if Keypad -> Broker -> TV
                elif target.device_type == Device.DeviceType.KEYPAD:
                    # Find ButtonMappings from this Keypad
                    keypad_mappings = ButtonMapping.objects.filter(source_device=target).select_related('target_device')
                    for kp_bm in keypad_mappings:
                        broker_or_tv = kp_bm.target_device
                        
                        # If Keypad -> TV directly
                        if broker_or_tv.device_type == Device.DeviceType.TV:
                            connected_tvs.add(broker_or_tv)
                        
                        # If Keypad -> Broker, check Broker -> TV
                        elif broker_or_tv.device_type == Device.DeviceType.BROKER:
                            broker_mappings = ButtonMapping.objects.filter(source_device=broker_or_tv).select_related('target_device')
                            for br_bm in broker_mappings:
                                if br_bm.target_device.device_type == Device.DeviceType.TV:
                                    connected_tvs.add(br_bm.target_device)
                
                # If target is a Broker, check Broker -> TV
                elif target.device_type == Device.DeviceType.BROKER:
                    broker_mappings = ButtonMapping.objects.filter(source_device=target).select_related('target_device')
                    for br_bm in broker_mappings:
                        if br_bm.target_device.device_type == Device.DeviceType.TV:
                            connected_tvs.add(br_bm.target_device)
        
        # Get all counters from all dispensers mapped to connected TVs
        # For each TV, get ALL dispensers mapped to it, then get counters from each dispenser
        for tv in connected_tvs:
            # Get ALL dispensers mapped to this TV (not just the current dispenser)
            tv_dispenser_mappings = TVDispenserMapping.objects.filter(
                tv=tv
            ).select_related('dispenser')
            
            # For each dispenser mapped to this TV, get all its counters
            for tv_dispenser_mapping in tv_dispenser_mappings:
                mapped_dispenser = tv_dispenser_mapping.dispenser
                button_index = tv_dispenser_mapping.button_index
                
                # Get all counters from this dispenser
                counter_mappings = CounterTokenDispenserMapping.objects.filter(
                    dispenser=mapped_dispenser
                ).select_related('counter')
                
                # Add all counters from this dispenser with the button_index from TVDispenserMapping
                for counter_mapping in counter_mappings:
                    counter = counter_mapping.counter
                    # Avoid duplicates by checking counter_id
                    if not any(tc['counter_id'] == counter.id for tc in tv_counters):
                        tv_counters.append({
                            'counter_id': counter.id,
                            'counter_name': counter.counter_name,
                            'counter_display_name': counter.counter_display_name,
                            'counter_prefix_code': counter.counter_prefix_code,
                            'max_token_number': counter.max_token_number,
                            'status': counter.status,
                            'button_index': button_index,  # Use button_index from TVDispenserMapping
                            'dispenser_s_no': mapped_dispenser.serial_number,
                        })
    except Exception as e:
        # If there's no TV mapping or any error, return empty list
        tv_counters = []
    
    # -------------------------------------------------------------------------
    # mapped_counters: counters from ALL dispensers in this dispenser's group.
    # If the dispenser has no group, returns only its own counters.
    # Each entry carries the dispenser's stored group button index.
    # -------------------------------------------------------------------------
    mapped_counters = []
    try:
        _gdm_self = GroupDispenserMapping.objects.filter(dispenser=dispenser).select_related('group').first()
        _group_obj = _gdm_self.group if _gdm_self else None

        # Build a lookup: counter_id → group-wide button_index from DB
        _gcbm_lookup = {}
        if _group_obj:
            for _gcbm in GroupCounterButtonMapping.objects.filter(group=_group_obj):
                _gcbm_lookup[_gcbm.counter_id] = _gcbm.button_index

        if _group_obj:
            _all_gdm = list(
                _group_obj.dispenser_slot_mappings
                           .select_related('dispenser')
                           .order_by('dispenser_button_index')
            )
        else:
            # Singleton: only this dispenser
            _all_gdm = [type('_FakeSlot', (), {
                'dispenser': dispenser,
                'dispenser_button_index': None,
            })()]

        _seen_mc = set()
        for _slot in _all_gdm:
            _disp = _slot.dispenser

            # Resolve TV-slot button_index and keypad_index for this dispenser.
            _this_button_index = None
            _this_keypad_index = None
            for tv in connected_tvs:
                _tv_disp_map = TVDispenserMapping.objects.filter(
                    tv=tv, dispenser=_disp
                ).first()
                if _tv_disp_map:
                    _this_button_index = _tv_disp_map.button_index

                _tv_keypad_map = TVKeypadMapping.objects.filter(
                    tv=tv, dispenser=_disp
                ).first()
                if _tv_keypad_map:
                    _this_keypad_index = _tv_keypad_map.keypad_index

                if _this_button_index or _this_keypad_index:
                    break

            # Use GCBM (true current state) for group dispensers; CTDM for singletons.
            if _group_obj:
                _disp_counter_mappings = GroupCounterButtonMapping.objects.filter(
                    group=_group_obj, dispenser=_disp
                ).select_related('counter').order_by('button_index')
            else:
                _disp_counter_mappings = CounterTokenDispenserMapping.objects.filter(
                    dispenser=_disp
                ).select_related('counter').order_by('id')

            for _fb_idx, _cm in enumerate(_disp_counter_mappings, start=1):
                _counter = _cm.counter
                if _counter.id in _seen_mc:
                    continue
                _seen_mc.add(_counter.id)

                # button_index: TV slot (same for all counters of this dispenser).
                if _this_button_index is not None:
                    _final_bi = str(_this_button_index)
                else:
                    try:
                        _final_bi = get_button_index_char(_fb_idx)
                    except ValueError:
                        _final_bi = chr(0x31)

                # dispenser_button_index: group-wide sequential index from DB only.
                # No fallback — the DB (GroupCounterButtonMapping) is the single
                # source of truth.  Returns None when no DB record exists.
                _dispenser_btn_idx = str(_gcbm_lookup[_counter.id]) if _counter.id in _gcbm_lookup else None

                mapped_counters.append({
                    'counter_id': _counter.id,
                    'counter_name': _counter.counter_name,
                    'counter_display_name': _counter.counter_display_name,
                    'counter_prefix_code': _counter.counter_prefix_code,
                    'max_token_number': _counter.max_token_number,
                    'status': _counter.status,
                    'button_index': _final_bi,
                    'dispenser_button_index': _dispenser_btn_idx,
                    'keypad_button_index': str(_this_keypad_index) if _this_keypad_index is not None else None,
                })
    except Exception:
        mapped_counters = []

    # -------------------------------------------------------------------------
    # group_counters: counters from ALL dispensers in the same FamilyMapping
    # group as this dispenser.  If this dispenser is not in any family, falls
    # back to only this dispenser's own counters (same as mapped_counters).
    # Each entry records which dispenser it came from and that dispenser's
    # button_index on the shared TV.
    # -------------------------------------------------------------------------
    group_counters = []
    try:
        gdm_self = GroupDispenserMapping.objects.filter(dispenser=dispenser).select_related('group').first()
        group_obj = gdm_self.group if gdm_self else None

        # Build a lookup: counter_id → group-wide button_index from DB
        gcbm_lookup = {}
        if group_obj:
            for gcbm in GroupCounterButtonMapping.objects.filter(group=group_obj):
                gcbm_lookup[gcbm.counter_id] = gcbm.button_index

        if group_obj:
            all_gdm = list(
                group_obj.dispenser_slot_mappings
                         .select_related('dispenser')
                         .order_by('dispenser_button_index')
            )
        else:
            all_gdm = [gdm_self] if gdm_self else []

        seen_counter_ids = set()
        for slot in all_gdm:
            fam_dispenser = slot.dispenser

            # Resolve TV-slot button_index and keypad_index for this sibling dispenser.
            fam_button_index = None
            fam_keypad_index = None
            for tv in connected_tvs:
                tv_disp_map = TVDispenserMapping.objects.filter(
                    tv=tv, dispenser=fam_dispenser
                ).first()
                if tv_disp_map:
                    fam_button_index = tv_disp_map.button_index

                tv_keypad_map = TVKeypadMapping.objects.filter(
                    tv=tv, dispenser=fam_dispenser
                ).first()
                if tv_keypad_map:
                    fam_keypad_index = tv_keypad_map.keypad_index

                if fam_button_index or fam_keypad_index:
                    break

            # Use GCBM (true current state) for group dispensers; CTDM for singletons.
            if group_obj:
                fam_counter_mappings = GroupCounterButtonMapping.objects.filter(
                    group=group_obj, dispenser=fam_dispenser
                ).select_related('counter').order_by('button_index')
            else:
                fam_counter_mappings = CounterTokenDispenserMapping.objects.filter(
                    dispenser=fam_dispenser
                ).select_related('counter').order_by('id')

            for fallback_index, cm in enumerate(fam_counter_mappings, start=1):
                counter = cm.counter
                if counter.id in seen_counter_ids:
                    continue
                seen_counter_ids.add(counter.id)

                # button_index: TV slot index for this dispenser (same for all its counters).
                if fam_button_index is not None:
                    final_bi = str(fam_button_index)
                else:
                    try:
                        final_bi = get_button_index_char(fallback_index)
                    except ValueError:
                        final_bi = chr(0x31)

                # dispenser_button_index: group-wide sequential index from DB only.
                # No fallback — the DB (GroupCounterButtonMapping) is the single
                # source of truth.  Returns None when no DB record exists.
                fam_dispenser_btn_idx = str(gcbm_lookup[counter.id]) if counter.id in gcbm_lookup else None

                group_counters.append({
                    'counter_id': counter.id,
                    'counter_name': counter.counter_name,
                    'counter_display_name': counter.counter_display_name,
                    'counter_prefix_code': counter.counter_prefix_code,
                    'max_token_number': counter.max_token_number,
                    'status': counter.status,
                    'button_index': final_bi,
                    'dispenser_button_index': fam_dispenser_btn_idx,
                    'keypad_button_index': str(fam_keypad_index) if fam_keypad_index is not None else None,
                    'dispenser_s_no': fam_dispenser.serial_number,
                })
    except Exception:
        group_counters = list(mapped_counters)

    # -------------------------------------------------------------------------
    # unmapped_counters: all CounterConfig records NOT mapped to ANY dispenser
    # -------------------------------------------------------------------------
    unmapped_counters = []
    try:
        mapped_counter_ids = CounterTokenDispenserMapping.objects.values_list(
            'counter_id', flat=True
        )
        unmapped_qs = CounterConfig.objects.exclude(
            id__in=mapped_counter_ids
        )
        # Scope to the dispenser's company so cross-company counters are not exposed
        if dispenser.company_id:
            unmapped_qs = unmapped_qs.filter(company_id=dispenser.company_id)
        unmapped_qs = unmapped_qs.order_by('counter_name')
        for uc in unmapped_qs:
            unmapped_counters.append({
                'counter_id': uc.id,
                'counter_name': uc.counter_name,
                'counter_display_name': uc.counter_display_name,
                'counter_prefix_code': uc.counter_prefix_code,
                'max_token_number': uc.max_token_number,
                'status': uc.status,
            })
    except Exception:
        unmapped_counters = []

    # -------------------------------------------------------------------------
    # own_counters: ONLY the counters whose GroupCounterButtonMapping row points
    # to THIS dispenser.  A counter swapped to another dispenser is excluded even
    # if a stale CTDM row still references this dispenser.
    # For dispensers not in any group, falls back to CTDM order.
    # -------------------------------------------------------------------------
    own_counters = []
    try:
        # Resolve TV-slot and keypad indices for this dispenser
        _own_btn_index = None
        _own_kp_index  = None
        for tv in connected_tvs:
            _tdm = TVDispenserMapping.objects.filter(tv=tv, dispenser=dispenser).first()
            if _tdm:
                _own_btn_index = _tdm.button_index
            _tkm = TVKeypadMapping.objects.filter(tv=tv, dispenser=dispenser).first()
            if _tkm:
                _own_kp_index = _tkm.keypad_index
            if _own_btn_index or _own_kp_index:
                break

        _own_gdm_row = GroupDispenserMapping.objects.filter(dispenser=dispenser).select_related('group').first()
        _own_group   = _own_gdm_row.group if _own_gdm_row else None

        if _own_group:
            # Source of truth: GCBM rows scoped to THIS dispenser only.
            # Counters swapped away will have their GCBM row pointing to the new
            # dispenser, so they won't appear here.
            own_gcbm_rows = (
                GroupCounterButtonMapping.objects
                .filter(group=_own_group, dispenser=dispenser)
                .select_related('counter')
                .order_by('button_index')
            )
            for _gcbm_row in own_gcbm_rows:
                _c   = _gcbm_row.counter
                _odbi = str(_gcbm_row.button_index)

                if _own_btn_index is not None:
                    _obi = str(_own_btn_index)
                else:
                    try:
                        _obi = get_button_index_char(int(_odbi))
                    except (ValueError, TypeError):
                        _obi = chr(0x31)

                own_counters.append({
                    'counter_id': _c.id,
                    'counter_name': _c.counter_name,
                    'counter_display_name': _c.counter_display_name,
                    'counter_prefix_code': _c.counter_prefix_code,
                    'max_token_number': _c.max_token_number,
                    'status': _c.status,
                    'button_index': _obi,
                    'dispenser_button_index': _odbi,
                    'keypad_button_index': str(_own_kp_index) if _own_kp_index is not None else None,
                })
        else:
            # No group: standalone dispenser — use CTDM order as before.
            own_mappings = CounterTokenDispenserMapping.objects.filter(
                dispenser=dispenser
            ).select_related('counter').order_by('id')

            for _oi, _cm in enumerate(own_mappings, start=1):
                _c = _cm.counter
                if _own_btn_index is not None:
                    _obi = str(_own_btn_index)
                else:
                    try:
                        _obi = get_button_index_char(_oi)
                    except ValueError:
                        _obi = chr(0x31)
                _odbi = None  # No group → no DB-stored dispenser_button_index
                own_counters.append({
                    'counter_id': _c.id,
                    'counter_name': _c.counter_name,
                    'counter_display_name': _c.counter_display_name,
                    'counter_prefix_code': _c.counter_prefix_code,
                    'max_token_number': _c.max_token_number,
                    'status': _c.status,
                    'button_index': _obi,
                    'dispenser_button_index': _odbi,
                    'keypad_button_index': str(_own_kp_index) if _own_kp_index is not None else None,
                })
    except Exception:
        own_counters = []

    # Build response
    response_data = {
        'status': 'success',
        'message': 'Token dispenser configuration fetched successfully',
        'device': {
            'id': dispenser.id,
            'serial_number': dispenser.serial_number,
            'display_name': dispenser.display_name,
            'device_type': dispenser.device_type,
            'token_type': dispenser.token_type,
            'random_number': dispenser.random_number
        },
        'config': config_json,
        'mapped_counters': mapped_counters,
        'own_counters': own_counters,
        'group_counters': group_counters,
        'unmapped_counters': unmapped_counters,
    }
    
    log_api_response('get_token_dispenser_config_api', 200, {
        'status': 'success',
        'dispenser_serial_number': dispenser.serial_number,
        'mapped_counters_count': len(mapped_counters)
    })
    action_logger.info(f"Token Dispenser Config Fetched | Dispenser: {dispenser.serial_number} | Mapped Counters: {len(mapped_counters)}")
    return Response(response_data)


@api_view(['POST'])
@permission_classes([AllowAny])
def external_device_counter_api(request):
    """
    External API endpoint for device-counter mapping.
    POST /api/external/device-counter
    
    Request Body:
    {
        "device_id": "TV001",
        "counter_name": "Counter A"
    }
    """
    device_id = request.data.get('device_id')
    counter_name = request.data.get('counter_name')
    
    if not device_id or not counter_name:
        return Response({
            'status': 'error',
            'message': 'device_id and counter_name are required'
        }, status=400)
    
    # Validate device exists (by serial_number or mac_address)
    device = None
    try:
        device = Device.objects.get(
            Q(serial_number=device_id) | Q(mac_address=device_id),
            device_type=Device.DeviceType.TV
        )
    except Device.DoesNotExist:
        return Response({
            'status': 'error',
            'message': 'Invalid device_id: Device not found'
        }, status=404)
    except Device.MultipleObjectsReturned:
        device = Device.objects.filter(
            Q(serial_number=device_id) | Q(mac_address=device_id),
            device_type=Device.DeviceType.TV
        ).first()
    
    # Validate counter exists
    counter = None
    try:
        counter = CounterConfig.objects.get(counter_name=counter_name, status=True)
    except CounterConfig.DoesNotExist:
        # Log even if counter doesn't exist
        try:
            log_entry = ExternalDeviceCounterLog.objects.create(
                device_id=device_id,
                counter_name=counter_name,
                api_payload=request.data
            )
        except Exception:
            pass  # Handle case where tables don't exist
        return Response({
            'status': 'error',
            'message': 'Invalid counter_name: Counter not found'
        }, status=404)
    except Exception:
        return Response({
            'status': 'error',
            'message': 'Database error: Counter tables not found. Please run migrations.'
        }, status=500)
    
    # Create log entry
    try:
        log_entry = ExternalDeviceCounterLog.objects.create(
            device_id=device_id,
            counter_name=counter_name,
            counter=counter,
            api_payload=request.data
        )
        
        # Auto-map counter to TV if not already mapped
        TVCounterMapping.objects.get_or_create(tv=device, counter=counter)
    except Exception:
        return Response({
            'status': 'error',
            'message': 'Database error: Counter mapping tables not found. Please run migrations.'
        }, status=500)
    
    return Response({
        'status': 'success',
        'message': 'Counter mapped successfully',
        'data': {
            'device_id': device_id,
            'counter_name': counter_name,
            'counter_id': counter.id
        }
    }, status=200)


# ============================================================================
# Counter Management UI Views
# ============================================================================

@login_required
def counter_list(request):
    """List counters for the current user's company."""
    company = getattr(request.user, 'company_relation', None)
    try:
        qs = CounterConfig.objects.all()
        if company:
            qs = qs.filter(company=company)
        counters = qs.order_by('counter_name')
    except Exception:
        counters = []
        messages.warning(request, 'Counter tables not found. Please run migrations first.')
    
    return render(request, 'configdetails/counter_list.html', {
        'counters': counters
    })


@login_required
def counter_create(request):
    """Create a new counter for the current user's company."""
    company = getattr(request.user, 'company_relation', None)
    if request.method == 'POST':
        try:
            counter = CounterConfig.objects.create(
                company=company,
                counter_name=request.POST.get('counter_name'),
                counter_prefix_code=request.POST.get('counter_prefix_code'),
                counter_display_name=request.POST.get('counter_display_name'),
                max_token_number=int(request.POST.get('max_token_number', 0)),
                status=request.POST.get('status') == 'on'
            )
            messages.success(request, f'Counter "{counter.counter_name}" created successfully!')
            return redirect('counter_list')
        except Exception as e:
            messages.error(request, f'Error creating counter: {str(e)}')
    
    return render(request, 'configdetails/counter_form.html', {
        'counter': None,
        'action': 'Create'
    })


@login_required
def counter_edit(request, counter_id):
    """Edit an existing counter – only if it belongs to the user's company."""
    company = getattr(request.user, 'company_relation', None)
    try:
        # Restrict to the user's company so one company cannot edit another's counter
        qs = CounterConfig.objects.all()
        if company:
            qs = qs.filter(company=company)
        counter = qs.filter(id=counter_id).first()
        if not counter:
            messages.error(request, 'Counter not found or you do not have permission to edit it.')
            return redirect('counter_list')
    except Exception:
        messages.error(request, 'Counter table not found. Please run migrations first.')
        return redirect('counter_list')
    
    if request.method == 'POST':
        try:
            counter.counter_name = request.POST.get('counter_name')
            counter.counter_prefix_code = request.POST.get('counter_prefix_code')
            counter.counter_display_name = request.POST.get('counter_display_name')
            counter.max_token_number = int(request.POST.get('max_token_number', 0))
            counter.status = request.POST.get('status') == 'on'
            counter.save()
            messages.success(request, f'Counter "{counter.counter_name}" updated successfully!')
            return redirect('counter_list')
        except Exception as e:
            messages.error(request, f'Error updating counter: {str(e)}')
    
    return render(request, 'configdetails/counter_form.html', {
        'counter': counter,
        'action': 'Edit'
    })


@login_required
def counter_delete(request, counter_id):
    """Delete a counter – only if it belongs to the user's company."""
    company = getattr(request.user, 'company_relation', None)
    try:
        qs = CounterConfig.objects.all()
        if company:
            qs = qs.filter(company=company)
        counter = qs.filter(id=counter_id).first()
        if not counter:
            messages.error(request, 'Counter not found or you do not have permission to delete it.')
            return redirect('counter_list')
    except Exception:
        messages.error(request, 'Counter table not found. Please run migrations first.')
        return redirect('counter_list')
    
    if request.method == 'POST':
        try:
            # Check if counter is mapped
            if TVCounterMapping.objects.filter(counter=counter).exists():
                messages.error(request, f'Cannot delete counter "{counter.counter_name}": It is mapped to one or more TV devices.')
                return redirect('counter_list')
            
            if CounterTokenDispenserMapping.objects.filter(counter=counter).exists():
                messages.error(request, f'Cannot delete counter "{counter.counter_name}": It is mapped to one or more token dispensers.')
                return redirect('counter_list')
            
            counter_name = counter.counter_name
            counter.delete()
            messages.success(request, f'Counter "{counter_name}" deleted successfully!')
        except Exception as e:
            messages.error(request, f'Error deleting counter: {str(e)}')
    
    return redirect('counter_list')


# ============================================================================
# Log Viewing Views
# ============================================================================

def log_list(request):
    """List available log dates organized by year/month/day with dropdown selection"""
    BASE_LOG_DIR = Path(settings.BASE_DIR) / 'CallQ_logs'
    
    if not BASE_LOG_DIR.exists():
        messages.warning(request, 'Log directory not found.')
        return render(request, 'configdetails/log_list.html', {
            'years': [],
            'selected_year': None,
            'selected_month': None,
            'selected_day': None,
            'months': [],
            'days': [],
            'log_info': None
        })
    
    # Organize logs by year/month/day structure
    log_structure = {}
    
    try:
        for year_dir in sorted(BASE_LOG_DIR.iterdir(), reverse=True):
            if not year_dir.is_dir():
                continue
            
            year = year_dir.name
            log_structure[year] = {}
            
            for month_dir in sorted(year_dir.iterdir(), reverse=True):
                if not month_dir.is_dir():
                    continue
                
                month = month_dir.name
                log_structure[year][month] = {}
                
                for day_dir in sorted(month_dir.iterdir(), reverse=True):
                    if not day_dir.is_dir():
                        continue
                    
                    day = day_dir.name
                    
                    # Check if log files exist
                    request_hits_log = day_dir / 'request_hits.log'
                    actions_log = day_dir / 'actions.log'
                    
                    has_request_hits = request_hits_log.exists() and request_hits_log.stat().st_size > 0
                    has_actions = actions_log.exists() and actions_log.stat().st_size > 0
                    
                    if has_request_hits or has_actions:
                        log_structure[year][month][day] = {
                            'has_request_hits': has_request_hits,
                            'has_actions': has_actions,
                            'request_hits_size': request_hits_log.stat().st_size if has_request_hits else 0,
                            'actions_size': actions_log.stat().st_size if has_actions else 0,
                        }
    except Exception as e:
        messages.error(request, f'Error reading log directory: {str(e)}')
        log_structure = {}
    
    # Get selected values from request
    selected_year = request.GET.get('year', '')
    selected_month = request.GET.get('month', '')
    selected_day = request.GET.get('day', '')
    
    # Prepare data for template
    years = sorted(log_structure.keys(), reverse=True)
    months = []
    days = []
    log_info = None
    
    if selected_year and selected_year in log_structure:
        months = sorted(log_structure[selected_year].keys(), reverse=True)
        
        if selected_month and selected_month in log_structure[selected_year]:
            days = sorted(log_structure[selected_year][selected_month].keys(), reverse=True)
            
            if selected_day and selected_day in log_structure[selected_year][selected_month]:
                log_info = log_structure[selected_year][selected_month][selected_day]
                log_info['year'] = selected_year
                log_info['month'] = selected_month
                log_info['day'] = selected_day
    
    return render(request, 'configdetails/log_list.html', {
        'years': years,
        'selected_year': selected_year,
        'selected_month': selected_month,
        'selected_day': selected_day,
        'months': months,
        'days': days,
        'log_info': log_info
    })


def log_view(request, year, month, day, log_type):
    """View log contents for a specific date and log type"""
    if log_type not in ['request_hits', 'actions']:
        messages.error(request, 'Invalid log type.')
        return redirect('log_list')
    
    BASE_LOG_DIR = Path(settings.BASE_DIR) / 'CallQ_logs'
    log_file_path = BASE_LOG_DIR / year / month / day / f'{log_type}.log'
    
    if not log_file_path.exists():
        messages.error(request, f'Log file not found for {year}/{month}/{day}.')
        return redirect('log_list')
    
    # Read log file
    try:
        with open(log_file_path, 'r', encoding='utf-8') as f:
            log_lines = f.readlines()
    except Exception as e:
        messages.error(request, f'Error reading log file: {str(e)}')
        return redirect('log_list')
    
    # Pagination
    page = request.GET.get('page', 1)
    paginator = Paginator(log_lines, 100)  # 100 lines per page
    
    try:
        log_page = paginator.page(page)
    except PageNotAnInteger:
        log_page = paginator.page(1)
    except EmptyPage:
        log_page = paginator.page(paginator.num_pages)
    
    # Get file size
    file_size = log_file_path.stat().st_size
    file_size_kb = file_size / 1024
    file_size_mb = file_size_kb / 1024
    
    # Calculate start index for line numbers
    start_index = (log_page.number - 1) * paginator.per_page + 1
    
    return render(request, 'configdetails/log_view.html', {
        'year': year,
        'month': month,
        'day': day,
        'log_type': log_type,
        'log_type_display': 'Request Hits' if log_type == 'request_hits' else 'Actions',
        'log_page': log_page,
        'start_index': start_index,
        'file_size': file_size,
        'file_size_kb': file_size_kb,
        'file_size_mb': file_size_mb,
        'total_lines': len(log_lines),
        'hide_sidebar': True,  # Hide sidebar for full-screen log view
    })


# =============================================================================
# External Device Registration API
# =============================================================================

@api_view(['POST'])
@permission_classes([AllowAny])
def register_device_api(request):
    """
    External API endpoint for registering a new device by serial number.
    POST /config/api/external/register-device

    Request Body:
    {
        "serial_number": "202505CAL000026",
        "device_type": "BROKER",              # TV | TOKEN_DISPENSER | KEYPAD | BROKER | LED
        "mac_address": "202505CAL000026",
        "time": "2026-02-26T11:51:20",        # optional timestamp from device
        "bluetooth_name": "Broker 1",         # stored in device_model column
        "display_name": "Broker 1"            # stored in display_name column
    }

    Responses:
        201 - Device created, waiting for licensing
        409 - Device already exists
        400 - Missing / invalid fields
    """
    log_api_request('register_device_api', request)

    serial_number = request.data.get('serial_number')
    device_type = request.data.get('device_type')
    mac_address = request.data.get('mac_address', '')
    device_time = request.data.get('time', '')
    bluetooth_name = request.data.get('bluetooth_name', '')  # stored in device_model
    display_name = request.data.get('display_name', '')      # stored in display_name

    # --- Validation ---
    if not serial_number or not device_type:
        log_api_response('register_device_api', 400, error='serial_number and device_type are required')
        return Response({
            'status': 'error',
            'message': 'serial_number and device_type are required'
        }, status=400)

    valid_types = [choice[0] for choice in Device.DeviceType.choices]
    if device_type not in valid_types:
        log_api_response('register_device_api', 400, error=f'Invalid device_type: {device_type}')
        return Response({
            'status': 'error',
            'message': f'Invalid device_type. Must be one of: {", ".join(valid_types)}'
        }, status=400)

    # --- Duplicate check ---
    if Device.objects.filter(serial_number=serial_number).exists():
        log_api_response('register_device_api', 409, error='Device already exists')
        return Response({
            'status': 'error',
            'message': 'Device already exists'
        }, status=409)

    # --- Create device + production serial entry inside a transaction ---
    try:
        with transaction.atomic():
            # 1. Create Device record (pending licensing)
            device = Device.objects.create(
                serial_number=serial_number,
                device_type=device_type,
                mac_address=mac_address if mac_address else None,
                device_model=bluetooth_name if bluetooth_name else 'PC',  # bluetooth_name -> device_model
                display_name=display_name if display_name else None,       # display_name -> display_name
                licence_status='Pending',
                is_active=True,
            )

            # 2. Get-or-create the "api" ProductionBatch
            api_batch, _ = ProductionBatch.objects.get_or_create(
                batch_id='api',
                defaults={'device_type': device_type},
            )

            # 3. Create ProductionSerialNumber entry
            ProductionSerialNumber.objects.create(
                batch=api_batch,
                serial_number=serial_number,
                device_type=device_type,
                is_registered=False,
            )

        log_api_response('register_device_api', 201, response_data={'serial_number': serial_number})
        return Response({
            'status': 'success',
            'message': 'Device is waiting for licensing',
            'data': {
                'serial_number': serial_number,
                'device_type': device_type,
                'mac_address': mac_address,
                'bluetooth_name': bluetooth_name,
                'display_name': display_name,
                'device_id': device.id,
            }
        }, status=201)

    except Exception as e:
        log_api_response('register_device_api', 500, error=str(e))
        return Response({
            'status': 'error',
            'message': f'Failed to register device: {str(e)}'
        }, status=500)


# ---------------------------------------------------------------------------
# Token Report API
# ---------------------------------------------------------------------------

@api_view(['POST'])
@permission_classes([AllowAny])
def token_report_api(request):
    """
    Store a token report entry.

    POST /api/external/token-report

    Request Body:
    {
        "received_message"  : "<raw token string>",
        "received_dateTime" : "2026-05-05T12:00:00",   // ISO-8601
        "displayed_dateTime": "2026-05-05T12:00:05",   // ISO-8601 (optional)
        "customerId"        : "CUST-001",
        "mac_address"       : "AA:BB:CC:DD:EE:FF"
    }

    Returns:
    {
        "status" : "success",
        "message": "Token report saved",
        "id"     : <pk of created record>
    }
    """
    from .models import TokenReport

    log_api_request('token_report_api', request)

    received_message   = request.data.get('received_message')
    received_dateTime  = request.data.get('received_dateTime')
    displayed_dateTime = request.data.get('displayed_dateTime')
    customerId         = request.data.get('customerId')
    mac_address        = request.data.get('mac_address')

    # If customerId is purely numeric (e.g. "0166"), convert to int to strip leading zeros
    if customerId:
        try:
            customerId = int(customerId)
        except (ValueError, TypeError):
            pass

    # --- Validate required fields ---
    missing = [
        field for field, val in [
            ('received_message',  received_message),
            ('received_dateTime', received_dateTime),
            ('customerId',        customerId),
            ('mac_address',       mac_address),
        ]
        if not val
    ]
    if missing:
        err = f"Missing required fields: {', '.join(missing)}"
        log_api_response('token_report_api', 400, error=err)
        return Response({'status': 'error', 'message': err}, status=400)

    # --- Parse datetimes ---
    from django.utils.dateparse import parse_datetime
    from django.utils import timezone

    def parse_dt(value, field_name):
        """Parse an ISO-8601 datetime string; return (dt, error_str)."""
        if not value:
            return None, None
        dt = parse_datetime(value)
        if dt is None:
            return None, f"Invalid datetime format for '{field_name}': {value!r}"
        # Make timezone-aware if naive
        if timezone.is_naive(dt):
            dt = timezone.make_aware(dt, timezone.get_current_timezone())
        return dt, None

    received_dt, err = parse_dt(received_dateTime, 'received_dateTime')
    if err:
        log_api_response('token_report_api', 400, error=err)
        return Response({'status': 'error', 'message': err}, status=400)

    displayed_dt, err = parse_dt(displayed_dateTime, 'displayed_dateTime')
    if err:
        log_api_response('token_report_api', 400, error=err)
        return Response({'status': 'error', 'message': err}, status=400)

    # --- Save ---
    try:
        report = TokenReport.objects.create(
            received_message=received_message,
            received_dateTime=received_dt,
            displayed_dateTime=displayed_dt,
            customerId=customerId,
            mac_address=mac_address,
        )
    except Exception as e:
        log_api_response('token_report_api', 500, error=str(e))
        return Response({
            'status': 'error',
            'message': f'Failed to save token report: {str(e)}'
        }, status=500)

    log_api_response('token_report_api', 201, response_data={'id': report.id})
    return Response({
        'status': 'success',
        'message': 'Token report saved',
        'id': report.id,
    }, status=201)


# ---------------------------------------------------------------------------
# Token Report - HTML Page View
# ---------------------------------------------------------------------------

@login_required
def token_report_list(request):
    """
    Paginated HTML page showing all TokenReport records.
    Supports GET filters: customer_id, mac_address, date_from, date_to.
    Scoped to the logged-in user's company / dealer customer.
    """
    from .models import TokenReport
    from companydetails.models import DealerCustomer
    from django.db.models import Q
    from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

    # --- Determine scope ---
    user = request.user
    qs = TokenReport.objects.all().order_by('-received_dateTime')

    # Scope by company / dealer customer if not super-admin / admin
    if hasattr(user, 'role'):
        if user.role not in ('SUPER_ADMIN', 'ADMIN'):
            scoped_ids = set()
            if hasattr(user, 'dealer_customer_relation') and user.dealer_customer_relation:
                dc = user.dealer_customer_relation
                raw = dc.customer_id
                scoped_ids.add(raw)
                try:
                    num = int(raw)
                    scoped_ids.add(str(num))         # strip leading zeros: "0166" → "166"
                    scoped_ids.add(str(num).zfill(3))# "166" → "166" or "016"
                    scoped_ids.add(str(num).zfill(4))# "166" → "0166"
                except (ValueError, TypeError):
                    pass
            elif hasattr(user, 'company_relation') and user.company_relation:
                c = user.company_relation
                # Collect raw candidate IDs (company_id CharField + DB primary key)
                raw_candidates = []
                if c.company_id:
                    raw_candidates.append(c.company_id)
                raw_candidates.append(str(c.id))
                # For each candidate, add raw value, int-normalised form, and zero-padded forms
                for raw in raw_candidates:
                    scoped_ids.add(raw)
                    try:
                        num = int(raw)
                        scoped_ids.add(str(num))         # "166"
                        scoped_ids.add(str(num).zfill(3))# "166" or "016"
                        scoped_ids.add(str(num).zfill(4))# "0166"
                    except (ValueError, TypeError):
                        pass
            if scoped_ids:
                qs = qs.filter(customerId__in=list(scoped_ids))

    # --- Filters ---
    filter_customer = request.GET.get('customer_id', '').strip()
    filter_mac      = request.GET.get('mac_address', '').strip()
    filter_date_from = request.GET.get('date_from', '').strip()
    filter_date_to   = request.GET.get('date_to', '').strip()

    if filter_customer:
        try:
            # If user searches for '0166', also search for '166'
            filter_customer_int = str(int(filter_customer))
            qs = qs.filter(
                Q(customerId__icontains=filter_customer) | 
                Q(customerId__icontains=filter_customer_int)
            )
        except ValueError:
            qs = qs.filter(customerId__icontains=filter_customer)

    if filter_mac:
        qs = qs.filter(mac_address__icontains=filter_mac)
    if filter_date_from:
        try:
            from django.utils.dateparse import parse_date
            d = parse_date(filter_date_from)
            if d:
                qs = qs.filter(received_dateTime__date__gte=d)
        except Exception:
            pass
    if filter_date_to:
        try:
            from django.utils.dateparse import parse_date
            d = parse_date(filter_date_to)
            if d:
                qs = qs.filter(received_dateTime__date__lte=d)
        except Exception:
            pass

    # --- Paginate ---
    paginator = Paginator(qs, 25)
    page_num  = request.GET.get('page', 1)
    try:
        reports = paginator.page(page_num)
    except PageNotAnInteger:
        reports = paginator.page(1)
    except EmptyPage:
        reports = paginator.page(paginator.num_pages)

    context = {
        'reports':        reports,
        'total_count':    paginator.count,
        'filter_customer': filter_customer,
        'filter_mac':      filter_mac,
        'filter_date_from': filter_date_from,
        'filter_date_to':   filter_date_to,
    }
    return render(request, 'configdetails/token_report_list.html', context)


# ============================================================================
# Group Device Edit APIs
# ============================================================================

@api_view(['GET'])
@login_required
def get_group_devices_api(request, group_id):
    """
    Returns:
      - current devices in the group (by type)
      - available (unallocated) devices of each type for the same company/branch
        (dispensers/keypads/LEDs must not already belong to another group;
         brokers/TVs can be in multiple groups so we show all company devices)
    """
    group = get_object_or_404(GroupMapping, id=group_id)

    # --- Current devices ---
    current_dispensers = list(
        GroupDispenserMapping.objects.filter(group=group)
        .select_related('dispenser')
        .order_by('dispenser_button_index')
    )
    current_keypads  = list(group.keypads.all().order_by('id'))
    current_leds     = list(group.leds.all().order_by('id'))
    current_brokers  = list(group.brokers.all().order_by('id'))
    current_tvs      = list(group.tvs.all().order_by('id'))

    def device_dict(d):
        return {
            'id': d.id,
            'serial_number': d.serial_number,
            'display_name': d.display_name or d.serial_number,
            'device_type': d.device_type,
            'token_type': getattr(d, 'token_type', None),
        }

    current_dispenser_ids = {gdm.dispenser.id for gdm in current_dispensers}
    current_keypad_ids    = {d.id for d in current_keypads}
    current_led_ids       = {d.id for d in current_leds}
    current_broker_ids    = {d.id for d in current_brokers}
    current_tv_ids        = {d.id for d in current_tvs}

    # Base queryset scoped to the group's company (and optionally branch)
    company = group.company
    branch  = group.branch
    base_qs = Device.objects.filter(company=company, is_active=True)
    if branch:
        base_qs = base_qs.filter(branch=branch)

    # Dispensers/Keypads/LEDs already claimed by ANY group
    taken_dispenser_ids = set(
        GroupDispenserMapping.objects.exclude(group=group)
        .values_list('dispenser_id', flat=True)
    )
    taken_keypad_ids = set(
        GroupMapping.objects.exclude(id=group_id)
        .values_list('keypads__id', flat=True)
    ) - {None}
    taken_led_ids = set(
        GroupMapping.objects.exclude(id=group_id)
        .values_list('leds__id', flat=True)
    ) - {None}

    # Available = correct type + not taken by another group + not already in current group
    avail_dispensers = list(
        base_qs.filter(device_type=Device.DeviceType.TOKEN_DISPENSER)
        .exclude(id__in=taken_dispenser_ids)
        .exclude(id__in=current_dispenser_ids)
        .order_by('display_name', 'serial_number')
    )
    avail_keypads = list(
        base_qs.filter(device_type=Device.DeviceType.KEYPAD)
        .exclude(id__in=taken_keypad_ids)
        .exclude(id__in=current_keypad_ids)
        .order_by('display_name', 'serial_number')
    )
    avail_leds = list(
        base_qs.filter(device_type=Device.DeviceType.LED)
        .exclude(id__in=taken_led_ids)
        .exclude(id__in=current_led_ids)
        .order_by('display_name', 'serial_number')
    )
    # Brokers/TVs — can be in multiple groups; show all not already in this group
    avail_brokers = list(
        base_qs.filter(device_type=Device.DeviceType.BROKER)
        .exclude(id__in=current_broker_ids)
        .order_by('display_name', 'serial_number')
    )
    avail_tvs = list(
        base_qs.filter(device_type=Device.DeviceType.TV)
        .exclude(id__in=current_tv_ids)
        .order_by('display_name', 'serial_number')
    )

    return Response({
        'group_id': group.id,
        'group_name': group.group_name,
        'current': {
            'dispensers': [
                {**device_dict(gdm.dispenser), 'dispenser_button_index': gdm.dispenser_button_index}
                for gdm in current_dispensers
            ],
            'keypads':  [device_dict(d) for d in current_keypads],
            'leds':     [device_dict(d) for d in current_leds],
            'brokers':  [device_dict(d) for d in current_brokers],
            'tvs':      [device_dict(d) for d in current_tvs],
        },
        'available': {
            'dispensers': [device_dict(d) for d in avail_dispensers],
            'keypads':    [device_dict(d) for d in avail_keypads],
            'leds':       [device_dict(d) for d in avail_leds],
            'brokers':    [device_dict(d) for d in avail_brokers],
            'tvs':        [device_dict(d) for d in avail_tvs],
        },
    })


@api_view(['POST'])
@login_required
def update_group_devices_api(request, group_id):
    """
    Accepts the DESIRED final device lists for the group:
      {
        "dispenser_ids": [...],
        "keypad_ids":    [...],
        "led_ids":       [...],
        "broker_ids":    [...],
        "tv_ids":        [...],
      }

    Validates exclusivity (dispenser/keypad/LED must not be in another group),
    then atomically:
      1. Removes devices that are no longer in the group
      2. Adds new devices
      3. Updates no_of_* counts
      4. Rebuilds all slot-index rows via _create_tv_slot_mappings()
    """
    from django.db import transaction

    group = get_object_or_404(GroupMapping, id=group_id)
    user  = request.user

    # Permission check
    if user.role == 'BRANCH_ADMIN' and group.branch != user.branch_relation:
        return Response({'status': 'error', 'message': 'Permission denied.'}, status=403)
    if user.role == 'COMPANY_ADMIN' and group.company != user.company_relation:
        return Response({'status': 'error', 'message': 'Permission denied.'}, status=403)

    new_dispenser_ids = [int(x) for x in request.data.get('dispenser_ids', []) if x]
    new_keypad_ids    = [int(x) for x in request.data.get('keypad_ids', []) if x]
    new_led_ids       = [int(x) for x in request.data.get('led_ids', []) if x]
    new_broker_ids    = [int(x) for x in request.data.get('broker_ids', []) if x]
    new_tv_ids        = [int(x) for x in request.data.get('tv_ids', []) if x]

    # ── Validate exclusivity for dispensers ─────────────────────────────────
    errors = []
    for d_id in new_dispenser_ids:
        conflict = GroupDispenserMapping.objects.filter(
            dispenser_id=d_id
        ).exclude(group=group).first()
        if conflict:
            d = Device.objects.filter(id=d_id).first()
            name = (d.display_name or d.serial_number) if d else str(d_id)
            errors.append(f"Dispenser '{name}' is already in group '{conflict.group.group_name}'.")

    for k_id in new_keypad_ids:
        conflict_qs = GroupMapping.objects.filter(keypads__id=k_id).exclude(id=group_id)
        if conflict_qs.exists():
            d = Device.objects.filter(id=k_id).first()
            name = (d.display_name or d.serial_number) if d else str(k_id)
            errors.append(f"Keypad '{name}' is already in group '{conflict_qs.first().group_name}'.")

    for l_id in new_led_ids:
        conflict_qs = GroupMapping.objects.filter(leds__id=l_id).exclude(id=group_id)
        if conflict_qs.exists():
            d = Device.objects.filter(id=l_id).first()
            name = (d.display_name or d.serial_number) if d else str(l_id)
            errors.append(f"LED '{name}' is already in group '{conflict_qs.first().group_name}'.")

    if errors:
        return Response({'status': 'error', 'message': ' '.join(errors)}, status=400)

    # ── Fetch device objects ─────────────────────────────────────────────────
    new_dispensers = list(Device.objects.filter(id__in=new_dispenser_ids, device_type=Device.DeviceType.TOKEN_DISPENSER))
    new_keypads    = list(Device.objects.filter(id__in=new_keypad_ids,    device_type=Device.DeviceType.KEYPAD))
    new_leds       = list(Device.objects.filter(id__in=new_led_ids,       device_type=Device.DeviceType.LED))
    new_brokers    = list(Device.objects.filter(id__in=new_broker_ids,    device_type=Device.DeviceType.BROKER))
    new_tvs        = list(Device.objects.filter(id__in=new_tv_ids,        device_type=Device.DeviceType.TV))

    if len(new_dispensers) != len(new_dispenser_ids):
        return Response({'status': 'error', 'message': 'One or more dispenser IDs are invalid.'}, status=400)
    if len(new_keypads) != len(new_keypad_ids):
        return Response({'status': 'error', 'message': 'One or more keypad IDs are invalid.'}, status=400)

    with transaction.atomic():
        # ── Determine removed devices ────────────────────────────────────────
        old_dispenser_ids = set(GroupDispenserMapping.objects.filter(group=group).values_list('dispenser_id', flat=True))
        old_keypad_ids    = set(group.keypads.values_list('id', flat=True))
        old_led_ids       = set(group.leds.values_list('id', flat=True))
        old_broker_ids    = set(group.brokers.values_list('id', flat=True))
        old_tv_ids        = set(group.tvs.values_list('id', flat=True))

        removed_dispenser_ids = old_dispenser_ids - set(new_dispenser_ids)
        removed_keypad_ids    = old_keypad_ids    - set(new_keypad_ids)
        removed_led_ids       = old_led_ids       - set(new_led_ids)
        removed_broker_ids    = old_broker_ids    - set(new_broker_ids)
        removed_tv_ids        = old_tv_ids        - set(new_tv_ids)

        removed_dispensers = Device.objects.filter(id__in=removed_dispenser_ids)
        removed_keypads    = Device.objects.filter(id__in=removed_keypad_ids)
        removed_tvs        = Device.objects.filter(id__in=removed_tv_ids)

        # ── Clean up slot index rows for removed devices ─────────────────────
        # Dispensers
        if removed_dispenser_ids:
            GroupDispenserMapping.objects.filter(group=group, dispenser__in=removed_dispensers).delete()
            GroupCounterButtonMapping.objects.filter(group=group, dispenser__in=removed_dispensers).delete()
            TVDispenserMapping.objects.filter(dispenser__in=removed_dispensers).delete()
            # Null-out keypad→dispenser references where that dispenser is removed
            TVKeypadMapping.objects.filter(dispenser__in=removed_dispensers).update(dispenser=None)
            # ButtonMappings connecting removed dispensers to other group devices
            all_group_device_ids = (
                list(old_dispenser_ids - removed_dispenser_ids) +
                list(old_keypad_ids) + list(old_led_ids) + list(old_broker_ids) + list(old_tv_ids)
            )
            ButtonMapping.objects.filter(
                Q(source_device__in=removed_dispensers) | Q(target_device__in=removed_dispensers)
            ).delete()

        # Keypads
        if removed_keypad_ids:
            TVKeypadMapping.objects.filter(keypad__in=removed_keypads).delete()
            ButtonMapping.objects.filter(
                Q(source_device__in=removed_keypads) | Q(target_device__in=removed_keypads)
            ).delete()
            group.keypads.remove(*removed_keypads)

        # LEDs
        if removed_led_ids:
            removed_leds = Device.objects.filter(id__in=removed_led_ids)
            ButtonMapping.objects.filter(
                Q(source_device__in=removed_leds) | Q(target_device__in=removed_leds)
            ).delete()
            group.leds.remove(*removed_leds)

        # Brokers
        if removed_broker_ids:
            removed_brokers = Device.objects.filter(id__in=removed_broker_ids)
            ButtonMapping.objects.filter(
                Q(source_device__in=removed_brokers) | Q(target_device__in=removed_brokers)
            ).delete()
            group.brokers.remove(*removed_brokers)

        # TVs
        if removed_tv_ids:
            TVDispenserMapping.objects.filter(tv__in=removed_tvs).delete()
            TVKeypadMapping.objects.filter(tv__in=removed_tvs).delete()
            ButtonMapping.objects.filter(
                Q(source_device__in=removed_tvs) | Q(target_device__in=removed_tvs)
            ).delete()
            group.tvs.remove(*removed_tvs)

        # ── Apply the new desired lists to the group ─────────────────────────
        # Dispensers — managed through GroupDispenserMapping (through model).
        # Assign proper sequential ASCII chars directly to avoid violating the
        # unique_together constraint on (group, dispenser_button_index) which
        # would IntegrityError when 2+ dispensers share the same default '1'.
        GroupDispenserMapping.objects.filter(group=group).delete()
        for pos, d in enumerate(new_dispensers, start=1):
            try:
                btn_char = get_button_index_char(pos)
            except ValueError:
                btn_char = chr(0x31)
            GroupDispenserMapping.objects.create(
                group=group,
                dispenser=d,
                dispenser_button_index=btn_char,
            )

        # Keypads (M2M direct)
        group.keypads.set(new_keypads)
        # LEDs
        group.leds.set(new_leds)
        # Brokers
        group.brokers.set(new_brokers)
        # TVs
        group.tvs.set(new_tvs)

        # ── Update no_of_* counts ────────────────────────────────────────────
        group.no_of_dispensers = len(new_dispensers)
        group.no_of_keypads    = len(new_keypads)
        group.no_of_leds       = len(new_leds)
        group.no_of_brokers    = len(new_brokers)
        group.no_of_tvs        = len(new_tvs)
        group.save(update_fields=[
            'no_of_dispensers', 'no_of_keypads', 'no_of_leds',
            'no_of_brokers', 'no_of_tvs', 'updated_at',
        ])

        # ── Rebuild all slot-index rows from scratch ─────────────────────────
        # _create_tv_slot_mappings rebuilds GroupDispenserMapping, GroupCounterButtonMapping,
        # TVDispenserMapping, and TVKeypadMapping compactly starting from slot '1'.
        _create_tv_slot_mappings(group)

        log_activity(
            request.user, "Group Devices Updated",
            f"Group '{group.group_name}' (ID: {group_id}) devices updated: "
            f"{len(new_dispensers)} dispensers, {len(new_keypads)} keypads, "
            f"{len(new_leds)} LEDs, {len(new_brokers)} brokers, {len(new_tvs)} TVs."
        )

    # ── Return updated slot assignments so UI can display them ───────────────
    updated_gdm = (
        GroupDispenserMapping.objects
        .filter(group=group)
        .select_related('dispenser')
        .order_by('dispenser_button_index')
    )
    updated_tkm = (
        TVKeypadMapping.objects
        .filter(tv__in=new_tvs)
        .select_related('keypad', 'tv')
        .order_by('tv', 'keypad_index')
    )

    return Response({
        'status': 'success',
        'message': f"Group '{group.group_name}' devices updated successfully.",
        'dispenser_slots': [
            {
                'dispenser_id': gdm.dispenser.id,
                'display_name': gdm.dispenser.display_name or gdm.dispenser.serial_number,
                'dispenser_button_index': gdm.dispenser_button_index,
            }
            for gdm in updated_gdm
        ],
        'keypad_slots': [
            {
                'keypad_id': tkm.keypad.id,
                'display_name': tkm.keypad.display_name or tkm.keypad.serial_number,
                'tv_id': tkm.tv.id,
                'keypad_index': tkm.keypad_index,
            }
            for tkm in updated_tkm
        ],
    })


# ============================================================================
# Group Button-Wise Mapping APIs
# ============================================================================

@api_view(['GET'])
@login_required
def get_group_button_mappings_api(request, group_id):
    """
    GET /api/mapping/group/<group_id>/button-mappings/

    Returns current explicit slot assignments for a group so the
    "Button Mapping" tab can pre-populate itself.
    """
    group = get_object_or_404(GroupMapping, id=group_id)

    def _d(dev):
        return {
            'id': dev.id,
            'display_name': dev.display_name or dev.serial_number,
            'serial_number': dev.serial_number,
            'device_type': dev.device_type,
        }

    group_dispensers = list(group.dispensers.all().order_by('id'))
    group_keypads    = list(group.keypads.all().order_by('id'))
    group_tvs        = list(group.tvs.all().order_by('id'))

    tv_name_map = {tv.id: (tv.display_name or tv.serial_number) for tv in group_tvs}

    dispenser_slots = []
    for tdm in (TVDispenserMapping.objects
                .filter(dispenser__in=group_dispensers, tv__in=group_tvs)
                .select_related('dispenser', 'tv')
                .order_by('tv', 'button_index')):
        dispenser_slots.append({
            'dispenser_id':  tdm.dispenser.id,
            'display_name':  tdm.dispenser.display_name or tdm.dispenser.serial_number,
            'tv_id':         tdm.tv.id,
            'tv_name':       tv_name_map.get(tdm.tv.id, tdm.tv.serial_number),
            'button_index':  tdm.button_index,
        })

    keypad_slots = []
    for tkm in (TVKeypadMapping.objects
                .filter(keypad__in=group_keypads, tv__in=group_tvs)
                .select_related('keypad', 'tv', 'dispenser')
                .order_by('tv', 'keypad_index')):
        keypad_slots.append({
            'keypad_id':      tkm.keypad.id,
            'display_name':   tkm.keypad.display_name or tkm.keypad.serial_number,
            'tv_id':          tkm.tv.id,
            'tv_name':        tv_name_map.get(tkm.tv.id, tkm.tv.serial_number),
            'keypad_index':   tkm.keypad_index,
            'dispenser_id':   tkm.dispenser.id if tkm.dispenser else None,
            'dispenser_name': (tkm.dispenser.display_name or tkm.dispenser.serial_number) if tkm.dispenser else None,
        })

    return Response({
        'group_id':         group.id,
        'group_name':       group.group_name,
        'dispenser_slots':  dispenser_slots,
        'keypad_slots':     keypad_slots,
        'group_dispensers': [_d(d) for d in group_dispensers],
        'group_keypads':    [_d(d) for d in group_keypads],
        'group_tvs':        [_d(d) for d in group_tvs],
    })


@api_view(['POST'])
@login_required
def save_group_button_mappings_api(request, group_id):
    """
    POST /api/mapping/group/<group_id>/button-mappings/save/

    Accepts explicit user-defined slot assignments and atomically rewrites
    TVDispenserMapping, TVKeypadMapping, GroupDispenserMapping, and
    GroupCounterButtonMapping.

    Request body:
    {
      "dispenser_slots": [
        { "dispenser_id": 5, "tv_id": 3, "button_index": "1" }
      ],
      "keypad_slots": [
        { "keypad_id": 7, "tv_id": 3, "keypad_index": "1", "dispenser_id": 5 }
      ]
    }
    button_index / keypad_index may be omitted → auto-assigned by position.
    """
    from collections import defaultdict

    group = get_object_or_404(GroupMapping, id=group_id)
    user  = request.user

    if user.role == 'BRANCH_ADMIN' and group.branch != user.branch_relation:
        return Response({'status': 'error', 'message': 'Permission denied.'}, status=403)
    if user.role == 'COMPANY_ADMIN' and group.company != user.company_relation:
        return Response({'status': 'error', 'message': 'Permission denied.'}, status=403)

    raw_disp_slots   = request.data.get('dispenser_slots', [])
    raw_keypad_slots = request.data.get('keypad_slots', [])

    group_dispenser_ids = set(group.dispensers.values_list('id', flat=True))
    group_keypad_ids    = set(group.keypads.values_list('id', flat=True))
    group_tv_ids        = set(group.tvs.values_list('id', flat=True))

    errors = []
    parsed_disp_slots = []
    seen_disp_tv_btn  = set()

    for i, row in enumerate(raw_disp_slots):
        try:
            d_id  = int(row.get('dispenser_id'))
            tv_id = int(row.get('tv_id'))
        except (TypeError, ValueError):
            errors.append(f"Dispenser slot {i+1}: invalid dispenser_id or tv_id.")
            continue
        btn = row.get('button_index')

        if d_id not in group_dispenser_ids:
            errors.append(f"Dispenser slot {i+1}: dispenser ID {d_id} is not in this group.")
        if tv_id not in group_tv_ids:
            errors.append(f"Dispenser slot {i+1}: TV ID {tv_id} is not in this group.")
        if btn is not None:
            btn = str(btn)
            if len(btn) != 1 or not (0x31 <= ord(btn) <= 0x7A):
                errors.append(f"Dispenser slot {i+1}: button_index '{btn}' is not valid ('1'-'z').")
            else:
                key = (tv_id, btn)
                if key in seen_disp_tv_btn:
                    errors.append(f"Dispenser slot {i+1}: duplicate button_index '{btn}' for TV {tv_id}.")
                seen_disp_tv_btn.add(key)

        parsed_disp_slots.append({'dispenser_id': d_id, 'tv_id': tv_id, 'button_index': btn})

    parsed_kp_slots = []
    seen_kp_tv_idx  = set()

    for i, row in enumerate(raw_keypad_slots):
        try:
            kp_id = int(row.get('keypad_id'))
            tv_id = int(row.get('tv_id'))
        except (TypeError, ValueError):
            errors.append(f"Keypad slot {i+1}: invalid keypad_id or tv_id.")
            continue
        kp_idx  = row.get('keypad_index')
        disp_id = row.get('dispenser_id')
        if disp_id is not None:
            try:
                disp_id = int(disp_id)
            except (TypeError, ValueError):
                disp_id = None

        if kp_id not in group_keypad_ids:
            errors.append(f"Keypad slot {i+1}: keypad ID {kp_id} is not in this group.")
        if tv_id not in group_tv_ids:
            errors.append(f"Keypad slot {i+1}: TV ID {tv_id} is not in this group.")
        if disp_id and disp_id not in group_dispenser_ids:
            errors.append(f"Keypad slot {i+1}: dispenser ID {disp_id} is not in this group.")
        if kp_idx is not None:
            kp_idx = str(kp_idx)
            if len(kp_idx) != 1 or not (0x31 <= ord(kp_idx) <= 0x7A):
                errors.append(f"Keypad slot {i+1}: keypad_index '{kp_idx}' is not valid.")
            else:
                key = (tv_id, kp_idx)
                if key in seen_kp_tv_idx:
                    errors.append(f"Keypad slot {i+1}: duplicate keypad_index '{kp_idx}' for TV {tv_id}.")
                seen_kp_tv_idx.add(key)

        parsed_kp_slots.append({
            'keypad_id': kp_id, 'tv_id': tv_id,
            'keypad_index': kp_idx, 'dispenser_id': disp_id,
        })

    if errors:
        return Response({'status': 'error', 'message': ' '.join(errors)}, status=400)

    # Auto-assign missing indices (sequential within each TV)
    def _auto_assign(slots, idx_key):
        by_tv = defaultdict(list)
        for row in slots:
            by_tv[row['tv_id']].append(row)
        result = []
        for tv_id, rows in by_tv.items():
            used = {row[idx_key] for row in rows if row[idx_key] is not None}
            pos = 1
            for row in rows:
                if row[idx_key] is None:
                    while True:
                        try:
                            ch = get_button_index_char(pos)
                        except ValueError:
                            ch = chr(0x31)
                            break
                        if ch not in used:
                            break
                        pos += 1
                    row = dict(row)
                    row[idx_key] = ch
                    used.add(ch)
                    pos += 1
                result.append(row)
        return result

    parsed_disp_slots = _auto_assign(parsed_disp_slots, 'button_index')
    parsed_kp_slots   = _auto_assign(parsed_kp_slots, 'keypad_index')

    tv_map     = {d.id: d for d in Device.objects.filter(id__in=group_tv_ids)}
    disp_map   = {d.id: d for d in Device.objects.filter(id__in=group_dispenser_ids)}
    keypad_map = {d.id: d for d in Device.objects.filter(id__in=group_keypad_ids)}

    with transaction.atomic():
        # Rewrite TVDispenserMapping
        TVDispenserMapping.objects.filter(dispenser__in=list(disp_map.values())).delete()
        for row in parsed_disp_slots:
            tv_obj   = tv_map.get(row['tv_id'])
            disp_obj = disp_map.get(row['dispenser_id'])
            if tv_obj and disp_obj:
                TVDispenserMapping.objects.create(
                    tv=tv_obj, dispenser=disp_obj, button_index=row['button_index'])

        # Rewrite TVKeypadMapping
        TVKeypadMapping.objects.filter(keypad__in=list(keypad_map.values())).delete()
        for row in parsed_kp_slots:
            tv_obj   = tv_map.get(row['tv_id'])
            kp_obj   = keypad_map.get(row['keypad_id'])
            disp_obj = disp_map.get(row['dispenser_id']) if row['dispenser_id'] else None
            if tv_obj and kp_obj:
                TVKeypadMapping.objects.create(
                    tv=tv_obj, keypad=kp_obj, dispenser=disp_obj,
                    keypad_index=row['keypad_index'])

        # Rebuild GroupDispenserMapping ordered by button_index
        GroupDispenserMapping.objects.filter(group=group).delete()
        all_disp_sorted = sorted(parsed_disp_slots, key=lambda r: r['button_index'])
        for pos, row in enumerate(all_disp_sorted, start=1):
            disp_obj = disp_map.get(row['dispenser_id'])
            if disp_obj:
                try:
                    grp_char = get_button_index_char(pos)
                except ValueError:
                    grp_char = chr(0x31)
                GroupDispenserMapping.objects.get_or_create(
                    group=group, dispenser=disp_obj,
                    defaults={'dispenser_button_index': grp_char})

        # Rebuild GroupCounterButtonMapping
        GroupCounterButtonMapping.objects.filter(group=group).delete()
        grp_btn_pos = 0
        for row in all_disp_sorted:
            disp_obj = disp_map.get(row['dispenser_id'])
            if not disp_obj:
                continue
            for cm in (CounterTokenDispenserMapping.objects
                       .filter(dispenser=disp_obj)
                       .select_related('counter')
                       .order_by('id')):
                grp_btn_pos += 1
                try:
                    btn_char = get_button_index_char(grp_btn_pos)
                except ValueError:
                    btn_char = chr(0x31)
                GroupCounterButtonMapping.objects.get_or_create(
                    group=group, dispenser=disp_obj, counter=cm.counter,
                    defaults={'button_index': btn_char})

        log_activity(
            request.user, "Group Button Mappings Saved",
            f"Group '{group.group_name}' (ID: {group_id}) — "
            f"{len(parsed_disp_slots)} dispenser slots, {len(parsed_kp_slots)} keypad slots configured."
        )

    return Response({
        'status': 'success',
        'message': f"Button mappings for group '{group.group_name}' saved successfully.",
        'dispenser_slots': parsed_disp_slots,
        'keypad_slots':    parsed_kp_slots,
    })


@api_view(['POST', 'GET'])
@permission_classes([AllowAny])
def get_android_mapped_counters(request):
    """
    API for fetching all mapped counters for a given customer_id for Android APK.
    Accepts: { "customer_id": "..." }
    """
    from companydetails.models import Company, DealerCustomer
    
    customer_id = request.data.get('customer_id') or request.GET.get('customer_id')
    
    if customer_id and str(customer_id).startswith('0'):
        customer_id = str(customer_id)[1:]
        
    if not customer_id:
        return Response({'error': 'customer_id is required'}, status=400)

    company = None
    if str(customer_id).isdigit():
        company = Company.objects.filter(id=customer_id).first()
    
    if not company:
        company = Company.objects.filter(company_id=customer_id).first()
        
    if not company:
        dealer_customer = DealerCustomer.objects.filter(customer_id=customer_id).first()
        if dealer_customer:
            company = dealer_customer.dealer
            
    if not company:
        return Response({'error': 'Invalid customer_id'}, status=404)
        
    counters = CounterConfig.objects.filter(company=company, status=True)
    
    mapped_counters = []
    for counter in counters:
        # Check mapping in GroupCounterButtonMapping
        mapping = GroupCounterButtonMapping.objects.filter(counter=counter).first()
        
        mapped_dispenser_sn = None
        button_index = None
        
        if mapping:
            mapped_dispenser_sn = mapping.dispenser.serial_number
            button_index = mapping.button_index
        else:
            # Fallback to old mapping
            old_mapping = CounterTokenDispenserMapping.objects.filter(counter=counter).first()
            if old_mapping:
                mapped_dispenser_sn = old_mapping.dispenser.serial_number
                button_index = '1'
        
        if mapped_dispenser_sn:
            mapped_counters.append({
                "id": counter.id,
                "counter_name": counter.counter_name,
                "counter_prefix_code": counter.counter_prefix_code,
                "counter_display_name": counter.counter_display_name,
                "max_token_number": counter.max_token_number,
                "mapped_dispenser_sn": mapped_dispenser_sn,
                "button_index": button_index
            })
            
    return Response({"counters": mapped_counters})
