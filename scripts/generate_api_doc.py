
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        if level == 1:
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
    
    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    def add_code_block(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)

    # --- Android Login API ---
    add_heading('Android Employee Login API Documentation', level=1)
    
    add_heading('Endpoint', level=2)
    add_paragraph('POST /CallQ/auth/api/android/login/', bold=True)
    add_paragraph('This API allows Android Employee devices to authenticate using their credentials, MAC address, and Customer ID.')

    add_heading('Description', level=2)
    add_paragraph('The Android Employee Login API is used to verify the employee\'s status and ensure the device is correctly mapped to a branch and customer before allowing login access.')

    add_heading('Request Body Example', level=2)
    add_code_block('{\n  "username": "employee_user",\n  "password": "secure_password",\n  "mac_address": "AA:BB:CC:DD:EE:FF",\n  "customer_id": "CUST123"\n}')

    add_heading('Request Fields', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    fields = [
        ('username', 'String', 'The employee\'s username or email.'),
        ('password', 'String', 'The employee\'s password.'),
        ('mac_address', 'String', 'The unique MAC address or serial number of the device.'),
        ('customer_id', 'String', 'The unique identifier of the customer company.')
    ]
    for field, ftype, desc in fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    add_heading('Success Response', level=2)
    add_code_block('{\n    "response": "Login Approved"\n}')
    add_paragraph('Note: If the device is awaiting admin approval, the response will be:')
    add_code_block('{\n    "response": "Waiting for Approval ! Contact Admin"\n}')

    add_heading('Error Responses', level=2)
    add_paragraph('Invalid or Missing Fields', bold=True)
    add_paragraph('400 BAD REQUEST')
    add_code_block('{\n  "error": "Invalid request data."\n}')
    
    add_paragraph('Unauthorized', bold=True)
    add_paragraph('401 UNAUTHORIZED')
    add_code_block('{\n  "detail": "Authentication credentials were not provided."\n}')

    doc.add_page_break()

    # --- Android Config API ---
    add_heading('Android TV Configuration API Documentation', level=1)
    
    add_heading('Endpoint', level=2)
    add_paragraph('POST /CallQ/config/api/android/config/', bold=True)
    add_paragraph('This API allows Android TV devices to fetch their specific configuration, advertisements, and mappings.')

    add_heading('Description', level=2)
    add_paragraph('The Android TV Configuration API provides comprehensive settings for TVs, including audio languages, advertisement files (URLs), layouts, and mappings to other devices like token dispensers and keypads.')

    add_heading('Request Body Example', level=2)
    add_code_block('{\n  "mac_address": "AA:BB:CC:DD:EE:FF",\n  "customer_id": "CUST123",\n  "flag": "EMBEDED"\n}')

    add_heading('Request Fields', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    fields = [
        ('mac_address', 'String', 'The unique MAC address or serial number of the device.'),
        ('customer_id', 'String', 'The unique identifier of the customer company.'),
        ('flag', 'String', '(Optional) Set to "EMBEDED" for full branch configuration.'),
        ('time', 'String', '(Optional) Request specific time (HH:MM) for shift profiles.'),
        ('date', 'String', '(Optional) Request specific date (YYYY-MM-DD) for shift profiles.')
    ]
    for field, ftype, desc in fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    add_heading('Success Response', level=2)
    add_code_block('{\n    "status": "success",\n    "message": "Configuration fetched successfully",\n    "device_id": 12,\n    "serial_number": "AA:BB:CC...",\n    "company_name": "Test Corp",\n    "config": { ... },\n    "tv_config": { \n        "audio_language": "en",\n        "show_ads": true,\n        "ad_files": ["http://.../media/ads/ad1.mp4"]\n    },\n    "mappings": [ ... ]\n}')

    add_heading('Error Responses', level=2)
    add_paragraph('Device Awaiting Approval', bold=True)
    add_paragraph('403 FORBIDDEN')
    add_code_block('{\n  "status": "pending",\n  "message": "Device awaiting approval"\n}')
    
    add_paragraph('Invalid Customer', bold=True)
    add_paragraph('404 NOT FOUND')
    add_code_block('{\n  "error": "Invalid customer_id"\n}')

    doc.save('/home/silpc-064/Desktop/CallQ/Android_API_Documentation.docx')

if __name__ == "__main__":
    create_doc()
