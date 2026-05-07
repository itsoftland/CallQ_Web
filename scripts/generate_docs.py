import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx():
    # Read the markdown file
    md_path = 'api_docs.md'
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r') as f:
        lines = f.readlines()

    doc = Document()
    
    # Title
    title = doc.add_heading('CallQ Android API Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_code_block = False
    
    for line in lines:
        line = line.strip('\n')
        
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('---'):
            doc.add_page_break()
            
        # Handle code blocks
        elif line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        elif in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 139) # Dark Blue
            p.paragraph_format.left_indent = Pt(20)
            
        # Handle lists
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
            
        # Handle plain text
        elif line.strip():
            # Basic bolding check
            if '**' in line:
                parts = line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 != 0:
                        run.bold = True
            else:
                doc.add_paragraph(line)

    output_path = 'CallQ_Android_API_Documentation.docx'
    doc.save(output_path)
    print(f"Documentation generated: {output_path}")

if __name__ == "__main__":
    generate_docx()
