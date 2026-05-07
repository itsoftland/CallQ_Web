"""
Generate Android Config Login API Documentation in DOCX format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)
    
    return table

def create_documentation():
    doc = Document()
    
    # Title
    title = doc.add_heading('Android Config App Login API Documentation', 0)
    
    # Overview
    add_heading(doc, 'Overview', 1)
    add_paragraph(doc, 'The Android Config App Login API allows Android devices running the config application to authenticate users and retrieve their accessible customers and devices based on their role hierarchy.')
    doc.add_paragraph()
    add_paragraph(doc, 'Base URL: /CallQ/auth/api/android/config-login', bold=True)
    add_paragraph(doc, 'Method: POST', bold=True)
    add_paragraph(doc, 'Content-Type: application/json', bold=True)
    
    # Authentication
    add_heading(doc, 'Authentication', 1)
    add_paragraph(doc, 'This API uses username/password authentication. Upon successful login, it returns the user\'s accessible resources based on their role.')
    
    # Request
    add_heading(doc, 'Request', 1)
    
    add_heading(doc, 'Headers', 2)
    add_table(doc, ['Header', 'Value', 'Required'], [
        ['Content-Type', 'application/json', 'Yes']
    ])
    doc.add_paragraph()
    
    add_heading(doc, 'Body Parameters', 2)
    add_table(doc, ['Parameter', 'Type', 'Required', 'Description'], [
        ['username', 'string', 'Yes', "User's username or email prefix"],
        ['password', 'string', 'Yes', "User's password"]
    ])
    doc.add_paragraph()
    
    add_heading(doc, 'Example Request', 2)
    add_code_block(doc, '''curl -X POST "https://your-domain.com/CallQ/auth/api/android/config-login" \\
  -H "Content-Type: application/json" \\
  -d '{
    "username": "admin",
    "password": "password123"
  }' ''')
    
    # Response Structure
    add_heading(doc, 'Response Structure', 1)
    
    add_heading(doc, 'Success Response Fields', 2)
    add_table(doc, ['Field', 'Type', 'Description'], [
        ['status', 'string', 'Always "success" for successful requests'],
        ['login_type', 'string', 'One of: "executive", "company", "dealer", "dealer_company"'],
        ['user_info', 'object', 'User details'],
        ['customers', 'array', 'List of customers (for executive/dealer login)'],
        ['customer', 'object', 'Single customer (for company/dealer_company login)'],
        ['devices', 'array', 'List of devices (for company/dealer_company login)'],
        ['dealer', 'object', 'Dealer info (for dealer login only)'],
        ['dealer_devices', 'array', 'Devices directly assigned to dealer (for dealer login only)']
    ])
    doc.add_paragraph()
    
    add_heading(doc, 'User Info Object', 2)
    add_table(doc, ['Field', 'Type', 'Description'], [
        ['id', 'integer', 'User ID'],
        ['username', 'string', 'Username'],
        ['email', 'string', 'User email'],
        ['role', 'string', 'User role (SUPER_ADMIN, ADMIN, COMPANY_ADMIN, etc.)']
    ])
    doc.add_paragraph()
    
    add_heading(doc, 'Device Object', 2)
    add_table(doc, ['Field', 'Type', 'Description'], [
        ['id', 'integer', 'Device ID'],
        ['serial_number', 'string', 'Device serial number'],
        ['mac_address', 'string', 'Device MAC address'],
        ['device_type', 'string', 'Type: TV, TOKEN_DISPENSER, KEYPAD, BROKER, LED'],
        ['licence_status', 'string', 'Status: Active, Inactive, Pending'],
        ['licence_active_to', 'string/null', 'License expiry date (YYYY-MM-DD)'],
        ['is_active', 'boolean', 'Whether device is enabled']
    ])
    
    # Login Types
    add_heading(doc, 'Login Types', 1)
    
    # Executive Login
    add_heading(doc, '1. Executive Login', 2)
    add_paragraph(doc, 'Roles: SUPER_ADMIN, ADMIN', bold=True)
    add_paragraph(doc, 'Description: Full access to all customers (companies and dealer customers) with their devices.')
    doc.add_paragraph()
    add_paragraph(doc, 'Response Example:', bold=True)
    add_code_block(doc, '''{
    "status": "success",
    "login_type": "executive",
    "user_info": {
        "id": 1,
        "username": "superadmin",
        "email": "admin@callq.com",
        "role": "SUPER_ADMIN"
    },
    "customers": [
        {
            "id": 1,
            "company_id": "COMP001",
            "name": "ABC Corporation",
            "type": "CUSTOMER",
            "devices": [
                {
                    "id": 10,
                    "serial_number": "TV-ABC-001",
                    "mac_address": "AA:BB:CC:DD:EE:01",
                    "device_type": "TV",
                    "licence_status": "Active",
                    "licence_active_to": "2026-12-31",
                    "is_active": true
                }
            ]
        }
    ]
}''')
    
    # Company Login
    add_heading(doc, '2. Company Login', 2)
    add_paragraph(doc, 'Roles: COMPANY_ADMIN', bold=True)
    add_paragraph(doc, "Description: Access to only the company's own devices.")
    doc.add_paragraph()
    add_paragraph(doc, 'Response Example:', bold=True)
    add_code_block(doc, '''{
    "status": "success",
    "login_type": "company",
    "user_info": {
        "id": 5,
        "username": "company_admin",
        "email": "admin@abccorp.com",
        "role": "COMPANY_ADMIN"
    },
    "customer": {
        "id": 1,
        "company_id": "COMP001",
        "name": "ABC Corporation",
        "type": "CUSTOMER"
    },
    "devices": [
        {
            "id": 10,
            "serial_number": "TV-ABC-001",
            "mac_address": "AA:BB:CC:DD:EE:01",
            "device_type": "TV",
            "licence_status": "Active",
            "licence_active_to": "2026-12-31",
            "is_active": true
        }
    ]
}''')
    
    # Dealer Login
    add_heading(doc, '3. Dealer Login', 2)
    add_paragraph(doc, 'Roles: DEALER_ADMIN', bold=True)
    add_paragraph(doc, 'Description: Access to all dealer customers and their devices, plus devices directly owned by the dealer.')
    doc.add_paragraph()
    add_paragraph(doc, 'Response Example:', bold=True)
    add_code_block(doc, '''{
    "status": "success",
    "login_type": "dealer",
    "user_info": {
        "id": 10,
        "username": "dealer_admin",
        "email": "admin@maindealer.com",
        "role": "DEALER_ADMIN"
    },
    "dealer": {
        "id": 2,
        "company_id": "DLR001",
        "name": "Main Dealer"
    },
    "customers": [
        {
            "id": 1,
            "customer_id": "DLR001-CUST0001",
            "name": "Customer Alpha",
            "dealer_id": 2,
            "dealer_name": "Main Dealer",
            "devices": [...]
        }
    ],
    "dealer_devices": [
        {
            "id": 30,
            "serial_number": "TV-DEALER-DEMO-001",
            "mac_address": "99:88:77:66:55:44",
            "device_type": "TV",
            "licence_status": "Active",
            "licence_active_to": "2027-01-01",
            "is_active": true
        }
    ]
}''')
    
    # Dealer Company Login
    add_heading(doc, '4. Dealer Company Login', 2)
    add_paragraph(doc, 'Roles: DEALER_CUSTOMER', bold=True)
    add_paragraph(doc, "Description: Access to only the specific dealer customer's devices.")
    doc.add_paragraph()
    add_paragraph(doc, 'Response Example:', bold=True)
    add_code_block(doc, '''{
    "status": "success",
    "login_type": "dealer_company",
    "user_info": {
        "id": 15,
        "username": "customer_alpha",
        "email": "admin@customeralpha.com",
        "role": "DEALER_CUSTOMER"
    },
    "customer": {
        "id": 1,
        "customer_id": "DLR001-CUST0001",
        "name": "Customer Alpha",
        "dealer_id": 2,
        "dealer_name": "Main Dealer"
    },
    "devices": [
        {
            "id": 20,
            "serial_number": "TV-ALPHA-001",
            "mac_address": "11:22:33:44:55:01",
            "device_type": "TV",
            "licence_status": "Active",
            "licence_active_to": "2026-06-30",
            "is_active": true
        }
    ]
}''')
    
    # Error Responses
    add_heading(doc, 'Error Responses', 1)
    add_table(doc, ['Status Code', 'Error', 'Description'], [
        ['400', 'Invalid request data', 'Missing username or password'],
        ['401', 'Invalid credentials', 'Wrong username or password'],
        ['401', 'Account disabled', 'User account is deactivated'],
        ['403', 'No company assigned', 'User not linked to company'],
        ['403', 'Unauthorized role', 'Role not allowed for config app'],
        ['404', 'Dealer customer not found', 'Could not find dealer customer record']
    ])
    doc.add_paragraph()
    
    add_heading(doc, 'Error Response Example', 2)
    add_code_block(doc, '''{
    "status": "error",
    "error": "Invalid credentials",
    "message": "The username or password provided is incorrect."
}''')
    
    # Role Hierarchy
    add_heading(doc, 'Role Hierarchy Reference', 1)
    add_table(doc, ['Role', 'Login Type', 'Access Level'], [
        ['SUPER_ADMIN', 'executive', 'All customers, all devices'],
        ['ADMIN', 'executive', 'All customers, all devices'],
        ['DEALER_ADMIN', 'dealer', 'Own dealer customers and their devices'],
        ['COMPANY_ADMIN', 'company', 'Own company devices only'],
        ['DEALER_CUSTOMER', 'dealer_company', 'Assigned dealer customer devices only'],
        ['BRANCH_ADMIN', 'Not supported', '-'],
        ['EMPLOYEE', 'Not supported', '-'],
        ['COMPANY_EMPLOYEE', 'Not supported', '-']
    ])
    
    # Notes
    add_heading(doc, 'Notes', 1)
    doc.add_paragraph('• Device Types: TV, TOKEN_DISPENSER, KEYPAD, BROKER, LED')
    doc.add_paragraph('• Licence Status Values: Active, Inactive, Pending')
    doc.add_paragraph('• Date Format: YYYY-MM-DD (ISO 8601)')
    doc.add_paragraph('• All API responses use UTF-8 encoding')
    
    # Version History
    add_heading(doc, 'Version History', 1)
    add_table(doc, ['Version', 'Date', 'Changes'], [
        ['1.0', '2026-01-30', 'Initial release with 4 login types']
    ])
    
    # Save document
    output_path = '/home/silpc-064/Desktop/CallQ/Android_Config_Login_API_Documentation.docx'
    doc.save(output_path)
    print(f'Documentation saved to: {output_path}')

if __name__ == '__main__':
    create_documentation()
