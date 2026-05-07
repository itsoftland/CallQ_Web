
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Android Login & Device API Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Helper functions
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    def add_code_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        
        # Split lines to handle basic formatting if needed
        lines = text.split('\n')
        for i, line in enumerate(lines):
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            if i < len(lines) - 1:
                run.add_break()

    def add_table(fields, headers=('Field', 'Type', 'Description')):
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for field, ftype, desc in fields:
            row = table.add_row().cells
            row[0].text = str(field)
            row[1].text = str(ftype)
            row[2].text = str(desc)
        
        doc.add_paragraph() # Spacing after table

    # ==========================================
    # API 1: Android Config Login (Step 1)
    # ==========================================
    add_heading('1. Android Config Login (Step 1)', level=1)
    add_paragraph('This is the modification implementation of the first step of the two-step login process for Android devices. It supports multiple logic types (Executive, Company, Dealer, Dealer Customer).')

    add_heading('Endpoint', level=2)
    add_paragraph('POST /CallQ/auth/api/android/Androidlogin', bold=True)

    add_heading('Request Body', level=2)
    add_table([
        ('username', 'String', 'Required. User username.'),
        ('password', 'String', 'Required. User password.'),
        ('mac_address', 'String', 'Required. Device MAC address or Serial Number.')
    ])

    add_heading('Request Example', level=2)
    add_code_block('{\n    "username": "admin",\n    "password": "password123",\n    "mac_address": "AA:BB:CC:DD:EE:FF"\n}')

    add_heading('Success Response', level=2)
    add_paragraph('The response structure varies slightly based on the user role (Executive, Company, Dealer, Dealer Customer), but generally follows this format:')
    add_code_block('{\n    "status": "success",\n    "login_type": "executive",\n    "Route": { ... },\n    "user_info": {\n        "id": 1,\n        "username": "admin",\n        "email": "admin@example.com",\n        "role": "SUPER_ADMIN",\n        "display_name": "Admin User"\n    },\n    "customers": [\n        {\n            "id": 101,\n            "company_name": "Example Company",\n            "devices": [ ... ]\n        }\n    ]\n}')
    
    add_heading('Key Response Fields', level=2)
    add_table([
        ('status', 'String', '"success" on successful login.'),
        ('login_type', 'String', 'Identifies the role context: "executive", "company", "dealer", or "dealer_company".'),
        ('Route', 'Object', 'Location tree (States -> Districts -> Cities) accessible to the user.'),
        ('user_info', 'Object', 'Details about the authenticated user.'),
        ('customers', 'Array', 'List of companies/customers the user can access. Includes "devices" list for Executive view.')
    ])

    add_heading('Error Responses', level=2)
    add_table([
        ('400 Bad Request', 'JSON', 'Missing required fields (username, password, mac_address).'),
        ('401 Unauthorized', 'JSON', 'Invalid credentials or account disabled.'),
        ('403 Forbidden', 'JSON', 'Android access denied, Device mapped to another customer, or User mapped to incorrect customer.'),
        ('Pending Approval', 'JSON', 'If status is "pending", the device is waiting for admin approval.')
    ], headers=('Status', 'Type', 'Description'))

    doc.add_page_break()

    # ==========================================
    # API 2: Get Device by Customer (Step 2)
    # ==========================================
    add_heading('2. Get Device by Customer (Step 2)', level=1)
    add_paragraph('This API is the second step of the login process, used to fetch specific device configurations for a selected customer.')

    add_heading('Endpoint', level=2)
    add_paragraph('POST /CallQ/auth/api/android/getDeviceByCustomer', bold=True)

    add_heading('Request Body', level=2)
    add_table([
        ('customer_id', 'String/Integer', 'Required. The ID of the Company or DealerCustomer.'),
        ('name', 'String', 'Optional. Name of the customer (echoed back in response).')
    ])

    add_heading('Request Example', level=2)
    add_code_block('{\n    "customer_id": "101",\n    "name": "Target Company"\n}')

    add_heading('Success Response', level=2)
    add_paragraph('Returns lists of devices separated by type (TVs vs others), with full configuration details.')
    add_code_block('{\n    "status": "success",\n    "customer_id": "101",\n    "customer_name": "Target Company",\n    "devices": [\n        {\n            "id": 50,\n            "serial_number": "DEV001",\n            "device_type": "KEYPAD",\n            "mappings_as_source": [ ... ],\n            "mappings_as_target": [ ... ]\n        }\n    ],\n    "tvs": [\n        {\n            "id": 60,\n            "serial_number": "TV001",\n            "device_type": "TV",\n            "tv_config": {\n                "orientation": "landscape",\n                "layout_type": "grid",\n                "token_font_size": 24,\n                "ads": [ "http://.../ad1.jpg" ],\n                ...\n            },\n            "tv_counters": [ ... ],\n            "tv_mappings": [ ... ],\n            "button_mappings": [ ... ]\n        }\n    ]\n}')

    add_heading('Key Response Fields', level=2)
    add_table([
        ('status', 'String', '"success" on valid request.'),
        ('devices', 'Array', 'List of non-TV devices (Keypads, Dispensers, etc.) including button mappings.'),
        ('tvs', 'Array', 'List of TV devices with comprehensive configuration (tv_config, tv_counters, ads) and mappings.')
    ])

    add_heading('Error Responses', level=2)
    add_table([
        ('400 Bad Request', 'JSON', 'Missing customer_id.'),
        ('404 Not Found', 'JSON', 'Customer ID not found for Company or DealerCustomer.')
    ], headers=('Status', 'Type', 'Description'))

    # Save Document
    filename = 'Android_Login_and_Device_API_Report.docx'
    output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'CallQ', filename)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_doc()
