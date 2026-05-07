from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    document = Document()

    # Title
    title = document.add_heading('CallQ Embedded Device Configuration API Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    document.add_heading('Overview', level=1)
    p = document.add_paragraph('This document details the API endpoint for retrieving configuration settings for embedded devices such as Keypads and Calling Units. The API returns a custom formatted string containing device-specific settings and mapping information.')

    # Endpoint Details
    document.add_heading('Endpoint Details', level=1)
    p = document.add_paragraph()
    p.add_run('URL: ').bold = True
    p.add_run('/CallQ/config/api/embedded/get-config/')
    
    p = document.add_paragraph()
    p.add_run('Method: ').bold = True
    p.add_run('GET or POST')

    p = document.add_paragraph()
    p.add_run('Content-Type: ').bold = True
    p.add_run('text/plain')

    # Request Parameters
    document.add_heading('Request Parameters', level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Required'
    hdr_cells[3].text = 'Description'

    row_cells = table.add_row().cells
    row_cells[0].text = 'mac_address'
    row_cells[1].text = 'String'
    row_cells[2].text = 'Yes'
    row_cells[3].text = 'The unique serial number or MAC address of the device.'

    row_cells = table.add_row().cells
    row_cells[0].text = 'device_type'
    row_cells[1].text = 'String'
    row_cells[2].text = 'No'
    row_cells[3].text = 'The type of device (e.g., KEYPAD).'

    # Response Format
    document.add_heading('Response Format', level=1)
    document.add_paragraph('The API returns a raw string with comma-separated values, terminated by XX#.')
    
    document.add_heading('Format Structure (Keypad/Standard):', level=2)
    p = document.add_paragraph('$1,SETTINGS,PASSWORD,TEXT,LOGO,COUNTER,MODE,SKIP,TRANSFER,VIP,VIP_FROM,VIP_TO,POOL,DISPENSER_SN,KEYPAD_COUNT,KP1,KP2,KP3,KP4,KP5,XX#')
    p.style = 'Quote'

    document.add_heading('Format Structure (Token Dispenser):', level=2)
    p = document.add_paragraph('$0,RANDOM_NUM,SERIAL_NUMBER,HEADER1,HEADER2,HEADER3,FOOTER1,FOOTER2,DAY_RESET,RESET_TKN,CUTTER,HALF_CUT,FEED,LOGO,MODE,LABEL,PAPER_OUT,TYPE,DUP_TKN,#')
    p.style = 'Quote'

    document.add_heading('Example Response (Keypad):', level=2)
    p = document.add_paragraph('$1,SETTINGS,XXXX,Welcome,1,1,1,1,1,0,0,0,0,DISP_SN,0,00000...0000,XX#')
    p.style = 'Quote'
    
    document.add_heading('Example Response (Dispenser):', level=2)
    p = document.add_paragraph('$0,4829,DISP_001,CallQ,Token,System,Thank You,Visit Again,1,1,1,0,1,1,1,Token,1,0,0,#')
    p.style = 'Quote'

    # Field Breakdown
    document.add_heading('Field Breakdown (Keypad)', level=2)
    
    fields = [
        ("Header", "$1,SETTINGS - Constant identifier."),
        ("PASSWORD", "XXXX - Default password (hardcoded security measure)."),
        ("TEXT", "Welcome - On-screen display text (configurable)."),
        ("LOGO", "1 or 0 - Enable/Disable logo display."),
        ("COUNTER", "Integer - The counter number assigned to this device."),
        ("MODE", "1 (Multiple) or 0 (Single) - Token calling mode."),
        ("SKIP", "1 or 0 - Enable/Disable Skip functionality."),
        ("TRANSFER", "1 or 0 - Enable/Disable Transfer functionality."),
        ("VIP", "1 or 0 - Enable/Disable VIP token calling."),
        ("VIP_FROM", "Integer - Start of VIP token range."),
        ("VIP_TO", "Integer - End of VIP token range."),
        ("POOL", "1 or 0 - Keypad Pool Mode status."),
        ("DISPENSER_SN", "String - Serial Number of the mapped Token Dispenser."),
        ("KEYPAD_COUNT", "Integer - Total number of keypads in the group."),
        ("KP1 to KP5", "String - Serial Numbers of up to 5 linked keypads. Unused slots are filled with zeros."),
        ("Trailer", "XX# - End of transmission marker.")
    ]

    for field, desc in fields:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(field + ": ").bold = True
        p.add_run(desc)

    document.add_heading('Field Breakdown (Dispenser)', level=2)
    disp_fields = [
        ("Header", "$0 - Constant identifier."),
        ("RANDOM_NUM", "4 Digit Random Number."),
        ("SERIAL_NUMBER", "Device Serial Number."),
        ("HEADER1-3", "Header Text Lines."),
        ("FOOTER1", "Footer Line 1."),
        ("FOOTER2", "Footer Line 2."),
        ("DAY_RESET", "1/0 - Auto Reset Daily."),
        ("RESET_TKN", "1/0 - Reset Token Count."),
        ("CUTTER", "1/0 - Enable Cutter."),
        ("HALF_CUT", "1/0 - Half Cut Mode."),
        ("FEED", "1/0 - Feed Paper."),
        ("LOGO", "1/0 - Enable Logo."),
        ("MODE", "1/0 - Single/Multiple Mode."),
        ("LABEL", "Token Label Text."),
        ("PAPER_OUT", "1/0 - Alert on Paper Out."),
        ("TYPE", "Device/Token Type."),
        ("DUP_TKN", "1/0 - Allow Duplicate Token.")
    ]

    for field, desc in disp_fields:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(field + ": ").bold = True
        p.add_run(desc)


    # Error Codes
    document.add_heading('Error Codes', level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'HTTP Status'
    hdr_cells[1].text = 'Response Body'
    hdr_cells[2].text = 'Description'

    row_cells = table.add_row().cells
    row_cells[0].text = '400 Bad Request'
    row_cells[1].text = 'ERROR: Missing mac_address'
    row_cells[2].text = 'The mac_address parameter was not provided.'

    row_cells = table.add_row().cells
    row_cells[0].text = '404 Not Found'
    row_cells[1].text = 'ERROR: Device not found'
    row_cells[2].text = 'No device exists with the provided Serial Number.'

    row_cells = table.add_row().cells
    row_cells[0].text = '403 Forbidden'
    row_cells[1].text = 'ERROR: Device inactive'
    row_cells[2].text = 'The device is registered but currently inactive or expired.'

    # Integration Notes
    document.add_heading('Integration Notes', level=1)
    document.add_paragraph('The PASSWORD field is consistently returned as XXXX and should not be used for authentication logic on the device side if security is a concern.', style='List Bullet')
    document.add_paragraph('Device configuration fields (Text, Logo, etc.) are managed via the CallQ Admin Portal.', style='List Bullet')
    document.add_paragraph('Ensure the device sends the mac_address exactly as registered in the portal (case-sensitive matching may apply depending on database collation, but uppercase is recommended).', style='List Bullet')

    document.save('/home/silpc-064/.gemini/antigravity/brain/1e36f249-d75e-4dba-abfe-45f5489a1b2c/CallQ_Embedded_API_Doc.docx')

if __name__ == "__main__":
    create_doc()
