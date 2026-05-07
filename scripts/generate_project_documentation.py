from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DATE = "February 10, 2026"
PROJECT_ROOT = "/home/silpc-064/Desktop/CallQ/CallQ"
OUTPUT_PATH = f"{PROJECT_ROOT}/CallQ_Project_Documentation.docx"


def add_title(doc):
    title = doc.add_heading("CallQ Project Documentation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Comprehensive system, workflow, and role documentation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p = doc.add_paragraph(f"Generated: {REPORT_DATE}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    root_p = doc.add_paragraph(f"Project root: {PROJECT_ROOT}")
    root_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    items = [
        "1. Overview",
        "2. Tech Stack",
        "3. Architecture and Project Structure",
        "4. Configuration and Environment",
        "5. Data Model",
        "6. User Roles and Responsibilities",
        "7. Authentication and Security",
        "8. External Integrations",
        "9. API Surface",
        "10. Page Wise Workflow",
        "11. Reporting and Exports",
        "12. Logging and Observability",
        "13. Operations and Maintenance",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def add_section(doc, title, level=1):
    doc.add_heading(title, level=level)


def add_kv(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_overview(doc):
    add_section(doc, "1. Overview", level=1)
    doc.add_paragraph(
        "CallQ is a Django-based web and API platform for queue and token management across customers, dealers, and branches. "
        "It supports device registration, configuration, and mapping for multiple device types (TV, token dispenser, keypad, broker, LED) "
        "and integrates with an external license portal for customer and device authentication. The system provides a web dashboard for "
        "administrators and company teams, and Android APIs for device and employee login workflows."
    )


def add_tech_stack(doc):
    add_section(doc, "2. Tech Stack", level=1)
    add_bullets(doc, [
        "Backend: Python, Django 5.2.9, Django REST Framework",
        "Frontend: Django templates, static JavaScript and CSS",
        "Database: Configured via environment variables (DB_ENGINE, DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT)",
        "Auth: Custom user model (userdetails.User) with email or username login backend",
        "Reporting: ReportLab for PDF exports, python-docx for Word exports",
        "Static and media: WhiteNoise for static files, media uploads served from MEDIA_ROOT",
        "Logging: File-based logs under CallQ_logs plus console output",
    ])


def add_architecture(doc):
    add_section(doc, "3. Architecture and Project Structure", level=1)
    doc.add_paragraph("Key directories in the project root:")
    add_bullets(doc, [
        "CallQ: Django project configuration (settings, urls, wsgi, middleware)",
        "callq_core: Core permissions, middleware, and external service integration",
        "companydetails: Company, branch, dealer, location, and activity logging",
        "configdetails: Device inventory, configuration, mapping, and production batches",
        "licenses: License batches, requests, validation, and purchase flows",
        "userdetails: Custom user model, role management, profiles, and Android login APIs",
        "templates: Shared templates and authentication templates",
        "static: Shared static assets (CSS, JS, images)",
        "media: Uploaded files (ads, audio, and other device assets)",
        "CallQ_logs: Log files grouped by date",
        "scripts: Utility scripts for documentation and admin tasks",
        "requirements.txt: Python dependencies",
    ])
    doc.add_paragraph("Core Django apps and responsibilities:")
    add_bullets(doc, [
        "companydetails: Customers, dealers, branches, locations, activity logs, and dashboard",
        "configdetails: Device lifecycle, device configuration, mapping, embedded profiles, approvals, and production batch tools",
        "licenses: License batches, requests, and validation",
        "userdetails: Custom user model, role-based access, profiles, and mobile authentication",
        "callq_core: Shared permissions, middleware, and external license portal integration",
    ])


def add_configuration(doc):
    add_section(doc, "4. Configuration and Environment", level=1)
    doc.add_paragraph("Key settings and environment variables:")
    add_bullets(doc, [
        "SECRET_KEY: Django secret key",
        "DEBUG: Debug mode (True/False)",
        "ALLOWED_HOSTS: Comma-separated hostnames",
        "CSRF_TRUSTED_ORIGINS: Comma-separated origins",
        "FORCE_SCRIPT_NAME: Optional subfolder hosting path",
        "APP_VERSION, PROJECT_NAME, PROJECT_DISPLAY_NAME: Branding and API metadata",
        "DB_ENGINE, DB_NAME, DB_USER, DB_PASSWORD, DB_HOST, DB_PORT: Database configuration",
        "LICENSE_PORTAL_URL: External license portal base URL",
        "EMAIL_HOST_USER, EMAIL_HOST_PASSWORD: SMTP credentials",
    ])
    doc.add_paragraph("Static and media paths:")
    add_bullets(doc, [
        "STATIC_URL: /CallQ/static/",
        "STATIC_ROOT: <BASE_DIR>/staticfiles",
        "STATICFILES_DIRS: <BASE_DIR>/static",
        "MEDIA_URL: /CallQ/media/",
        "MEDIA_ROOT: <BASE_DIR>/media",
    ])
    doc.add_paragraph("Time zone and localization:")
    add_bullets(doc, [
        "TIME_ZONE: Asia/Kolkata",
        "LANGUAGE_CODE: en-us",
        "USE_TZ: True",
    ])


def add_data_model(doc):
    add_section(doc, "5. Data Model", level=1)
    doc.add_paragraph("The following models represent the core data structures. Fields are listed as name and type.")

    add_section(doc, "userdetails", level=2)
    add_section(doc, "User", level=3)
    add_bullets(doc, [
        "Inherits AbstractUser fields (password, last_login, is_staff, etc.)",
        "username: CharField",
        "display_name: CharField",
        "email: EmailField (unique)",
        "role: CharField (choices: SUPER_ADMIN, ADMIN, DEALER_ADMIN, COMPANY_ADMIN, BRANCH_ADMIN, DEALER_CUSTOMER, PRODUCTION_ADMIN, EMPLOYEE, COMPANY_EMPLOYEE)",
        "zone: CharField",
        "assigned_state: JSONField",
        "is_web_user: BooleanField",
        "is_android_user: BooleanField",
        "company_relation: ForeignKey to Company",
        "branch_relation: ForeignKey to Branch",
    ])
    add_section(doc, "AppLoginHistory", level=3)
    add_bullets(doc, [
        "user: ForeignKey to User",
        "company: ForeignKey to Company",
        "mac_address: CharField",
        "version: CharField",
        "timestamp: DateTimeField (auto_now_add)",
    ])

    add_section(doc, "companydetails", level=2)
    add_section(doc, "Country", level=3)
    add_bullets(doc, ["name: CharField", "code: CharField"])
    add_section(doc, "State", level=3)
    add_bullets(doc, ["country: ForeignKey to Country", "name: CharField"])
    add_section(doc, "District", level=3)
    add_bullets(doc, ["state: ForeignKey to State", "name: CharField"])
    add_section(doc, "Company", level=3)
    add_bullets(doc, [
        "company_id: CharField",
        "company_name: CharField",
        "company_type: CharField (CUSTOMER or DEALER)",
        "parent_company: ForeignKey to Company (self)",
        "is_dealer_created: BooleanField",
        "branch_configuration: CharField (SINGLE or MULTIPLE)",
        "company_email: EmailField",
        "gst_number: CharField",
        "contact_person: CharField",
        "contact_number: CharField",
        "address: TextField",
        "address_2: TextField",
        "city: CharField",
        "district: CharField",
        "state: CharField",
        "country: CharField",
        "zip_code: CharField",
        "number_of_licence: IntegerField",
        "noof_broker_devices: IntegerField",
        "noof_token_dispensors: IntegerField",
        "noof_keypad_devices: IntegerField",
        "noof_television_devices: IntegerField",
        "noof_led_devices: IntegerField",
        "authentication_status: CharField",
        "product_registration_id: IntegerField",
        "unique_identifier: CharField",
        "product_from_date: DateField",
        "product_to_date: DateField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "Branch", level=3)
    add_bullets(doc, [
        "company: ForeignKey to Company",
        "branch_name: CharField",
        "address: TextField",
        "city: CharField",
        "state: CharField",
        "country: CharField",
        "zip_code: CharField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "AuthenticationLog", level=3)
    add_bullets(doc, [
        "company: ForeignKey to Company",
        "authentication_status: CharField",
        "product_registration_id: IntegerField",
        "from_date: DateField",
        "to_date: DateField",
        "number_of_licence: IntegerField",
        "response_json: JSONField",
        "created_at: DateTimeField",
    ])
    add_section(doc, "ActivityLog", level=3)
    add_bullets(doc, [
        "user: ForeignKey to User",
        "action: CharField",
        "details: TextField",
        "timestamp: DateTimeField",
    ])
    add_section(doc, "DealerCustomer", level=3)
    add_bullets(doc, [
        "dealer: ForeignKey to Company (dealer)",
        "customer_id: CharField",
        "company_name: CharField",
        "company_email: EmailField",
        "gst_number: CharField",
        "contact_person: CharField",
        "contact_number: CharField",
        "address: TextField",
        "address_2: TextField",
        "city: CharField",
        "district: CharField",
        "state: CharField",
        "country: CharField",
        "zip_code: CharField",
        "is_active: BooleanField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])

    add_section(doc, "configdetails", level=2)
    add_section(doc, "Device", level=3)
    add_bullets(doc, [
        "serial_number: CharField",
        "display_name: CharField",
        "device_type: CharField (TV, TOKEN_DISPENSER, KEYPAD, BROKER, LED)",
        "token_type: CharField (1_BUTTON to 4_BUTTON)",
        "company: ForeignKey to Company",
        "branch: ForeignKey to Branch",
        "device_registration_id: CharField",
        "mac_address: CharField",
        "batch: ForeignKey to Batch",
        "dealer_customer: ForeignKey to DealerCustomer",
        "device_model: CharField",
        "licence_status: CharField",
        "licence_active_to: DateField",
        "project_name: CharField",
        "apk_version: CharField",
        "product_type_id: IntegerField",
        "embedded_profile: ForeignKey to EmbeddedProfile",
        "is_active: BooleanField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "ProductionBatch", level=3)
    add_bullets(doc, [
        "batch_id: CharField",
        "device_type: CharField",
        "created_at: DateTimeField",
    ])
    add_section(doc, "ProductionSerialNumber", level=3)
    add_bullets(doc, [
        "batch: ForeignKey to ProductionBatch",
        "serial_number: CharField",
        "device_type: CharField",
        "is_registered: BooleanField",
    ])
    add_section(doc, "Mapping", level=3)
    add_bullets(doc, [
        "company: ForeignKey to Company",
        "branch: ForeignKey to Branch",
        "dealer_customer: ForeignKey to DealerCustomer",
        "token_dispenser: ForeignKey to Device",
        "tv: ForeignKey to Device",
        "keypad: ForeignKey to Device",
        "broker: ForeignKey to Device",
        "keypad_2: ForeignKey to Device",
        "keypad_3: ForeignKey to Device",
        "keypad_4: ForeignKey to Device",
        "led: ForeignKey to Device",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "ButtonMapping", level=3)
    add_bullets(doc, [
        "company: ForeignKey to Company",
        "branch: ForeignKey to Branch",
        "dealer_customer: ForeignKey to DealerCustomer",
        "source_device: ForeignKey to Device",
        "source_button: CharField",
        "target_device: ForeignKey to Device",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "DeviceConfig", level=3)
    add_bullets(doc, [
        "device: OneToOneField to Device",
        "config_json: JSONField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "TVConfig", level=3)
    add_bullets(doc, [
        "tv: OneToOneField to Device",
        "token_audio_file: FileField",
        "token_music_file: FileField",
        "audio_language: CharField",
        "save_audio_external: BooleanField",
        "enable_counter_announcement: BooleanField",
        "enable_token_announcement: BooleanField",
        "show_ads: BooleanField",
        "ad_interval: IntegerField",
        "orientation: CharField",
        "layout_type: CharField",
        "display_rows: IntegerField",
        "display_columns: IntegerField",
        "counter_text_color: CharField",
        "token_text_color: CharField",
        "token_font_size: IntegerField",
        "counter_font_size: IntegerField",
        "tokens_per_counter: IntegerField",
        "no_of_counters: IntegerField",
        "ad_placement: CharField",
        "current_token_color: CharField",
        "previous_token_color: CharField",
        "blink_current_token: BooleanField",
        "token_format: CharField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "TVAd", level=3)
    add_bullets(doc, [
        "tv_config: ForeignKey to TVConfig",
        "file: FileField",
        "created_at: DateTimeField",
    ])
    add_section(doc, "LedConfig", level=3)
    add_bullets(doc, [
        "led: OneToOneField to Device",
        "led_identifier_name: CharField",
        "voice_announcement: BooleanField",
        "counter_number: IntegerField",
        "token_calling: CharField",
        "counter_voice: BooleanField",
        "token_voice: CharField",
        "linked_calling_device: ManyToManyField to Device",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "TVCounter", level=3)
    add_bullets(doc, [
        "tv_config: ForeignKey to TVConfig",
        "counter_id: CharField",
        "counter_name: CharField",
        "counter_code: CharField",
        "row_span: IntegerField",
        "col_span: IntegerField",
        "counter_audio_file: FileField",
        "is_enabled: BooleanField",
    ])
    add_section(doc, "Counter", level=3)
    add_bullets(doc, [
        "branch: ForeignKey to Branch",
        "counter_name: CharField",
        "counter_number: IntegerField",
        "assigned_device: ForeignKey to Device",
    ])
    add_section(doc, "ConfigProfile", level=3)
    add_bullets(doc, [
        "name: CharField",
        "device: ForeignKey to Device",
        "company: ForeignKey to Company",
        "config_json: JSONField",
        "day: CharField",
        "start_time: TimeField",
        "end_time: TimeField",
        "is_active: BooleanField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "EmbeddedProfile", level=3)
    add_bullets(doc, [
        "name: CharField",
        "device_type: CharField",
        "company: ForeignKey to Company",
        "branch: ForeignKey to Branch",
        "config_json: JSONField",
        "day: JSONField",
        "start_time: TimeField",
        "end_time: TimeField",
        "is_active: BooleanField",
        "is_api_sourced: BooleanField",
        "external_id: CharField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])

    add_section(doc, "licenses", level=2)
    add_section(doc, "Batch", level=3)
    add_bullets(doc, [
        "name: CharField",
        "customer: ForeignKey to Company",
        "dealer_customer: ForeignKey to DealerCustomer",
        "max_tvs: IntegerField",
        "max_dispensers: IntegerField",
        "max_keypads: IntegerField",
        "max_brokers: IntegerField",
        "max_leds: IntegerField",
        "status: CharField (PENDING, ACTIVE, REJECTED)",
        "created_at: DateTimeField",
    ])
    add_section(doc, "License", level=3)
    add_bullets(doc, [
        "batch: ForeignKey to Batch",
        "device_type: CharField",
        "license_key: CharField",
        "status: CharField (ACTIVE, INACTIVE, REVOKED)",
        "device_uid: CharField",
        "activated_at: DateTimeField",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])
    add_section(doc, "BatchRequest", level=3)
    add_bullets(doc, [
        "requester: ForeignKey to Company",
        "dealer_customer: ForeignKey to DealerCustomer",
        "requester_type: CharField (CUSTOMER, DEALER, DEALER_CUSTOMER)",
        "requested_tvs: IntegerField",
        "requested_dispensers: IntegerField",
        "requested_keypads: IntegerField",
        "requested_brokers: IntegerField",
        "requested_leds: IntegerField",
        "reason: TextField",
        "status: CharField (PENDING, APPROVED, REJECTED)",
        "reviewed_by: ForeignKey to User",
        "reviewed_at: DateTimeField",
        "admin_notes: TextField",
        "approved_batch: OneToOneField to Batch",
        "created_at: DateTimeField",
        "updated_at: DateTimeField",
    ])


def add_user_roles(doc):
    add_section(doc, "6. User Roles and Responsibilities", level=1)
    doc.add_paragraph("Roles are defined in userdetails.User.Role and enforced through callq_core.permissions and view logic.")

    roles = [
        ("SUPER_ADMIN", "Full system access across all companies, dealers, devices, users, logs, and licenses."),
        ("ADMIN", "State-scoped administrator. Manages companies, devices, users, logs, and license requests within assigned states."),
        ("DEALER_ADMIN", "Manages dealer company, dealer customers, child companies, device assignments, and dealer-scoped license requests."),
        ("COMPANY_ADMIN", "Manages a single customer company, its branches, devices, mappings, and users."),
        ("BRANCH_ADMIN", "Manages a single branch and its devices and mappings."),
        ("DEALER_CUSTOMER", "Limited access to the specific dealer customer devices and batches."),
        ("PRODUCTION_ADMIN", "Manages production batches, serial uploads, and batch reporting."),
        ("EMPLOYEE", "System-level employee with visibility limited to assigned states."),
        ("COMPANY_EMPLOYEE", "Company-level employee focused on Android app usage and device access."),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Role"
    hdr[1].text = "Primary Responsibilities"

    for role, desc in roles:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = desc

    doc.add_paragraph("Platform access flags:")
    add_bullets(doc, [
        "is_web_user: Controls access to the web dashboard (enforced in CustomLoginView)",
        "is_android_user: Controls access to Android login APIs",
    ])


def add_auth_security(doc):
    add_section(doc, "7. Authentication and Security", level=1)
    add_bullets(doc, [
        "CustomLoginView enforces dealer license approval and expiry checks before allowing login for DEALER_ADMIN accounts.",
        "EmailOrUsernameBackend allows login with email or username.",
        "LicenseCheckMiddleware enforces company license expiry and warning banners after login.",
        "GlobalExceptionMiddleware captures exceptions and returns JSON for API calls or redirects to dashboard for web.",
        "Role checks are enforced using callq_core.permissions and view-level decorators.",
        "Password reset flows are available under /CallQ/accounts/password_reset/ and /CallQ/auth/password-reset/.",
    ])


def add_external_integrations(doc):
    add_section(doc, "8. External Integrations", level=1)
    doc.add_paragraph("License portal integration (callq_core.services.LicenseManagementService):")
    add_bullets(doc, [
        "ProductRegistration: Registers customer/company details in the external license portal.",
        "ProductAuthentication: Validates customer license and retrieves device counts.",
        "DeviceRegistration: Registers devices with the external portal.",
        "CheckDeviceStatus: Retrieves device license status, expiry, and version info.",
        "Requests and responses are logged in request_hits and actions logs.",
    ])


def add_api_surface(doc):
    add_section(doc, "9. API Surface", level=1)

    add_section(doc, "Company APIs", level=2)
    add_bullets(doc, [
        "GET/POST /CallQ/api/companies/",
        "POST /CallQ/api/customer_register/",
        "PATCH /CallQ/api/customer-registration/save/<pk>/",
        "POST /CallQ/api/customer_authentication/",
        "PATCH /CallQ/api/save_product_authentication/<pk>/",
        "GET /CallQ/api/company-branches/<company_id>/",
        "GET /CallQ/api/get-states/",
        "GET /CallQ/api/get-districts/",
    ])

    add_section(doc, "Config and Device APIs", level=2)
    add_bullets(doc, [
        "GET/POST /CallQ/config/api/devices/",
        "GET/POST /CallQ/config/api/mappings/",
        "GET/POST /CallQ/config/api/profiles/",
        "POST /CallQ/config/api/android/config",
        "POST /CallQ/config/api/device/<device_id>/check-status/",
        "POST /CallQ/config/api/device/<device_id>/authenticate/",
        "GET /CallQ/config/api/branch/<branch_id>/devices/",
        "GET /CallQ/config/api/branch/<branch_id>/mappings/",
        "POST /CallQ/config/api/mapping/button/save/",
        "POST /CallQ/config/api/mapping/button/<mapping_id>/delete/",
        "GET /CallQ/config/api/dealer-customer/<customer_id>/devices/",
        "GET /CallQ/config/api/dealer-customer/<customer_id>/mappings/",
        "GET /CallQ/config/api/available-serial-numbers/",
    ])

    add_section(doc, "User and Android APIs", level=2)
    add_bullets(doc, [
        "POST /CallQ/auth/api/android/login",
        "POST /CallQ/auth/api/android/Androidlogin",
        "GET /CallQ/auth/api/android/getDeviceByCustomer",
    ])

    add_section(doc, "License APIs", level=2)
    add_bullets(doc, [
        "GET/POST /CallQ/licenses/batches/",
        "GET/POST /CallQ/licenses/licenses/",
        "POST /CallQ/licenses/licenses/validate/",
    ])


def add_page_workflow(doc):
    add_section(doc, "10. Page Wise Workflow", level=1)

    page_groups = [
        {
            "name": "Authentication and Base Layout",
            "pages": [
                {
                    "name": "Login",
                    "url": "/CallQ/login/",
                    "template": "templates/login.html",
                    "roles": "Unauthenticated users",
                    "view": "companydetails.views.CustomLoginView",
                    "models": "User, Company",
                    "workflow": [
                        "User submits username or email and password.",
                        "CustomLoginView authenticates and checks is_web_user.",
                        "Dealer Admins are blocked if license status is not approved or expired.",
                        "User is redirected to dashboard on success.",
                    ],
                    "actions": ["Login", "Error messaging on failure"],
                },
                {
                    "name": "Password Reset (Account Flow)",
                    "url": "/CallQ/accounts/password_reset/ and related /done, /confirm, /complete",
                    "template": "templates/registration/*.html",
                    "roles": "All users",
                    "view": "CustomPasswordResetView and Django auth views",
                    "models": "User",
                    "workflow": [
                        "User submits email address.",
                        "System sends password reset email if address exists.",
                        "User follows link to set new password.",
                    ],
                    "actions": ["Submit email", "Reset password", "View success messages"],
                },
            ],
        },
        {
            "name": "Dashboard and Activity Logs",
            "pages": [
                {
                    "name": "Dashboard",
                    "url": "/CallQ/ or /CallQ/dashboard/",
                    "template": "companydetails/dashboard.html",
                    "roles": "All authenticated roles",
                    "view": "companydetails.views.dashboard",
                    "models": "Company, Device, License, ActivityLog",
                    "workflow": [
                        "User lands on dashboard after login.",
                        "Counts and stats are scoped by role and state.",
                        "Recent activity is shown for permitted scope.",
                    ],
                    "actions": ["View counts and alerts", "Navigate to modules"],
                },
                {
                    "name": "Activity Logs",
                    "url": "/CallQ/logs/",
                    "template": "companydetails/activity_log_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "companydetails.views.activity_log_list",
                    "models": "ActivityLog",
                    "workflow": [
                        "User opens logs list scoped by role.",
                        "Filters by pagination and optional export format.",
                        "Exports available as CSV, PDF, or DOCX.",
                    ],
                    "actions": ["Search and paginate", "Export logs"],
                },
            ],
        },
        {
            "name": "Customer, Dealer, and Branch Management",
            "pages": [
                {
                    "name": "Customer List",
                    "url": "/CallQ/customer-list/",
                    "template": "companydetails/customer_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin",
                    "view": "companydetails.views.customer_list",
                    "models": "Company, Device",
                    "workflow": [
                        "User views customers scoped by role.",
                        "License status is calculated and displayed.",
                        "Actions include authenticate, delete, and export.",
                    ],
                    "actions": ["View license status", "Authenticate customer", "Delete customer"],
                },
                {
                    "name": "Dealer List",
                    "url": "/CallQ/dealers/",
                    "template": "companydetails/dealer_list.html",
                    "roles": "Super Admin, Admin",
                    "view": "companydetails.views.dealer_list",
                    "models": "Company",
                    "workflow": [
                        "User lists dealer companies.",
                        "Pagination and search applied.",
                    ],
                    "actions": ["View dealer records"],
                },
                {
                    "name": "Customer Registration",
                    "url": "/CallQ/customer-register/",
                    "template": "companydetails/customer_register.html",
                    "roles": "Dealer Admin, Super Admin, Admin",
                    "view": "companydetails.views.customer_registration",
                    "models": "Company, DealerCustomer",
                    "workflow": [
                        "User fills registration form.",
                        "Dealer Admin creates dealer-customer contact or child company as needed.",
                        "External license registration and authentication flows can be triggered.",
                    ],
                    "actions": ["Create customer", "Register and authenticate with license portal"],
                },
                {
                    "name": "Branch List",
                    "url": "/CallQ/branches/",
                    "template": "companydetails/branch_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "companydetails.views.branch_list",
                    "models": "Branch",
                    "workflow": [
                        "User views branches within their scope.",
                        "Branches can be created or edited.",
                    ],
                    "actions": ["View branches", "Navigate to create or edit"],
                },
                {
                    "name": "Branch Form",
                    "url": "/CallQ/branches/create/ or /CallQ/branches/<pk>/edit/",
                    "template": "companydetails/branch_form.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin",
                    "view": "companydetails.views.branch_create / branch_edit",
                    "models": "Branch",
                    "workflow": [
                        "User fills branch details and submits.",
                        "Branch is created or updated for the selected company.",
                    ],
                    "actions": ["Create branch", "Edit branch"],
                },
                {
                    "name": "Dealer Customers",
                    "url": "/CallQ/dealer-customers/",
                    "template": "companydetails/dealer_customer_list.html",
                    "roles": "Dealer Admin",
                    "view": "companydetails.dealer_customer_views.dealer_customer_list",
                    "models": "DealerCustomer",
                    "workflow": [
                        "Dealer Admin views their customer contacts.",
                        "Search and status filters applied.",
                    ],
                    "actions": ["Create, edit, or delete dealer customer"],
                },
                {
                    "name": "Dealer Customer Form",
                    "url": "/CallQ/dealer-customers/create/ or /CallQ/dealer-customers/<pk>/edit/",
                    "template": "companydetails/dealer_customer_form.html",
                    "roles": "Dealer Admin",
                    "view": "companydetails.dealer_customer_views.dealer_customer_create / dealer_customer_edit",
                    "models": "DealerCustomer",
                    "workflow": [
                        "Dealer Admin submits customer contact details.",
                        "Customer contact record is created or updated.",
                    ],
                    "actions": ["Create or update dealer customer"],
                },
                {
                    "name": "Location Management",
                    "url": "/CallQ/location-management/",
                    "template": "companydetails/location_management.html",
                    "roles": "Super Admin, Admin",
                    "view": "companydetails.views.location_management",
                    "models": "Country, State, District",
                    "workflow": [
                        "Admin manages country, state, and district records.",
                        "Used for customer and branch location selection.",
                    ],
                    "actions": ["Add or update location records"],
                },
                {
                    "name": "Device Report",
                    "url": "/CallQ/reports/devices/",
                    "template": "companydetails/device_report.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "companydetails.views.device_report",
                    "models": "Device, Company, Branch",
                    "workflow": [
                        "User views device inventory report.",
                        "Filters applied by role scope.",
                    ],
                    "actions": ["View report", "Export if enabled"],
                },
            ],
        },
        {
            "name": "Device and Configuration Management",
            "pages": [
                {
                    "name": "Device List",
                    "url": "/CallQ/config/devices/",
                    "template": "configdetails/device_list.html",
                    "roles": "All roles with device access",
                    "view": "configdetails.views.device_list",
                    "models": "Device, Branch",
                    "workflow": [
                        "User views devices scoped by role and state.",
                        "Pagination applied and actions available per device.",
                    ],
                    "actions": ["Configure device", "Assign branch", "Check status", "Delete device"],
                },
                {
                    "name": "Device Registration",
                    "url": "/CallQ/config/device/register/",
                    "template": "configdetails/device_register.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin, Production Admin",
                    "view": "configdetails.views.device_register",
                    "models": "Device, Company, Branch",
                    "workflow": [
                        "User selects company and branch.",
                        "Devices are registered with external license portal.",
                        "Device status and license info are stored locally.",
                    ],
                    "actions": ["Register device", "Bulk register via form"],
                },
                {
                    "name": "Device Configuration",
                    "url": "/CallQ/config/device/<device_id>/config/",
                    "template": "configdetails/device_config.html",
                    "roles": "Company Admin, Dealer Admin, Super Admin, Admin",
                    "view": "configdetails.views.device_config",
                    "models": "Device, DeviceConfig, TVConfig, TVCounter, TVAd, LedConfig",
                    "workflow": [
                        "User opens device config page.",
                        "Configuration is stored as JSON and device-specific records.",
                        "Ads and counters are managed for TVs; LED settings saved for LED devices.",
                    ],
                    "actions": ["Save configuration", "Upload ads", "Manage counters"],
                },
                {
                    "name": "TV Config (Legacy Page)",
                    "url": "/CallQ/config/tv/<tv_id>/config/",
                    "template": "configdetails/tv_config.html",
                    "roles": "Company Admin, Dealer Admin, Super Admin, Admin",
                    "view": "configdetails.views.tv_config",
                    "models": "TVConfig",
                    "workflow": [
                        "User edits basic TV settings.",
                        "TVConfig record is updated.",
                    ],
                    "actions": ["Update orientation and ads"],
                },
                {
                    "name": "Mapping",
                    "url": "/CallQ/config/mapping/",
                    "template": "configdetails/mapping.html",
                    "roles": "Company Admin, Dealer Admin, Branch Admin, Super Admin, Admin",
                    "view": "configdetails.views.mapping_view",
                    "models": "Mapping, ButtonMapping, Device",
                    "workflow": [
                        "User views existing mappings.",
                        "Selects devices for token, keypad, broker, TV, and LED.",
                        "Mapping and button-level mapping are saved.",
                    ],
                    "actions": ["Create or update mappings", "View mapping coverage"],
                },
                {
                    "name": "Mapping List",
                    "url": "/CallQ/config/mapping/list/",
                    "template": "configdetails/mapping_list.html",
                    "roles": "Company Admin, Dealer Admin, Branch Admin, Super Admin, Admin",
                    "view": "configdetails.views.mapping_list_view",
                    "models": "Mapping",
                    "workflow": [
                        "User views mappings in list format.",
                        "Pagination and filters are applied.",
                    ],
                    "actions": ["Review mappings"],
                },
                {
                    "name": "Assign Devices to Branch",
                    "url": "/CallQ/config/devices/assign-branch/",
                    "template": "configdetails/assign_devices_to_branch.html",
                    "roles": "Company Admin, Super Admin, Admin",
                    "view": "configdetails.views.assign_devices_to_branch",
                    "models": "Device, Branch",
                    "workflow": [
                        "User selects a branch and devices.",
                        "Devices are updated with the selected branch.",
                    ],
                    "actions": ["Bulk assign devices to branch"],
                },
                {
                    "name": "Map Devices to Dealer Customer",
                    "url": "/CallQ/config/mapping/map-device/",
                    "template": "configdetails/map_device_to_customer.html",
                    "roles": "Dealer Admin",
                    "view": "configdetails.views.map_device_to_customer",
                    "models": "Device, DealerCustomer",
                    "workflow": [
                        "Dealer Admin selects dealer customer and devices.",
                        "Devices are assigned to the dealer customer and branch cleared.",
                    ],
                    "actions": ["Assign and unassign devices"],
                },
                {
                    "name": "Device Approvals",
                    "url": "/CallQ/config/approvals/",
                    "template": "configdetails/device_approval_list.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin, Branch Admin",
                    "view": "configdetails.views.device_approval_list",
                    "models": "Device",
                    "workflow": [
                        "User filters devices by licence_status (Pending, Active, Rejected).",
                        "Approval actions update device status and company assignment where allowed.",
                    ],
                    "actions": ["Approve or reject device requests"],
                },
                {
                    "name": "Embedded Profiles",
                    "url": "/CallQ/config/embedded-profiles/",
                    "template": "configdetails/embedded_profile_list.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin",
                    "view": "configdetails.views.embedded_profile_list",
                    "models": "EmbeddedProfile",
                    "workflow": [
                        "User views reusable embedded profiles.",
                        "Profiles can be created, edited, deleted, or allocated to devices.",
                    ],
                    "actions": ["Create/edit profile", "Allocate profile"],
                },
                {
                    "name": "Embedded Profile Form",
                    "url": "/CallQ/config/embedded-profiles/create/ and /edit/",
                    "template": "configdetails/embedded_profile_form.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin",
                    "view": "configdetails.views.embedded_profile_create / embedded_profile_edit",
                    "models": "EmbeddedProfile",
                    "workflow": [
                        "User enters profile name, device type, and JSON config.",
                        "Optional schedule is applied.",
                    ],
                    "actions": ["Save profile"],
                },
                {
                    "name": "Embedded Profile Delete",
                    "url": "/CallQ/config/embedded-profiles/<pk>/delete/",
                    "template": "configdetails/embedded_profile_confirm_delete.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin",
                    "view": "configdetails.views.embedded_profile_delete",
                    "models": "EmbeddedProfile",
                    "workflow": [
                        "User confirms delete.",
                        "Profile is removed if allowed.",
                    ],
                    "actions": ["Delete profile"],
                },
                {
                    "name": "Embedded Profile Allocation",
                    "url": "/CallQ/config/embedded-profiles/<pk>/allocate/",
                    "template": "configdetails/embedded_profile_allocate.html",
                    "roles": "Super Admin, Admin, Company Admin, Dealer Admin",
                    "view": "configdetails.views.embedded_profile_allocate",
                    "models": "EmbeddedProfile, Device",
                    "workflow": [
                        "User selects devices to allocate a profile.",
                        "Devices reference the embedded profile.",
                    ],
                    "actions": ["Allocate profile to devices"],
                },
                {
                    "name": "Production Batch Upload",
                    "url": "/CallQ/config/production-batch/upload/",
                    "template": "configdetails/production_batch_upload.html",
                    "roles": "Super Admin, Admin, Production Admin",
                    "view": "configdetails.views.production_batch_upload",
                    "models": "ProductionBatch, ProductionSerialNumber",
                    "workflow": [
                        "User creates or selects a production batch.",
                        "Serial numbers are uploaded via CSV or XLSX.",
                        "System stores serial numbers and device types.",
                    ],
                    "actions": ["Upload serial numbers", "Review batch list"],
                },
                {
                    "name": "Production Report",
                    "url": "/CallQ/config/production/report/",
                    "template": "configdetails/production_report.html",
                    "roles": "Super Admin, Admin, Production Admin",
                    "view": "configdetails.views.production_report_view",
                    "models": "ProductionBatch, ProductionSerialNumber",
                    "workflow": [
                        "User views past production batches.",
                        "Summary counts are computed by device type.",
                    ],
                    "actions": ["Download batch reports"],
                },
            ],
        },
        {
            "name": "License and Batch Management",
            "pages": [
                {
                    "name": "Batch List",
                    "url": "/CallQ/licenses/list/",
                    "template": "licenses/batch_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin, Dealer Customer",
                    "view": "licenses.views.batch_page",
                    "models": "Batch, Device",
                    "workflow": [
                        "User views batches scoped by role.",
                        "Usage statistics and remaining counts are computed.",
                    ],
                    "actions": ["View batch usage", "Download batch report"],
                },
                {
                    "name": "Purchase Batch",
                    "url": "/CallQ/licenses/purchase/",
                    "template": "licenses/purchase_batch.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "licenses.views.purchase_batch",
                    "models": "Batch",
                    "workflow": [
                        "User enters requested counts.",
                        "Batch is created with PENDING status.",
                        "Admin approval is required.",
                    ],
                    "actions": ["Submit purchase request"],
                },
                {
                    "name": "Request Batch",
                    "url": "/CallQ/licenses/request/",
                    "template": "licenses/request_batch.html",
                    "roles": "Company Admin, Dealer Admin, Dealer Customer",
                    "view": "licenses.views.request_batch",
                    "models": "BatchRequest",
                    "workflow": [
                        "User submits request for additional device counts.",
                        "Request is marked PENDING for approval.",
                    ],
                    "actions": ["Submit request"],
                },
                {
                    "name": "Batch Requests List",
                    "url": "/CallQ/licenses/requests/",
                    "template": "licenses/batch_requests_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin",
                    "view": "licenses.views.batch_requests_list",
                    "models": "BatchRequest",
                    "workflow": [
                        "Approvers view pending, approved, or rejected requests.",
                        "Requests can be approved or rejected with notes.",
                    ],
                    "actions": ["Approve or reject requests"],
                },
            ],
        },
        {
            "name": "User Management",
            "pages": [
                {
                    "name": "User List",
                    "url": "/CallQ/auth/users/",
                    "template": "userdetails/user_list.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin, Branch Admin",
                    "view": "userdetails.views.user_list",
                    "models": "User",
                    "workflow": [
                        "User sees list of users scoped by role and branch/company.",
                        "Filters by search, role, and branch.",
                    ],
                    "actions": ["Create, edit, or delete users"],
                },
                {
                    "name": "User Form",
                    "url": "/CallQ/auth/users/create/ or /CallQ/auth/users/<pk>/edit/",
                    "template": "userdetails/user_form.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "userdetails.views.user_create / user_edit",
                    "models": "User, Company, Branch",
                    "workflow": [
                        "User enters email, role, and branch/company relationships.",
                        "System assigns role-specific relations and permissions.",
                    ],
                    "actions": ["Create user", "Update user"],
                },
                {
                    "name": "Profile",
                    "url": "/CallQ/auth/profile/",
                    "template": "userdetails/profile.html",
                    "roles": "All authenticated users",
                    "view": "userdetails.views.profile",
                    "models": "User, Company",
                    "workflow": [
                        "User updates email and profile details.",
                        "Company Admin and Dealer Admin may update company profile fields.",
                    ],
                    "actions": ["Update profile", "Trigger password reset"],
                },
                {
                    "name": "App Login History",
                    "url": "/CallQ/auth/app/login-history/",
                    "template": "userdetails/app_login_history.html",
                    "roles": "Super Admin, Admin, Dealer Admin, Company Admin",
                    "view": "userdetails.views.app_login_history",
                    "models": "AppLoginHistory",
                    "workflow": [
                        "User views Android login history scoped by role.",
                        "Search by username, email, MAC address, or version.",
                    ],
                    "actions": ["Review Android login activity"],
                },
            ],
        },
    ]

    for group in page_groups:
        add_section(doc, group["name"], level=2)
        for page in group["pages"]:
            add_section(doc, page["name"], level=3)
            add_kv(doc, "URL", page["url"])
            add_kv(doc, "Template", page["template"])
            add_kv(doc, "Primary Roles", page["roles"])
            add_kv(doc, "Backend View", page["view"])
            if page.get("models"):
                add_kv(doc, "Models", page["models"])
            add_kv(doc, "Workflow", "")
            add_numbers(doc, page["workflow"])
            add_kv(doc, "Key Actions", "")
            add_bullets(doc, page["actions"])



def add_reports(doc):
    add_section(doc, "11. Reporting and Exports", level=1)
    add_bullets(doc, [
        "Activity Logs: CSV, PDF, and DOCX export options.",
        "License Batch Reports: PDF download with device usage.",
        "Production Batch Reports: CSV, PDF, and DOCX export formats.",
        "Device Report: Web report filtered by role and scope.",
    ])


def add_logging(doc):
    add_section(doc, "12. Logging and Observability", level=1)
    add_bullets(doc, [
        "Actions log: CallQ_logs/<year>/<month>/<day>/actions.log",
        "Request hits log: CallQ_logs/<year>/<month>/<day>/request_hits.log",
        "GlobalExceptionMiddleware records errors and returns safe responses.",
    ])


def add_operations(doc):
    add_section(doc, "13. Operations and Maintenance", level=1)
    add_bullets(doc, [
        "Run migrations after model changes: python manage.py makemigrations and python manage.py migrate",
        "Create superuser accounts: python manage.py createsuperuser",
        "Collect static files for production: python manage.py collectstatic",
        "Back up database and media files regularly, especially before batch operations.",
        "Monitor license expiry warnings and external portal availability.",
    ])


def generate_document():
    doc = Document()
    add_title(doc)
    add_toc(doc)
    add_overview(doc)
    add_tech_stack(doc)
    add_architecture(doc)
    add_configuration(doc)
    add_data_model(doc)
    add_user_roles(doc)
    add_auth_security(doc)
    add_external_integrations(doc)
    add_api_surface(doc)
    add_page_workflow(doc)
    add_reports(doc)
    add_logging(doc)
    add_operations(doc)

    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_document()
