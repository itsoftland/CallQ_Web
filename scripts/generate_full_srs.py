
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def setup_document_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

def add_title_page(doc):
    doc.add_heading('CallQ', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Token Management System\nSoftware Requirements Specification (SRS)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_page_break()

def create_srs_docx():
    doc = Document()
    setup_document_styles(doc)
    add_title_page(doc)
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph('The purpose of this document is to define the software requirements for the CallQ system, a comprehensive Token Management and Queue System designed to streamline customer flow.')
    
    doc.add_heading('1.2 Scope', level=2)
    p = doc.add_paragraph()
    p.add_run('The system encompasses:').bold = True
    items = [
        'Web Administration Portal for configuration and management.',
        'Android Application for mobile access.',
        'Embedded Device Integration (IoT) for hardware control.',
        'RESTful and Raw Socket APIs for communication.'
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.3 Definitions & Acronyms', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    
    definitions = [
        ('SRS', 'Software Requirements Specification'),
        ('Super Admin', 'Top-level system administrator with global access.'),
        ('Dealer', 'Intermediary entity managing multiple customer companies.'),
        ('Keypad', 'Input device used by staff to call the next token.'),
        ('Dispenser', 'Device that prints or issues tokens to customers.'),
        ('Mapping', 'Logical association between input and output devices.')
    ]
    
    for term, definition in definitions:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition

    # 2. System Overview
    doc.add_heading('2. System Overview', level=1)
    doc.add_heading('2.1 User Roles', level=2)
    roles = [
        'Super Admin: Full system control.',
        'Admin: System maintenance and oversight.',
        'Production Admin: Manages hardware inventory (Batches/Serial Numbers).',
        'Dealer Admin: Manages assigned customer companies.',
        'Company Admin: Manages a specific company instance.',
        'Branch Admin: Manages a specific physical location.',
        'Employee: Standard operational user.'
    ]
    for role in roles:
        doc.add_paragraph(role, style='List Bullet')

    doc.add_heading('2.2 Technology Stack', level=2)
    doc.add_paragraph('Backend: Python Django\nDatabase: SQLite / PostgreSQL\nFrontend: Django Templates (HTML/JS)\nMobile: Android (Kotlin)\nEmbedded: C/C++')

    # 3. Functional Requirements
    doc.add_heading('3. Functional Requirements', level=1)
    
    doc.add_heading('3.1 Company & Location Management', level=2)
    doc.add_paragraph('The system maintains a hierarchical structure of graphical locations (Country > State > District) and organizational entities (Company > Branch). It supports Single and Multiple branch configurations.')
    
    doc.add_heading('3.2 Device Management', level=2)
    doc.add_paragraph('Core module for tracking hardware lifecycle:')
    device_features = [
        'Inventory tracking via Serial Number and MAC Address.',
        'Device Types: TV, Token Dispenser, Keypad, Broker, LED.',
        'License validity tracking with auto-expiry calculation.',
        'Production Batch management for manufacturing control.'
    ]
    for feat in device_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_heading('3.3 Configuration & Mapping', level=2)
    doc.add_paragraph('Allows linking of input devices to output displays. Features include:')
    config_features = [
        'TV Config: Layouts, Ad scheduling, Audio files, Colors.',
        'Device Mapping: Linking Keypads to specific Counters and Displays.',
        'Profiles: Time-based configuration profiles (Morning/Evening modes).'
    ]
    for feat in config_features:
        doc.add_paragraph(feat, style='List Bullet')

    # 4. API Specification Overview
    doc.add_heading('4. API Specification Overview', level=1)
    doc.add_paragraph('The system exposes endpoints for embedded devices to fetch configuration booting up.')
    p = doc.add_paragraph()
    p.add_run('Endpoint: ').bold = True
    p.add_run('/CallQ/config/api/embedded/get-config/')
    
    p = doc.add_paragraph()
    p.add_run('Format: ').bold = True
    p.add_run('Custom comma-separated raw string (e.g., $1,SETTINGS,...)')

    # Save
    output_path = '/home/silpc-064/.gemini/antigravity/brain/00ecc9a8-9deb-4601-8b45-67d5da539207/CallQ_SRS.docx'
    doc.save(output_path)
    print(f"SRS Document saved to: {output_path}")

if __name__ == "__main__":
    create_srs_docx()
